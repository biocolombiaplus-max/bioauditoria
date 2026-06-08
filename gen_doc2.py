
# -*- coding: utf-8 -*-
"""Genera PRO-PP-002 Proceso de Procedimientos Estéticos No Invasivos"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = "/home/user/bioauditoria/documentos_habilitacion/CARPETA_5_PROCESOS_PRIORITARIOS/PRO-PP-002_Proceso_Procedimientos_Esteticos_No_Invasivos.docx"

AZUL_OSCURO = RGBColor(0, 51, 102)
AZUL_MEDIO = RGBColor(21, 101, 192)
ROJO = RGBColor(180, 0, 0)
BLANCO = RGBColor(255, 255, 255)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_header_footer(doc, code, title):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = f"{code} | {title}"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.runs[0]
    run.font.size = Pt(9)
    run.font.color.rgb = AZUL_OSCURO
    run.font.name = 'Calibri'
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)

def add_page_number(paragraph):
    paragraph.clear()
    run = paragraph.add_run("Versión 2.0 | Junio 2025    Página ")
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run2 = paragraph.add_run()
    run2.font.size = Pt(9)
    run2.font.name = 'Calibri'
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Calibri'
    run.font.color.rgb = AZUL_OSCURO

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    run.font.color.rgb = AZUL_MEDIO

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = AZUL_OSCURO

def aviso(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run("⚠ AVISO LEGAL: " + text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = ROJO
    set_cell_bg_para(p, 'FFF3CD')

def set_cell_bg_para(para, color_hex):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

def table_header_row(table, headers, bg='003366'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        set_cell_bg(cell, bg)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = BLANCO
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table_row(table, values, shaded=False):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        if shaded:
            set_cell_bg(cell, 'F2F2F2')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

add_header_footer(doc, "PRO-PP-002", "Proceso de Procedimientos Estéticos No Invasivos")

# PORTADA
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[LOGO DEL CONSULTORIO]")
run.font.size = Pt(14); run.font.name = 'Calibri'; run.font.color.rgb = AZUL_OSCURO; run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PROCESO DE PROCEDIMIENTOS ESTÉTICOS NO INVASIVOS")
run.font.size = Pt(18); run.font.name = 'Calibri'; run.font.color.rgb = AZUL_OSCURO; run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[NOMBRE DEL CONSULTORIO]")
run.font.size = Pt(14); run.font.name = 'Calibri'; run.font.color.rgb = AZUL_MEDIO; run.bold = True

doc.add_paragraph()
aviso(doc, "Los procedimientos estéticos descritos en este documento son ACTOS MÉDICOS que por ley colombiana solo pueden ser realizados por un MÉDICO LEGALMENTE HABILITADO con tarjeta profesional vigente, en un establecimiento habilitado por la Secretaría de Salud. El ejercicio de estos procedimientos por personas sin título médico o en establecimientos no habilitados constituye ejercicio ilegal de la medicina (Ley 23 de 1981, Art. 1), delito tipificado en el Código Penal colombiano. Este documento es de uso exclusivo del personal médico del [NOMBRE DEL CONSULTORIO].")

t = doc.add_table(rows=7, cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
portada_data = [
    ("Código:", "PRO-PP-002"), ("Versión:", "2.0"), ("Fecha de Aprobación:", "Junio 2025"),
    ("Elaborado por:", "Dra. [NOMBRE DE LA MÉDICA]"), ("Revisado por:", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("Aprobado por:", "Dra. [NOMBRE DE LA MÉDICA] - Directora Médica"), ("NIT:", "[NIT DEL CONSULTORIO]"),
]
for i, (k, v) in enumerate(portada_data):
    row = t.rows[i]; row.cells[0].text = k; row.cells[1].text = v
    set_cell_bg(row.cells[0], 'E8EEF4')
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11); run.font.name = 'Calibri'
                run.bold = (cell == row.cells[0])

doc.add_page_break()

# CONTROL DE VERSIONES
h1(doc, "CONTROL DE VERSIONES")
t_ver = doc.add_table(rows=3, cols=5)
t_ver.style = 'Table Grid'
table_header_row(t_ver, ["Versión", "Fecha", "Descripción del Cambio", "Elaboró", "Aprobó"])
add_table_row(t_ver, ["1.0", "Enero 2024", "Versión inicial. Inclusión de toxina botulínica tipo A y ácido hialurónico.", "Dra. [NOMBRE DE LA MÉDICA]", "Dra. [NOMBRE DE LA MÉDICA]"])
add_table_row(t_ver, ["2.0", "Junio 2025", "Ampliación a mesoterapia, PRP y peelings químicos. Incorporación de protocolo de emergencia vascular. Actualización normativa según Res. 3100/2019.", "Dra. [NOMBRE DE LA MÉDICA]", "Dra. [NOMBRE DE LA MÉDICA]"], shaded=True)

doc.add_paragraph()

# 1. INTRODUCCIÓN
h1(doc, "1. INTRODUCCIÓN")
body(doc, "La medicina estética constituye una rama especializada de la medicina clínica que aplica conocimientos de medicina interna, dermatología, cirugía menor y farmacología para el diagnóstico y tratamiento de condiciones estéticas que afectan la apariencia física de los pacientes, sin comprometer su salud general. En las últimas dos décadas, el campo de la medicina estética no invasiva ha experimentado un crecimiento exponencial a nivel mundial y en Colombia, impulsado por el desarrollo de nuevas técnicas, el refinamiento de los productos inyectables, el aumento de la demanda de procedimientos con tiempos de recuperación mínimos y la mayor conciencia de los pacientes sobre las opciones disponibles.")
body(doc, "En Colombia, la regulación de la medicina estética ha sido objeto de atención normativa creciente, especialmente tras la promulgación de la Ley 711 de 2001 sobre el ejercicio de la cosmetología y la estética integral, y las sucesivas resoluciones del Ministerio de Salud que han delimitado con precisión creciente cuáles procedimientos pueden ser realizados por cosmetólogos y cuáles constituyen ACTOS MÉDICOS que solo pueden ser ejecutados por médicos con tarjeta profesional vigente. La Resolución 3100 de 2019 establece los estándares de habilitación para el servicio de 'medicina estética' como servicio habilitado diferente a la consulta de medicina general, aunque en centros de primer nivel los procedimientos estéticos no invasivos realizados por médico general pueden estar incluidos dentro del portafolio habilitado del consultorio.")
body(doc, "El [NOMBRE DEL CONSULTORIO], bajo la dirección médica de la Dra. [NOMBRE DE LA MÉDICA], ofrece procedimientos estéticos no invasivos en el marco del portafolio de servicios habilitados, con el más alto estándar de seguridad, calidad y ética médica. Todos los procedimientos son realizados exclusivamente por la Dra. [NOMBRE DE LA MÉDICA], médica con formación específica y certificada en medicina estética. El presente proceso estandariza la evaluación previa, la técnica de ejecución, el manejo de complicaciones y el seguimiento post-procedimiento para cada uno de los procedimientos ofrecidos.")
body(doc, "La seguridad del paciente es el principio rector de este proceso. En el campo de la medicina estética, la seguridad implica: una evaluación médica completa y honesta del paciente antes de cada procedimiento para determinar indicaciones y contraindicaciones; un consentimiento informado exhaustivo que garantice la autonomía del paciente con información real sobre beneficios, riesgos, alternativas y resultados esperados y no esperados; una técnica depurada y basada en el conocimiento de la anatomía relevante, especialmente en procedimientos inyectables donde el riesgo vascular y de necrosis tisular es una posibilidad real; disponibilidad de protocolos de emergencia y materiales para el manejo de complicaciones graves; y un seguimiento post-procedimiento estructurado para la detección temprana de complicaciones.")

# 2. MARCO LEGAL CON AVISO PROMINENTE
h1(doc, "2. MARCO LEGAL, NORMATIVO Y ÉTICO")
aviso(doc, "ADVERTENCIA LEGAL FUNDAMENTAL: Los procedimientos descritos en este documento son ACTOS MÉDICOS exclusivos de médicos habilitados. La aplicación de toxina botulínica, ácido hialurónico u otros rellenos dérmicos por personas sin título médico constituye ejercicio ilegal de la medicina según la Ley 23 de 1981, y puede generar consecuencias penales, civiles y disciplinarias. El médico que delega o permite la realización de estos procedimientos a personal no médico es igualmente responsable por omisión. Ninguna franquicia, spa, centro de estética o salón de belleza, por más que utilice el término 'médico' en su nombre comercial, está autorizado para realizar estas intervenciones sin el médico presente y ejecutando el procedimiento.")

body(doc, "El marco legal específico aplicable a los procedimientos estéticos no invasivos en Colombia incluye las siguientes normas, que la Dra. [NOMBRE DE LA MÉDICA] conoce y cumple en su totalidad:")
t_legal = doc.add_table(rows=1, cols=3)
t_legal.style = 'Table Grid'
table_header_row(t_legal, ["Norma", "Descripción", "Aplicabilidad"])
normas_est = [
    ("Ley 711 de 2001", "Regula el ejercicio de la cosmetología y estética integral. Delimita claramente que los cosmetólogos solo pueden realizar procedimientos superficiales sobre piel sana, y prohíbe expresamente la realización de actos médicos.", "Delimita el campo de acción del médico vs. cosmetólogo. Todo procedimiento inyectable, invasivo o que requiera diagnóstico médico es exclusivo del médico."),
    ("Resolución 2378 de 2008 del INVIMA", "Regula la publicidad de productos y servicios de salud. Prohíbe la publicidad engañosa de procedimientos estéticos médicos.", "Obliga al consultorio a no hacer publicidad que prometa resultados garantizados o minimice riesgos de los procedimientos estéticos."),
    ("Resolución 3100 de 2019 - MINSALUD", "Estándares de habilitación para procedimientos estéticos. Incluye requisitos específicos de talento humano (médico), dotación y procesos para procedimientos estéticos habilitados.", "Norma principal que rige los estándares del servicio de procedimientos estéticos en el consultorio."),
    ("Ley 23 de 1981 y Decreto 3380/1981", "Código de Ética Médica. Art. 16: el médico no puede garantizar resultados. Art. 15: consentimiento informado. Art. 10: historia clínica obligatoria.", "El médico no puede garantizar resultados estéticos ni prometer cambios específicos. Toda intervención requiere HC y CI."),
    ("Decreto 4725 de 2005 y sus modificaciones", "Regulación de dispositivos médicos en Colombia. Los rellenos dérmicos y la toxina botulínica son dispositivos médicos/medicamentos sujetos a registro INVIMA.", "Solo se pueden utilizar en el consultorio productos con registro sanitario INVIMA vigente. El uso de productos sin registro es ilegal y constituye responsabilidad del médico."),
    ("Circular 01 de 2019 - INVIMA", "Alerta sanitaria sobre productos de relleno dérmico, toxina botulínica y biopolímeros sin registro INVIMA. Restricciones de uso y comercialización.", "Prohíbe el uso de biopolímeros (polimetilmetacrilato, silicona líquida inyectable) para rellenos dérmicos por su alta tasa de complicaciones tardías irreversibles."),
    ("Ley 1164 de 2007 - Talento Humano en Salud", "Regula las ocupaciones del área de la salud. Establece que el médico general puede realizar procedimientos estéticos no invasivos si cuenta con la formación específica.", "La Dra. [NOMBRE DE LA MÉDICA] debe acreditar formación certificada en los procedimientos estéticos que realiza."),
    ("Resolución 8430 de 1993 (parcialmente)", "Normas para investigación en salud. Cuando se utilizan productos en indicaciones diferentes a las aprobadas (uso off-label), se deben seguir principios éticos de investigación.", "El uso off-label de toxina botulínica o ácido hialurónico (en indicaciones no aprobadas por INVIMA) requiere consentimiento informado específico que lo mencione."),
    ("Ley 1581 de 2012 - Habeas Data", "Protección de datos personales. Los datos de salud y las fotografías de los pacientes son datos sensibles con máxima protección.", "Las fotografías clínicas del paciente son datos sensibles. Requieren consentimiento escrito específico para su uso y almacenamiento. No pueden publicarse en redes sociales sin consentimiento explícito adicional."),
    ("Resolución 13437 de 1991 - Derechos del Paciente", "Derechos del paciente incluyendo recibir información suficiente para la toma de decisiones. En estética: derecho a conocer riesgos reales, alternativas y a no ser sometido a presión comercial.", "El médico debe presentar la información de manera neutral, sin presión de venta, dando tiempo al paciente para reflexionar antes de firmar el consentimiento."),
]
for i, row_data in enumerate(normas_est):
    add_table_row(t_legal, list(row_data), shaded=(i % 2 == 1))

doc.add_paragraph()

# 3. OBJETIVOS
h1(doc, "3. OBJETIVOS")
h2(doc, "3.1 Objetivo General")
body(doc, "Estandarizar el proceso de atención para la realización de procedimientos estéticos no invasivos en el [NOMBRE DEL CONSULTORIO], garantizando la seguridad del paciente, la excelencia técnica, el cumplimiento normativo y los más altos estándares éticos de la práctica médica, desde la evaluación previa hasta el seguimiento post-procedimiento.")
h2(doc, "3.2 Objetivos Específicos")
bullet(doc, "Establecer los criterios de selección de pacientes, incluyendo indicaciones y contraindicaciones absolutas y relativas para cada procedimiento estético ofrecido en el consultorio.")
bullet(doc, "Garantizar que el consentimiento informado de cada procedimiento sea un proceso real, completo, comprensible y libre de presión comercial, que respete la autonomía del paciente.")
bullet(doc, "Estandarizar la técnica de ejecución de cada procedimiento (toxina botulínica, rellenos dérmicos, peelings químicos, mesoterapia, PRP) mediante protocolos basados en la mejor evidencia disponible y la anatomía de riesgo relevante.")
bullet(doc, "Garantizar la disponibilidad permanente de materiales e insumos de emergencia para el manejo de complicaciones graves, especialmente oclusión vascular por relleno dérmico.")
bullet(doc, "Capacitar y entrenar al equipo del consultorio en la identificación temprana y el manejo inmediato de las complicaciones de los procedimientos estéticos.")
bullet(doc, "Establecer un proceso de seguimiento post-procedimiento estructurado que permita la detección temprana de complicaciones y la evaluación de resultados.")

# 4. ALCANCE Y EXCLUSIONES
h1(doc, "4. ALCANCE Y EXCLUSIONES")
h2(doc, "4.1 Procedimientos Incluidos en Este Proceso")
bullet(doc, "Aplicación de toxina botulínica tipo A (BTX-A): frente, entrecejo (glabela), patas de gallo, bandas platismales del cuello, hiperhidrosis axilar, hiperhidrosis palmar/plantar, bruxismo (músculo masetero), tratamiento de migraña.")
bullet(doc, "Rellenos dérmicos con ácido hialurónico (AH): labios (aumento y definición), surcos nasolabiales, código de barras peribucal, ojeras (surco lagrimal), proyección de pómulos, corrección de asimetrías faciales menores, hidratación dérmica profunda (skinboosters).")
bullet(doc, "Peelings químicos superficiales: ácido glicólico (10-50%), ácido salicílico (10-30%), ácido mandélico, ácido tricloroacético (TCA) hasta 15% en peeling superficial, solución de Jessner.")
bullet(doc, "Mesoterapia facial: aplicación de cócteles vitamínicos, ácido hialurónico no reticulado, péptidos bioestimulantes mediante microinyecciones intradérmicas.")
bullet(doc, "Plasma Rico en Plaquetas (PRP): para rejuvenecimiento facial, alopecia androgenética, cicatrices de acné.")
h2(doc, "4.2 Exclusiones")
aviso(doc, "ESTÁ ABSOLUTAMENTE PROHIBIDO en este consultorio: la aplicación de biopolímeros (polimetilmetacrilato, silicona líquida inyectable, aceite mineral, parafina), el uso de productos sin registro INVIMA vigente, la realización de cualquier procedimiento invasivo que requiera sedación, la liposucción, los implantes, la cirugía plástica, y cualquier procedimiento que supere la habilitación del establecimiento. El médico que aplica biopolímeros en Colombia comete un delito tipificado como lesiones personales con consecuencias culposas (Código Penal, Art. 112), pues estas sustancias generan complicaciones graves, crónicas e irreversibles.")

# 5. DEFINICIONES
h1(doc, "5. DEFINICIONES ESPECÍFICAS DE MEDICINA ESTÉTICA")
definiciones_est = [
    ("Procedimiento Estético No Invasivo", "Intervención médica que modifica la apariencia estética mediante la aplicación de agentes físicos o químicos en la superficie cutánea, o mediante inyecciones percutáneas de agentes biológicos o farmacológicos, sin requerir incisiones quirúrgicas ni anestesia general. Son 'no invasivos' en el sentido de que no requieren hospitalización, pero sí constituyen actos médicos que implican penetración de la integridad cutánea cuando se usan agujas."),
    ("Toxina Botulínica Tipo A (BTX-A)", "Neurotoxina proteica producida por la bacteria Clostridium botulinum que actúa en la unión neuromuscular inhibiendo la liberación de acetilcolina, produciendo paresia o parálisis temporal del músculo tratado. Su efecto es dosis-dependiente, localizado y reversible, con duración de 3 a 6 meses. En medicina estética se utiliza para el tratamiento de líneas de expresión dinámicas (arrugas causadas por movimiento muscular repetitivo)."),
    ("Ácido Hialurónico (AH)", "Glucosaminoglicano lineal de alto peso molecular, componente natural de la matriz extracelular de la dermis humana, articulaciones y humor vítreo. Los rellenos dérmicos de AH se producen mediante fermentación bacteriana seguida de reticulación (cross-linking) con 1,4-butanodiol diglicidil éter (BDDE), que aumenta la durabilidad del producto. Son biodegradables (degradados por hialuronidasa endógena) y su efecto dura entre 6 y 18 meses dependiendo del grado de reticulación, el área tratada y el metabolismo del paciente."),
    ("Oclusión Vascular por Relleno Dérmico", "Complicación grave y potencialmente catastrófica de la inyección de rellenos dérmicos, producida por la inyección inadvertida de AH (u otro relleno) dentro de un vaso sanguíneo o por compresión vascular extrínseca, que resulta en isquemia tisular aguda. Puede causar necrosis cutánea, ceguera irreversible (por oclusión de la arteria central de la retina o sus ramas) o accidente cerebrovascular. Requiere reconocimiento inmediato y tratamiento de emergencia con hialuronidasa."),
    ("Hialuronidasa", "Enzima que cataliza la hidrólisis del ácido hialurónico, utilizada como antídoto para la oclusión vascular por rellenos dérmicos de AH. Su aplicación inmediata en el área de isquemia o en el trayecto vascular comprometido puede revertir la obstrucción y restaurar la perfusión. Es el tratamiento de primera línea ante cualquier sospecha de oclusión vascular por AH. El consultorio debe tener hialuronidasa disponible en todo momento en que se realicen procedimientos con AH."),
    ("Peeling Químico", "Aplicación controlada de un agente cáustico sobre la superficie cutánea con el propósito de inducir exfoliación epidérmica y/o dérmica, estimulando la renovación celular y la síntesis de colágeno. Se clasifica según la profundidad alcanzada: superficial (epidermis y dermis papilar superficial), medio (dermis papilar profunda) y profundo (dermis reticular). Los peelings superficiales y medio-superficiales son los indicados en primer nivel."),
    ("Mesoterapia", "Técnica de administración de pequeños volúmenes de sustancias activas (vitaminas, minerales, aminoácidos, ácido hialurónico no reticulado, péptidos) mediante múltiples microinyecciones intradérmicas o subdérmicas superficiales, con el propósito de actuar localmente en el tejido diana. El término fue acuñado por el médico francés Michel Pistor en 1952. En Colombia, la mesoterapia es un acto médico."),
    ("Plasma Rico en Plaquetas (PRP)", "Concentrado autólogo de plaquetas obtenido mediante la centrifugación de la sangre periférica del paciente. Contiene factores de crecimiento plaquetarios (PDGF, TGF-β, VEGF, EGF, IGF-1) que estimulan la regeneración tisular, la síntesis de colágeno y la angiogénesis. El proceso de obtención incluye extracción de sangre venosa, centrifugación en tubo especializado, extracción de la fracción rica en plaquetas y reinyección en el área tratada."),
    ("Escala GAIS (Global Aesthetic Improvement Scale)", "Escala de evaluación de resultados estéticos de 5 puntos: 5 = aspecto excepcional, mucho mejor que el estado ideal; 4 = mejoría marcada, respecto al estado basal; 3 = mejoría, mejora obvia del aspecto inicial; 2 = sin cambio, aspecto equivalente al inicial; 1 = empeoramiento, estado peor que el inicial. Se aplica antes y después de los procedimientos para documentar el resultado."),
    ("Consentimiento Informado Estético", "Proceso y documento específico para procedimientos estéticos que, además de los elementos del consentimiento informado general, incluye: descripción precisa del procedimiento, resultados esperados REALISTAS (no garantizados), posibilidad de necesitar tratamientos repetidos, variabilidad individual de resultados, fotografías de antes y después representativas (con resultados promedio, no solo los mejores), todas las complicaciones posibles incluyendo las graves, costo del tratamiento de complicaciones si las hubiera, y tiempo de recuperación real."),
]
for term, defi in definiciones_est:
    p = doc.add_paragraph()
    run_term = p.add_run(f"{term}: ")
    run_term.bold = True; run_term.font.size = Pt(11); run_term.font.name = 'Calibri'; run_term.font.color.rgb = AZUL_OSCURO
    run_def = p.add_run(defi)
    run_def.font.size = Pt(11); run_def.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)

# 6. EVALUACIÓN PREVIA
h1(doc, "6. EVALUACIÓN MÉDICA PREVIA AL PROCEDIMIENTO")
body(doc, "Todo paciente que solicite un procedimiento estético en el [NOMBRE DEL CONSULTORIO] debe ser evaluado por la Dra. [NOMBRE DE LA MÉDICA] antes de la realización del procedimiento. Esta evaluación no es un mero trámite administrativo, sino una consulta médica de primera vez o de control con los mismos estándares que cualquier consulta de medicina general, adaptada al contexto de la medicina estética. La evaluación previa tiene el propósito de: verificar que el paciente es candidato adecuado para el procedimiento solicitado, identificar contraindicaciones absolutas o relativas, establecer expectativas realistas, obtener un consentimiento informado real, y planificar la técnica más adecuada para cada caso individual.")

h2(doc, "6.1 Historia Clínica de Medicina Estética")
body(doc, "La historia clínica para procedimientos estéticos incluye todos los componentes de la historia clínica de medicina general más los siguientes elementos específicos:")
bullet(doc, "Motivo de consulta estético: área de preocupación, tiempo de evolución, tratamientos previos en el área (incluyendo otros rellenos, cirugías previas, radiofrecuencia, láser), resultados y complicaciones de tratamientos previos.")
bullet(doc, "Antecedentes específicos de relevancia estética: trastornos de la coagulación, uso de anticoagulantes o antiagregantes (AAS, clopidogrel, warfarina, heparina, DACO), uso de suplementos con efecto anticoagulante (vitamina E >400 UI/día, omega-3, ginkgo biloba, ajo), enfermedades autoinmunes activas, antecedente de herpes simple peribucal (contraindicación relativa para rellenos de labios), historia de queloides o cicatrización hipertrófica, antecedente de reacciones alérgicas a anestésicos locales o a rellenos previos, antecedentes de procedimientos estéticos previos (tipo, producto utilizado, cuándo, dónde, resultado), uso actual de isotretinoína (contraindicación para peelings durante el tratamiento y 6 meses después), embarazo o lactancia (contraindicación absoluta para la mayoría de procedimientos estéticos).")
bullet(doc, "Evaluación de la motivación y expectativas del paciente: exploración de las expectativas reales del paciente sobre el procedimiento, identificación de señales de trastorno dismórfico corporal (preocupación desproporcionada por defectos mínimos o inexistentes, historia de múltiples procedimientos sin satisfacción, solicitudes de cambios radicales). El trastorno dismórfico corporal es una contraindicación absoluta para cualquier procedimiento estético.")
bullet(doc, "Fotografía clínica estandarizada: toma de fotografías de frente, perfil derecho, perfil izquierdo, ¾ derecho y ¾ izquierdo en condiciones estandarizadas (misma iluminación, fondo neutro, sin maquillaje) antes de cada procedimiento. Las fotografías forman parte de la historia clínica y requieren consentimiento escrito específico del paciente.")

h2(doc, "6.2 Contraindicaciones Absolutas (aplican a todos los procedimientos)")
t_contra = doc.add_table(rows=1, cols=2)
t_contra.style = 'Table Grid'
table_header_row(t_contra, ["Contraindicación Absoluta", "Justificación"])
contras_abs = [
    ("Embarazo (cualquier trimestre)", "Seguridad no establecida en el feto para ninguno de los productos o agentes utilizados."),
    ("Lactancia", "Posible transferencia de toxinas o agentes al lactante a través de la leche materna."),
    ("Trastorno dismórfico corporal diagnosticado o sospechado", "El procedimiento no mejorará la percepción del paciente y puede empeorar el cuadro psiquiátrico."),
    ("Infección activa en el área a tratar (herpes activo, acné inflamatorio severo, infección bacteriana)", "Riesgo de diseminación de la infección y complicaciones graves."),
    ("Alergia conocida al producto a utilizar o a cualquiera de sus componentes", "Riesgo de reacción anafiláctica potencialmente fatal."),
    ("Coagulopatía no controlada o uso de anticoagulantes que no pueden suspenderse", "Riesgo de hematomas extensos, equimosis severa o hemorragia."),
    ("Enfermedades autoinmunes en fase activa con afectación sistémica", "Riesgo aumentado de efectos adversos y posible desencadenamiento de brote."),
    ("Expectativas irreales o no aceptación de los riesgos explicados", "El procedimiento no puede realizarse sin consentimiento informado real y completo."),
]
for i, row_data in enumerate(contras_abs):
    add_table_row(t_contra, list(row_data), shaded=(i % 2 == 1))

doc.add_paragraph()

# 7. PROCEDIMIENTOS ESPECÍFICOS
h1(doc, "7. PROTOCOLOS TÉCNICOS DE PROCEDIMIENTOS ESTÉTICOS")

# 7.1 TOXINA BOTULÍNICA
h2(doc, "7.1 TOXINA BOTULÍNICA TIPO A (BTX-A)")
h3(doc, "7.1.1 Farmacología y Mecanismo de Acción")
body(doc, "La toxina botulínica tipo A (BTX-A) es una proteína de cadena pesada (~100 kDa) y cadena ligera (~50 kDa) unidas por un puente disulfuro, producida por Clostridium botulinum. En medicina estética, las formulaciones comerciales disponibles en Colombia con registro INVIMA incluyen: Botox® (onabotulinumtoxinA, Allergan/AbbVie), Dysport® (abobotulinumtoxinA, Ipsen), Xeomin® (incobotulinumtoxinA, Merz). Estas formulaciones NO son bioequivalentes: las unidades de BTX-A de Botox no son intercambiables con las unidades de Dysport (ratio aproximado 1:2.5-3) ni con las de Xeomin (ratio aproximado 1:1, aunque con ligeras variaciones). El médico debe conocer la formulación que utiliza y las dosis específicas para cada una.")
body(doc, "Mecanismo de acción: La BTX-A bloquea la liberación de acetilcolina en la unión neuromuscular mediante la endocitosis de la toxina en la terminación nerviosa presináptica, seguida de la escisión proteolítica de la proteína SNAP-25 (Synaptosomal-Associated Protein 25 kDa), que es esencial para el proceso de fusión de las vesículas de acetilcolina con la membrana plasmática. El resultado es una paresia o parálisis temporal y reversible del músculo tratado, sin efecto sobre la condición del nervio motor ni sobre la fibra muscular. El efecto se inicia entre los 2 y 5 días post-inyección, alcanza su máximo entre los 7 y 14 días, y persiste entre 3 y 6 meses, momento en el cual las terminaciones nerviosas restablecen nuevas sinapsis funcionantes y el músculo recupera gradualmente su actividad.")
body(doc, "Almacenamiento y reconstitución: La BTX-A se comercializa liofilizada en frascos de 50U o 100U (Botox), o en presentaciones equivalentes para otras marcas. Se almacena entre 2-8°C (refrigerador), protegida de la luz, antes de su reconstitución. Se reconstituye con solución salina normal (SSN 0.9%) estéril sin preservante. El volumen de reconstitución varía según la formulación y la concentración deseada: para Botox, la reconstitución habitual es con 2.5 mL de SSN por frasco de 100U (40U/mL), aunque puede variar entre 1 y 4 mL según la preferencia del médico. Una vez reconstituida, debe refrigerarse y utilizarse dentro de las 4 a 24 horas (máximo 72 horas con conservantes) según las especificaciones del fabricante.")
body(doc, "Propiedades de difusión: La BTX-A difunde radialmente desde el punto de inyección en un radio proporcional al volumen inyectado y a la concentración. Una mayor dilución (mayor volumen por unidad) produce mayor difusión pero menor precisión. En áreas donde se desea efecto localizado (por ejemplo, elevador del párpado superior, músculos periorbitarios), se prefieren volúmenes pequeños y alta concentración para minimizar la difusión inadvertida a músculos no diana.")

h3(doc, "7.1.2 Anatomía Relevante y Músculos Diana")
body(doc, "El conocimiento preciso de la anatomía facial es indispensable para la aplicación segura y eficaz de la BTX-A. Los principales músculos tratados en medicina estética y su anatomía relevante son:")
bullet(doc, "FRENTE (músculo frontal): Músculo par, de fibras verticales, que eleva las cejas y produce las arrugas horizontales de la frente. La inyección debe realizarse en la mitad superior del músculo (>2 cm por encima de la ceja) para evitar la paresia del elevador del párpado superior. Se distribuyen 10-20U de Botox en 4-6 puntos de inyección a lo largo del músculo. PRECAUCIÓN: el frontal es el único elevador de la ceja; un tratamiento excesivo produce ptosis de la ceja con aspecto 'pesado' y puede comprometer el campo visual superior.")
bullet(doc, "GLABELA (corrugador superciliar y prócer): Complejo muscular que produce las arrugas verticales y oblicuas del entrecejo ('número 11'). El corrugador se origina en el arco superciliar medial y se inserta en la piel de la porción media de la ceja. El prócer se origina en el hueso nasal y se inserta en la piel del entrecejo. Se inyectan 10-25U de Botox distribuidos en 3-5 puntos. PRECAUCIÓN ANATÓMICA: la arteria supratroclear y la arteria supraorbitaria cursan en la región medial de la ceja y el entrecejo. La inyección intravascular inadvertida puede producir oclusión vascular y necrosis de la glabela o, en el peor escenario, oclusión retrográdamente hasta la arteria oftálmica causando ceguera. La inyección debe ser intramuscular, no intravascular. Aspirar antes de inyectar en esta área.")
bullet(doc, "PATAS DE GALLO (músculo orbicular del ojo, fibras laterales): Las arrugas del ángulo lateral externo del ojo son producidas por el orbicular del ojo, que rodea circunferencialmente el ojo. Las inyecciones se realizan lateralmente al ojo, a >1 cm del margen orbital lateral, en la piel subcutánea o intramuscular superficial. Se inyectan 8-15U de Botox distribuidos en 2-4 puntos. PRECAUCIÓN: inyecciones muy próximas al ojo o en posición inferior-medial pueden difundir al recto inferior causando diplopía transitoria.")
bullet(doc, "BANDAS PLATISMALES (platisma): Músculo de fibras verticales del cuello que produce las bandas verticales prominentes del cuello ('cuello de pavo'). Requiere inyecciones directas sobre las bandas visibles en la piel del cuello, 3-6 puntos por banda, 2-4U por punto. PRECAUCIÓN: el platisma subyace a estructuras importantes del cuello (carótida, yugular, nervio espinal accesorio). La inyección debe ser superficial-intramuscular.")
bullet(doc, "MÚSCULO MASETERO (bruxismo/hipertrofia): El masetero es el principal músculo masticador, de gran potencia. Su inyección con BTX-A reduce el tono muscular y el volumen del masetero hipertrófico, con resultados tanto funcionales (reducción del bruxismo y dolor orofacial) como estéticos (contorno facial más delgado). Dosis: 20-40U por lado. El punto de inyección es el tercio inferior y posterior del masetero, palpable como una prominencia al apretar los dientes. PRECAUCIÓN: evitar la parótida, que es contigua al masetero, para prevenir xerostomía y cambios en la secreción salival.")

h3(doc, "7.1.3 Técnica de Aplicación - Protocolo Paso a Paso")
body(doc, "PREPARACIÓN Y VERIFICACIÓN PRE-PROCEDIMIENTO:")
pasos_btx = [
    "Verificar que el consentimiento informado FOR-PP-002 esté firmado y fechado, después de haber explicado todos los riesgos al paciente.",
    "Verificar que la historia clínica de medicina estética esté completamente diligenciada, incluyendo todas las contraindicaciones investigadas.",
    "Confirmar la ausencia de contraindicaciones en la consulta del día: no embarazo, no infección activa, no medicamentos anticoagulantes activos (se recomienda suspender AAS 5-7 días antes, omega-3 y vitamina E 3 días antes, previa autorización del médico tratante si los toma con indicación médica).",
    "Tomar fotografías clínicas estandarizadas de frente, perfil y ¾ sin maquillaje, en condiciones de iluminación estandarizadas. Guardar en la historia clínica con fecha y código del paciente.",
    "Realizar el análisis facial completo: simetría de cejas, posición del párpado superior (distancia del margen palpebral al reflejo corneal, normal >3.5 mm), tono de los músculos diana en reposo y en movimiento, distribución y profundidad de las arrugas. Registrar los hallazgos en la historia clínica.",
    "Calcular la dosis total según la formulación utilizada y las áreas a tratar. Documentar dosis por área en la historia clínica.",
    "Reconstituir el frasco de BTX-A si no estaba previamente reconstituido, siguiendo las instrucciones del fabricante. Anotar el lote, fecha de vencimiento y volumen de reconstitución en la historia clínica.",
    "Preparar la jeringa de insulina (1 mL, aguja 30G o 31G) con la dosis calculada para el área a tratar. Cambiar la aguja antes de cada inyección para garantizar que esté perfectamente afilada y no contaminada.",
    "Limpiar el área a tratar con solución antiséptica (gel antibacterial de alcohol) y dejar secar 30 segundos. NO usar povidona yodada porque inactiva la toxina si hay contacto.",
    "Posicionar al paciente: sentado o reclinado a 45°, con la cabeza en posición neutra, iluminación adecuada sobre el área a tratar.",
]
for i, paso in enumerate(pasos_btx):
    p = doc.add_paragraph()
    run_num = p.add_run(f"{i+1}. ")
    run_num.bold = True; run_num.font.size = Pt(11); run_num.font.name = 'Calibri'
    run_text = p.add_run(paso)
    run_text.font.size = Pt(11); run_text.font.name = 'Calibri'
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)

body(doc, "TÉCNICA DE INYECCIÓN:")
pasos_btx2 = [
    "Identificar visualmente y palpar los músculos diana con el paciente en movimiento (pedirle que frunza el ceño, eleve las cejas, sonría, etc.) y en reposo.",
    "Marcar los puntos de inyección con un lápiz dérmico blanco si se requiere, especialmente en tratamientos complejos o en pacientes nuevos.",
    "Insertar la aguja en el punto de inyección con bisel hacia arriba, a la profundidad requerida según el músculo (intramuscular para frontal, glabela y masetero; subcutáneo-intramuscular superficial para orbicular lateral).",
    "SIEMPRE aspirar levemente el émbolo antes de inyectar para verificar ausencia de sangre en la jeringa (confirmar que no está en posición intravascular). Si aparece sangre: retirar la aguja, comprimir el punto por 2 minutos y reintroducir en nuevo punto.",
    "Inyectar el volumen calculado de manera lenta y controlada (no inyectar a presión). Si se siente resistencia, verificar la posición de la aguja.",
    "Retirar la aguja con movimiento firme y aplicar presión suave con gasa estéril durante 30-60 segundos. NO frotar el área, pues esto promueve la difusión no deseada.",
    "Repetir el procedimiento en cada punto de inyección planificado, cambiando la aguja si se nota pérdida de filo.",
    "Al terminar todos los puntos, evaluar visualmente el resultado inmediato: simetría, hematomas, signos de inyección intravascular.",
    "Descartar el frasco vial parcialmente utilizado si no va a usarse en otro paciente en la misma sesión, o conservarlo en refrigeración con las especificaciones del fabricante si se va a usar en menos de 24 horas.",
    "Registrar en la historia clínica: producto utilizado (nombre comercial, lote, fecha de vencimiento), dosis total aplicada y distribución por puntos de inyección, técnica utilizada, tolerancia del procedimiento, y cualquier incidencia.",
]
for i, paso in enumerate(pasos_btx2):
    p = doc.add_paragraph()
    run_num = p.add_run(f"{i+11}. ")
    run_num.bold = True; run_num.font.size = Pt(11); run_num.font.name = 'Calibri'
    run_text = p.add_run(paso)
    run_text.font.size = Pt(11); run_text.font.name = 'Calibri'
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)

h3(doc, "7.1.4 Cuidados Post-Procedimiento e Instrucciones al Paciente")
bullet(doc, "Permanecer erguido (no acostarse) durante las primeras 4 horas post-inyección para minimizar la difusión gravitacional de la toxina.")
bullet(doc, "No frotar, masajear ni presionar el área tratada durante 24 horas.")
bullet(doc, "Evitar actividad física intensa, saunas, baños calientes, exposición solar intensa y alcohol durante 24-48 horas.")
bullet(doc, "No realizar procedimientos de calor (ultrasonido, radiofrecuencia) en el área tratada durante 2 semanas.")
bullet(doc, "El efecto completo se aprecia a los 7-14 días. Cita de revisión a los 14 días para evaluación de resultados y ajustes si se requieren.")
bullet(doc, "Señales de alarma que requieren consulta inmediata: ptosis palpebral (párpado caído), diplopía (visión doble), dificultad para deglutir, alteraciones del habla, debilidad generalizada.")

h3(doc, "7.1.5 Complicaciones y su Manejo")
t_comp_btx = doc.add_table(rows=1, cols=4)
t_comp_btx.style = 'Table Grid'
table_header_row(t_comp_btx, ["Complicación", "Causa", "Prevalencia", "Manejo"])
comp_btx = [
    ("Equimosis / hematoma", "Punción de vaso sanguíneo superficial", "10-25%", "Presión local 2-5 min, compresas frías, arnica tópica. Resuelve espontáneamente en 7-10 días."),
    ("Cefalea post-inyección", "Mecanismo no completamente establecido. Posiblemente tensión muscular reflexiva.", "5-15%", "Analgésico oral (acetaminofén 1 g). Resuelve en 24-48 horas."),
    ("Ptosis palpebral (párpado caído)", "Difusión de BTX-A al músculo elevador del párpado superior. Principalmente en tratamiento de la glabela.", "1-3%", "Colirio de apraclonidina 0.5% o fenilefrina 2.5% (estimula el músculo de Müller). Resolución espontánea en 4-6 semanas. Informar al paciente, usar gafas de sol si es necesario."),
    ("Ptosis de la ceja", "Tratamiento excesivo del músculo frontal sin compensar con glabela.", "5-10%", "No tiene tratamiento farmacológico. Resolución espontánea en 6-12 semanas. Evitar en próxima sesión con menor dosis en frontal."),
    ("Diplopía (visión doble)", "Difusión al músculo recto inferior u otros músculos extraoculares en tratamiento periorbitario.", "<0.5%", "Parche ocular, evaluación por oftalmología. Resolución espontánea en 4-8 semanas."),
    ("Asimetría", "Distribución desigual de la toxina, anatómica asimétrica del paciente.", "5-10%", "Evaluación a los 14 días. Si persiste, ajuste con pequeña dosis adicional en el lado con menor efecto."),
    ("Resistencia al tratamiento (anticuerpos neutralizantes)", "Formación de anticuerpos anti-BTX-A tras exposiciones repetidas. Más frecuente con altas dosis y cortos intervalos.", "<1% con dosis bajas", "Cambio a otra formulación de BTX-A, espaciado de las sesiones. Considerar toxina tipo B si disponible."),
]
for i, row_data in enumerate(comp_btx):
    add_table_row(t_comp_btx, list(row_data), shaded=(i % 2 == 1))

doc.add_paragraph()

# 7.2 RELLENOS DÉRMICOS
h2(doc, "7.2 RELLENOS DÉRMICOS CON ÁCIDO HIALURÓNICO (AH)")
h3(doc, "7.2.1 Farmacología, Propiedades Físicas y Selección de Producto")
body(doc, "El ácido hialurónico (AH) utilizado en medicina estética es producido mediante fermentación bacteriana de Streptococcus equi y posterior reticulación con BDDE. Las propiedades físicas del relleno que determinan su comportamiento clínico son: a) Concentración de AH (mg/mL): varía entre 15 y 25 mg/mL. Mayor concentración = mayor volumen de tejido desplazado. b) Grado de reticulación: determina la dureza (cohesividad) del gel y su resistencia a la degradación enzimática. Rellenos altamente reticulados son más firmes, duran más y son apropiados para proyección de estructuras profundas (pómulos, mandíbula). Rellenos de baja reticulación son más suaves, más apropiados para labios y superficies delgadas. c) Fuerza de elevación (G'): módulo de elasticidad del gel. Alto G' = relleno firme, buena proyección. Bajo G' = relleno suave, mejor para labios y áreas de movimiento. d) Higroscopicidad: la capacidad del AH de atraer agua puede producir edema tardío en áreas de acumulación de agua (ojeras, labios) que puede interpretarse erróneamente como exceso de producto.")
body(doc, "Marcas disponibles en Colombia con registro INVIMA (ejemplos): Juvederm® (Allergan/AbbVie), Restylane® (Galderma), Belotero® (Merz), Sculptra® (Galderma, ácido poli-L-láctico - diferente al AH). El médico debe conocer el producto específico que utiliza, su perfil de reología, las indicaciones aprobadas y las contraindicaciones del fabricante. NUNCA usar productos sin registro INVIMA o de origen desconocido.")
body(doc, "Anestesia tópica y local: Los rellenos dérmicos producen dolor en el momento de la inyección. Para minimizar el dolor se utilizan: a) Crema anestésica tópica de aplicación previa (mezcla eutéctica de lidocaína 2.5% y prilocaína 2.5%, EMLA® o equivalente), aplicada 45-60 minutos antes del procedimiento cubierta con film de polietileno. b) Bloqueos nerviosos periféricos con lidocaína sin vasoconstrictor para tratamiento de labios (bloqueo del nervio infraorbitario y mentoniano). c) Muchos rellenos modernos incorporan lidocaína 0.3% en su formulación (Juvederm® XC, Restylane® con lidocaína), lo que reduce el dolor de forma significativa.")

h3(doc, "7.2.2 Anatomía de Riesgo Vascular - CONOCIMIENTO ESENCIAL DE SEGURIDAD")
aviso(doc, "Esta sección es DE LECTURA OBLIGATORIA antes de realizar cualquier procedimiento con rellenos dérmicos. El desconocimiento de la anatomía vascular facial es la causa principal de complicaciones catastróficas (necrosis, ceguera) con rellenos dérmicos. El médico debe actualizarse periódicamente en la anatomía de zonas de alto riesgo.")
body(doc, "Zonas de ALTO RIESGO vascular para inyección de rellenos dérmicos (en orden de mayor a menor riesgo de ceguera):")
bullet(doc, "GLABELA Y ENTRECEJO: La arteria supratroclear y la arteria supraorbitaria (ramas de la arteria oftálmica, que es rama de la carótida interna) cursan profundas bajo la musculatura del entrecejo. La inyección de AH en esta área puede ocluir estos vasos retrógradamente hasta la arteria oftálmica y producir CEGUERA IRREVERSIBLE. Esta es la zona de más alto riesgo de ceguera en medicina estética. Técnica segura: inyectar superficialmente (subdérmico), volúmenes mínimos, sin presión excesiva, con aguja o cánula.")
bullet(doc, "REGIÓN NASAL: La arteria dorsal de la nariz (rama terminal de la arteria angular, de la arteria facial) y las arterias laterales de la nariz pueden comunicarse con la arteria oftálmica. Relleno en la nariz puede producir ceguera por flujo retrógrado. La rinoplastia no quirúrgica ('nariz de botox') tiene uno de los perfiles de riesgo más altos de ceguera en medicina estética. Requiere formación avanzada específica.")
bullet(doc, "REGIÓN TEMPORAL: La arteria temporal superficial es superficial en la sien. La inyección en el temporal puede comprometer la irrigación del músculo temporal y producir necrosis del colgajo temporal.")
bullet(doc, "LABIOS Y ÁREA PERIBUCAL: La arteria labial superior e inferior son ramas de la arteria facial. Cursan en el plano submucoso del labio, profundas respecto a la submucosa. El riesgo de oclusión labial es elevado en manos inexpertas.")
bullet(doc, "SURCO NASOLABIAL: La arteria facial transcurre en profundidad en el surco nasolabial. Inyección profunda puede comprometer la arteria facial.")
bullet(doc, "SURCO LAGRIMAL (ojeras): La arteria infraorbitaria emerge del foramen infraorbitario en esta zona. Área de alta complejidad que requiere experiencia avanzada.")

h3(doc, "7.2.3 Técnica de Aplicación - Protocolo Paso a Paso")
body(doc, "PREPARACIÓN PRE-PROCEDIMIENTO (pasos 1-10 similares a BTX-A). Adicionalmente:")
pasos_ah = [
    "Verificar el consentimiento informado FOR-PP-003 firmado, con énfasis en que el paciente comprende el riesgo de oclusión vascular.",
    "Tener la hialuronidasa preparada y disponible en el área de procedimientos ANTES de iniciar cualquier inyección de AH. Verificar la fecha de vencimiento. No iniciar el procedimiento si no se cuenta con hialuronidasa.",
    "Seleccionar el relleno adecuado para la indicación: relleno suave (bajo G') para labios; relleno medio para surcos nasolabiales y ojeras; relleno duro (alto G', alto G') para proyección de pómulos, mandíbula; relleno de AH no reticulado para skinbooster.",
    "Aplicar anestesia tópica (EMLA® u otra) 45-60 minutos antes si el paciente no tiene contraindicaciones. Cubrir con film transparente.",
    "Retirar la anestesia tópica y limpiar el área con solución antiséptica. Dejar actuar 30 segundos. NO usar povidona yodada.",
    "Seleccionar la herramienta de inyección: aguja (30G para labios, surcos, ojeras) o cánula (23G-25G según el área). Las cánulas reducen el riesgo de punción vascular directa pero no eliminan el riesgo de oclusión por compresión extrínseca.",
    "Identificar los puntos de entrada y los planos de inyección planificados para cada área.",
    "Insertar la aguja o cánula en el plano planificado (subdérmico, supraperióstico o intramuscular según el área y el producto).",
    "SIEMPRE aspirar antes de inyectar (aunque la evidencia sobre su sensibilidad para detectar posición intravascular es discutida, es una medida de seguridad mínima). Si aparece sangre: retirar completamente, comprimir 2 min, usar nuevo punto de entrada.",
    "Inyectar en retroceso (técnica de retroinyección): mover la aguja en retroceso mientras se inyecta el producto, o depositar pequeños bolos. Nunca inyectar a alta presión.",
    "Verificar el resultado inmediato: simetría, homogeneidad, ausencia de nódulos, ausencia de blanqueamiento cutáneo (signo de oclusión vascular).",
    "Modelar suavemente el producto con los dedos enguantados para distribuirlo uniformemente si es necesario.",
    "Registrar en la historia clínica: producto utilizado (nombre comercial, lote, fecha de vencimiento, volumen utilizado por área), técnica empleada, tolerancia al procedimiento, resultado inmediato.",
]
for i, paso in enumerate(pasos_ah):
    p = doc.add_paragraph()
    run_num = p.add_run(f"{i+1}. ")
    run_num.bold = True; run_num.font.size = Pt(11); run_num.font.name = 'Calibri'
    run_text = p.add_run(paso)
    run_text.font.size = Pt(11); run_text.font.name = 'Calibri'
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)

h3(doc, "7.2.4 PROTOCOLO DE EMERGENCIA POR OCLUSIÓN VASCULAR")
aviso(doc, "La oclusión vascular por relleno dérmico es una EMERGENCIA MÉDICA. El tiempo de actuación desde el reconocimiento hasta el inicio del tratamiento con hialuronidasa es CRÍTICO. Cada minuto de isquemia aumenta el daño tisular. El médico debe conocer este protocolo de memoria y tener todos los materiales necesarios disponibles en el consultorio en todo momento en que se realicen procedimientos con AH.")

body(doc, "SIGNOS DE ALARMA DE OCLUSIÓN VASCULAR (reconocerlos DURANTE y después del procedimiento):")
bullet(doc, "SIGNO PRINCIPAL: Blanqueamiento cutáneo súbito en el área inyectada o áreas adyacentes (piel blanca o gris por isquemia). Este es el signo más temprano y el que debe activar INMEDIATAMENTE el protocolo de emergencia.")
bullet(doc, "Dolor intenso y desproporcionado al procedimiento durante la inyección.")
bullet(doc, "Cambio de coloración progresivo: de blanco a azul-violáceo (cianosis por isquemia venosa) o moteado.")
bullet(doc, "Pérdida visual súbita o parcial, visión borrosa, escotoma (urgencia oftalmológica absoluta).")
bullet(doc, "Dolor ocular ipsilateral al área inyectada.")

body(doc, "PROTOCOLO DE EMERGENCIA VASCULAR - SECUENCIA DE PASOS INMEDIATOS:")
pasos_emergency = [
    "DETENER inmediatamente la inyección al primer signo de alarma. No inyectar ni un microlito más de producto.",
    "LLAMAR AL 123 (SISTEMA DE EMERGENCIAS MÉDICAS) si el paciente presenta pérdida visual, dolor ocular, alteración neurológica o compromiso vascular extenso. Indicar: 'Paciente con oclusión vascular por relleno dérmico, requiere traslado urgente a urgencias con oftalmología y cirugía vascular.'",
    "HIALURONIDASA - INYECCIÓN INMEDIATA EN EL ÁREA AFECTADA: Preparar hialuronidasa (200-300 UI en 1-2 mL de solución salina estéril para el área isquémica). Inyectar EN y alrededor del área de blanqueamiento, distribuida en 5-10 puntos de inyección en el área isquémica y a lo largo del trayecto vascular comprometido.",
    "DOSIS DE HIALURONIDASA: Para oclusión cutánea (sin compromiso ocular): 200-300 UI en el área afectada, repitiendo cada 60 minutos si no hay mejoría visible. Para oclusión con riesgo de ceguera (área de la glabela, frente, nariz): 500-1500 UI en el área inyectada, distribuidas en múltiples puntos, de manera INMEDIATA. Informar a la literatura que dosis >1500 UI en área periocular pueden ser necesarias.",
    "CALOR LOCAL: Aplicar compresas tibias (no calientes) sobre el área afectada para promover vasodilatación y aumentar la perfusión.",
    "MASAJE SUAVE: Masaje suave sobre el área isquémica para ayudar a distribuir la hialuronidasa y disolver el AH.",
    "NITROGLICERINA TÓPICA: Aplicar pasta de nitroglicerina 2% (si está disponible) sobre el área isquémica para producir vasodilatación local.",
    "ÁCIDO ACETILSALICÍLICO (AAS): Administrar AAS 300 mg VO si el paciente no es alérgico, como antiagregante plaquetario para reducir el riesgo de extensión de la trombosis.",
    "MONITORIZAR al paciente: signos vitales, estado de la isquemia (mejoría o progresión), estado visual, nivel de conciencia. Documentar el seguimiento.",
    "Si hay mejoría evidente del color cutáneo en 30-60 minutos: continuar monitorización, citar al día siguiente para control. Si no hay mejoría: REFERENCIA URGENTE a urgencias hospitalarias.",
    "Para compromiso visual: REFERENCIA OFTALMOLÓGICA DE EMERGENCIA dentro de la primera hora del evento. No esperar. Llamar directamente a la clínica de referencia para recibir al paciente.",
    "Documentar COMPLETAMENTE el evento en la historia clínica: hora de inicio, signos observados, dosis de hialuronidasa, pasos realizados, respuesta del paciente, hora de llamado al 123 si aplica, condición del paciente al alta o traslado.",
    "Reportar el evento adverso en el formato FOR-PP-005 (Reporte de Evento Adverso) dentro de las 24 horas siguientes, según el protocolo de seguridad del paciente.",
]
for i, paso in enumerate(pasos_emergency):
    p = doc.add_paragraph()
    run_num = p.add_run(f"PASO {i+1}: ")
    run_num.bold = True; run_num.font.size = Pt(11); run_num.font.name = 'Calibri'; run_num.font.color.rgb = ROJO
    run_text = p.add_run(paso)
    run_text.font.size = Pt(11); run_text.font.name = 'Calibri'
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)

body(doc, "MATERIALES DE EMERGENCIA QUE DEBEN ESTAR SIEMPRE DISPONIBLES EN EL CONSULTORIO:")
t_mat_em = doc.add_table(rows=1, cols=3)
t_mat_em.style = 'Table Grid'
table_header_row(t_mat_em, ["Material", "Especificación", "Cantidad Mínima en Stock"])
materiales_em = [
    ("Hialuronidasa", "Presentación ≥150 UI/ampolla o 1500 UI/vial. Marca de referencia: Hyaluronidase® o equivalente con registro INVIMA.", "2 viales o ampollas disponibles, vigentes"),
    ("Jeringas de 1 mL con aguja 30G", "Para dilución y aplicación de hialuronidasa", "10 unidades"),
    ("Solución salina normal (SSN 0.9%) estéril", "Para dilución de hialuronidasa", "50 mL"),
    ("AAS (ácido acetilsalicílico) 300 mg", "Comprimidos vía oral", "4 comprimidos"),
    ("Pasta de nitroglicerina tópica 2%", "Vasodilatador local", "1 tubo"),
    ("Compresas estériles", "Para compresas tibias y apósitos", "10 unidades"),
    ("Adrenalina 1:1000 (0.1%)", "Para manejo de reacción anafiláctica (emergencia adicional)", "2 ampollas"),
    ("Antihistamínico parenteral (clorfenamina o difenhidramina)", "Para manejo de reacciones alérgicas leves a moderadas", "2 ampollas"),
    ("Hidrocortisona 100 mg IV/IM", "Para manejo de reacciones alérgicas severas", "2 ampollas"),
    ("Esfigmomanómetro y oxímetro de pulso", "Monitoreo de signos vitales durante emergencia", "Equipos de la consulta"),
]
for i, row_data in enumerate(materiales_em):
    add_table_row(t_mat_em, list(row_data), shaded=(i % 2 == 1))

doc.add_paragraph()

# 7.3 PEELINGS QUIMICOS
h2(doc, "7.3 PEELINGS QUÍMICOS SUPERFICIALES")
h3(doc, "7.3.1 Mecanismo de Acción y Clasificación")
body(doc, "Los agentes de peeling químico actúan produciendo lesión química controlada de la epidermis y/o dermis papilar, induciendo la renovación epitelial y estimulando la síntesis de colágeno tipo I y tipo III por los fibroblastos dérmicos. El grado de penetración (y por tanto la profundidad del peeling) depende de: el agente utilizado y su concentración, el pH de la solución, el número de capas aplicadas, el tiempo de contacto, la preparación previa de la piel del paciente (uso de ácido retinoico previo que adelgaza el estrato córneo y uniformiza la penetración), y las características individuales del paciente (grosor de piel, presencia de queratosis, hidratación).")
body(doc, "CLASIFICACIÓN DE LOS AGENTES DE PEELING y sus indicaciones principales en primer nivel:")
t_peel_class = doc.add_table(rows=1, cols=4)
t_peel_class.style = 'Table Grid'
table_header_row(t_peel_class, ["Agente", "Concentración Superficial", "Mecanismo", "Indicaciones Principales"])
peel_agents = [
    ("Ácido glicólico (alfa-hidroxiácido)", "20-50% para superficial; 50-70% para medio-superficial", "Disrupciona los desmosomas del estrato córneo, promueve la descamación. A mayor concentración, penetra más profundo.", "Hiperpigmentación, melasma leve, acné comedogénico, piel fotoenvejecida, poros dilatados. Indicado para pieles Fitzpatrick I-IV."),
    ("Ácido salicílico (beta-hidroxiácido)", "10-30%", "Liposoluble: penetra el folículo pilosebáceo. Efecto queratolítico, comedolítico, antiinflamatorio y antiséptico.", "Acné inflamatorio y comedogénico, piel grasa, foliculitis. Indicado para Fitzpatrick I-VI. NO usar en embarazadas (salicilismo en altas concentraciones)."),
    ("Ácido mandélico", "10-30%", "Alfa-hidroxiácido de molécula grande: penetración más lenta y uniforme, menor irritación.", "Acné, hiperpigmentación, piel sensible. Buena tolerabilidad en pieles oscuras (Fitzpatrick IV-VI) con menor riesgo de hiperpigmentación post-inflamatoria."),
    ("Ácido tricloroacético (TCA)", "10-15% para superficial", "Precipitación de proteínas epidérmicas y dérmicas ('escarcha blanca' como punto final de aplicación).", "Fotoenvejecimiento, hiperpigmentación, cicatrices superficiales. Requiere mayor experiencia por el mayor riesgo de complicaciones."),
    ("Solución de Jessner", "Fórmula: ácido resorcinol 14g + ácido salicílico 14g + ácido láctico 14g + etanol 95% csp 100 mL", "Efecto combinado de los tres ácidos. Utilizada sola (superficial) o como primer paso antes de TCA (medio).", "Hiperpigmentación, melasma, acné, fotoenvejecimiento. Versatilidad alta."),
]
for i, row_data in enumerate(peel_agents):
    add_table_row(t_peel_class, list(row_data), shaded=(i % 2 == 1))

doc.add_paragraph()

h3(doc, "7.3.2 Protocolo de Preparación Pre-Peeling (4-6 semanas antes)")
body(doc, "La preparación pre-peeling es fundamental para garantizar resultados uniformes, reducir el riesgo de hiperpigmentación post-inflamatoria (especialmente en pieles oscuras Fitzpatrick IV-VI) y acelerar la recuperación. El esquema de preparación estándar incluye:")
bullet(doc, "Ácido retinoico 0.025-0.05% tópico nocturno: comenzar 4-6 semanas antes del peeling para adelgazar el estrato córneo, uniformizar la penetración del agente de peeling y estimular la renovación epidérmica. Suspender 5-7 días antes del peeling para reducir la irritación el día del procedimiento.")
bullet(doc, "Hidroquinona 4% tópica matutina (en pacientes con hiperpigmentación o pieles Fitzpatrick IV-VI): inhibidor de la tirosinasa para reducir el riesgo de hiperpigmentación post-inflamatoria. Puede combinarse con ácido azelaico 15-20%.")
bullet(doc, "Protector solar de amplio espectro SPF ≥50: uso diario estricto durante la preparación y durante los 3 meses post-peeling. La fotoproteección es el factor más importante para el mantenimiento de los resultados.")
bullet(doc, "Hidratación cutánea diaria: con emoliente sin fragancia para mantener la barrera cutánea íntegra.")

h3(doc, "7.3.3 Técnica de Aplicación de Peeling Químico - Paso a Paso")
pasos_peel = [
    "Verificar consentimiento informado FOR-PP-004 firmado. Explicar al paciente qué sensaciones puede esperar (ardor, calor, eritema) y el proceso de descamación post-peeling.",
    "Limpiar el rostro con limpiador suave sin aceite. Retirar maquillaje completamente. Secar con palmaditas (no frotar).",
    "Desengrase la piel con acetona o alcohol isopropílico 70% para remover el exceso de sebo, que interfiere con la penetración uniforme del agente. Secar completamente.",
    "Proteger zonas sensibles: cantus medial de los ojos, orificios nasales, comisuras labiales con vaselina en pequeña cantidad con aplicador de punta fina.",
    "Aplicar el agente de peeling con brocha, gasa o hisopo de algodón, en movimientos rápidos y uniformes desde la frente hacia abajo, empezando por las áreas más resistentes (frente, mentón) y terminando en las más sensibles (área periorbitaria, peribucal). NO aplicar en párpados, mucosa labial ni en lesiones abiertas.",
    "Observar el proceso de penetración: eritema progresivo (signo de reacción inflamatoria cutánea), nivel de ardor referido por el paciente, y para TCA: aparición de 'escarcha blanca' (precipitación de proteínas = punto final de aplicación para TCA).",
    "Para peelings superficiales con ácido glicólico: neutralizar con solución de bicarbonato de sodio al 5-10% o agua abundante al alcanzar el tiempo de exposición estipulado (2-6 minutos según concentración y respuesta). El ácido salicílico se autolimita (cristaliza); el TCA no requiere neutralización (se autolimita con la escarcha).",
    "Aplicar compresas de agua fría para calmar la sensación de ardor post-peeling.",
    "Aplicar crema calmante y anti-inflamatoria (pantenol, centella asiática, niacinamida) en capa generosa.",
    "Aplicar protector solar físico de amplio espectro (dióxido de titanio o zinc) SPF 50+ como última capa.",
    "Registrar en la historia clínica: agente utilizado, concentración, número de capas, tiempo de exposición, punto final observado, tolerancia del procedimiento, apariencia de la piel al finalizar.",
]
for i, paso in enumerate(pasos_peel):
    p = doc.add_paragraph()
    run_num = p.add_run(f"{i+1}. ")
    run_num.bold = True; run_num.font.size = Pt(11); run_num.font.name = 'Calibri'
    run_text = p.add_run(paso)
    run_text.font.size = Pt(11); run_text.font.name = 'Calibri'
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)

# 7.4 MESOTERAPIA
h2(doc, "7.4 MESOTERAPIA FACIAL")
h3(doc, "7.4.1 Fundamentos y Evidencia Científica")
body(doc, "La mesoterapia es una técnica de administración de medicamentos y sustancias bioactivas mediante microinyecciones intradérmicas o subdérmicas superficiales. El principio fundamental es la administración local de sustancias activas en el tejido diana, maximizando la concentración local y minimizando los efectos sistémicos. Las sustancias utilizadas en mesoterapia facial incluyen: cócteles de vitaminas (vitamina C, complejo B, biotina), aminoácidos, ácido hialurónico no reticulado (skinbooster), péptidos bioestimulantes (Matrixyl, Argireline en baja concentración), coenzima Q10, glutatión. La evidencia científica para mesoterapia es limitada (estudios de bajo nivel de evidencia, series de casos) comparada con la evidencia de BTX-A y AH. El médico debe informar al paciente sobre este aspecto de manera transparente.")
body(doc, "TÉCNICAS DE APLICACIÓN en mesoterapia: a) Técnica de papula: inyección intradérmica de 0.01-0.05 mL creando una pápula visible en la piel, espaciadas 0.5-1 cm entre sí, en toda la zona de tratamiento. b) Técnica lineal: avance de la aguja en el plano intradérmico o subdérmico con depósito continuo de producto al retirar. c) Técnica de nappage (ducha): múltiples microinyecciones rápidas y superficiales. Las agujas utilizadas son de calibre 30G-32G, 4 mm de longitud para intradérmico.")
h3(doc, "7.4.2 Protocolo de Mesoterapia Facial")
bullet(doc, "Verificar consentimiento informado específico para mesoterapia.")
bullet(doc, "Aplicar anestesia tópica 30-45 minutos antes.")
bullet(doc, "Limpiar y desinfectar el área con antiséptico.")
bullet(doc, "Preparar el cóctel de mesoterapia según el protocolo indicado para la condición del paciente (envejecimiento, deshidratación, manchas, caída del cabello).")
bullet(doc, "Aplicar las microinyecciones de manera sistemática en toda el área de tratamiento, espaciando 0.5-1 cm entre puntos.")
bullet(doc, "Post-procedimiento: aplicar crema calmante y fotoprotector. Instruir al paciente sobre el eritema transitorio esperado (1-3 días).")
bullet(doc, "Las sesiones de mesoterapia se realizan en series de 3-6 sesiones con intervalos de 1-4 semanas, seguidas de sesiones de mantenimiento cada 1-3 meses.")

# 7.5 PRP
h2(doc, "7.5 PLASMA RICO EN PLAQUETAS (PRP)")
h3(doc, "7.5.1 Fundamento Biológico y Proceso de Obtención")
body(doc, "El PRP es un concentrado autólogo de plaquetas que se obtiene de la sangre periférica del paciente mediante centrifugación. El protocolo de preparación estándar incluye: a) EXTRACCIÓN DE SANGRE: extracción de 10-30 mL de sangre venosa del paciente (usualmente vena antecubital) en tubos especiales para PRP con anticoagulante (citrato de sodio) y activador o sin activador según el kit utilizado. b) CENTRIFUGACIÓN: protocolo bifásico: primera centrifugación a baja velocidad (1000-1200 rpm, 10 min) para separar eritrocitos de plasma y plaquetas; segunda centrifugación del plasma a mayor velocidad (2500-3000 rpm, 10 min) para concentrar las plaquetas. c) EXTRACCIÓN DEL PRP: La fracción rica en plaquetas se ubica en la interfase entre el plasma pobre en plaquetas (superior, se descarta) y los eritrocitos (inferior). Se extrae con jeringa estéril. d) ACTIVACIÓN (opcional): algunos protocolos activan el PRP con cloruro de calcio 10% o trombina bovina justo antes de la inyección, para iniciar la liberación de factores de crecimiento. e) APLICACIÓN: dentro de los 30 minutos siguientes a la preparación, mediante inyecciones intradérmicas o subdérmicas en el área a tratar, o mediante microagujación (microneedling) superficial.")
body(doc, "INDICACIONES PRINCIPALES del PRP en medicina estética: rejuvenecimiento facial (arrugas finas, pérdida de luminosidad, textura irregular), cicatrices de acné superficiales, alopecia androgenética (cuero cabelludo, inyecciones perifoliculares). Para alopecia, series de 3-4 sesiones mensuales con mantenimiento trimestral han mostrado resultados en estudios controlados. CONTRAINDICACIONES ESPECÍFICAS del PRP: trombocitopenia (<150.000 plaquetas/μL), enfermedades hematológicas activas, infección activa sistémica o en el área a tratar, uso de anticoagulantes, hepatitis activa, neoplasias activas.")

# 8. SEGUIMIENTO POST-PROCEDIMIENTO
h1(doc, "8. SEGUIMIENTO POST-PROCEDIMIENTO")
body(doc, "El seguimiento post-procedimiento es una parte esencial del proceso estético que garantiza la seguridad del paciente y la evaluación de resultados. El esquema de seguimiento para cada procedimiento es:")
t_seguim = doc.add_table(rows=1, cols=4)
t_seguim.style = 'Table Grid'
table_header_row(t_seguim, ["Procedimiento", "Control a los 14 días", "Evaluación de Resultados", "Mantenimiento"])
seguim_data = [
    ("Toxina Botulínica", "Control obligatorio para evaluar simetría, resultado y ajustes si se requieren (touch-up). No cobrar consulta de control.", "Fotografías post vs. pre. Escala GAIS.", "Cada 4-6 meses o al retorno de la movilidad muscular."),
    ("Rellenos Dérmicos (AH)", "Control a los 14 días para evaluar integración del producto, edema resuelto, simetría y resultado. Fotografías comparativas.", "Fotografías post vs. pre. Escala GAIS. Volumen estimado retenido.", "Cada 6-18 meses según el producto y el área tratada."),
    ("Peeling Químico", "Control a los 7 días para evaluar proceso de descamación y cicatrización. Indicar crema emoliente y fotoprotector.", "Fotografías a los 30 días. Evaluación de hiperpigmentación, textura, luminosidad.", "Series de 3-6 sesiones mensuales. Mantenimiento cada 3 meses."),
    ("Mesoterapia", "Control entre sesiones. Evaluación de tolerancia y respuesta.", "Fotografías al final de la serie (3-6 sesiones).", "Sesiones de mantenimiento cada 1-3 meses."),
    ("PRP", "Control entre sesiones. Evaluación de tolerancia y respuesta.", "Fotografías al final de la serie. Evaluación de densidad capilar si es para alopecia.", "Mantenimiento cada 3-6 meses."),
]
for i, row_data in enumerate(seguim_data):
    add_table_row(t_seguim, list(row_data), shaded=(i % 2 == 1))

doc.add_paragraph()

# 9. INDICADORES
h1(doc, "9. INDICADORES DE CALIDAD Y SEGURIDAD")
indicadores_est = [
    ("IND-PP2-001: Tasa de Complicaciones por Procedimiento Estético", "(N° de complicaciones reportadas / N° total de procedimientos realizados) × 100. Meta: <3% para complicaciones menores (equimosis, asimetría leve) y 0% para complicaciones graves (oclusión vascular, necrosis, ceguera). Frecuencia: mensual."),
    ("IND-PP2-002: Tasa de Satisfacción del Paciente con Procedimientos Estéticos", "(N° de pacientes con calificación ≥4/5 en encuesta de satisfacción / N° total de pacientes encuestados) × 100. Meta: ≥90%. Frecuencia: mensual."),
    ("IND-PP2-003: Proporción de Procedimientos con Consentimiento Informado Completo", "(N° de procedimientos con CI firmado, fechado y completo en HC / N° total de procedimientos) × 100. Meta: 100%. Frecuencia: mensual."),
    ("IND-PP2-004: Proporción de Procedimientos con Fotografía Clínica Pre-procedimiento", "(N° de procedimientos con fotografía pre-procedimiento en HC / N° total de procedimientos) × 100. Meta: 100%. Frecuencia: mensual."),
    ("IND-PP2-005: Disponibilidad de Materiales de Emergencia", "Verificación mensual de la disponibilidad y vigencia de: hialuronidasa, adrenalina, antihistamínico parenteral, corticoide parenteral. Meta: 100% disponibilidad en toda sesión de procedimientos. Frecuencia: antes de cada jornada de procedimientos."),
]
for nombre, desc in indicadores_est:
    h3(doc, nombre)
    body(doc, desc)
    doc.add_paragraph()

# 10. BIBLIOGRAFÍA
h1(doc, "10. BIBLIOGRAFÍA")
refs_est = [
    "Rohrich RJ, Bartlett EL, Dayan E. Practical Approach and Safety of Hyaluronic Acid Fillers. Plast Reconstr Surg Glob Open. 2019;7(6):e2172.",
    "DeLorenzi C. Complications of injectable fillers, part 2: vascular complications. Aesthet Surg J. 2014;34(4):584-600.",
    "Carruthers J, Carruthers A. Botulinum Toxin A in the Mid and Lower Face and Neck. Dermatol Clin. 2004;22:151-158.",
    "Ministerio de Salud y Protección Social de Colombia. Ley 711 de 2001: reglamenta el ejercicio de la cosmetología y la estética integral. Bogotá: MINSALUD; 2001.",
    "INVIMA Colombia. Circular 01 de 2019: alerta sanitaria sobre rellenos dérmicos y biopolímeros. Bogotá: INVIMA; 2019.",
    "Monheit G, Coleman KM. Hyaluronic Acid Fillers. Dermatol Ther. 2006;19:141-150.",
    "Kontis TC, Rivkin A. The history of injectable facial fillers. Facial Plast Surg. 2009;25(2):67-72.",
    "Alam M, Gladstone H, Kramer EM, et al. ASDS guidelines of care: injectable fillers. Dermatol Surg. 2008;34(suppl 1):S115-148.",
    "Rohrich RJ, Ghavami A, Crosby MA. The role of hyaluronic acid fillers (Restylane) in facial cosmetic surgery: review and technical considerations. Plast Reconstr Surg. 2007;120(6 suppl):41S-54S.",
    "Sclafani AP, Fagien S. Treatment of injectable soft tissue filler complications. Dermatol Surg. 2009;35(suppl 2):1672-1680.",
    "Fabbrocini G, Padova MP, Tosti A. Chemical Peels: What's New and What Isn't New but Still Works Well. Facial Plast Surg. 2009;25(5):329-336.",
    "Giordano CN, Matarasso SL, Ozog DM. Injectable and topical neurotoxins in dermatology: basic science, anatomy, and therapeutic agents. J Am Acad Dermatol. 2017;76(6):1013-1024.",
]
for i, ref in enumerate(refs_est):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    run = p.add_run(f"{i+1}. {ref}")
    run.font.size = Pt(10); run.font.name = 'Calibri'

doc.save(PATH)
print(f"DOC2 guardado: {PATH}")
