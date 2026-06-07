#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera PRO-PP-007 Plan Emergencias y PRO-GC-002 Gestión de Riesgos"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

DARK_BLUE = RGBColor(0, 51, 102)
OUT_DIR_5 = "/home/user/bioauditoria/documentos_habilitacion/CARPETA_5_PROCESOS_PRIORITARIOS"
OUT_DIR_8 = "/home/user/bioauditoria/documentos_habilitacion/CARPETA_8_GESTION_CALIDAD"

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14) if level == 1 else (Pt(12) if level == 2 else Pt(11))
    run.font.color.rgb = DARK_BLUE
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'Calibri'; run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table_header(table, headers):
    row = table.rows[0]
    for i, h in enumerate(headers):
        if i >= len(row.cells): break
        cell = row.cells[i]
        cell.text = h
        set_cell_bg(cell, '003366')
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True; run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(10); run.font.name = 'Calibri'
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_portada(doc, titulo, codigo):
    doc.add_paragraph()
    doc.add_paragraph()
    for txt, sz, bold in [("[NOMBRE DEL CONSULTORIO]",18,True),
                           ("Consultorio Médico General y Procedimientos Estéticos No Invasivos",13,False),
                           (titulo,16,True)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.bold = bold; r.font.size = Pt(sz); r.font.color.rgb = DARK_BLUE; r.font.name = 'Calibri'
        if not bold: doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_table(rows=5, cols=2)
    t.style = 'Table Grid'
    datos = [("Código:",codigo),("Versión:","1.0"),("Fecha:","Junio 2025"),
             ("Responsable:","[NOMBRE DE LA MÉDICA]"),("Dirección:","[DIRECCIÓN DEL CONSULTORIO]")]
    for i, (k,v) in enumerate(datos):
        t.rows[i].cells[0].text = k; t.rows[i].cells[1].text = v
        set_cell_bg(t.rows[i].cells[0], 'BDD7EE')
        for cell in t.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(11)
    doc.add_page_break()

def add_control_versiones(doc):
    add_heading(doc, "CONTROL DE VERSIONES", 1)
    t = doc.add_table(rows=2, cols=5)
    t.style = 'Table Grid'
    add_table_header(t, ["Versión","Fecha","Descripción","Elaboró","Aprobó"])
    t.rows[1].cells[0].text = "1.0"; t.rows[1].cells[1].text = "Junio 2025"
    t.rows[1].cells[2].text = "Creación del documento"
    t.rows[1].cells[3].text = "[NOMBRE DE LA MÉDICA]"; t.rows[1].cells[4].text = "[NOMBRE DE LA MÉDICA]"
    for cell in t.rows[1].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

def setup_doc_with_header(titulo, codigo):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3); sec.right_margin = Cm(2.5)
    header = sec.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear(); hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(f"{titulo} | {codigo}")
    run.font.size = Pt(9); run.font.name = 'Calibri'; run.font.color.rgb = DARK_BLUE
    footer = sec.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = fp.add_run("Versión 1.0 | Junio 2025 | Página ")
    run2.font.size = Pt(9); run2.font.name = 'Calibri'
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run3 = fp.add_run()
    run3._r.append(fldChar1); run3._r.append(instrText); run3._r.append(fldChar2)
    run3.font.size = Pt(9); run3.font.name = 'Calibri'
    return doc

def bullet(doc, text):
    pb = doc.add_paragraph(text, style='List Bullet')
    for run in pb.runs:
        run.font.name='Calibri'; run.font.size=Pt(11)

# =====================================================================
# DOCUMENTO 3: PRO-PP-007 Plan Emergencias
# =====================================================================
def gen_emergencias():
    titulo = "PLAN DE EMERGENCIAS Y CONTINGENCIAS"
    codigo = "PRO-PP-007"
    doc = setup_doc_with_header(titulo, codigo)
    add_portada(doc, titulo, codigo)
    add_control_versiones(doc)

    add_heading(doc, "TABLA DE CONTENIDO", 1)
    for item in ["1. Objetivo","2. Alcance","3. Marco Legal","4. Definiciones",
                 "5. Análisis de Vulnerabilidad y Amenazas","6. Organización para Emergencias",
                 "7. Procedimientos por Tipo de Emergencia","8. Rutas de Evacuación y Señalización",
                 "9. Punto de Encuentro","10. Brigada de Emergencias","11. Simulacros",
                 "12. Recursos para Emergencias","13. Responsables","14. Registros"]:
        add_body(doc, item)
    doc.add_page_break()

    add_heading(doc, "1. OBJETIVO", 1)
    add_body(doc, "Establecer los procedimientos, responsabilidades y recursos para prevenir, preparar la respuesta y recuperarse ante emergencias y contingencias que puedan presentarse en el consultorio [NOMBRE DEL CONSULTORIO], garantizando la protección de la vida e integridad física del personal, los usuarios y visitantes, así como la continuidad de los servicios de salud.")

    add_heading(doc, "2. ALCANCE", 1)
    add_body(doc, "El presente plan aplica a todas las instalaciones del consultorio [NOMBRE DEL CONSULTORIO] ubicado en [DIRECCIÓN DEL CONSULTORIO], a todo el personal que labora en él, a los usuarios y visitantes que se encuentren en las instalaciones en el momento de una emergencia.")

    add_heading(doc, "3. MARCO LEGAL", 1)
    normas = [
        ("Ley 1523 de 2012","Por la cual se adopta la política nacional de gestión del riesgo de desastres y se establece el Sistema Nacional de Gestión del Riesgo de Desastres."),
        ("Decreto 2157 de 2017","Por medio del cual se adoptan directrices generales para la elaboración del plan de gestión del riesgo de desastres de las entidades públicas y privadas."),
        ("NSR-10","Reglamento Colombiano de Construcción Sismo Resistente. Normas de diseño y construcción sismorresistente."),
        ("Resolución 0256 de 2014 (Min. Trabajo)","Por la cual se dictan disposiciones sobre el Sistema de Gestión de la Seguridad y Salud en el Trabajo."),
        ("Decreto 1072 de 2015","Decreto Único Reglamentario del Sector Trabajo. Incluye disposiciones sobre planes de emergencia."),
        ("NFPA 101","Código de Seguridad Humana (referencia técnica para rutas de evacuación y sistemas contra incendios)."),
        ("NTC 1867","Sistemas de señales de evacuación, emergencia y contra incendios."),
    ]
    t = doc.add_table(rows=len(normas)+1, cols=2)
    t.style = 'Table Grid'
    add_table_header(t, ["Norma","Descripción"])
    for i, (n, d) in enumerate(normas):
        t.rows[i+1].cells[0].text = n; t.rows[i+1].cells[1].text = d
        set_cell_bg(t.rows[i+1].cells[0], 'EBF3FB')
        for cell in t.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc, "4. DEFINICIONES", 1)
    defs = [
        ("Emergencia","Situación que pone en peligro inmediato la integridad de las personas, los bienes o el medio ambiente y que requiere de una respuesta inmediata y coordinada."),
        ("Contingencia","Evento incierto o condición que, si ocurre, tendrá un efecto en los objetivos de la organización."),
        ("Evacuación","Desplazamiento planeado, ordenado y protegido de personas en peligro, desde un lugar amenazado hasta un lugar seguro."),
        ("Punto de encuentro","Lugar seguro, predeterminado, exterior a la instalación, donde se reúne el personal evacuado para verificar su integridad."),
        ("Brigada de emergencias","Equipo de personas entrenadas para actuar durante una emergencia."),
        ("Análisis de vulnerabilidad","Proceso mediante el cual se determina el nivel de exposición y predisposición a la pérdida de una instalación frente a una amenaza específica."),
        ("Ruta de evacuación","Camino continuo y sin obstáculos que une cualquier punto de la edificación con el exterior o a un lugar seguro."),
    ]
    for term, defi in defs:
        p = doc.add_paragraph()
        r = p.add_run(f"{term}: "); r.bold = True; r.font.name='Calibri'; r.font.size=Pt(11)
        r2 = p.add_run(defi); r2.font.name='Calibri'; r2.font.size=Pt(11)

    add_heading(doc, "5. ANÁLISIS DE VULNERABILIDAD Y AMENAZAS", 1)
    add_body(doc, "El análisis de vulnerabilidad identifica las amenazas a las que está expuesto el consultorio y determina la capacidad de respuesta ante ellas.")

    add_heading(doc, "5.1 Identificación y Valoración de Amenazas", 2)
    t2 = doc.add_table(rows=9, cols=5)
    t2.style = 'Table Grid'
    add_table_header(t2, ["Amenaza","Origen","Probabilidad\n(1-3)","Impacto\n(1-3)","Nivel de Riesgo"])
    amenazas = [
        ("Incendio","Interno - eléctrico o materiales inflamables","2","3","ALTO"),
        ("Sismo","Natural","2","3","ALTO"),
        ("Derrame de sustancias químicas","Interno - productos estéticos","2","2","MEDIO"),
        ("Emergencia médica dentro del consultorio","Interno - estado del usuario","3","3","ALTO"),
        ("Corte de energía eléctrica","Externo","3","2","MEDIO"),
        ("Inundación","Natural/externo","1","2","BAJO"),
        ("Hurto/seguridad","Externo","2","2","MEDIO"),
        ("Accidente con residuos biológicos","Interno","2","2","MEDIO"),
    ]
    colores_riesgo = {'ALTO': 'FF0000', 'MEDIO': 'FFC000', 'BAJO': '70AD47'}
    for i, row in enumerate(amenazas):
        for j, val in enumerate(row):
            t2.rows[i+1].cells[j].text = val
            for para in t2.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
        nivel = row[4]
        if nivel in colores_riesgo:
            set_cell_bg(t2.rows[i+1].cells[4], colores_riesgo[nivel])
    doc.add_paragraph()

    add_heading(doc, "6. ORGANIZACIÓN PARA EMERGENCIAS", 1)
    add_body(doc, "Dado que el consultorio es una entidad pequeña, el personal asume múltiples roles en caso de emergencia. La organización para emergencias es la siguiente:")

    t3 = doc.add_table(rows=5, cols=3)
    t3.style = 'Table Grid'
    add_table_header(t3, ["Rol","Responsable","Funciones Principales"])
    roles = [
        ("Director de Emergencias","[NOMBRE DE LA MÉDICA]","Coordinar la respuesta ante la emergencia. Declarar la emergencia y activar el plan. Autorizar la evacuación. Establecer comunicación con servicios de emergencia (123, bomberos, Cruz Roja). Informar a las autoridades competentes."),
        ("Jefe de Evacuación","[NOMBRE DE LA MÉDICA] / Auxiliar designado","Dirigir la evacuación ordenada de usuarios y personal. Verificar que todas las personas hayan evacuado. Confirmar el total de personas en el punto de encuentro. Impedir el reingreso hasta que sea seguro."),
        ("Primeros Auxilios","[NOMBRE DE LA MÉDICA] (médica)","Atender a los lesionados en caso de emergencia. Estabilizar a los pacientes hasta la llegada de servicios de emergencia. Coordinar el traslado si es necesario."),
        ("Control de Riesgos","Personal disponible","Intentar controlar la emergencia si es seguro hacerlo (extintores en incendios pequeños). Aislar el área del riesgo. No poner en riesgo la propia vida."),
    ]
    for i, row in enumerate(roles):
        for j, val in enumerate(row):
            t3.rows[i+1].cells[j].text = val
            for para in t3.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
        if i % 2 == 0:
            for j in range(3):
                set_cell_bg(t3.rows[i+1].cells[j], 'EBF3FB')
    doc.add_paragraph()

    add_heading(doc, "7. PROCEDIMIENTOS POR TIPO DE EMERGENCIA", 1)

    emergencias_proc = [
        ("7.1 INCENDIO", [
            ("Detección","Cualquier persona que detecte humo, llamas u olor a quemado debe activar la alarma de emergencia (voz, timbre o alarma instalada)."),
            ("Alerta","Notificar de inmediato al Director de Emergencias. Llamar al Cuerpo de Bomberos: 119."),
            ("Evaluación","Evaluar si el incendio está en etapa inicial y es controlable con extintor. Si no es controlable, evacuar inmediatamente."),
            ("Control (si es posible)","Usar el extintor más cercano apuntando a la base del fuego. Nunca poner en riesgo la vida propia. Si el extintor no controla el incendio en 30 segundos, evacuar."),
            ("Evacuación","Activar la ruta de evacuación. Cerrar puertas y ventanas al salir (sin llave). Apagar equipos eléctricos si es seguro. Nunca usar el ascensor en caso de incendio."),
            ("Verificación en punto de encuentro","Verificar que todas las personas estén en el punto de encuentro. Informar a bomberos sobre posibles personas atrapadas."),
            ("Regreso","Solo reingresar cuando los bomberos lo autoricen."),
        ]),
        ("7.2 SISMO", [
            ("Durante el sismo","Mantener la calma. Alejarse de ventanas, estantes y equipos pesados. Ubicarse bajo escritorios o mesas resistentes, o junto a columnas estructurales. No correr hacia las puertas o escaleras. Proteger la cabeza con los brazos."),
            ("Al terminar el sismo","Verificar el estado de las personas. Cortar el suministro de gas (si aplica) y energía eléctrica. Evaluar daños estructurales antes de moverse."),
            ("Evacuación post-sismo","Salir de forma ordenada si hay daños estructurales, riesgo de incendio o fugas de gas. Usar las rutas de evacuación establecidas. Tener cuidado con escombros y vidrios rotos."),
            ("En el punto de encuentro","Verificar todas las personas. Esperar instrucciones de las autoridades. No reingresar hasta que un especialista certifique la seguridad de la edificación."),
        ]),
        ("7.3 DERRAME DE SUSTANCIAS", [
            ("Identificación","Identificar la sustancia derramada (revisar hojas de seguridad SDS disponibles en el consultorio)."),
            ("Aislamiento","Aislar el área del derrame. Evitar el contacto con la sustancia. Usar los EPP adecuados antes de intervenir."),
            ("Confinamiento","Contener el derrame con material absorbente (aserrín, papel absorbente, kit de derrames). Evitar que llegue a alcantarillas."),
            ("Limpieza","Recoger el material contaminado. Disponer como residuo peligroso (bolsa roja). Limpiar y desinfectar el área."),
            ("Reporte","Registrar el incidente en el formato de reporte de eventos adversos/accidentes. Notificar si hay lesionados."),
        ]),
        ("7.4 EMERGENCIA MÉDICA DENTRO DEL CONSULTORIO", [
            ("Detección","Identificar que un usuario o visitante presenta una emergencia médica (síncope, reacción alérgica, paro cardiorrespiratorio, etc.)."),
            ("Evaluación inicial","La médica evalúa el estado del paciente (estado de consciencia, respiración, pulso)."),
            ("Activación del plan","Si es necesario, llamar al número de emergencias 123 o a la línea de la EPS del paciente para remisión."),
            ("Atención de emergencia","Aplicar los protocolos de atención de urgencias: RCP si hay paro cardiorrespiratorio, administrar oxígeno si hay disponible, iniciar tratamiento de reacciones alérgicas graves (epinefrina si hay anafilaxia)."),
            ("Remisión","Coordinar el traslado al servicio de urgencias más cercano si el caso lo requiere. Entregar información clínica del paciente."),
            ("Registro","Diligenciar el evento en la historia clínica y en el formato de eventos adversos."),
        ]),
        ("7.5 CORTE DE ENERGÍA ELÉCTRICA", [
            ("Inmediato","Mantener la calma. Usar linternas de emergencia (deben estar disponibles y cargadas)."),
            ("Evaluación","Verificar si el corte es interno (falla de circuito) o externo (falla de la red pública). Revisar el tablero eléctrico."),
            ("Procedimiento","Si hay equipos en uso, finalizar de forma segura los procedimientos en curso. Suspender nuevas atenciones que requieran equipos eléctricos. Informar a los usuarios en sala de espera."),
            ("Duración prolongada","Si el corte se prolonga más de 30 minutos y hay procedimientos que no pueden continuar, reprogramar las citas. Verificar la cadena de frío de los biológicos almacenados."),
            ("Reporte","Reportar a la empresa de energía eléctrica. Registrar el evento y sus consecuencias."),
        ]),
    ]

    for titulo_em, pasos in emergencias_proc:
        add_heading(doc, titulo_em, 2)
        t_em = doc.add_table(rows=len(pasos)+1, cols=2)
        t_em.style = 'Table Grid'
        add_table_header(t_em, ["Etapa", "Procedimiento"])
        for i, (etapa, proc) in enumerate(pasos):
            t_em.rows[i+1].cells[0].text = etapa
            t_em.rows[i+1].cells[1].text = proc
            set_cell_bg(t_em.rows[i+1].cells[0], 'BDD7EE' if i%2==0 else 'EBF3FB')
            for cell in t_em.rows[i+1].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name='Calibri'; run.font.size=Pt(9)
        doc.add_paragraph()

    add_heading(doc, "8. RUTAS DE EVACUACIÓN Y SEÑALIZACIÓN", 1)
    add_body(doc, "Las rutas de evacuación deben estar claramente señalizadas con pictogramas normalizados (NTC 1867) y garantizar la salida segura de todas las personas desde cualquier punto de las instalaciones hasta el exterior o punto de encuentro.")

    add_heading(doc, "8.1 Principios de las Rutas de Evacuación", 2)
    for item in [
        "Las rutas de evacuación deben estar siempre despejadas y sin obstáculos.",
        "El ancho mínimo de las rutas de evacuación debe ser de 0.90 metros.",
        "Las puertas de salida deben abrir hacia el exterior o en el sentido de la evacuación.",
        "La señalización debe ser visible en condiciones normales y de emergencia (señales fotoluminiscentes o con iluminación de emergencia).",
        "Deben existir señales de 'SALIDA' en todas las salidas de emergencia.",
        "Las rutas deben estar señalizadas en el plano del consultorio publicado en lugar visible.",
    ]:
        bullet(doc, item)

    add_heading(doc, "8.2 Señalización Requerida", 2)
    t4 = doc.add_table(rows=9, cols=3)
    t4.style = 'Table Grid'
    add_table_header(t4, ["Señal","Ubicación","Norma"])
    senales = [
        ("Salida de emergencia","Sobre puertas de salida y en la ruta de evacuación","NTC 1867"),
        ("Ruta de evacuación (flecha)","A lo largo de la ruta de evacuación, en cada cambio de dirección","NTC 1867"),
        ("Extintor","Junto a cada extintor y en la pared a 1.20 m de altura","NTC 1867"),
        ("Botiquín de primeros auxilios","Junto al botiquín","NTC 1867"),
        ("No correr","En pasillos y escaleras","NTC 1867"),
        ("Punto de encuentro","En el punto de encuentro exterior","NTC 1867"),
        ("Peligro biológico","En el área de almacenamiento de residuos peligrosos","NTC 1867"),
        ("Área restringida","En áreas de acceso restringido","NTC 1867"),
    ]
    for i, row in enumerate(senales):
        for j, val in enumerate(row):
            t4.rows[i+1].cells[j].text = val
            for para in t4.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
        if i%2==0:
            for j in range(3):
                set_cell_bg(t4.rows[i+1].cells[j], 'EBF3FB')
    doc.add_paragraph()

    add_heading(doc, "9. PUNTO DE ENCUENTRO", 1)
    add_body(doc, "El punto de encuentro es el lugar seguro exterior al consultorio donde se reúne todo el personal y los usuarios en caso de evacuación. Para el consultorio [NOMBRE DEL CONSULTORIO], el punto de encuentro se establece en:")
    add_body(doc, "[DEFINIR PUNTO DE ENCUENTRO - Ejemplo: frente al edificio sobre la acera, esquina de [DIRECCIÓN DEL CONSULTORIO], alejado de la fachada del edificio]")
    add_body(doc, "Características del punto de encuentro:")
    for item in [
        "Ubicado en lugar visible y de fácil acceso desde el interior del consultorio.",
        "Alejado de la fachada del edificio para evitar riesgos por caída de objetos.",
        "Señalizado con el pictograma correspondiente.",
        "Conocido por todo el personal y practicado en los simulacros.",
        "Con espacio suficiente para reunir a todo el personal y los usuarios.",
    ]:
        bullet(doc, item)

    add_heading(doc, "10. BRIGADA DE EMERGENCIAS", 1)
    add_body(doc, "Dado el tamaño del consultorio, la brigada de emergencias está conformada por el mismo personal del establecimiento, capacitado en los procedimientos básicos de respuesta ante emergencias.")

    t5 = doc.add_table(rows=4, cols=4)
    t5.style = 'Table Grid'
    add_table_header(t5, ["Cargo en la Brigada","Nombre","Capacitaciones Requeridas","Teléfono de Contacto"])
    brigada = [
        ("Jefe de Brigada","[NOMBRE DE LA MÉDICA]","Primeros auxilios, uso de extintor, evacuación","[TELÉFONO]"),
        ("Brigadista de Evacuación","[Auxiliar/Recepcionista]","Evacuación, uso de extintor","[TELÉFONO]"),
        ("Brigadista de Primeros Auxilios","[NOMBRE DE LA MÉDICA]","Primeros auxilios avanzados (ACLS/ATLS)","[TELÉFONO]"),
    ]
    for i, row in enumerate(brigada):
        for j, val in enumerate(row):
            t5.rows[i+1].cells[j].text = val
            for para in t5.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc, "11. SIMULACROS", 1)
    add_body(doc, "Los simulacros son ejercicios prácticos que permiten evaluar la efectividad del plan de emergencias y entrenar al personal en los procedimientos de respuesta. Se establece el siguiente cronograma de simulacros:")

    t6 = doc.add_table(rows=5, cols=5)
    t6.style = 'Table Grid'
    add_table_header(t6, ["N°","Tipo de Simulacro","Mes Programado","Responsable","Objetivo"])
    simulacros = [
        ("1","Evacuación por incendio","Marzo","[NOMBRE DE LA MÉDICA]","Verificar la ruta de evacuación y el tiempo de salida"),
        ("2","Sismo y evacuación post-sismo","Junio","[NOMBRE DE LA MÉDICA]","Practicar la respuesta ante sismo y posterior evacuación"),
        ("3","Emergencia médica (RCP)","Septiembre","[NOMBRE DE LA MÉDICA]","Practicar el protocolo de atención de paro cardiorrespiratorio"),
        ("4","Evacuación general (sorpresa)","Noviembre","[NOMBRE DE LA MÉDICA]","Evaluar la respuesta real del personal y usuarios"),
    ]
    for i, row in enumerate(simulacros):
        for j, val in enumerate(row):
            t6.rows[i+1].cells[j].text = val
            for para in t6.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
        if i%2==0:
            for j in range(5):
                set_cell_bg(t6.rows[i+1].cells[j], 'EBF3FB')
    doc.add_paragraph()

    add_heading(doc, "12. RECURSOS PARA EMERGENCIAS", 1)
    t7 = doc.add_table(rows=8, cols=4)
    t7.style = 'Table Grid'
    add_table_header(t7, ["Recurso","Cantidad","Ubicación","Responsable de Mantenimiento"])
    recursos = [
        ("Extintor multipropósito ABC (5 kg)","2","Consultorio principal y área de procedimientos","[NOMBRE DE LA MÉDICA]"),
        ("Botiquín de primeros auxilios","1","Área de procedimientos/consultorio","[NOMBRE DE LA MÉDICA]"),
        ("Linterna de emergencia","2","Recepción y consultorio","[NOMBRE DE LA MÉDICA]"),
        ("Camilla rígida o silla de ruedas","1","Área de procedimientos","[NOMBRE DE LA MÉDICA]"),
        ("Directorio de emergencias (visible)","1","Recepción","[NOMBRE DE LA MÉDICA]"),
        ("Plano de evacuación","2","Recepción y consultorio","[NOMBRE DE LA MÉDICA]"),
        ("Kit de derrames","1","Área de almacenamiento de sustancias","[NOMBRE DE LA MÉDICA]"),
    ]
    for i, row in enumerate(recursos):
        for j, val in enumerate(row):
            t7.rows[i+1].cells[j].text = val
            for para in t7.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()
    add_body(doc, "DIRECTORIO DE EMERGENCIAS (debe estar visible en recepción):")
    numeros = [("Emergencias nacionales","123"),("Bomberos","119"),("Cruz Roja","132"),
               ("Defensa Civil","144"),("CRUE (Centro Regulador de Urgencias)","125"),("Policía","112")]
    t_dir = doc.add_table(rows=len(numeros)+1, cols=2)
    t_dir.style = 'Table Grid'
    add_table_header(t_dir, ["Entidad","Número"])
    for i, (ent, num) in enumerate(numeros):
        t_dir.rows[i+1].cells[0].text = ent; t_dir.rows[i+1].cells[1].text = num
        for cell in t_dir.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(11)
    doc.add_paragraph()

    add_heading(doc, "14. REGISTROS", 1)
    t8 = doc.add_table(rows=5, cols=3)
    t8.style = 'Table Grid'
    add_table_header(t8, ["Registro","Código","Tiempo de Retención"])
    regs = [
        ("Acta de simulacro","FOR-PP-007-01","5 años"),
        ("Informe de emergencia","FOR-PP-007-02","5 años"),
        ("Lista de asistencia a capacitación","FOR-PP-007-03","5 años"),
        ("Inspección de recursos de emergencia","FOR-PP-007-04","2 años"),
    ]
    for i, row in enumerate(regs):
        for j, val in enumerate(row):
            t8.rows[i+1].cells[j].text = val
            for para in t8.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)

    path = os.path.join(OUT_DIR_5, "PRO-PP-007_Plan_Emergencias_Contingencias.docx")
    doc.save(path)
    print(f"Guardado: {path}")

# =====================================================================
# DOCUMENTO 4: PRO-GC-002 Gestión de Riesgos
# =====================================================================
def gen_riesgos():
    titulo = "PROCESO DE GESTIÓN DE RIESGOS EN SALUD"
    codigo = "PRO-GC-002"
    doc = setup_doc_with_header(titulo, codigo)
    add_portada(doc, titulo, codigo)
    add_control_versiones(doc)

    add_heading(doc, "TABLA DE CONTENIDO", 1)
    for item in ["1. Objetivo","2. Alcance","3. Marco Legal","4. Definiciones",
                 "5. Metodología de Gestión de Riesgos","6. Identificación de Riesgos",
                 "7. Valoración de Riesgos","8. Mapa de Riesgos del Consultorio",
                 "9. Controles y Barreras de Seguridad","10. Riesgos en Procedimientos Estéticos",
                 "11. Monitoreo y Seguimiento","12. Responsables","13. Registros"]:
        add_body(doc, item)
    doc.add_page_break()

    add_heading(doc, "1. OBJETIVO", 1)
    add_body(doc, "Establecer la metodología para la identificación, análisis, valoración, tratamiento, monitoreo y comunicación de los riesgos en salud del consultorio [NOMBRE DEL CONSULTORIO], con el fin de prevenir la ocurrencia de eventos adversos, proteger la seguridad de los pacientes, el personal y la organización, y mejorar continuamente la calidad de los servicios prestados.")

    add_heading(doc, "2. ALCANCE", 1)
    add_body(doc, "El presente proceso aplica a todos los riesgos clínicos, administrativos, de infraestructura y tecnológicos identificados en el consultorio [NOMBRE DEL CONSULTORIO], ubicado en [DIRECCIÓN DEL CONSULTORIO]. Incluye especialmente los riesgos asociados a los procedimientos de medicina estética no invasiva.")

    add_heading(doc, "3. MARCO LEGAL", 1)
    normas = [
        ("Resolución 2003 de 2014","Define los procedimientos y condiciones de inscripción de prestadores. Incluye criterios de seguridad del paciente."),
        ("Resolución 0256 de 2016","Establece indicadores para el monitoreo de la calidad, incluyendo eventos adversos."),
        ("GTC 45 (ICONTEC)","Guía para la identificación de los peligros y la valoración de los riesgos en seguridad y salud en el trabajo."),
        ("ISO 31000:2018","Norma internacional de gestión de riesgos. Principios y directrices."),
        ("Decreto 1011 de 2006","Sistema Obligatorio de Garantía de Calidad - incluye gestión de riesgos."),
        ("Resolución 3100 de 2019","Estándares de habilitación que incluyen criterios de seguridad del paciente y gestión de riesgos."),
    ]
    t = doc.add_table(rows=len(normas)+1, cols=2)
    t.style = 'Table Grid'
    add_table_header(t, ["Norma","Descripción"])
    for i, (n, d) in enumerate(normas):
        t.rows[i+1].cells[0].text = n; t.rows[i+1].cells[1].text = d
        set_cell_bg(t.rows[i+1].cells[0], 'EBF3FB')
        for cell in t.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc, "4. DEFINICIONES", 1)
    defs = [
        ("Riesgo","Combinación de la probabilidad de que ocurra un evento y sus consecuencias. En salud, posibilidad de que un paciente sufra daño innecesario asociado a la atención sanitaria."),
        ("Evento adverso","Resultado clínico adverso que se deriva de la atención en salud y no de la enfermedad subyacente del paciente."),
        ("Incidente","Evento o circunstancia que podría haber ocasionado u ocasionó daño innecesario a un paciente pero que no llegó a producir consecuencias."),
        ("Barrera de seguridad","Acción o circunstancia que reduce la probabilidad de presentación del incidente o del evento adverso."),
        ("Mapa de riesgos","Representación gráfica de los riesgos identificados, su probabilidad e impacto, que sirve para priorizar las intervenciones."),
        ("Probabilidad","Frecuencia con la que se espera que ocurra el riesgo."),
        ("Impacto","Consecuencias o gravedad del daño que puede ocasionar el riesgo si se materializa."),
    ]
    for term, defi in defs:
        p = doc.add_paragraph()
        r = p.add_run(f"{term}: "); r.bold = True; r.font.name='Calibri'; r.font.size=Pt(11)
        r2 = p.add_run(defi); r2.font.name='Calibri'; r2.font.size=Pt(11)

    add_heading(doc, "5. METODOLOGÍA DE GESTIÓN DE RIESGOS", 1)
    add_body(doc, "La metodología adoptada se basa en los principios de la GTC 45 (ICONTEC) e ISO 31000:2018, adaptada al contexto de los prestadores de servicios de salud del nivel primario. Consta de las siguientes etapas:")

    etapas = [
        ("Identificación","Reconocimiento de todos los riesgos que pueden afectar el cumplimiento de los objetivos del consultorio."),
        ("Análisis","Determinación de la probabilidad de ocurrencia e impacto de cada riesgo identificado."),
        ("Valoración","Combinación de probabilidad e impacto para determinar el nivel del riesgo y definir prioridades."),
        ("Tratamiento","Definición de acciones para eliminar, reducir, transferir o aceptar el riesgo."),
        ("Monitoreo","Seguimiento periódico a los riesgos y a la efectividad de los controles implementados."),
        ("Comunicación","Difusión de los resultados de la gestión de riesgos a los involucrados."),
    ]
    t2 = doc.add_table(rows=len(etapas)+1, cols=2)
    t2.style = 'Table Grid'
    add_table_header(t2, ["Etapa","Descripción"])
    for i, (e, d) in enumerate(etapas):
        t2.rows[i+1].cells[0].text = e; t2.rows[i+1].cells[1].text = d
        set_cell_bg(t2.rows[i+1].cells[0], 'BDD7EE')
        for cell in t2.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc, "5.1 Escala de Probabilidad", 2)
    t3 = doc.add_table(rows=4, cols=3)
    t3.style = 'Table Grid'
    add_table_header(t3, ["Nivel","Descripción","Criterio"])
    prob = [("1 - Baja","El evento puede ocurrir solo en circunstancias excepcionales","Menos de 1 vez al año"),
            ("2 - Media","El evento puede ocurrir en algunas ocasiones","1-6 veces al año"),
            ("3 - Alta","El evento ocurre con frecuencia","Más de 6 veces al año o mensualmente")]
    for i, row in enumerate(prob):
        for j, val in enumerate(row):
            t3.rows[i+1].cells[j].text = val
            for para in t3.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc, "5.2 Escala de Impacto", 2)
    t4 = doc.add_table(rows=4, cols=3)
    t4.style = 'Table Grid'
    add_table_header(t4, ["Nivel","Descripción","Consecuencias"])
    imp = [("1 - Leve","Daño mínimo, sin consecuencias permanentes","Molestia menor, no requiere tratamiento adicional"),
           ("2 - Moderado","Daño que requiere tratamiento pero sin secuelas permanentes","Requiere intervención médica, incapacidad temporal"),
           ("3 - Grave","Daño grave con secuelas permanentes o muerte","Hospitalización, secuela permanente, muerte")]
    for i, row in enumerate(imp):
        for j, val in enumerate(row):
            t4.rows[i+1].cells[j].text = val
            for para in t4.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc, "5.3 Matriz de Valoración del Riesgo", 2)
    add_body(doc, "Nivel de Riesgo = Probabilidad x Impacto")
    t5 = doc.add_table(rows=4, cols=4)
    t5.style = 'Table Grid'
    add_table_header(t5, ["Probabilidad / Impacto","Leve (1)","Moderado (2)","Grave (3)"])
    matriz = [
        ("Alta (3)","3 - MEDIO","6 - ALTO","9 - MUY ALTO"),
        ("Media (2)","2 - BAJO","4 - MEDIO","6 - ALTO"),
        ("Baja (1)","1 - BAJO","2 - BAJO","3 - MEDIO"),
    ]
    nivel_colores = {"1 - BAJO":"70AD47","2 - BAJO":"70AD47","3 - MEDIO":"FFC000",
                     "4 - MEDIO":"FFC000","6 - ALTO":"FF0000","9 - MUY ALTO":"C00000"}
    for i, row in enumerate(matriz):
        set_cell_bg(t5.rows[i+1].cells[0], 'BDD7EE')
        t5.rows[i+1].cells[0].text = row[0]
        for j in range(1, 4):
            t5.rows[i+1].cells[j].text = row[j]
            color = nivel_colores.get(row[j], 'FFFFFF')
            set_cell_bg(t5.rows[i+1].cells[j], color)
        for cell in t5.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10); run.bold=True
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_heading(doc, "6. IDENTIFICACIÓN DE RIESGOS", 1)

    categorias_riesgos = [
        ("6.1 Riesgos Clínicos", [
            ("RC-001","Error diagnóstico","Diagnóstico incorrecto o tardío","2","2","4 - MEDIO"),
            ("RC-002","Error en prescripción","Medicamento o dosis incorrectos","1","3","3 - MEDIO"),
            ("RC-003","Reacción adversa a medicamentos","Reacción no prevista a medicamento administrado","2","2","4 - MEDIO"),
            ("RC-004","Caída del paciente","Caída dentro del consultorio","2","2","4 - MEDIO"),
            ("RC-005","Infección asociada a procedimiento","Infección post-procedimiento","1","3","3 - MEDIO"),
            ("RC-006","Queja o reclamación no resuelta","Insatisfacción del usuario","2","2","4 - MEDIO"),
        ]),
        ("6.2 Riesgos Administrativos", [
            ("RA-001","Incumplimiento de normas de habilitación","Hallazgos en visita de verificación","2","3","6 - ALTO"),
            ("RA-002","Pérdida o alteración de historias clínicas","Daño a registros clínicos","1","3","3 - MEDIO"),
            ("RA-003","Incumplimiento de obligaciones tributarias/legales","Sanciones por incumplimiento","1","3","3 - MEDIO"),
            ("RA-004","Rotación o ausencia de personal clave","Interrupción del servicio","2","2","4 - MEDIO"),
        ]),
        ("6.3 Riesgos de Infraestructura", [
            ("RI-001","Incendio","Daño a instalaciones y personas","1","3","3 - MEDIO"),
            ("RI-002","Sismo","Daño estructural","1","3","3 - MEDIO"),
            ("RI-003","Falla del sistema eléctrico","Interrupción del servicio","2","2","4 - MEDIO"),
            ("RI-004","Corte del servicio de agua","Imposibilidad de cumplir estándares de higiene","2","2","4 - MEDIO"),
        ]),
        ("6.4 Riesgos Tecnológicos", [
            ("RT-001","Falla del sistema de información/historia clínica","Pérdida de datos clínicos","2","2","4 - MEDIO"),
            ("RT-002","Falla de equipos médicos","Imposibilidad de prestar servicio","2","2","4 - MEDIO"),
            ("RT-003","Violación de confidencialidad de datos","Fuga de información de pacientes","1","3","3 - MEDIO"),
        ]),
    ]

    for cat_titulo, riesgos in categorias_riesgos:
        add_heading(doc, cat_titulo, 2)
        t_r = doc.add_table(rows=len(riesgos)+1, cols=6)
        t_r.style = 'Table Grid'
        add_table_header(t_r, ["Código","Riesgo","Descripción","Prob.\n(1-3)","Impacto\n(1-3)","Nivel de Riesgo"])
        for i, row in enumerate(riesgos):
            for j, val in enumerate(row):
                t_r.rows[i+1].cells[j].text = val
                for para in t_r.rows[i+1].cells[j].paragraphs:
                    for run in para.runs:
                        run.font.name='Calibri'; run.font.size=Pt(9)
            nivel = row[5]
            if "ALTO" in nivel:
                set_cell_bg(t_r.rows[i+1].cells[5], 'FF0000')
            elif "MEDIO" in nivel:
                set_cell_bg(t_r.rows[i+1].cells[5], 'FFC000')
            else:
                set_cell_bg(t_r.rows[i+1].cells[5], '70AD47')
        doc.add_paragraph()

    add_heading(doc, "8. MAPA DE RIESGOS DEL CONSULTORIO", 1)
    add_body(doc, "El mapa de riesgos resume todos los riesgos identificados y su nivel de valoración. Los riesgos de nivel ALTO requieren intervención inmediata y controles prioritarios.")

    t_mapa = doc.add_table(rows=14, cols=5)
    t_mapa.style = 'Table Grid'
    add_table_header(t_mapa, ["Código","Riesgo","Categoría","Nivel","Prioridad de Atención"])
    todos_riesgos = [
        ("RA-001","Incumplimiento normas habilitación","Administrativo","6 - ALTO","1 - INMEDIATA"),
        ("RC-003","Reacción adversa a medicamentos","Clínico","4 - MEDIO","2 - CORTO PLAZO"),
        ("RC-001","Error diagnóstico","Clínico","4 - MEDIO","2 - CORTO PLAZO"),
        ("RI-003","Falla del sistema eléctrico","Infraestructura","4 - MEDIO","2 - CORTO PLAZO"),
        ("RT-001","Falla sistema de información","Tecnológico","4 - MEDIO","2 - CORTO PLAZO"),
        ("RC-002","Error en prescripción","Clínico","3 - MEDIO","3 - MEDIANO PLAZO"),
        ("RC-005","Infección asociada a procedimiento","Clínico","3 - MEDIO","3 - MEDIANO PLAZO"),
        ("RA-002","Pérdida de historias clínicas","Administrativo","3 - MEDIO","3 - MEDIANO PLAZO"),
        ("RI-001","Incendio","Infraestructura","3 - MEDIO","3 - MEDIANO PLAZO"),
        ("RI-002","Sismo","Infraestructura","3 - MEDIO","3 - MEDIANO PLAZO"),
        ("RT-003","Violación confidencialidad datos","Tecnológico","3 - MEDIO","3 - MEDIANO PLAZO"),
        ("RC-004","Caída del paciente","Clínico","4 - MEDIO","2 - CORTO PLAZO"),
        ("RA-004","Rotación personal clave","Administrativo","4 - MEDIO","3 - MEDIANO PLAZO"),
    ]
    for i, row in enumerate(todos_riesgos):
        for j, val in enumerate(row):
            t_mapa.rows[i+1].cells[j].text = val
            for para in t_mapa.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
        if "ALTO" in row[3]:
            set_cell_bg(t_mapa.rows[i+1].cells[3], 'FF0000')
        elif "MEDIO" in row[3]:
            set_cell_bg(t_mapa.rows[i+1].cells[3], 'FFC000')
    doc.add_paragraph()

    add_heading(doc, "9. CONTROLES Y BARRERAS DE SEGURIDAD", 1)
    t_ctrl = doc.add_table(rows=9, cols=4)
    t_ctrl.style = 'Table Grid'
    add_table_header(t_ctrl, ["Código Riesgo","Riesgo","Control/Barrera de Seguridad","Tipo de Control"])
    controles = [
        ("RA-001","Incumpl. habilitación","Autoevaluación semestral + lista de chequeo Res.3100","Preventivo"),
        ("RC-001","Error diagnóstico","Uso de guías de práctica clínica + doble verificación","Preventivo"),
        ("RC-002","Error prescripción","Protocolo de prescripción segura + revisión doble","Preventivo"),
        ("RC-003","Reacción adversa","Anamnesis completa de alergias + consentimiento informado","Preventivo"),
        ("RC-004","Caída paciente","Piso antideslizante + acompañamiento en procedimientos","Preventivo"),
        ("RC-005","Infección procedimiento","Protocolo de asepsia y antisepsia + esterilización","Preventivo"),
        ("RI-003","Falla eléctrica","UPS o generador + protocolos de contingencia","Contingencia"),
        ("RT-001","Falla sistema info","Backup diario + manual de contingencia","Contingencia"),
    ]
    for i, row in enumerate(controles):
        for j, val in enumerate(row):
            t_ctrl.rows[i+1].cells[j].text = val
            for para in t_ctrl.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
        if i%2==0:
            for j in range(4):
                set_cell_bg(t_ctrl.rows[i+1].cells[j], 'EBF3FB')
    doc.add_paragraph()

    add_heading(doc, "10. RIESGOS ESPECÍFICOS EN PROCEDIMIENTOS ESTÉTICOS NO INVASIVOS", 1)
    add_body(doc, "Los procedimientos de medicina estética no invasiva que realiza el consultorio presentan riesgos específicos que requieren controles especiales:")

    t_est = doc.add_table(rows=8, cols=5)
    t_est.style = 'Table Grid'
    add_table_header(t_est, ["Procedimiento","Riesgo Específico","Prob.","Impacto","Control/Barrera"])
    riesgos_est = [
        ("Toxina botulínica","Ptosis palpebral post-aplicación","2","2","Técnica correcta, dosis adecuada, consentimiento informado"),
        ("Toxina botulínica","Reacción alérgica anafiláctica","1","3","Anamnesis de alergias, kit de emergencia con epinefrina disponible"),
        ("Ácido hialurónico","Oclusión vascular","1","3","Técnica aspirativa, uso de cánula, consentimiento informado detallado"),
        ("Ácido hialurónico","Necrosis cutánea","1","3","Conocimiento anatómico, hialuronidasa disponible"),
        ("Mesoterapia","Infección local","2","2","Técnica aséptica, material estéril, seguimiento post-procedimiento"),
        ("Peelings químicos","Quemadura química","2","2","Concentración controlada, tiempo de exposición, protocolo de neutralización"),
        ("Todos los procedimientos","Insatisfacción del resultado","3","2","Consentimiento informado con fotografías, expectativas realistas"),
    ]
    for i, row in enumerate(riesgos_est):
        for j, val in enumerate(row):
            t_est.rows[i+1].cells[j].text = val
            for para in t_est.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
    doc.add_paragraph()

    add_heading(doc, "11. MONITOREO Y SEGUIMIENTO", 1)
    t_mon = doc.add_table(rows=5, cols=4)
    t_mon.style = 'Table Grid'
    add_table_header(t_mon, ["Actividad de Monitoreo","Descripción","Periodicidad","Responsable"])
    monitoreo = [
        ("Revisión del mapa de riesgos","Actualizar la valoración de los riesgos y verificar la efectividad de los controles","Semestral","[NOMBRE DE LA MÉDICA]"),
        ("Análisis de eventos adversos","Revisión de todos los eventos adversos e incidentes reportados, identificación de causas raíz","Mensual","[NOMBRE DE LA MÉDICA]"),
        ("Auditoría de controles","Verificar que los controles establecidos estén implementados y funcionando","Trimestral","[NOMBRE DE LA MÉDICA]"),
        ("Informe de gestión de riesgos","Consolidado de riesgos, eventos, controles e indicadores","Semestral","[NOMBRE DE LA MÉDICA]"),
    ]
    for i, row in enumerate(monitoreo):
        for j, val in enumerate(row):
            t_mon.rows[i+1].cells[j].text = val
            for para in t_mon.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
        if i%2==0:
            for j in range(4):
                set_cell_bg(t_mon.rows[i+1].cells[j], 'EBF3FB')

    path = os.path.join(OUT_DIR_8, "PRO-GC-002_Proceso_Gestion_Riesgos.docx")
    doc.save(path)
    print(f"Guardado: {path}")

if __name__ == "__main__":
    gen_emergencias()
    gen_riesgos()
    print("Documentos 3 y 4 completados.")
