
# -*- coding: utf-8 -*-
"""Genera PRO-PP-001 Proceso de Consulta de Medicina General"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

PATH = "/home/user/bioauditoria/documentos_habilitacion/CARPETA_5_PROCESOS_PRIORITARIOS/PRO-PP-001_Proceso_Consulta_Medicina_General.docx"

AZUL_OSCURO = RGBColor(0, 51, 102)
AZUL_MEDIO = RGBColor(21, 101, 192)
BLANCO = RGBColor(255, 255, 255)
GRIS_CLARO = RGBColor(242, 242, 242)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_header_footer(doc, code, title):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = f"{code} | {title}"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.runs[0]
    run.font.size = Pt(9)
    run.font.color.rgb = AZUL_OSCURO
    run.font.name = 'Calibri'
    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)

def add_page_number(paragraph):
    paragraph.clear()
    run = paragraph.add_run("Versión 2.0 | Junio 2025    Página ")
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run2 = paragraph.add_run()
    run2.font.size = Pt(9)
    run2.font.name = 'Calibri'
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    run3 = paragraph.add_run(" de ")
    run3.font.size = Pt(9)
    run3.font.name = 'Calibri'
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = 'NUMPAGES'
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run4 = paragraph.add_run()
    run4.font.size = Pt(9)
    run4.font.name = 'Calibri'
    run4._r.append(fldChar3)
    run4._r.append(instrText2)
    run4._r.append(fldChar4)

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Calibri'
    run.font.color.rgb = AZUL_OSCURO
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    run.font.color.rgb = AZUL_MEDIO
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = AZUL_OSCURO
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p

def table_header_row(table, headers, bg='003366'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        set_cell_bg(cell, bg)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = BLANCO
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table_row(table, values, shaded=False):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        if shaded:
            set_cell_bg(cell, 'F2F2F2')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

doc = Document()

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

add_header_footer(doc, "PRO-PP-001", "Proceso de Consulta de Medicina General")

# ============================================================
# PORTADA
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[LOGO DEL CONSULTORIO]")
run.font.size = Pt(14)
run.font.name = 'Calibri'
run.font.color.rgb = AZUL_OSCURO
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PROCESO DE CONSULTA DE MEDICINA GENERAL")
run.font.size = Pt(20)
run.font.name = 'Calibri'
run.font.color.rgb = AZUL_OSCURO
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[NOMBRE DEL CONSULTORIO]")
run.font.size = Pt(16)
run.font.name = 'Calibri'
run.font.color.rgb = AZUL_MEDIO
run.bold = True

doc.add_paragraph()

# Tabla de datos de portada
t = doc.add_table(rows=7, cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
portada_data = [
    ("Código:", "PRO-PP-001"),
    ("Versión:", "2.0"),
    ("Fecha de Aprobación:", "Junio 2025"),
    ("Elaborado por:", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("Revisado por:", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("Aprobado por:", "Dra. [NOMBRE DE LA MÉDICA] - Directora Médica"),
    ("NIT:", "[NIT DEL CONSULTORIO]"),
]
for i, (k, v) in enumerate(portada_data):
    row = t.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    set_cell_bg(row.cells[0], 'E8EEF4')
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
                run.bold = (cell == row.cells[0])

doc.add_page_break()

# ============================================================
# CONTROL DE VERSIONES
# ============================================================
h1(doc, "CONTROL DE VERSIONES")
t_ver = doc.add_table(rows=3, cols=5)
t_ver.style = 'Table Grid'
table_header_row(t_ver, ["Versión", "Fecha", "Descripción del Cambio", "Elaboró", "Aprobó"])
add_table_row(t_ver, ["1.0", "Enero 2024", "Versión inicial del proceso de consulta de medicina general. Estructuración del flujo de atención, definición de responsabilidades y tiempos estándar.", "Dra. [NOMBRE DE LA MÉDICA]", "Dra. [NOMBRE DE LA MÉDICA]"])
add_table_row(t_ver, ["2.0", "Junio 2025", "Actualización integral según Resolución 3100 de 2019. Incorporación de indicadores de calidad, matriz de riesgos, sección de gestión del consentimiento informado y actualización del marco legal.", "Dra. [NOMBRE DE LA MÉDICA]", "Dra. [NOMBRE DE LA MÉDICA]"], shaded=True)

doc.add_paragraph()

# ============================================================
# 1. INTRODUCCIÓN
# ============================================================
h1(doc, "1. INTRODUCCIÓN")
body(doc, "La consulta médica de primer nivel de atención constituye la puerta de entrada al Sistema General de Seguridad Social en Salud (SGSSS) colombiano. En el marco del modelo de atención integral en salud establecido por el Ministerio de Salud y Protección Social mediante la Política de Atención Integral en Salud (PAIS) y la Resolución 429 de 2016, la consulta de medicina general representa el primer contacto formal y sistemático del paciente con el sistema de salud institucionalizado, siendo históricamente responsable de resolver aproximadamente el 80% de los problemas de salud de la población general, de acuerdo con las estadísticas del Observatorio Nacional de Salud del Instituto Nacional de Salud.")
body(doc, "En Colombia, el médico general que ejerce en el primer nivel de atención cumple un rol fundamental e irremplazable dentro de la arquitectura del sistema de salud. Su función no se limita a la atención curativa de enfermedades agudas, sino que abarca la promoción de la salud, la prevención de la enfermedad, el diagnóstico temprano, el tratamiento oportuno, la rehabilitación básica y la paliación del sufrimiento en condiciones que no requieren manejo especializado. Esta integralidad del quehacer del médico general es reconocida en la Ley 1751 de 2015 (Ley Estatutaria de Salud) como componente esencial del derecho fundamental a la salud.")
body(doc, "El modelo de atención integral en salud, consagrado en la Resolución 429 de 2016, propone un abordaje sistémico e integrado que supera la visión fragmentada y centrada exclusivamente en la enfermedad. Este modelo postula que la atención en salud debe organizarse en torno a las personas, las familias y las comunidades, reconociendo sus determinantes sociales, económicos, culturales y ambientales. Para el consultorio [NOMBRE DEL CONSULTORIO], este enfoque implica que cada consulta de medicina general es una oportunidad no solo de resolver el motivo de consulta inmediato, sino de evaluar el estado de salud global del paciente, identificar factores de riesgo modificables, aplicar actividades preventivas pertinentes y fortalecer la autonomía del paciente en el cuidado de su propia salud.")
body(doc, "Desde la perspectiva de la medicina basada en la evidencia (MBE), la práctica clínica del médico general en primer nivel debe fundamentarse en la mejor evidencia científica disponible, integrada con la experiencia clínica del profesional y las preferencias informadas del paciente. En Colombia, el Ministerio de Salud y Protección Social ha desarrollado Guías de Práctica Clínica (GPC) basadas en evidencia para las principales causas de consulta en primer nivel, como infección de vías urinarias, hipertensión arterial, diabetes mellitus tipo 2, asma bronquial, depresión, entre otras. El presente proceso incorpora el uso sistemático de estas guías como herramienta de toma de decisiones clínicas estandarizadas.")
body(doc, "El contexto institucional del [NOMBRE DEL CONSULTORIO] es el de un establecimiento habilitado de primer nivel de atención, con el propósito de ofrecer servicios de medicina general y procedimientos estéticos no invasivos a la población de [CIUDAD] y su área de influencia. En cumplimiento de los estándares de habilitación establecidos en la Resolución 3100 de 2019, el consultorio cuenta con los recursos humanos, físicos, tecnológicos y de información necesarios para brindar una atención médica segura, oportuna, pertinente, accesible, continua y satisfactoria. La Directora Médica, Dra. [NOMBRE DE LA MÉDICA], con tarjeta profesional N° [N° TP] vigente, es la responsable directa de la atención médica y del cumplimiento de los estándares de calidad establecidos en el presente proceso.")
body(doc, "La relación del consultorio con el Sistema General de Seguridad Social en Salud se establece a través de contratos con Entidades Promotoras de Salud (EPS) del régimen contributivo y subsidiado, así como mediante la atención de pacientes particulares. En todos los casos, el proceso de consulta de medicina general descrito en el presente documento aplica de manera uniforme, garantizando la equidad e igualdad en la atención independientemente del tipo de aseguramiento del paciente. El abordaje biopsicosocial, propuesto por el modelo de Engel y adaptado al contexto colombiano, permite reconocer que los problemas de salud tienen dimensiones biológicas, psicológicas y sociales que deben ser evaluadas e intervenidas de manera integral en cada consulta médica.")
body(doc, "El presente proceso ha sido elaborado con el rigor técnico, normativo y ético que exige la práctica médica responsable en Colombia, incorporando los requisitos de la Resolución 3100 de 2019 en materia de habilitación, los lineamientos del Sistema Obligatorio de Garantía de Calidad en Salud (SOGCS) establecido en el Decreto 1011 de 2006, y las mejores prácticas de gestión de procesos en salud. Su implementación contribuye al mejoramiento continuo de la calidad de la atención en el [NOMBRE DEL CONSULTORIO] y al fortalecimiento de la confianza de los pacientes y la comunidad en sus servicios de salud.")

# ============================================================
# 2. OBJETIVOS
# ============================================================
h1(doc, "2. OBJETIVOS")
h2(doc, "2.1 Objetivo General")
body(doc, "Estandarizar el proceso de atención en consulta de medicina general del [NOMBRE DEL CONSULTORIO], garantizando una atención integral, segura, oportuna y de calidad, fundamentada en la medicina basada en la evidencia y los principios éticos de la práctica médica, en cumplimiento de los estándares de habilitación establecidos en la Resolución 3100 de 2019 del Ministerio de Salud y Protección Social y demás normatividad aplicable vigente.")

h2(doc, "2.2 Objetivos Específicos")
bullet(doc, "Garantizar la atención oportuna del paciente en los tiempos establecidos por la normatividad vigente, con un tiempo máximo de espera para consulta programada de primera vez no superior a tres (3) días hábiles, y de manera inmediata para urgencias menores, en concordancia con las metas establecidas por el Ministerio de Salud y el indicador de oportunidad en la asignación de cita en consulta médica.")
bullet(doc, "Asegurar la elaboración completa, veraz, legible y oportuna de la historia clínica en todos sus componentes obligatorios, en cumplimiento estricto de la Resolución 1995 de 1999 del Ministerio de Salud, garantizando que constituya un documento médico-legal válido y un instrumento efectivo de continuidad asistencial.")
bullet(doc, "Aplicar sistemáticamente el modelo de anamnesis estructurada y el examen físico completo y pertinente en cada consulta, asegurando que la evaluación clínica sea integral, metodológicamente correcta y fundamentada en la semiología médica, con el fin de minimizar errores diagnósticos y garantizar la seguridad del paciente.")
bullet(doc, "Formular diagnósticos basados en la Clasificación Internacional de Enfermedades (CIE-10) con la precisión y el nivel de especificidad requeridos, facilitando el análisis epidemiológico, la trazabilidad clínica y el reporte a las autoridades sanitarias cuando sea procedente según la normatividad de enfermedades de notificación obligatoria.")
bullet(doc, "Elaborar planes de manejo terapéuticos integrales que incluyan componentes farmacológicos y no farmacológicos, sustentados en guías de práctica clínica nacionales vigentes, con prescripción médica que cumpla en su totalidad los requisitos establecidos en la Resolución 1478 de 2006, garantizando el uso racional de medicamentos.")
bullet(doc, "Implementar un proceso sistemático de educación al paciente y su familia en cada consulta, proporcionando información comprensible sobre diagnóstico, pronóstico, plan de manejo, señales de alarma y medidas de autocuidado, contribuyendo al fortalecimiento de la autonomía del paciente y la adherencia terapéutica.")
bullet(doc, "Gestionar eficientemente los procesos de referencia y contrarreferencia hacia niveles superiores de complejidad cuando el estado de salud del paciente lo requiera, garantizando la continuidad y la integralidad de la atención en red, con tiempos de referencia acordes al grado de urgencia de cada caso.")
bullet(doc, "Medir, analizar y mejorar continuamente los indicadores de calidad asociados al proceso de consulta de medicina general, en el marco del Programa de Auditoría para el Mejoramiento de la Calidad (PAMEC) del consultorio, con el propósito de alcanzar y mantener los estándares óptimos de atención en salud.")

# ============================================================
# 3. ALCANCE
# ============================================================
h1(doc, "3. ALCANCE")
h2(doc, "3.1 Población y Servicios Cubiertos")
body(doc, "El presente proceso aplica a la totalidad de las consultas de medicina general que se realicen en el [NOMBRE DEL CONSULTORIO], ubicado en [DIRECCIÓN DEL CONSULTORIO], [CIUDAD], Colombia. Comprende todas las modalidades de atención médica de primer nivel incluyendo:")
bullet(doc, "Consultas de primera vez: paciente que acude por primera vez al consultorio o que no ha tenido consulta en los últimos doce (12) meses.")
bullet(doc, "Consultas de control y seguimiento: paciente con diagnóstico establecido que acude para evaluación de evolución, ajuste terapéutico o verificación de resultados.")
bullet(doc, "Consultas de urgencia menor: paciente con condición aguda que no compromete la vida pero requiere atención médica dentro de las primeras horas.")
bullet(doc, "Valoraciones para certificados médicos: examen médico con fines laborales, deportivos, de estudio o similares.")
bullet(doc, "Valoración inicial previa a procedimientos: evaluación médica previa a la realización de procedimientos estéticos no invasivos en el marco del servicio habilitado.")
bullet(doc, "Consultas de medicina preventiva: aplicación de actividades de detección temprana y protección específica según normas técnicas del Ministerio de Salud.")

h2(doc, "3.2 Exclusiones")
body(doc, "El presente proceso NO aplica en las siguientes situaciones:")
bullet(doc, "Procedimientos estéticos no invasivos (regulados por PRO-PP-002 Proceso de Procedimientos Estéticos No Invasivos).")
bullet(doc, "Emergencias médicas mayores que comprometan la vida del paciente (paro cardiorrespiratorio, politrauma grave, evento cerebrovascular agudo, entre otros), las cuales requieren activación inmediata del Sistema de Emergencias Médicas (SEM) mediante llamado al número 123 y remisión inmediata al servicio de urgencias de la red hospitalaria de mayor complejidad.")
bullet(doc, "Procedimientos diagnósticos invasivos o terapéuticos que requieran habilitación específica del servicio correspondiente.")
bullet(doc, "Atención domiciliaria de pacientes (que cuenta con su propio protocolo específico si aplica).")

# ============================================================
# 4. MARCO LEGAL
# ============================================================
h1(doc, "4. MARCO LEGAL Y NORMATIVO")
body(doc, "El presente proceso se enmarca en el conjunto de disposiciones constitucionales, legales y reglamentarias que regulan la prestación de servicios de salud en Colombia. A continuación se presenta el marco normativo exhaustivo aplicable, con descripción de su relevancia específica para el proceso de consulta de medicina general:")

t_legal = doc.add_table(rows=1, cols=3)
t_legal.style = 'Table Grid'
t_legal.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(t_legal, ["Norma", "Descripción", "Aplicabilidad al Proceso"])

normas = [
    ("Constitución Política de Colombia, Art. 49 (1991)", "Consagra la atención en salud como servicio público a cargo del Estado y establece los principios de eficiencia, universalidad y solidaridad del SGSSS.", "Fundamento constitucional del derecho a la salud y la obligación de prestar atención de calidad."),
    ("Ley 23 de 1981 y Decreto 3380/1981", "Código de Ética Médica. Regula los principios éticos de la profesión médica en Colombia: beneficencia, no maleficencia, autonomía, justicia. Arts. 1, 2, 5, 10, 15, 16, 34 sobre relación médico-paciente, secreto médico y responsabilidad profesional.", "Rige el comportamiento ético del médico en cada acto médico, incluyendo la consulta de medicina general."),
    ("Ley 100 de 1993 (Arts. 153, 156, 185, 186)", "Crea el Sistema General de Seguridad Social en Salud. Establece principios de calidad, eficiencia, integralidad y libre escogencia. Define los planes de beneficios.", "Marco institucional dentro del cual opera el consultorio como prestador de servicios de salud."),
    ("Ley 1122 de 2007", "Reforma al SGSSS. Fortalece el papel de la Superintendencia Nacional de Salud, regula la inspección, vigilancia y control de IPS.", "Establece las reglas de operación como IPS habilitada y los mecanismos de control."),
    ("Ley 1438 de 2011 (Arts. 2, 3, 65, 66)", "Segunda reforma al SGSSS. Introduce el modelo de Atención Primaria en Salud con enfoque familiar y comunitario. Regula el Sistema de Referencia y Contrarreferencia.", "Orienta el modelo de atención del consultorio hacia el enfoque de APS y la coordinación en red."),
    ("Ley 1751 de 2015 - Ley Estatutaria de Salud", "Eleva la salud a derecho fundamental autónomo. Define sus elementos: disponibilidad, aceptabilidad, accesibilidad, calidad. Arts. 2, 6, 10, 17 sobre obligaciones del Estado y los prestadores.", "Obliga al consultorio a garantizar atención oportuna, pertinente, de calidad y sin discriminación."),
    ("Decreto 1011 de 2006 - SOGCS", "Establece el Sistema Obligatorio de Garantía de Calidad en Salud: habilitación, acreditación, auditoría para el mejoramiento y sistema de información.", "Marco general del sistema de calidad dentro del cual se inscribe el presente proceso."),
    ("Resolución 2003 de 2014 (derogada parcialmente)", "Definió los procedimientos y condiciones de inscripción de los Prestadores de Servicios de Salud. Referente histórico.", "Antecedente normativo de la habilitación del consultorio."),
    ("Resolución 3100 de 2019", "Vigente. Define los procedimientos y condiciones de habilitación de servicios de salud. Estándares para consulta de medicina general y procedimientos estéticos.", "Norma principal que rige los estándares mínimos de habilitación que el consultorio debe cumplir."),
    ("Resolución 1995 de 1999", "Establece las normas para el manejo de la historia clínica en Colombia: características, componentes, custodia, confidencialidad y tiempo de conservación.", "Regula la elaboración, gestión y archivo de la historia clínica en cada consulta."),
    ("Resolución 13437 de 1991", "Adopta el Decálogo de los Derechos de los Pacientes: dignidad, información, confidencialidad, autonomía, acceso a segunda opinión.", "Obliga al consultorio a garantizar el ejercicio pleno de los derechos de los pacientes."),
    ("Resolución 1478 de 2006", "Establece los requisitos de la prescripción médica: datos obligatorios, uso de denominación común internacional, prohibición de inducción a prescripción.", "Regula la emisión de fórmulas médicas en el consultorio."),
    ("Resolución 0256 de 2016", "Define los indicadores de calidad del SOGCS de reporte obligatorio por los prestadores de servicios de salud al SISPRO.", "Establece los indicadores de calidad que el consultorio debe medir y reportar."),
    ("Resolución 8430 de 1993", "Normas científicas, técnicas y administrativas para la investigación en salud. Clasifica la investigación según riesgo.", "Aplica cuando se realicen actividades de investigación o cuando se utilicen datos de pacientes con fines estadísticos."),
    ("Resolución 5061 de 1997", "Define los actos médico-quirúrgicos que pueden realizar los médicos según su especialidad y nivel de entrenamiento.", "Delimita el campo de acción del médico general en el consultorio."),
    ("Resolución 1552 de 2013", "Regula la atención a usuarios del Sistema de Salud: tiempos de respuesta a quejas y reclamos, información a usuarios.", "Obliga al consultorio a establecer mecanismos de atención de PQRSF."),
    ("Ley 1581 de 2012 - Habeas Data", "Protección de datos personales. Regula el tratamiento de información sensible, incluyendo datos de salud.", "Rige el manejo confidencial de la información de salud de los pacientes."),
    ("Resolución 429 de 2016 - PAIS", "Adopta la Política de Atención Integral en Salud con enfoque de APS, atención primaria, gestión del riesgo y modelo integral de atención.", "Orienta el modelo de atención del consultorio hacia la integralidad y la gestión del riesgo en salud."),
]

for i, (norma, desc, aplic) in enumerate(normas):
    add_table_row(t_legal, [norma, desc, aplic], shaded=(i % 2 == 1))

doc.add_paragraph()

# ============================================================
# 5. DEFINICIONES
# ============================================================
h1(doc, "5. DEFINICIONES Y GLOSARIO DE TÉRMINOS")
body(doc, "Para efectos del presente proceso y en el marco de la práctica médica en Colombia, se adoptan las siguientes definiciones, basadas en la normatividad vigente, la literatura médica reconocida y las directrices del Ministerio de Salud y Protección Social:")

definiciones = [
    ("Acto Médico", "Toda actividad realizada por un profesional de la medicina legalmente habilitado, tendiente a la curación, rehabilitación o paliación de la enfermedad, o a la promoción de la salud y prevención de la enfermedad. El acto médico está regido por los principios de beneficencia, no maleficencia, autonomía y justicia, y genera responsabilidad ética, legal y civil para el profesional que lo ejecuta."),
    ("Anamnesis", "Interrogatorio clínico sistemático y estructurado que realiza el médico al paciente (o a sus acompañantes cuando el paciente no puede suministrar la información) con el propósito de recopilar información sobre el motivo de consulta, la historia de la enfermedad actual, los antecedentes patológicos personales y familiares, los hábitos de vida y la revisión por sistemas. Constituye el primer y más importante paso del proceso diagnóstico."),
    ("Atención en Salud", "Conjunto de servicios que se prestan al usuario en el marco de los procesos propios del aseguramiento, así como de las actividades, procedimientos e intervenciones asistenciales en las fases de promoción y prevención, diagnóstico, tratamiento y rehabilitación que se prestan a toda la población. (Decreto 1011 de 2006)."),
    ("Auditoría Médica", "Mecanismo sistemático y continuo de evaluación y mejoramiento de la calidad observada, respecto de la calidad esperada, de la atención de salud que reciben los usuarios. Tiene como objetivos el análisis de la práctica clínica, la identificación de oportunidades de mejora y el cumplimiento de los estándares de calidad establecidos. (Decreto 1011 de 2006)."),
    ("Calidad en Salud", "Provisión de servicios de salud a los usuarios individuales y colectivos de manera accesible y equitativa, a través de un nivel profesional óptimo, teniendo en cuenta el balance entre beneficios, riesgos y costos, con el propósito de lograr la adhesión y satisfacción de dichos usuarios. (Decreto 1011 de 2006)."),
    ("Certificado Médico", "Documento expedido por un médico legalmente habilitado que da fe del estado de salud de una persona, de la existencia de una enfermedad o lesión, o de la capacidad o incapacidad para desarrollar determinadas actividades. Tiene valor probatorio ante autoridades administrativas y judiciales y genera responsabilidad médica y legal para quien lo expide."),
    ("CIE-10", "Clasificación Internacional de Enfermedades, décima revisión, de la Organización Mundial de la Salud (OMS). Sistema alfanumérico de codificación diagnóstica utilizado internacionalmente para el registro, análisis y comparación de datos de morbilidad y mortalidad. En Colombia es de uso obligatorio en la elaboración de historias clínicas y reportes estadísticos en salud."),
    ("Consentimiento Informado", "Proceso comunicativo entre el médico y el paciente mediante el cual el paciente recibe información comprensible, veraz y suficiente sobre su diagnóstico, el procedimiento o tratamiento propuesto, sus beneficios, riesgos, alternativas y consecuencias de no realizarlo, y manifiesta libre y voluntariamente su decisión de aceptar o rechazar la intervención propuesta. Es un derecho del paciente y una obligación ética y legal del médico."),
    ("Consulta de Primera Vez", "Atención médica brindada a un paciente que acude por primera vez a la institución prestadora de servicios de salud, o cuyo último contacto con la misma data de más de doce (12) meses. Requiere una evaluación integral que incluye anamnesis completa, examen físico exhaustivo, formulación diagnóstica y plan de manejo."),
    ("Consulta de Control o Seguimiento", "Atención médica brindada a un paciente que ya cuenta con diagnóstico establecido y plan de manejo previo, con el propósito de evaluar la evolución clínica, valorar la adherencia y respuesta al tratamiento, ajustar el plan terapéutico según sea necesario, y reforzar la educación en salud."),
    ("Diagnóstico Presuntivo", "Hipótesis diagnóstica formulada por el médico con base en la información recopilada en la anamnesis y el examen físico, antes de contar con los resultados de los exámenes complementarios. Es el diagnóstico de trabajo inicial que orienta las decisiones clínicas inmediatas."),
    ("Diagnóstico Definitivo", "Conclusión diagnóstica establecida por el médico una vez integrada la información clínica completa (anamnesis, examen físico, resultados de exámenes complementarios). Debe ser expresado en términos de la CIE-10 con el nivel de especificidad que permita la información disponible."),
    ("Diagnóstico Diferencial", "Lista razonada de los diagnósticos posibles que comparten características clínicas similares al cuadro del paciente, ordenados por probabilidad y gravedad. Su formulación explícita en la historia clínica refleja el rigor del proceso diagnóstico y protege al médico y al paciente ante errores de omisión diagnóstica."),
    ("Enfermedad Actual", "Descripción cronológica, ordenada y detallada del problema de salud que motiva la consulta del paciente, incluyendo: fecha de inicio, forma de inicio (súbita o gradual), características semiológicas del síntoma principal, evolución en el tiempo, factores modificadores, síntomas acompañantes, consultas y tratamientos previos para el mismo problema."),
    ("Examen Físico", "Exploración clínica sistemática del cuerpo del paciente realizada por el médico mediante las técnicas de inspección, palpación, percusión y auscultación, con el propósito de identificar signos clínicos que, junto con los síntomas referidos en la anamnesis, permitan formular hipótesis diagnósticas y orientar el plan de estudio y manejo."),
    ("Formulación o Prescripción Médica", "Orden escrita emitida por un médico legalmente habilitado que contiene la indicación de uno o más medicamentos para un paciente específico, con la información necesaria para su dispensación y administración correcta. Debe cumplir los requisitos establecidos en la Resolución 1478 de 2006."),
    ("Historia Clínica", "Documento privado, obligatorio y sometido a reserva, en el cual se registran cronológicamente las condiciones de salud del paciente, los actos médicos y los demás procedimientos ejecutados por el equipo de salud que interviene en su atención. (Resolución 1995 de 1999)."),
    ("Incapacidad Médica", "Documento expedido por el médico tratante que certifica la imposibilidad temporal o permanente del paciente para desarrollar sus actividades laborales habituales como consecuencia de una enfermedad o lesión, e indica el número de días de reposo recomendados. Su expedición debe estar respaldada por hallazgos clínicos documentados en la historia clínica."),
    ("Medicina Basada en la Evidencia (MBE)", "Uso consciente, explícito y juicioso de la mejor evidencia clínica disponible en la toma de decisiones sobre el cuidado de cada paciente individual. Integra la experiencia clínica del médico, la evidencia de la investigación clínica y las preferencias y valores del paciente. (Sackett et al., 1996)."),
    ("Motivo de Consulta", "Razón principal expresada por el paciente, en sus propias palabras, que lo llevó a buscar atención médica en el momento actual. Debe registrarse de manera literal, entre comillas, en la historia clínica. Constituye el punto de partida de la anamnesis."),
    ("Oportunidad en la Consulta", "Posibilidad que tiene el usuario de obtener los servicios que requiere, sin que se presenten retrasos que pongan en riesgo su vida o su salud. Se mide como el tiempo transcurrido entre la solicitud de la cita y la fecha de la consulta efectiva. Meta establecida por el Ministerio de Salud: ≤3 días hábiles para consulta de primera vez de medicina general."),
    ("Plan de Manejo", "Conjunto de decisiones clínicas adoptadas por el médico tras el proceso diagnóstico, que incluye: tratamiento farmacológico, medidas no farmacológicas, solicitud de exámenes complementarios, interconsultas a especialistas, plan educativo para el paciente y programación de seguimiento. Debe ser coherente con el diagnóstico y fundamentado en la mejor evidencia disponible."),
    ("Primer Nivel de Atención", "Primer nivel de contacto de los individuos, la familia y la comunidad con el sistema de salud, llevando lo más cerca posible la atención de salud al lugar donde residen y trabajan las personas. Constituye el primer elemento de un proceso permanente de asistencia sanitaria. (Declaración de Alma-Ata, OMS, 1978). En Colombia corresponde a la atención de baja complejidad."),
    ("Referencia y Contrarreferencia", "Conjunto de procesos, procedimientos y actividades técnicas y administrativas que permiten prestar adecuadamente los servicios de salud al paciente, garantizando la calidad, accesibilidad, oportunidad, continuidad e integralidad de los servicios, en función de la organización de la red de prestadores y las necesidades de salud de la población. (Resolución 1438 de 2011)."),
    ("Seguridad del Paciente", "Conjunto de elementos estructurales, procesos, instrumentos y metodologías basadas en evidencias científicamente probadas que propenden por minimizar el riesgo de sufrir un evento adverso en el proceso de atención de salud o de mitigar sus consecuencias. (Decreto 1011 de 2006)."),
    ("SGSSS", "Sistema General de Seguridad Social en Salud. Sistema establecido por la Ley 100 de 1993 que regula la prestación de servicios de salud en Colombia mediante la articulación de entidades públicas y privadas, aseguradoras (EPS) y prestadoras (IPS), bajo la dirección y vigilancia del Estado."),
    ("Signos Vitales", "Mediciones fisiológicas básicas que indican el estado funcional de los sistemas vitales del organismo: temperatura corporal, frecuencia cardíaca (pulso), frecuencia respiratoria, presión arterial y, según disponibilidad y pertinencia clínica, saturación de oxígeno y glucemia capilar. Su medición y registro son obligatorios en toda consulta médica."),
    ("Triage", "Proceso de valoración clínica preliminar que ordena a los pacientes antes de la atención médica definitiva, según la urgencia que su situación clínica requiere, con el fin de garantizar la atención de las condiciones que son amenaza inmediata para la vida. En primer nivel, se aplica una clasificación simplificada de urgencia vs. consulta electiva."),
    ("Usuario/Paciente", "Toda persona que utiliza, requiere o recibe servicios de salud. El término 'usuario' enfatiza la perspectiva del sistema de salud y los derechos de participación, mientras 'paciente' enfatiza la relación clínica de atención. En el presente documento se usan indistintamente con el mismo sentido."),
    ("Evento Adverso", "Resultado no deseado, no intencional, que pudo haber sido evitado mediante el cumplimiento de los estándares del cuidado asistencial disponibles en un momento determinado. Los eventos adversos pueden ser prevenibles o no prevenibles, y su reporte es un mecanismo fundamental de aprendizaje organizacional en seguridad del paciente."),
]

for term, defi in definiciones:
    p = doc.add_paragraph()
    run_term = p.add_run(f"{term}: ")
    run_term.bold = True
    run_term.font.size = Pt(11)
    run_term.font.name = 'Calibri'
    run_term.font.color.rgb = AZUL_OSCURO
    run_def = p.add_run(defi)
    run_def.font.size = Pt(11)
    run_def.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)

# ============================================================
# 6. DESCRIPCIÓN DEL PROCESO
# ============================================================
h1(doc, "6. DESCRIPCIÓN DETALLADA DEL PROCESO")
body(doc, "El proceso de consulta de medicina general en el [NOMBRE DEL CONSULTORIO] se desarrolla en cuatro fases secuenciales e interrelacionadas: Agendamiento y Programación, Recepción y Verificación Administrativa, Atención Médica, y Cierre de la Consulta. Cada fase comprende pasos específicos con tiempos estándar, responsables definidos y criterios de calidad medibles. A continuación se describe en detalle cada uno de los pasos del proceso:")

h2(doc, "FLUJOGRAMA GENERAL DEL PROCESO")
body(doc, "INICIO → [1.1] Solicitud de cita (presencial/telefónica/digital) → [1.2] Agendamiento y confirmación → [2.1] Recepción del paciente y verificación de identidad → [2.2] Verificación de afiliación al SGSSS → [2.3] Medición de signos vitales → [3.1] Anamnesis completa → [3.2] Examen físico sistemático → [3.3] Análisis clínico y formulación diagnóstica → [3.4] Plan de manejo integral → [3.5] Prescripción médica (si aplica) → [3.6] Solicitud de ayudas diagnósticas (si aplica) → [3.7] Educación al paciente → [4.1] Cierre y firma de historia clínica → [4.2] Entrega de documentos al paciente → [4.3] Programación de seguimiento → FIN")
body(doc, "Puntos de decisión en el flujograma: ¿Requiere atención inmediata? → Sí: activar protocolo de urgencia / No: continuar flujo normal. ¿Diagnóstico resuelto en primer nivel? → Sí: continuar / No: iniciar proceso de referencia. ¿Requiere procedimiento adicional? → Sí: derivar a PRO-PP-002 / No: cerrar consulta.")

h2(doc, "FASE 1: AGENDAMIENTO Y PROGRAMACIÓN")
h3(doc, "Paso 1.1: Solicitud de Cita Médica")
body(doc, "La solicitud de cita médica puede realizarse por tres canales: presencial en el consultorio, telefónica al número [TELÉFONO], o mediante plataformas digitales habilitadas. En todos los casos, la persona que recibe la solicitud (médica o personal administrativo) debe recopilar los siguientes datos mínimos obligatorios para el agendamiento:")
bullet(doc, "Nombre completo del paciente (primer apellido, segundo apellido, primer nombre, segundo nombre).")
bullet(doc, "Número y tipo de documento de identidad (CC, TI, CE, PA, RC, NUIP).")
bullet(doc, "Fecha de nacimiento y edad.")
bullet(doc, "Número de contacto telefónico principal y secundario.")
bullet(doc, "Tipo de aseguramiento: EPS contributivo, EPS subsidiado, particular, otro régimen.")
bullet(doc, "Nombre de la EPS si aplica.")
bullet(doc, "Motivo de la consulta (en términos generales, sin requerir diagnóstico previo).")
bullet(doc, "Tipo de consulta solicitada: primera vez, control, urgencia menor, certificado médico.")
bullet(doc, "Si es consulta de control: fecha de la última consulta y diagnóstico conocido.")
body(doc, "Con base en la información recopilada, se clasifica el tipo de consulta y se asigna la franja horaria correspondiente. Cada consulta de medicina general tiene asignado un tiempo mínimo de 30 minutos para garantizar una atención integral. Para urgencias menores, se reservan espacios de atención prioritaria dentro de la agenda del día. La confirmación de la cita se realiza verbalmente en el momento del agendamiento y, si el paciente lo solicita, mediante mensaje de texto al número registrado. El tiempo estándar para el proceso de agendamiento es de 5 minutos.")

h3(doc, "Paso 1.2: Preparación de la Consulta")
body(doc, "En el día de la consulta programada, y previo al inicio de la jornada laboral, la Dra. [NOMBRE DE LA MÉDICA] o quien corresponda realiza la preparación de la consulta, que incluye: revisión de la agenda del día con los nombres de los pacientes citados, identificación de los pacientes de control para los cuales debe recuperarse la historia clínica previa, verificación de la disponibilidad y correcto funcionamiento de los equipos médicos básicos (tensiómetro calibrado, oxímetro, termómetro, glucómetro, fonendoscopio, linterna, martillo de reflejos), disponibilidad de los formatos de historia clínica impresos si el sistema es físico, o correcto funcionamiento del sistema de historia clínica electrónica si aplica, y verificación del stock de insumos básicos de la consulta (guantes de examen, tapabocas, gel antibacterial, gasas, torundas). El tiempo estándar para la preparación de la consulta es de 10 minutos al inicio de la jornada.")

h2(doc, "FASE 2: RECEPCIÓN Y VERIFICACIÓN ADMINISTRATIVA")
h3(doc, "Paso 2.1: Recepción del Paciente")
body(doc, "Al llegar el paciente al consultorio, la persona que lo recibe (médica o personal administrativo) lo saluda de manera cordial y respetuosa, presentándose por su nombre y cargo. Se verifica la identidad del paciente solicitando la exhibición del documento de identidad original y comparándolo con los datos registrados en la agenda. En caso de discrepancia en los datos, se corrige el registro antes de continuar. Se solicita la tarjeta de afiliación a la EPS o el documento que acredite el tipo de aseguramiento. Si el paciente no cuenta con aseguramiento activo, se le informa sobre las opciones de atención como particular y el costo de la consulta antes de iniciar la atención, respetando su autonomía de decisión.")
body(doc, "La verificación de la afiliación activa al SGSSS se realiza en tiempo real a través de la plataforma ADRES en línea (www.adres.gov.co) o mediante el aplicativo de la EPS correspondiente, cuando el paciente cuenta con orden de atención o autorización. Se registra en el sistema el número de la autorización, si aplica. En caso de que el sistema de verificación en línea no esté disponible, se acepta la tarjeta de afiliación vigente y el número de documento como evidencia provisional de aseguramiento, dejando nota en el registro del motivo por el cual no se pudo verificar en línea. El tiempo estándar para la recepción y verificación administrativa es de 5 a 8 minutos.")

h3(doc, "Paso 2.2: Medición de Signos Vitales")
body(doc, "Antes de ingresar al consultorio para la atención médica, se realiza la medición y registro de los signos vitales básicos. Esta actividad puede ser realizada por la auxiliar de enfermería (si el consultorio cuenta con este recurso) o por la propia médica al inicio de la consulta. Los signos vitales que se miden de manera rutinaria en toda consulta de medicina general son:")
bullet(doc, "Presión Arterial (PA): se mide con tensiómetro aneroide o electrónico calibrado, con el paciente sentado y en reposo por al menos 5 minutos, con el brazo apoyado a la altura del corazón, manguito correctamente posicionado 2-3 cm por encima del pliegue del codo. Se registra como sistólica/diastólica en mmHg. En pacientes hipertensos conocidos o con primera toma elevada, se realizan dos mediciones con 5 minutos de diferencia y se registra el promedio.")
bullet(doc, "Frecuencia Cardíaca (FC): se mide por palpación del pulso radial o cubital durante 60 segundos completos, o mediante oxímetro de pulso. Se registra en latidos por minuto (lpm) y se describe el ritmo (regular/irregular).")
bullet(doc, "Frecuencia Respiratoria (FR): se mide por observación de los movimientos respiratorios durante 60 segundos, idealmente sin que el paciente lo advierta para evitar modificación voluntaria. Se registra en respiraciones por minuto (rpm).")
bullet(doc, "Temperatura corporal: se mide con termómetro digital en axila (tiempo mínimo 60 segundos), oral (60 segundos) o timpánico. Se registra en grados Celsius (°C) y se especifica el sitio de medición.")
bullet(doc, "Saturación de oxígeno (SpO2): se mide con oxímetro de pulso en dedo índice, con el paciente en reposo y sin esmalte de uñas. Se registra como porcentaje (%). Valor normal ≥95%.")
bullet(doc, "Peso: se mide en báscula calibrada con el paciente de pie, sin zapatos, con ropa ligera, en exhalación. Se registra en kilogramos (kg) con un decimal.")
bullet(doc, "Talla: se mide con tallímetro fijo o portátil calibrado, con el paciente de pie, sin zapatos, talones juntos, espalda recta, mirada al frente (plano de Frankfurt). Se registra en centímetros (cm).")
bullet(doc, "Índice de Masa Corporal (IMC): se calcula automáticamente como peso (kg) / talla² (m²). Se registra con un decimal e interpretación según clasificación OMS: bajo peso (<18.5), normal (18.5-24.9), sobrepeso (25-29.9), obesidad grado I (30-34.9), II (35-39.9), III (≥40).")
body(doc, "Todos los signos vitales se registran de manera inmediata en la historia clínica, con la hora de medición. El tiempo estándar para la medición completa de signos vitales es de 5 minutos. Cualquier valor anormal identificado debe ser informado de inmediato a la médica para su evaluación prioritaria.")

h2(doc, "FASE 3: ATENCIÓN MÉDICA")
h3(doc, "Paso 3.1: Anamnesis Completa")
body(doc, "La anamnesis constituye el pilar del proceso diagnóstico en medicina clínica. Se estima que el 70-80% de los diagnósticos pueden establecerse o sospecharse fuertemente con base únicamente en una anamnesis bien realizada. La Dra. [NOMBRE DE LA MÉDICA] realiza la anamnesis de manera sistemática, en un ambiente de privacidad y respeto, utilizando técnicas de entrevista motivacional y comunicación efectiva que favorezcan la confianza del paciente y la revelación completa de su problema de salud. La anamnesis comprende los siguientes componentes obligatorios:")
body(doc, "MOTIVO DE CONSULTA: Se registra textualmente, entre comillas, la razón principal que expresa el paciente para buscar atención médica, junto con el tiempo de evolución. Ejemplos: 'Me duele la cabeza desde hace 3 días', 'Vengo a control de la presión', 'Necesito un certificado médico para el trabajo'. Este dato no se parafrasea ni se interpreta en este punto.")
body(doc, "ENFERMEDAD ACTUAL: Descripción cronológica y detallada del problema de salud principal. Para cada síntoma o conjunto de síntomas se investigan sistemáticamente: a) Inicio: fecha y modo de inicio (súbito, gradual, insidioso); b) Características: calidad del síntoma (tipo de dolor, características del sangrado, etc.); c) Localización: sitio anatómico preciso e irradiación; d) Intensidad: cuantificada mediante la Escala Visual Analógica (EVA) 0-10 para el dolor; e) Duración y periodicidad: continuo, intermitente, episódico; f) Factores que modifican el síntoma: qué lo aumenta, qué lo disminuye; g) Síntomas acompañantes: fiebre, náuseas, vómito, cambios en el apetito, trastornos del sueño, etc.; h) Consultas y tratamientos previos para el mismo problema: medicamentos, dosis, respuesta.")
body(doc, "ANTECEDENTES PERSONALES PATOLÓGICOS: Registro sistemático de todas las enfermedades crónicas, agudas recurrentes o episódicas que el paciente ha padecido a lo largo de su vida, con especial énfasis en: hipertensión arterial, diabetes mellitus, enfermedades cardiovasculares (cardiopatía isquémica, arritmias, insuficiencia cardíaca), enfermedades pulmonares (asma, EPOC), enfermedades oncológicas (tipo, tratamiento, estado actual), enfermedades autoinmunes (lupus, artritis reumatoide), enfermedades neurológicas (epilepsia, migraña, enfermedad cerebrovascular), enfermedades psiquiátricas (depresión, trastorno de ansiedad, esquizofrenia), enfermedades renales, hepáticas, tiroideas. Para cada antecedente se registra la fecha de diagnóstico, el tratamiento actual y el estado de control.")
body(doc, "ANTECEDENTES QUIRÚRGICOS: Lista de todas las intervenciones quirúrgicas realizadas, con el tipo de cirugía, año de realización, hospital o clínica donde se realizó y si hubo complicaciones. La información quirúrgica es relevante para evaluar anatomía alterada, posibles adherencias, y riesgo anestésico-quirúrgico en caso de requerirse futuras intervenciones.")
body(doc, "ANTECEDENTES FARMACOLÓGICOS: Lista completa y detallada de los medicamentos que el paciente consume actualmente, incluyendo: nombre genérico y/o comercial, dosis, vía de administración, frecuencia y tiempo de uso. Se incluyen medicamentos de prescripción médica, automedicación, suplementos vitamínicos, productos herbales y homeopáticos, ya que pueden interactuar con los medicamentos que se van a prescribir.")
body(doc, "ANTECEDENTES ALÉRGICOS: Registro de todas las alergias conocidas a medicamentos, alimentos, materiales (látex, metales) o agentes ambientales (pólenes, ácaros). Para cada alergia se especifica el tipo de reacción presentada (urticaria, angioedema, anafilaxia, exantema, síntomas gastrointestinales) y la gravedad. Se diferencia claramente entre alergia verdadera (mediada por IgE) e intolerancia o efecto adverso no alérgico. Esta información debe estar resaltada en la historia clínica con una señal visual prominente (recuadro rojo o similar) para prevenir errores de medicación.")
body(doc, "ANTECEDENTES GINECO-OBSTÉTRICOS (en pacientes de sexo femenino): Menarquia (edad de inicio de la menstruación), fecha de última menstruación (FUM), características del ciclo menstrual (duración, frecuencia, cantidad), historia obstétrica resumida como G_P_A_V (gestaciones, partos, abortos, vivos), tipo de partos (vaginal/cesárea), complicaciones obstétricas, método de planificación familiar actual, fecha de última citología cervicovaginal y resultado, síntomas climatéricos si aplica, uso de hormonas (anticonceptivos, terapia hormonal de reemplazo).")
body(doc, "ANTECEDENTES FAMILIARES: Enfermedades en familiares de primer grado (padres, hermanos, hijos) y segundo grado (abuelos, tíos) con relevancia genética o epidemiológica: hipertensión arterial, diabetes mellitus tipo 2, cardiopatía isquémica prematura (hombres <55 años, mujeres <65 años), dislipidemias, cáncer (tipo y familiar afectado), enfermedades autoinmunes, enfermedades mentales, muerte súbita, enfermedad cerebrovascular.")
body(doc, "ANTECEDENTES LABORALES: Ocupación actual y previa, tipo de trabajo (sedentario, actividad física), riesgos ocupacionales conocidos (exposición a ruido, vibraciones, sustancias químicas, posiciones ergonómicas inadecuadas, estrés laboral), accidentes de trabajo previos, enfermedades laborales reconocidas.")
body(doc, "ANTECEDENTES SOCIALES: Escolaridad (nivel educativo alcanzado), estrato socioeconómico, condiciones de vivienda (tipo, número de personas que conviven, acceso a servicios públicos), hábitos (tabaquismo: paquetes/año; consumo de alcohol: patrón, frecuencia, cantidad; uso de sustancias psicoactivas), actividad física (tipo, frecuencia, duración), alimentación (patrón general, restricciones dietarias), red de apoyo social (familia, pareja, amigos).")
body(doc, "REVISIÓN POR SISTEMAS: Interrogatorio sistemático sobre síntomas en cada sistema orgánico, independientemente del motivo de consulta, que permite identificar problemas de salud no referidos espontáneamente por el paciente. Se evalúan: sistema cardiovascular (palpitaciones, disnea de esfuerzo, ortopnea, edemas, dolor precordial, síncope); sistema respiratorio (tos, expectoración, hemoptisis, disnea, sibilancias, dolor pleurítico); sistema digestivo (náuseas, vómito, disfagia, pirosis, dolor abdominal, cambios en el hábito intestinal, rectorragia, melenas, ictericia); sistema genitourinario (disuria, polaquiuria, hematuria, secreción uretral o vaginal, disfunción eréctil); sistema osteomuscular (dolor articular, rigidez, inflamación articular, limitación funcional); sistema neurológico (cefalea, mareo, vértigo, parestesias, déficit motor o sensitivo focal, alteraciones del lenguaje, convulsiones); sistema dermatológico (lesiones cutáneas, cambios en piel y mucosas, prurito); sistema endocrino (poliuria, polidipsia, polifagia, intolerancia al calor o frío, cambios de peso involuntarios); dimensión psicológica (estado de ánimo, ansiedad, trastornos del sueño, cambios cognitivos, ideación suicida o autolesiva).")

h3(doc, "Paso 3.2: Examen Físico Sistemático")
body(doc, "El examen físico se realiza de manera sistemática, de céfalo a caudal, empleando las técnicas semiológicas de inspección, palpación, percusión y auscultación. Se realiza con el pleno consentimiento del paciente, garantizando en todo momento su privacidad, dignidad y comodidad. Para el examen de áreas corporales que requieren exposición, se utiliza bata de paciente y se descubre únicamente la región a examinar. El examen físico completo en una consulta de primera vez en adulto comprende:")
body(doc, "INSPECCIÓN GENERAL: Evaluación del aspecto general del paciente: estado de conciencia y orientación (persona, lugar, tiempo), actitud y postura, estado nutricional (normal, caquexia, obesidad), hidratación (mucosas húmedas/secas, signo del pliegue), coloración de piel y mucosas (palidez, ictericia, cianosis central/periférica, eritema), facies (expresión de dolor, ansiedad, rasgos dismórficos), marcha y movimiento espontáneo. Se describe si el paciente luce en buen, regular o mal estado general, con la justificación semiológica de la impresión.")
body(doc, "CABEZA Y CARA: Inspección del cráneo (forma, simetría, deformidades), cuero cabelludo (alopecia, lesiones, seborrea), cara (simetría, parálisis facial, edema periorbitario, lesiones cutáneas). Ojos: conjuntivas (palidez, inyección conjuntival, ictericia escleral), pupilas (isocoria/anisocoria, mióticas/midriáticas, reflejo fotomotor directo y consensual), motilidad ocular extrínseca, agudeza visual si se requiere. Oídos: pabellón auricular, conducto auditivo externo (cerumen, secreción). Nariz: permeabilidad de fosas nasales, tabique (desviaciones), mucosa nasal (congestión, secreción, pólipos). Boca: labios, mucosa bucal, encías, lengua (saburra, ulceraciones), paladar, amígdalas (tamaño, hiperemia, exudado), orofaringe posterior, dentición.")
body(doc, "CUELLO: Inspección y palpación de tiroides (tamaño, consistencia, nódulos, movilidad con la deglución), cadenas ganglionares cervicales (submandibulares, yugulares, posteriores, supraclaviculares: tamaño, consistencia, movilidad, sensibilidad), pulsos carotídeos (simetría, amplitud), ingurgitación yugular (evaluada a 45°), movilidad cervical (flexión, extensión, rotación, lateralización).")
body(doc, "TÓRAX ANTERIOR: Inspección: forma del tórax (normolíneo, pectus excavatum, pectus carinatum, tórax en barril), frecuencia y patrón respiratorio, uso de músculos accesorios, asimetrías. Palpación: expansibilidad (manos en base torácica anterior), frémito vocal (repetición de 'treinta y tres'). Percusión: sonoridad pulmonar bilateral (sonoridad, matidez, timpanismo, hipersonoridad). Auscultación pulmonar: campos anterior, lateral y posterior; murmullo vesicular (presente, disminuido, ausente), ruidos sobreañadidos (sibilancias, roncus, crépitos, frote pleural). Auscultación cardíaca: en los cuatro focos clásicos: aórtico (2° EID), pulmonar (2° EII), tricuspídeo (4° EII) y mitral (5° EII - ápex); se evalúa ritmo (regular/irregular), frecuencia, intensidad y calidad de los ruidos cardíacos (R1 y R2), presencia de ruidos extra (R3, R4, clics), soplos (sistólicos, diastólicos, continuos; localización, irradiación, intensidad en escala de Levine I-VI/VI).")
body(doc, "MAMAS (en pacientes de sexo femenino, con consentimiento explícito): Inspección en reposo y con maniobras (brazos elevados, manos en cadera): simetría, cambios en piel (retracción, edema en piel de naranja, eritema), cambios en el pezón (inversión, eccema). Palpación sistemática en cuatro cuadrantes y región retroareolar en decúbito supino con brazo ipsilateral elevado, búsqueda de nódulos (localización, tamaño, forma, consistencia, límites, movilidad, sensibilidad), expresión del pezón (secreción: tipo, color). Palpación de ganglios axilares, supraclaviculares e infraclaviculares.")
body(doc, "ABDOMEN: Inspección: forma (plano, excavado, globoso, distendido), cicatrices quirúrgicas (localización), herniaciones (umbilical, inguinal, incisional), peristaltismo visible, circulación colateral. Auscultación (antes de palpación): ruidos intestinales (presentes/ausentes, normales, aumentados, metálicos), soplos vasculares (aórtico, renal, ilíaco). Percusión: timpanismo (predominante), matidez hepática y esplénica, matidez desplazable (ascitis). Palpación superficial: temperatura, tono muscular, defensa voluntaria e involuntaria, sensibilidad generalizada. Palpación profunda: masas, organomegalias (hepatomegalia: tamaño en cm por debajo del reborde costal, consistencia, borde, sensibilidad; esplenomegalia: maniobra de Schuster y Middleton), puntos dolorosos específicos (epigastrio, Murphy, McBurney, puntos ureterales, hipogastrio). Signos de irritación peritoneal: Blumberg (rebote), Rovsing. Evaluación de región inguinal: hernias, adenopatías.")
body(doc, "EXTREMIDADES: Inspección: simetría, forma, deformidades articulares (nódulos de Heberden/Bouchard, desviación en ulnar, hallux valgus), atrofia muscular, edemas (localización, simetría, consistencia: blando/duro; escala: +/++/+++/++++, godet positivo/negativo). Palpación: pulsos periféricos (radial, cubital, femoral, poplíteo, tibial posterior, pedio, en escala 0-4+), llenado capilar (normal <2 segundos), temperatura cutánea, varices (localización, extensión). Fuerza muscular de grupos principales en escala de Medical Research Council (0-5). Sensibilidad: táctil, vibratoria (diapasón), propioceptiva. Movilidad articular activa y pasiva de las principales articulaciones.")
body(doc, "SISTEMA NEUROLÓGICO BÁSICO: Estado mental: nivel de conciencia (Escala de Glasgow: apertura ocular 1-4, respuesta verbal 1-5, respuesta motora 1-6; total 3-15), orientación (persona, lugar, tiempo), atención, memoria (eventos recientes y remotos), lenguaje (fluencia, comprensión, denominación). Nervios craneales: evaluación básica de pares II, III, IV, VI (agudeza visual, campo visual por confrontación, movimientos oculares, pupilas), V (sensibilidad facial, músculos masticadores), VII (simetría facial en reposo y movimiento), VIII (acumetría de Weber y Rinne básica), IX y X (úvula, deglución), XII (protrusión lingual). Función motora: tono muscular, fuerza de grupos musculares principales. Reflejos osteotendinosos: bicipital, tricipital, rotuliano, aquileano (escala 0-4+). Cerebelo básico: dedo-nariz, marcha en tándem. Sensibilidad: táctil, vibratoria. Signos meníngeos si se sospecha meningitis: rigidez de nuca, Kernig, Brudzinski.")
body(doc, "PIEL Y FANERAS: Color y homogeneidad. Hidratación (turgencia, elasticidad). Temperatura (distribución, asimetría). Lesiones cutáneas: descripción morfológica precisa (tipo: mácula, pápula, vesícula, ampolla, pústula, nódulo, tumor, úlcera, costra, escama, liquenificación, atrofia, cicatriz; color, tamaño en mm/cm, bordes, distribución: localizada/diseminada/generalizada; número: única, múltiple; configuración: aislada, agrupada, lineal, en diana). Uñas: forma, color, textura (coiloniquia, onicólisis, leuconiquia, clubbing). Cabello: distribución, cantidad, textura.")

h3(doc, "Paso 3.3: Análisis Clínico e Impresión Diagnóstica")
body(doc, "Una vez completada la anamnesis y el examen físico, la Dra. [NOMBRE DE LA MÉDICA] realiza el análisis clínico que integra toda la información recopilada para formular la impresión diagnóstica. Este proceso cognitivo incluye: identificación del problema principal y los problemas secundarios, correlación entre síntomas referidos y signos objetivos encontrados, formulación del diagnóstico principal con el código CIE-10 correspondiente al nivel de especificidad que permita la información disponible, identificación de diagnósticos secundarios activos que requieran manejo simultáneo, elaboración del diagnóstico diferencial razonado con mínimo tres diagnósticos alternativos ordenados por probabilidad, identificación de criterios de gravedad que puedan modificar el plan de manejo o requerir derivación urgente, y evaluación del estado de control de las enfermedades crónicas conocidas.")
body(doc, "La codificación CIE-10 debe realizarse al nivel más específico posible: se prefieren los códigos de cuatro caracteres (ejemplo: J06.9 Infección aguda de las vías respiratorias superiores, no especificada) sobre los de tres caracteres cuando la información clínica lo permite. Para diagnósticos que no cumplen criterios de certeza clínica, se utilizan los modificadores 'presuntivo' o 'en estudio' en el registro de la historia clínica, aunque el código CIE-10 se asigna al diagnóstico más probable. La claridad en el proceso diagnóstico es un elemento de seguridad del paciente y de calidad de la atención médica.")

h3(doc, "Paso 3.4: Plan de Manejo Integral")
body(doc, "El plan de manejo se elabora de manera individualizada para cada paciente, considerando su diagnóstico, comorbilidades, medicamentos actuales, alergias, contexto socioeconómico, preferencias y capacidad de adherencia. Comprende los siguientes componentes:")
body(doc, "PLAN FARMACOLÓGICO: Selección de medicamentos con base en las guías de práctica clínica nacionales vigentes (GPC del Ministerio de Salud), priorizando el uso de medicamentos genéricos del listado del Plan de Beneficios en Salud (PBS). Para cada medicamento se especifica: nombre genérico (denominación común internacional - DCI), concentración, forma farmacéutica, dosis unitaria, vía de administración, frecuencia de administración, duración del tratamiento, e indicaciones especiales de administración (con o sin alimentos, precauciones). Se verifica la ausencia de interacciones farmacológicas relevantes entre los medicamentos prescritos y los medicamentos actuales del paciente, y la ausencia de contraindicaciones por alergias o comorbilidades. El principio del uso racional de medicamentos implica prescribir solo lo necesario, a la dosis correcta, durante el tiempo adecuado.")
body(doc, "PLAN NO FARMACOLÓGICO: Medidas complementarias al tratamiento farmacológico con alto nivel de evidencia: recomendaciones dietarias específicas según el diagnóstico (dieta hiposódica en HTA, dieta para diabéticos, dieta de fácil digestión en gastroenteritis, aumento de líquidos en infección urinaria), indicaciones de actividad física (tipo, intensidad, frecuencia, duración), reposo relativo o absoluto si está indicado (con duración específica), medidas locales (compresas frías/calientes, elevación de miembros, vendaje), y otras medidas higiénico-dietéticas pertinentes según el diagnóstico.")
body(doc, "PLAN EDUCATIVO: Información al paciente y su familia sobre: nombre y naturaleza del diagnóstico en términos comprensibles (sin tecnicismos innecesarios), mecanismo y propósito de cada medicamento prescrito, importancia de la adherencia al tratamiento completo, señales de alarma que requieren consulta urgente antes de la cita programada (para cada diagnóstico se definen señales específicas), medidas preventivas para evitar recurrencia o complicaciones, y cambios en el estilo de vida recomendados. La educación debe ser bidireccional: se verifica la comprensión del paciente mediante preguntas de confirmación ('¿Puede decirme cuántas veces al día tomará este medicamento?').")
body(doc, "AYUDAS DIAGNÓSTICAS: Solicitud de exámenes de laboratorio o imágenes diagnósticas cuando estén indicados clínicamente. Para cada examen solicitado se registra en la historia clínica la justificación clínica (sospecha diagnóstica que sustenta la solicitud). Los exámenes de laboratorio de mayor uso en primer nivel incluyen: hemograma completo con diferencial, glicemia en ayunas, HbA1c, perfil lipídico, creatinina, BUN, ácido úrico, TSH, uroanálisis con sedimento, coproscópico, prueba de embarazo, serología para VIH (con consentimiento informado previo), hepatitis B y C. Los estudios de imagen más frecuentes incluyen: radiografía de tórax PA y lateral, radiografía de extremidades, ecografía abdominal total. Todas las órdenes de laboratorio o imagen se emiten en el formato del consultorio o en el formato de la EPS, con los datos del paciente, el médico, el diagnóstico y la justificación.")
body(doc, "REFERENCIA A NIVEL SUPERIOR: Cuando el diagnóstico o la evolución del paciente supera la capacidad resolutiva del primer nivel o requiere la intervención de un especialista, se inicia el proceso de referencia formal según el PRO-RC-001 (Proceso de Referencia y Contrarreferencia). Se especifica: especialidad a la que se refiere, motivo clínico de la referencia, criterio de urgencia (urgente: referencia en el día o al día siguiente; prioritario: dentro de 3-5 días hábiles; electivo: dentro de 30 días), información clínica que acompaña al paciente (resumen de la historia clínica, resultados de exámenes previos). En caso de que el estado del paciente sea de riesgo vital, se activa el sistema de emergencias (123) y se acompaña al paciente hasta la llegada del servicio de emergencias.")

h3(doc, "Paso 3.5: Prescripción Médica")
body(doc, "La prescripción médica se elabora en cumplimiento estricto de la Resolución 1478 de 2006 del Ministerio de Salud. Los datos obligatorios que debe contener toda prescripción médica son:")
bullet(doc, "Nombre completo del paciente y número de documento de identidad.")
bullet(doc, "Fecha de la prescripción (día, mes, año).")
bullet(doc, "Nombre del medicamento en denominación común internacional (DCI) o nombre genérico. Se permite adicionar el nombre comercial entre paréntesis.")
bullet(doc, "Concentración del medicamento (mg, mcg, UI, %).")
bullet(doc, "Forma farmacéutica (tableta, cápsula, suspensión, crema, solución inyectable, etc.).")
bullet(doc, "Vía de administración (oral, tópica, intramuscular, intravenosa, inhalatoria, etc.).")
bullet(doc, "Dosis unitaria a administrar.")
bullet(doc, "Frecuencia de administración (cada X horas o N veces al día).")
bullet(doc, "Duración del tratamiento (N días, N semanas, hasta nueva orden, de por vida).")
bullet(doc, "Número de unidades a dispensar.")
bullet(doc, "Indicaciones especiales de administración si aplican.")
bullet(doc, "Diagnóstico que justifica la prescripción (código CIE-10 y descripción).")
bullet(doc, "Nombre completo del médico prescriptor.")
bullet(doc, "Número de tarjeta profesional: [N° TP].")
bullet(doc, "Sello del médico con número de tarjeta profesional.")
bullet(doc, "Firma del médico.")
body(doc, "NOTA IMPORTANTE: Está absolutamente prohibida la prescripción de medicamentos de control especial (psicotrópicos, estupefacientes) sin la receta oficial de color establecida por el INVIMA y el MINSALUD. La Dra. [NOMBRE DE LA MÉDICA] podrá prescribir medicamentos de control especial únicamente en las circunstancias clínicas y con los formatos establecidos por la normatividad vigente.")

h2(doc, "FASE 4: CIERRE DE LA CONSULTA")
h3(doc, "Paso 4.1: Cierre y Firma de la Historia Clínica")
body(doc, "Al finalizar la atención médica, la Dra. [NOMBRE DE LA MÉDICA] realiza la revisión final de la historia clínica para verificar la completitud y coherencia de todos los campos registrados: motivo de consulta, anamnesis completa, signos vitales, examen físico, diagnóstico(s) con código(s) CIE-10, plan de manejo completo, y prescripción médica. La historia clínica debe ser legible, clara y no contener borrones, tachones sin enmienda o espacios en blanco. Si se identifica un error de registro, se sigue el procedimiento establecido en el MAN-HC-001: tachar con una línea horizontal, escribir 'error' encima, inicial del médico y fecha, y registrar la información correcta a continuación. NUNCA se utiliza corrector de texto (liquid paper, cinta correctora) en historia clínica. La historia clínica se cierra con la firma y sello del médico, número de tarjeta profesional y fecha y hora de la atención.")

h3(doc, "Paso 4.2: Entrega de Documentos al Paciente")
body(doc, "Al finalizar la consulta, el médico entrega al paciente todos los documentos generados durante la consulta: fórmula médica original (una copia queda en la historia clínica), órdenes de laboratorio y/o imágenes diagnósticas, incapacidad médica si fue otorgada (con la justificación clínica debidamente documentada en la historia clínica), certificado médico si fue solicitado, y material educativo impreso relevante para el diagnóstico (si disponible). El médico verifica verbalmente que el paciente comprende el plan de manejo, los medicamentos prescritos y las señales de alarma para consulta urgente. Se entrega información de contacto del consultorio para consultas urgentes o dudas sobre el tratamiento.")

h3(doc, "Paso 4.3: Programación del Seguimiento")
body(doc, "Al final de cada consulta se establece y comunica al paciente el plan de seguimiento: fecha y hora de la próxima cita de control (se agenda en ese momento si es posible), criterios clínicos para consultar antes de la cita programada (señales de alarma específicas del diagnóstico), y opciones de contacto para consultas entre citas (teléfono del consultorio, si aplica). El tiempo estándar para el cierre y entrega de documentos es de 5 a 8 minutos.")

# ============================================================
# 7. TIEMPOS ESTÁNDAR
# ============================================================
h1(doc, "7. TIEMPOS ESTÁNDAR DEL PROCESO")
t_tiempos = doc.add_table(rows=1, cols=4)
t_tiempos.style = 'Table Grid'
table_header_row(t_tiempos, ["Actividad", "Tiempo Mínimo", "Tiempo Óptimo", "Responsable"])
tiempos = [
    ("Agendamiento y confirmación de cita", "3 min", "5 min", "Médica / Auxiliar Administrativa"),
    ("Preparación de la consulta (inicio de jornada)", "5 min", "10 min", "Médica"),
    ("Recepción del paciente y verificación administrativa", "5 min", "8 min", "Médica / Auxiliar"),
    ("Verificación de afiliación ADRES", "2 min", "3 min", "Médica / Auxiliar"),
    ("Medición de signos vitales completos", "5 min", "7 min", "Médica / Auxiliar de Enfermería"),
    ("Anamnesis completa (primera vez)", "10 min", "15 min", "Médica"),
    ("Anamnesis de control", "5 min", "8 min", "Médica"),
    ("Examen físico completo (primera vez)", "10 min", "15 min", "Médica"),
    ("Examen físico focalizado (control)", "5 min", "8 min", "Médica"),
    ("Análisis clínico y formulación diagnóstica", "5 min", "8 min", "Médica"),
    ("Elaboración del plan de manejo", "5 min", "8 min", "Médica"),
    ("Prescripción médica y órdenes", "3 min", "5 min", "Médica"),
    ("Educación al paciente", "3 min", "5 min", "Médica"),
    ("Cierre y firma de historia clínica", "3 min", "5 min", "Médica"),
    ("Entrega de documentos y programación de control", "3 min", "5 min", "Médica"),
    ("TOTAL CONSULTA DE PRIMERA VEZ", "45 min", "60 min", "Médica"),
    ("TOTAL CONSULTA DE CONTROL", "25 min", "35 min", "Médica"),
    ("TIEMPO MÁXIMO DE ESPERA EN SALA (meta)", "—", "≤20 min", "Médica"),
    ("OPORTUNIDAD ASIGNACIÓN DE CITA (meta)", "—", "≤3 días hábiles", "Médica"),
]
for i, row_data in enumerate(tiempos):
    add_table_row(t_tiempos, list(row_data), shaded=(i % 2 == 1))

doc.add_paragraph()

# ============================================================
# 8. ROLES Y RESPONSABILIDADES
# ============================================================
h1(doc, "8. ROLES Y RESPONSABILIDADES")
t_roles = doc.add_table(rows=1, cols=3)
t_roles.style = 'Table Grid'
table_header_row(t_roles, ["Cargo", "Responsabilidades Específicas en el Proceso", "Requisitos del Cargo"])
roles = [
    ("Médica General - Directora Médica\nDra. [NOMBRE DE LA MÉDICA]\nT.P.: [N° TP]",
     "• Realizar la totalidad del acto médico con estándares de calidad.\n• Elaborar la historia clínica completa en cada consulta.\n• Formular diagnósticos con codificación CIE-10.\n• Elaborar prescripciones médicas según Res. 1478/2006.\n• Solicitar ayudas diagnósticas pertinentes con justificación.\n• Gestionar referencias oportunas cuando se requieran.\n• Educar al paciente en cada consulta.\n• Gestionar el consentimiento informado cuando aplique.\n• Reportar eventos adversos e incidentes de seguridad.\n• Diligenciar el registro de notificación de enfermedades de declaración obligatoria (SIVIGILA) cuando corresponda.\n• Medir y reportar los indicadores de calidad del proceso.\n• Mantener actualizados sus conocimientos clínicos.\n• Verificar la calibración y funcionamiento de los equipos médicos.",
     "• Título de médico y cirujano legalmente reconocido.\n• Tarjeta profesional vigente del Tribunal Ético de Medicina.\n• Certificado de BLS/RCP vigente (no mayor a 2 años).\n• Certificado de curso de urgencias médicas básicas.\n• Inscripción vigente en el RETHUS.\n• Ausencia de sanciones disciplinarias activas ante el Tribunal Ético de Medicina."),
    ("Auxiliar Administrativa (si aplica)",
     "• Recibir y gestionar las solicitudes de citas (presencial, telefónica, digital).\n• Verificar la identidad del paciente al ingreso.\n• Verificar la afiliación activa al SGSSS en plataforma ADRES.\n• Registrar datos de ingreso en el sistema de información.\n• Medir signos vitales básicos según entrenamiento.\n• Archivar la documentación clínica según protocolo.\n• Gestionar las autorizaciones con las EPS cuando aplique.\n• Atender y registrar las PQRSF de los usuarios.\n• Mantener el área de recepción en condiciones adecuadas.",
     "• Técnico o tecnólogo en auxiliar administrativo en salud o área afín.\n• Curso de recepción y atención al usuario en servicios de salud.\n• Manejo básico de sistemas de información.\n• Conocimiento de normatividad básica del SGSSS."),
]
for i, row_data in enumerate(roles):
    add_table_row(t_roles, list(row_data), shaded=(i % 2 == 1))

doc.add_paragraph()

# ============================================================
# 9. RECURSOS NECESARIOS
# ============================================================
h1(doc, "9. RECURSOS NECESARIOS PARA EL PROCESO")
h2(doc, "9.1 Recurso Humano")
body(doc, "El proceso de consulta de medicina general requiere como mínimo un (1) médico general con tarjeta profesional vigente expedida por el Tribunal Ético de Medicina de Colombia y registrado en el Registro Único Nacional del Talento Humano en Salud (RETHUS). De conformidad con la Resolución 3100 de 2019, el médico que presta el servicio de consulta de medicina general debe acreditar los siguientes requisitos adicionales:")
bullet(doc, "Certificado vigente de Soporte Vital Básico (BLS/RCP) de la American Heart Association (AHA) o equivalente reconocido, con vigencia no mayor a dos (2) años.")
bullet(doc, "Certificado de asistencia a cursos de actualización en medicina de urgencias básica o atención inicial del paciente en urgencias (AIAU) dentro de los últimos cinco (5) años.")
bullet(doc, "Educación médica continua documentada: mínimo 20 horas anuales en actividades de actualización clínica pertinentes al servicio prestado.")

h2(doc, "9.2 Equipos Médicos Obligatorios")
body(doc, "De conformidad con los estándares de la Resolución 3100 de 2019 para el servicio de consulta de medicina general, el consultorio debe contar con los siguientes equipos en condiciones de funcionamiento y con mantenimiento y calibración documentados:")
t_equipos = doc.add_table(rows=1, cols=4)
t_equipos.style = 'Table Grid'
table_header_row(t_equipos, ["Equipo", "Especificación Mínima", "Frecuencia de Calibración/Mantenimiento", "Normativa de Referencia"])
equipos = [
    ("Tensiómetro aneroide o electrónico", "Validado clínicamente, tallas de manguito para adulto, adulto obeso y pediátrico", "Calibración anual certificada por proveedor autorizado", "Res. 3100/2019"),
    ("Fonendoscopio", "Biauricular, doble campana, membrana de diafragma", "Revisión mensual de estado físico y sello", "Res. 3100/2019"),
    ("Termómetro digital", "Precisión ±0.1°C, tiempo de respuesta <1 min, con protección de cubierta desechable", "Calibración según protocolo del fabricante (anual)", "Res. 3100/2019"),
    ("Oxímetro de pulso", "Rango SpO2 70-99%, precisión ±2%, con indicador de perfusión", "Verificación mensual con paciente de referencia conocida", "Res. 3100/2019"),
    ("Báscula con tallímetro", "Capacidad mínima 150 kg, precisión 100 g. Tallímetro 0-200 cm, precisión 1 mm", "Calibración semestral con pesas patrón certificadas", "Res. 3100/2019"),
    ("Linterna clínica / Otoscopio básico", "Luz LED de alta intensidad, espéculos desechables o esterilizables", "Verificación mensual de funcionamiento", "Res. 3100/2019"),
    ("Martillo de reflejos", "Tipo Taylor o triangular, mango de goma antideslizante", "Revisión semestral de estado físico", "Res. 3100/2019"),
    ("Glucómetro", "Rango 10-600 mg/dL, control de calidad con soluciones estándar, tiras reactivas vigentes", "Control de calidad mensual con solución estándar", "Res. 3100/2019"),
    ("Camilla de examen", "Tapizado lavable, papel desechable de camilla, ajustable en altura o plegable", "Revisión mensual de estado físico y tapizado", "Res. 3100/2019"),
    ("Electrocardiógrafo (recomendado)", "12 derivaciones, impresión de trazado, interpretación automática", "Calibración anual y prueba de señal mensual", "Recomendado por buenas prácticas clínicas"),
]
for i, row_data in enumerate(equipos):
    add_table_row(t_equipos, list(row_data), shaded=(i % 2 == 1))

doc.add_paragraph()
h2(doc, "9.3 Insumos y Materiales")
bullet(doc, "Guantes de examen de nitrilo en tallas S, M, L (al menos 100 unidades en stock permanente).")
bullet(doc, "Mascarillas quirúrgicas tipo IIR (mínimo 50 unidades en stock).")
bullet(doc, "Gel antibacterial para manos con concentración de alcohol ≥70% (dispensador en cada punto de atención).")
bullet(doc, "Papel desechable para camilla de examen (rollos).")
bullet(doc, "Formatos de historia clínica impresos (si el sistema es físico) o sistema de historia clínica electrónica operativo.")
bullet(doc, "Formatos de prescripción médica membretados con todos los datos del consultorio y del médico.")
bullet(doc, "Formatos de orden de laboratorio e imágenes diagnósticas.")
bullet(doc, "Formatos de incapacidad médica y certificado médico.")
bullet(doc, "Tiras reactivas para glucómetro con fecha de vencimiento vigente.")
bullet(doc, "Lancetas desechables para glucometría.")
bullet(doc, "Torundas de algodón y alcohol antiséptico.")

h2(doc, "9.4 Sistemas de Información")
bullet(doc, "Agenda de pacientes: física o digital, con registro de nombre, documento, tipo de consulta, fecha y hora asignada.")
bullet(doc, "Sistema de historia clínica: físico (archivo organizado por documento de identidad, con acceso restringido) o electrónico (con respaldo periódico, acceso con contraseña, registro de auditoría de accesos).")
bullet(doc, "Conectividad a internet para verificación de afiliación en plataforma ADRES y consulta de guías de práctica clínica.")
bullet(doc, "Teléfono fijo o celular para comunicación con pacientes, EPS y servicios de referencia.")

# ============================================================
# 10. INDICADORES DE CALIDAD
# ============================================================
h1(doc, "10. INDICADORES DE CALIDAD DEL PROCESO")
body(doc, "En cumplimiento de la Resolución 0256 de 2016 y los lineamientos del SOGCS, el presente proceso cuenta con los siguientes indicadores de calidad para su monitorización continua:")

indicadores = [
    {
        "nombre": "IND-PP-001: Oportunidad en la Asignación de Cita en Consulta de Medicina General de Primera Vez",
        "objetivo": "Medir el tiempo transcurrido entre la solicitud de la cita de primera vez y la fecha de la consulta efectiva, para garantizar el acceso oportuno al servicio.",
        "formula": "Promedio de [Fecha de consulta efectiva - Fecha de solicitud de cita] en días hábiles, sobre el total de consultas de primera vez en el período.",
        "fuente": "Registro de agenda del consultorio. Se registra la fecha de solicitud y la fecha de consulta en cada agendamiento.",
        "frecuencia": "Mensual",
        "responsable": "Dra. [NOMBRE DE LA MÉDICA]",
        "meta": "≤3 días hábiles (promedio mensual)",
        "umbral": "Alerta cuando el promedio supere 4 días hábiles en cualquier semana del mes.",
        "acciones": "Si se supera la meta: análisis de causa (sobredemanda, fallas de agenda, ausencia del médico), ajuste de agenda para ampliar disponibilidad, comunicación a usuarios sobre tiempos de espera, consideración de extensión de horario de atención.",
    },
    {
        "nombre": "IND-PP-002: Tiempo Promedio de Espera en Sala de Pacientes",
        "objetivo": "Medir el tiempo que el paciente espera desde su llegada al consultorio hasta el inicio efectivo de la atención médica, para garantizar la satisfacción y el respeto del tiempo del usuario.",
        "formula": "Promedio de [Hora de inicio de atención médica - Hora de llegada al consultorio] en minutos, sobre el total de consultas en el período.",
        "fuente": "Registro de hora de llegada (firmado por el paciente en la hoja de recepción) y hora de inicio de consulta (registrada en la historia clínica).",
        "frecuencia": "Mensual",
        "responsable": "Dra. [NOMBRE DE LA MÉDICA]",
        "meta": "≤20 minutos (promedio mensual)",
        "umbral": "Alerta cuando el promedio supere 25 minutos en cualquier semana.",
        "acciones": "Análisis de causas de demora (problemas administrativos, equipos, agenda sobrecargada), optimización del proceso de recepción, ajuste de la agenda con mayor holgura entre consultas, implementación de mensajes de cortesía a pacientes en espera prolongada.",
    },
    {
        "nombre": "IND-PP-003: Satisfacción Global del Usuario con el Servicio de Consulta",
        "objetivo": "Medir el grado de satisfacción de los usuarios con la totalidad del proceso de consulta de medicina general, como indicador de la calidad percibida del servicio.",
        "formula": "(Número de usuarios con calificación 'Satisfecho' o 'Muy Satisfecho' / Total de usuarios encuestados) × 100. Se aplica encuesta de satisfacción validada al finalizar la consulta o mediante formulario digital.",
        "fuente": "Encuesta de satisfacción FOR-GC-005, aplicada al 100% de los pacientes o muestra representativa ≥30% mensual.",
        "frecuencia": "Mensual",
        "responsable": "Dra. [NOMBRE DE LA MÉDICA]",
        "meta": "≥90% de usuarios satisfechos o muy satisfechos",
        "umbral": "Alerta cuando la satisfacción caiga por debajo del 85% en cualquier mes.",
        "acciones": "Análisis detallado de las causas de insatisfacción por componente (trato, tiempo de espera, información, instalaciones), plan de mejora con acciones concretas y responsables, seguimiento mensual hasta recuperar la meta.",
    },
    {
        "nombre": "IND-PP-004: Tasa de Historias Clínicas con Todos los Campos Obligatorios Diligenciados",
        "objetivo": "Garantizar que el 100% de las historias clínicas generadas en el proceso de consulta cumplan con los requisitos de completitud establecidos en la Resolución 1995 de 1999.",
        "formula": "(Número de HC con todos los campos obligatorios completos en la muestra auditada / Total de HC auditadas) × 100.",
        "fuente": "Auditoría de historias clínicas mediante instrumento FOR-HC-003, muestra mínima de 10 HC por mes seleccionadas aleatoriamente.",
        "frecuencia": "Mensual",
        "responsable": "Dra. [NOMBRE DE LA MÉDICA]",
        "meta": "100% de cumplimiento de campos obligatorios",
        "umbral": "Cualquier HC con campos obligatorios incompletos activa alerta inmediata.",
        "acciones": "Identificación del campo o campos faltantes, retroalimentación inmediata a la médica, análisis del patrón de incumplimiento, refuerzo del proceso de cierre de HC, registro del hallazgo en el PAMEC.",
    },
    {
        "nombre": "IND-PP-005: Proporción de Prescripciones que Cumplen los Requisitos de la Resolución 1478 de 2006",
        "objetivo": "Verificar que el 100% de las prescripciones médicas emitidas en el consultorio cumplan con todos los requisitos legales establecidos en la Resolución 1478 de 2006.",
        "formula": "(Número de prescripciones con todos los requisitos legales completos en la muestra auditada / Total de prescripciones auditadas) × 100.",
        "fuente": "Auditoría de fórmulas médicas archivadas, muestra de 10 prescripciones por mes seleccionadas aleatoriamente.",
        "frecuencia": "Mensual",
        "responsable": "Dra. [NOMBRE DE LA MÉDICA]",
        "meta": "100% de cumplimiento",
        "umbral": "Cualquier prescripción incompleta activa alerta.",
        "acciones": "Identificación del requisito incumplido, retroalimentación inmediata, revisión del formato de prescripción, capacitación si es necesario, seguimiento en el siguiente ciclo de auditoría.",
    },
]

for ind in indicadores:
    h3(doc, ind["nombre"])
    t_ind = doc.add_table(rows=8, cols=2)
    t_ind.style = 'Table Grid'
    campos = [
        ("Objetivo del Indicador", ind["objetivo"]),
        ("Fórmula de Cálculo", ind["formula"]),
        ("Fuente de Datos", ind["fuente"]),
        ("Frecuencia de Medición", ind["frecuencia"]),
        ("Responsable de Medición", ind["responsable"]),
        ("Meta", ind["meta"]),
        ("Umbral de Alerta", ind["umbral"]),
        ("Acciones ante Incumplimiento", ind["acciones"]),
    ]
    for i, (campo, valor) in enumerate(campos):
        row = t_ind.rows[i]
        row.cells[0].text = campo
        row.cells[1].text = valor
        set_cell_bg(row.cells[0], 'E8EEF4')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
        row.cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

# ============================================================
# 11. GESTIÓN DE RIESGOS
# ============================================================
h1(doc, "11. GESTIÓN DE RIESGOS DEL PROCESO")
body(doc, "La siguiente matriz identifica los principales riesgos asociados al proceso de consulta de medicina general, su valoración y los controles establecidos, en el marco de la política de seguridad del paciente del [NOMBRE DEL CONSULTORIO]:")

t_riesgos = doc.add_table(rows=1, cols=8)
t_riesgos.style = 'Table Grid'
table_header_row(t_riesgos, ["Riesgo", "Causa Probable", "Consecuencia", "P (1-5)", "I (1-5)", "Nivel", "Control Existente", "Control Propuesto"])

riesgos = [
    ("Error diagnóstico por anamnesis incompleta", "Presión de tiempo, fatiga del médico, paciente con dificultad para comunicar síntomas", "Diagnóstico erróneo, tratamiento inadecuado, daño al paciente, responsabilidad médica", "3", "5", "ALTO", "Lista de verificación mental de componentes de anamnesis", "Formato estructurado de anamnesis; tiempo mínimo de 30 min por consulta en agenda"),
    ("Error en prescripción médica", "Letra ilegible, nombre incorrecto del medicamento, dosis incorrecta, omisión de datos obligatorios", "Dispensación incorrecta, daño al paciente por medicamento equivocado o dosis incorrecta", "3", "4", "ALTO", "Revisión antes de entregar la fórmula al paciente", "Prescripción electrónica; verificación verbal con el paciente del nombre y dosis del medicamento prescrito"),
    ("Reacción alérgica a medicamento prescrito", "Alergias no registradas en la historia clínica, paciente no informa alergias", "Reacción anafiláctica, urticaria, angioedema, daño grave al paciente", "2", "5", "ALTO", "Interrogatorio rutinario de alergias en anamnesis; registro prominente en HC", "Campo de alergias con fondo de color rojo en la HC; preguntar alergias en cada consulta antes de prescribir"),
    ("Demora en el inicio de la atención (tiempo de espera prolongado)", "Sobredemanda, consultas más largas de lo previsto, agenda mal organizada", "Insatisfacción del usuario, abandono de la consulta, deterioro de la condición del paciente en urgencia menor", "3", "3", "MODERADO", "Agenda con espacios de 30 minutos por consulta; gestión de urgencias con prioridad", "Monitorización diaria del indicador IND-PP-002; comunicación proactiva al paciente sobre el tiempo de espera estimado"),
    ("Falla en la gestión de referencia oportuna", "No reconocimiento de criterios de referencia urgente, dificultad para conseguir citas con especialistas", "Progresión de la enfermedad, complicaciones evitables, daño grave o muerte del paciente", "2", "5", "ALTO", "Conocimiento de criterios de referencia según diagnóstico; contacto directo con servicios de urgencias en casos graves", "Protocolo escrito de criterios de referencia por diagnóstico; directorio actualizado de servicios de referencia y urgencias de la red"),
    ("Pérdida de confidencialidad de información clínica", "Acceso no autorizado a la historia clínica, conversaciones en área pública, divulgación inadvertida", "Violación del derecho a la privacidad del paciente, daño reputacional, sanciones legales", "2", "4", "MODERADO", "Historia clínica en archivo bajo llave; consulta en área privada", "Política de protección de datos firmada por todo el personal; cifrado de información digital; protocolo de acceso con contraseña"),
    ("Accidente biológico del personal de salud", "Pinchazo con aguja, exposición a fluidos corporales durante el examen físico o procedimiento", "Exposición a VIH, hepatitis B, hepatitis C, otras infecciones transmisibles por sangre", "2", "4", "MODERADO", "Uso de guantes en procedimientos con riesgo de exposición a fluidos", "Protocolo de accidente biológico escrito y accesible; esquema de vacunación completo del personal (hepatitis B, influenza, tétanos)"),
    ("Deterioro o falla de equipos médicos durante la consulta", "Falta de mantenimiento preventivo, batería agotada, uso inadecuado del equipo", "Datos clínicos incorrectos (diagnóstico erróneo), demora en la atención, imagen de falta de profesionalismo", "2", "3", "MODERADO", "Verificación diaria de equipos al inicio de la jornada", "Cronograma de mantenimiento preventivo documentado; registro de calibración de equipos; equipos de respaldo para los de mayor criticidad (tensiómetro, glucómetro)"),
]

for i, row_data in enumerate(riesgos):
    add_table_row(t_riesgos, list(row_data), shaded=(i % 2 == 1))

doc.add_paragraph()

# ============================================================
# 12. DOCUMENTOS RELACIONADOS
# ============================================================
h1(doc, "12. DOCUMENTOS RELACIONADOS")
t_docs = doc.add_table(rows=1, cols=3)
t_docs.style = 'Table Grid'
table_header_row(t_docs, ["Código", "Nombre del Documento", "Tipo"])
docs_rel = [
    ("PRO-PP-002", "Proceso de Procedimientos Estéticos No Invasivos", "Proceso"),
    ("PRO-PP-003", "Proceso de Seguridad del Paciente", "Proceso"),
    ("PRO-PP-004", "Proceso de Prevención de Infecciones Asociadas a la Atención en Salud", "Proceso"),
    ("PRO-PP-005", "PAMEC - Programa de Auditoría para el Mejoramiento de la Calidad", "Proceso"),
    ("PRO-RC-001", "Proceso de Referencia y Contrarreferencia", "Proceso"),
    ("MAN-TH-001", "Manual de Gestión del Talento Humano", "Manual"),
    ("MAN-HC-001", "Manual de Historia Clínica", "Manual"),
    ("MAN-PP-001", "Manual de Atención al Paciente", "Manual"),
    ("FOR-PP-001", "Consentimiento Informado para Consulta de Medicina General", "Formato"),
    ("FOR-HC-001", "Historia Clínica de Medicina General", "Formato"),
    ("FOR-HC-003", "Instrumento de Auditoría de Historias Clínicas", "Formato"),
    ("FOR-PP-005", "Formato de Reporte de Evento Adverso", "Formato"),
    ("FOR-GC-005", "Encuesta de Satisfacción del Usuario", "Formato"),
    ("PER-TH-001", "Perfil del Cargo de Médico General", "Perfil"),
    ("PRO-TH-001", "Proceso de Selección y Contratación de Personal", "Proceso"),
    ("FOR-TH-001", "Formato de Evaluación de Desempeño", "Formato"),
]
for i, row_data in enumerate(docs_rel):
    add_table_row(t_docs, list(row_data), shaded=(i % 2 == 1))

doc.add_paragraph()

# ============================================================
# 13. BIBLIOGRAFÍA
# ============================================================
h1(doc, "13. BIBLIOGRAFÍA")
referencias = [
    "Ministerio de Salud y Protección Social de Colombia. Resolución 3100 de 2019: Por la cual se definen los procedimientos y condiciones de inscripción de los Prestadores de Servicios de Salud y de habilitación de los servicios de salud. Bogotá: MINSALUD; 2019.",
    "Ministerio de Salud y Protección Social de Colombia. Resolución 1995 de 1999: Por la cual se establecen normas para el manejo de la Historia Clínica. Bogotá: MINSALUD; 1999.",
    "Ministerio de Salud y Protección Social de Colombia. Resolución 1478 de 2006: Por la cual se expiden normas para el control, seguimiento y vigilancia de la importación, exportación, procesamiento, síntesis, producción, distribución, dispensación, compra, venta, destrucción y uso de sustancias sometidas a fiscalización. Bogotá: MINSALUD; 2006.",
    "Ministerio de Salud y Protección Social de Colombia. Política de Atención Integral en Salud (PAIS). Resolución 429 de 2016. Bogotá: MINSALUD; 2016.",
    "Ministerio de Salud y Protección Social de Colombia. Decreto 1011 de 2006: Sistema Obligatorio de Garantía de Calidad en Salud. Bogotá: MINSALUD; 2006.",
    "Sackett DL, Rosenberg WM, Gray JA, Haynes RB, Richardson WS. Evidence based medicine: what it is and what it isn't. BMJ. 1996;312(7023):71-2.",
    "Organización Mundial de la Salud (OMS). Declaración de Alma-Ata sobre Atención Primaria en Salud. Conferencia Internacional sobre Atención Primaria de Salud. Alma-Ata, URSS: OMS; 1978.",
    "Bickley LS, Szilagyi PG. Bates' Guide to Physical Examination and History Taking. 12th ed. Philadelphia: Wolters Kluwer; 2017.",
    "Swartz MH. Tratado de Semiología: Anamnesis y Exploración. 7ª ed. Barcelona: Elsevier; 2015.",
    "Congreso de la República de Colombia. Ley 1751 de 2015: Ley Estatutaria de Salud. Por medio de la cual se regula el derecho fundamental a la salud. Bogotá: Congreso de la República; 2015.",
    "Instituto Nacional de Salud de Colombia. Guía para la Notificación de Eventos de Interés en Salud Pública (EISP). Bogotá: INS; 2023.",
    "Ministerio de Salud y Protección Social de Colombia. Guías de Práctica Clínica basadas en evidencia. Disponibles en: https://www.minsalud.gov.co/salud/publica/Paginas/guias-practica-clinica.aspx",
]
for i, ref in enumerate(referencias):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    run = p.add_run(f"{i+1}. {ref}")
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

# ============================================================
# 14. FIRMAS DE APROBACIÓN
# ============================================================
h1(doc, "14. APROBACIÓN Y FIRMAS")
doc.add_paragraph()
t_firmas = doc.add_table(rows=3, cols=3)
t_firmas.style = 'Table Grid'
table_header_row(t_firmas, ["Elaboró", "Revisó", "Aprobó"])
firma_row = t_firmas.rows[1]
for cell in firma_row.cells:
    cell.text = "\n\n_______________________\n"
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
firma_row2 = t_firmas.rows[2]
nombres = [
    "Dra. [NOMBRE DE LA MÉDICA]\nT.P.: [N° TP]\nMédica General",
    "Dra. [NOMBRE DE LA MÉDICA]\nT.P.: [N° TP]\nMédica General",
    "Dra. [NOMBRE DE LA MÉDICA]\nT.P.: [N° TP]\nDirectora Médica",
]
for i, nombre in enumerate(nombres):
    firma_row2.cells[i].text = nombre
    firma_row2.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for para in firma_row2.cells[i].paragraphs:
        for run in para.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri'

doc.save(PATH)
print(f"DOC1 guardado: {PATH}")
