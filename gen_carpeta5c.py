#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Carpeta 5 - Parte C: PRO-PP-003, PRO-PP-004, Consentimientos y Reporte de Eventos"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

BASE5="/home/user/bioauditoria/documentos_habilitacion/CARPETA_5_PROCESOS_PRIORITARIOS"

def margins(doc):
    for s in doc.sections:
        s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5); s.left_margin=Cm(3); s.right_margin=Cm(2.5)

def hf(doc,code,title):
    for s in doc.sections:
        h=s.header; h.paragraphs[0].clear()
        r=h.paragraphs[0].add_run(f"CONSULTORIO [NOMBRE DEL CONSULTORIO]  |  {code}  |  {title}")
        r.font.size=Pt(8); r.font.color.rgb=RGBColor(64,64,64)
        h.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        f=s.footer; f.paragraphs[0].clear()
        r2=f.paragraphs[0].add_run(f"{code} - V1.0 - {datetime.date.today().strftime('%d/%m/%Y')} | Pág. ")
        r2.font.size=Pt(8)
        fld=OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'),'begin')
        f.paragraphs[0].runs[-1]._r.append(fld)
        it=OxmlElement('w:instrText'); it.text='PAGE'
        f.paragraphs[0].runs[-1]._r.append(it)
        fld2=OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'),'end')
        f.paragraphs[0].runs[-1]._r.append(fld2)
        f.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

def portada(doc,code,title):
    doc.add_paragraph(); doc.add_paragraph()
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=t.add_run("CONSULTORIO MÉDICO\n[NOMBRE DEL CONSULTORIO]")
    r.font.size=Pt(18); r.font.bold=True; r.font.color.rgb=RGBColor(31,73,125)
    doc.add_paragraph()
    t2=doc.add_paragraph(); t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=t2.add_run(title); r2.font.size=Pt(14); r2.font.bold=True
    doc.add_paragraph()
    tb=doc.add_table(rows=6,cols=2); tb.style='Table Grid'; tb.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(l,v) in enumerate([("Código:",code),("Versión:","1.0"),("Fecha:",datetime.date.today().strftime('%d/%m/%Y')),
           ("Elaboró:","[NOMBRE DE LA MÉDICA], Médica General"),
           ("Revisó:","[NOMBRE DE LA MÉDICA]"),("Aprobó:","[NOMBRE DE LA MÉDICA]")]):
        tb.rows[i].cells[0].text=l; tb.rows[i].cells[1].text=v
        tb.rows[i].cells[0].paragraphs[0].runs[0].font.bold=True
    doc.add_page_break()
    tb2=doc.add_table(rows=3,cols=5); tb2.style='Table Grid'
    hrow(tb2,["Versión","Fecha","Cambio","Elaboró","Aprobó"])
    tb2.rows[1].cells[0].text="1.0"; tb2.rows[1].cells[1].text=datetime.date.today().strftime('%d/%m/%Y')
    tb2.rows[1].cells[2].text="Versión inicial"; tb2.rows[1].cells[3].text="[NOMBRE DE LA MÉDICA]"
    tb2.rows[1].cells[4].text="[NOMBRE DE LA MÉDICA]"
    doc.add_page_break()

def hrow(table,headers,bg="1F497D"):
    row=table.rows[0]
    for i,h in enumerate(headers):
        cell=row.cells[i]; cell.text=""
        run=cell.paragraphs[0].add_run(h)
        run.font.bold=True; run.font.color.rgb=RGBColor(255,255,255); run.font.size=Pt(10)
        cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),bg)
        shd.set(qn('w:color'),'auto'); shd.set(qn('w:val'),'clear'); tcPr.append(shd)

def H(doc,text,lv=1):
    p=doc.add_heading(text,level=lv)
    colors={1:RGBColor(31,73,125),2:RGBColor(46,116,181),3:RGBColor(68,114,196)}
    sizes={1:14,2:12,3:11}
    for run in p.runs:
        run.font.size=Pt(sizes[lv]); run.font.bold=True; run.font.color.rgb=colors[lv]
    return p

def B(doc,text,bold=False):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=p.add_run(text); r.font.size=Pt(11); r.font.bold=bold; r.font.name='Calibri'; return p

def BL(doc,text):
    p=doc.add_paragraph(style='List Bullet'); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=p.add_run(text); r.font.size=Pt(11); r.font.name='Calibri'; return p

def firmas(doc):
    doc.add_paragraph()
    tb=doc.add_table(rows=2,cols=3); tb.style='Table Grid'
    hrow(tb,["ELABORÓ","REVISÓ","APROBÓ"])
    for j in range(3):
        tb.rows[1].cells[j].text="[NOMBRE DE LA MÉDICA]\nMédica General - Propietaria\n\nFirma: _________________________\n\nFecha: "+datetime.date.today().strftime('%d/%m/%Y')
        tb.rows[1].cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

# ===== PRO-PP-003: Seguridad del Paciente =====
def crear_pro_pp003():
    doc=Document(); margins(doc); hf(doc,"PRO-PP-003","Proceso de Seguridad del Paciente")
    portada(doc,"PRO-PP-003","PROCESO DE SEGURIDAD DEL PACIENTE")

    H(doc,"1. OBJETIVO")
    B(doc,"Establecer las políticas, estrategias y herramientas para garantizar la seguridad de los pacientes atendidos en el Consultorio [NOMBRE DEL CONSULTORIO], previniendo la ocurrencia de eventos adversos e incidentes, y promoviendo una cultura de seguridad basada en el aprendizaje continuo y la mejora de los procesos asistenciales, en cumplimiento de la Resolución 0256 de 2016 y los lineamientos del Sistema Obligatorio de Garantía de Calidad.")

    H(doc,"2. MARCO LEGAL")
    for n in ["Resolución 0256 de 2016 - Sistema de Información para la Calidad, indicadores de seguridad",
              "Decreto 1011 de 2006 - SOGCS, seguridad del paciente como componente",
              "Resolución 2003 de 2014 - Habilitación, estándares de seguridad",
              "Resolución 3100 de 2019 - Actualización estándares habilitación",
              "OMS - Nueve soluciones para la seguridad del paciente (2007)",
              "Ley 1751 de 2015 - Ley Estatutaria de Salud"]: BL(doc,n)

    H(doc,"3. POLÍTICA DE SEGURIDAD DEL PACIENTE")
    B(doc,"El Consultorio [NOMBRE DEL CONSULTORIO], bajo la dirección de la Dra. [NOMBRE DE LA MÉDICA], se compromete a:")
    pol_seg=["Brindar atención médica con los más altos estándares de seguridad",
             "Prevenir activamente la ocurrencia de eventos adversos a través de la implementación de buenas prácticas de seguridad",
             "Promover una cultura de reporte abierto y no punitivo de los incidentes y eventos adversos",
             "Aprender de los errores para mejorar continuamente los procesos de atención",
             "Informar a los pacientes sobre los riesgos inherentes a los procedimientos y los eventos adversos que ocurran",
             "Implementar las metas internacionales de seguridad del paciente adaptadas al contexto del consultorio"]
    for item in pol_seg: BL(doc,item)

    H(doc,"4. METAS DE SEGURIDAD DEL PACIENTE")
    H(doc,"4.1 Meta 1: Identificación Correcta del Paciente",2)
    B(doc,"Se verificará la identidad del paciente con MÍNIMO DOS IDENTIFICADORES antes de cualquier procedimiento o administración de medicamentos:")
    ident=["Nombre completo del paciente (tal como está en su documento de identidad)",
           "Número de documento de identidad (cédula, TI, pasaporte, etc.)",
           "Los identificadores se verifican consultando el documento físico del paciente o la historia clínica",
           "NUNCA se usa el número de habitación o cama como identificador (no aplica en consultorio, pero el principio se mantiene)",
           "Para procedimientos estéticos: verificar adicionalmente el área anatómica a tratar (coincide con lo descrito en el CI y la HC)"]
    for item in ident: BL(doc,item)

    H(doc,"4.2 Meta 2: Comunicación Efectiva",2)
    B(doc,"La comunicación efectiva entre el médico y el paciente, y entre el médico y otros prestadores de salud, es fundamental para la seguridad:")
    com_ef=["Técnica SBAR (Situación, Background/Antecedentes, Valoración, Recomendación) para comunicar información clínica a otros prestadores",
            "Comunicación completa del diagnóstico y plan al paciente antes de que abandone el consultorio",
            "Verificar comprensión del paciente: pedir al paciente que repita las instrucciones más importantes (teach-back)",
            "Documentar en la historia clínica todas las comunicaciones relevantes, incluyendo indicaciones dadas al paciente",
            "Para remisiones: incluir en el formulario de remisión toda la información necesaria para la continuidad del cuidado"]
    for item in com_ef: BL(doc,item)

    H(doc,"4.3 Meta 3: Medicamentos de Alto Riesgo",2)
    B(doc,"Los medicamentos de alto riesgo disponibles en el consultorio requieren manejo especial:")
    med_alto_riesgo=[
        ("Toxina botulínica tipo A","Verificar concentración/dilución antes de administrar. Verificar sitio de inyección en el mapa anatómico. No intercambiar entre marcas sin ajustar dosis (unidades NO equivalentes entre marcas)."),
        ("Ácido hialurónico (rellenos)","Verificar que es el producto correcto (densidad adecuada para la zona a tratar). Tener hialuronidasa disponible siempre."),
        ("Adrenalina del botiquín","Verificar concentración (1:1000 para IM en anafilaxia) vs (1:10000 para IV - NO usar IV sin entrenamiento)."),
        ("Agentes para peeling","Verificar concentración antes de aplicar. Nunca confundir TCA 10% con TCA 30% - etiquetar claramente."),
    ]
    tb=doc.add_table(rows=len(med_alto_riesgo)+1,cols=2); tb.style='Table Grid'
    hrow(tb,["Medicamento/dispositivo de alto riesgo","Medida de seguridad específica"])
    for i,(m,s) in enumerate(med_alto_riesgo):
        tb.rows[i+1].cells[0].text=m; tb.rows[i+1].cells[1].text=s
        for c in tb.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"4.4 Meta 4: Prevención de Infecciones Asociadas",2)
    B(doc,"La prevención de infecciones es una prioridad de seguridad. Las acciones clave son:")
    prev_inf=["Lavado de manos con técnica OMS de 6 pasos ANTES y DESPUÉS de cada paciente (ver PRO-PP-004)",
              "Uso de guantes estériles para procedimientos invasivos (inyecciones, peelings, PRP)",
              "Antisepsia adecuada de la piel antes de cada inyección",
              "Material estéril de un solo uso para cada paciente (agujas, jeringas, cánulas)",
              "Desinfección adecuada de superficies de trabajo entre pacientes",
              "Gestión correcta de residuos cortopunzantes (guardianes)"]
    for item in prev_inf: BL(doc,item)

    H(doc,"5. REPORTE Y ANÁLISIS DE EVENTOS ADVERSOS")
    H(doc,"5.1 Definiciones",2)
    defs=[("Evento adverso:","Daño involuntario o complicación que ocurre como resultado de la atención de salud y no de la enfermedad subyacente del paciente, que resulta en incapacidad, hospitalización o muerte."),
          ("Incidente:","Evento que no llega a ocasionar daño al paciente pero que podría haberlo hecho. También se llama 'near miss' o casi-falla."),
          ("Evento centinela:","Evento adverso grave, inesperado, que requiere investigación inmediata (muerte inesperada, pérdida permanente de función, accidente quirúrgico grave)."),
          ("Causa raíz:","Factores fundamentales que dan origen a un problema, cuya eliminación previene la recurrencia del evento.")]
    for term,defi in defs:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r1=p.add_run(term+" "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(defi); r2.font.size=Pt(11)

    H(doc,"5.2 Proceso de Reporte",2)
    B(doc,"Todo el personal del consultorio (incluyendo la médica propietaria) tiene el DEBER de reportar incidentes y eventos adversos utilizando el formato FOR-PP-005. El reporte es de cultura no punitiva: el objetivo es aprender y mejorar, no sancionar.")
    reporte=[("Paso 1","Identificación","Cualquier persona del consultorio identifica un incidente o evento adverso"),
             ("Paso 2","Atención inmediata","Si el paciente sufrió daño, se atiende primero al paciente, luego se hace el reporte"),
             ("Paso 3","Reporte","Diligenciar el FOR-PP-005 en las primeras 24 horas"),
             ("Paso 4","Análisis","La médica propietaria analiza el evento: qué pasó, por qué pasó (causa raíz)"),
             ("Paso 5","Plan de mejora","Se establecen acciones correctivas para evitar la recurrencia"),
             ("Paso 6","Seguimiento","Se hace seguimiento a la implementación de las acciones correctivas"),
             ("Paso 7","Reporte externo","Los eventos adversos graves se reportan al INVIMA (si relacionados con dispositivos médicos o medicamentos) y a la SDS o entidad territorial de salud si así lo requiere la normativa")]
    tb2=doc.add_table(rows=len(reporte)+1,cols=3); tb2.style='Table Grid'
    hrow(tb2,["Paso","Actividad","Descripción"])
    for i,(n,a,d) in enumerate(reporte):
        tb2.rows[i+1].cells[0].text=n; tb2.rows[i+1].cells[1].text=a; tb2.rows[i+1].cells[2].text=d
        for c in tb2.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"6. INDICADORES DE SEGURIDAD")
    tb3=doc.add_table(rows=5,cols=4); tb3.style='Table Grid'
    hrow(tb3,["Indicador","Fórmula","Meta","Frecuencia"])
    for i,row in enumerate([
        ("Tasa de eventos adversos reportados","N° eventos adversos / 100 atenciones","Tendencia a cero","Mensual"),
        ("% incidentes analizados con plan de mejora","(Incidentes con plan/Total incidentes)×100","100%","Mensual"),
        ("% cumplimiento protocolo identificación paciente","Auditoria observacional (N° verificaciones correctas/Total)×100","100%","Trimestral"),
        ("% cumplimiento lavado de manos","(Lavados correctos/Total oportunidades observadas)×100","≥90%","Trimestral")]):
        for j,v in enumerate(row): tb3.rows[i+1].cells[j].text=v
        for c in tb3.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph(); firmas(doc)
    path=os.path.join(BASE5,"PRO-PP-003_Proceso_Seguridad_Paciente.docx")
    doc.save(path); print(f"✓ Creado: {path}")

# ===== PRO-PP-004: Prevención de Infecciones =====
def crear_pro_pp004():
    doc=Document(); margins(doc); hf(doc,"PRO-PP-004","Proceso de Prevención de Infecciones")
    portada(doc,"PRO-PP-004","PROCESO DE PREVENCIÓN Y CONTROL DE INFECCIONES")

    H(doc,"1. OBJETIVO")
    B(doc,"Establecer las precauciones estándar, los protocolos de esterilización/desinfección y las medidas de bioseguridad necesarias para prevenir la transmisión de microorganismos patógenos durante la atención de los pacientes en el Consultorio [NOMBRE DEL CONSULTORIO], protegiendo tanto a los pacientes como al personal de salud.")

    H(doc,"2. MARCO LEGAL")
    for n in ["Decreto 351 de 2014 - Gestión de residuos generados en atención en salud",
              "Resolución 1164 de 2002 - Manual de procedimientos PGIRH",
              "OMS - Directrices sobre higiene de manos en la atención sanitaria (2009)",
              "Resolución 2003 de 2014 y 3100 de 2019 - Bioseguridad como estándar de habilitación",
              "Decreto 1011 de 2006 - SOGCS",
              "Normas internacionales CDC sobre precauciones estándar"]: BL(doc,n)

    H(doc,"3. PRECAUCIONES ESTÁNDAR")
    B(doc,"Las precauciones estándar se aplican a TODOS los pacientes independientemente de su diagnóstico o estado infeccioso conocido. Incluyen:")

    H(doc,"3.1 Protocolo de Higiene de Manos (OMS 6 Pasos)",2)
    B(doc,"La higiene de manos es la medida más importante para prevenir infecciones. Se realiza con agua y jabón (cuando las manos están visiblemente sucias) o con alcohol glicerinado al 70% (para manos sin suciedad visible).")
    B(doc,"TÉCNICA DE LAVADO CON AGUA Y JABÓN (mínimo 40-60 segundos):", bold=True)
    pasos_lm=["Paso 1: Mojar las manos con agua corriente",
              "Paso 2: Aplicar suficiente jabón para cubrir toda la superficie de las manos",
              "Paso 3: Frotar las palmas entre sí",
              "Paso 4: Frotar la palma derecha sobre el dorso izquierdo con los dedos entrelazados y viceversa",
              "Paso 5: Frotar las palmas entre sí con los dedos entrelazados",
              "Paso 6: Frotar el dorso de los dedos contra la palma opuesta, manteniendo unidos los dedos",
              "Paso 7: Frotar con movimiento rotatorio el pulgar izquierdo atrapándolo con la palma derecha y viceversa",
              "Paso 8: Frotar la punta de los dedos de la mano derecha contra la palma de la mano izquierda haciendo movimiento rotatorio y viceversa",
              "Paso 9: Enjuagar las manos con agua",
              "Paso 10: Secar con toalla desechable",
              "Paso 11: Usar la toalla para cerrar el grifo (no contaminar manos limpias)"]
    for item in pasos_lm: BL(doc,item)
    B(doc,"LAS 5 OCASIONES PARA EL LAVADO DE MANOS (OMS):", bold=True)
    cinco_momentos=["1. ANTES de tocar al paciente",
                    "2. ANTES de realizar una tarea limpia/aséptica",
                    "3. DESPUÉS del riesgo de exposición a fluidos corporales",
                    "4. DESPUÉS de tocar al paciente",
                    "5. DESPUÉS de tocar el entorno del paciente"]
    for item in cinco_momentos: BL(doc,item)

    H(doc,"3.2 Elementos de Protección Personal (EPP)",2)
    epp=[("Guantes","Usar guantes para cualquier procedimiento que implique contacto con piel no intacta, mucosas, sangre, fluidos corporales. Cambiar entre pacientes. No tocar superficies con guantes contaminados.",
          "Latex o nitrilo (talla S, M, L disponibles). Estériles para procedimientos invasivos."),
         ("Tapabocas","Para procedimientos con riesgo de salpicadura de fluidos. Para la atención de pacientes con síntomas respiratorios. Para procedimientos en zonas próximas a mucosas.",
          "Tapabocas quirúrgico (fluidos) o N95 (aerosoles)"),
         ("Gafas de protección","Para procedimientos con alto riesgo de salpicadura: extracción de PRP, peelings (vapores químicos), inyecciones en zonas con presión.",
          "Gafas de seguridad o visores"),
         ("Bata/delantal","Para procedimientos con riesgo de salpicadura significativa.",
          "Bata desechable o lavable (lavado a 60°C+)")]
    tb=doc.add_table(rows=len(epp)+1,cols=3); tb.style='Table Grid'
    hrow(tb,["EPP","¿Cuándo usar?","Tipo disponible en el consultorio"])
    for i,(e,c,t) in enumerate(epp):
        tb.rows[i+1].cells[0].text=e; tb.rows[i+1].cells[1].text=c; tb.rows[i+1].cells[2].text=t
        for c2 in tb.rows[i+1].cells:
            for p in c2.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"4. LIMPIEZA, DESINFECCIÓN Y ESTERILIZACIÓN")
    H(doc,"4.1 Niveles de Procesamiento",2)
    B(doc,"Según la clasificación de Spaulding, los materiales e instrumentos del consultorio requieren diferentes niveles de procesamiento:")
    tb2=doc.add_table(rows=4,cols=4); tb2.style='Table Grid'
    hrow(tb2,["Clasificación Spaulding","Descripción","Nivel requerido","Ejemplos en el consultorio"])
    spaulding=[("Críticos","Penetran tejidos estériles o el sistema vascular","Esterilización","Agujas, jeringas, cánulas (todos desechables), bisturís"),
               ("Semicríticos","Contactan mucosas o piel no intacta","Desinfección de alto nivel o esterilización","Espéculos (si aplica), laringoscopio"),
               ("No críticos","Contactan piel intacta","Desinfección de bajo nivel","Tensiómetro, fonendoscopio, cinta métrica, camilla")]
    for i,row in enumerate(spaulding):
        for j,v in enumerate(row): tb2.rows[i+1].cells[j].text=v
        for c in tb2.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"4.2 Desinfección de Superficies",2)
    B(doc,"La desinfección de superficies del área de atención se realiza con los siguientes agentes y frecuencias:")
    desinfec=[("Camilla de examen","Limpiar y desinfectar entre cada paciente","Hipoclorito de sodio al 0.5% o alcohol 70%"),
              ("Mesa/escritorio médico","Limpieza al inicio y final de la jornada, y cuando hay contaminación visible","Alcohol 70% o solución desinfectante de superficies"),
              ("Superficies de trabajo (procedimientos estéticos)","Desinfección antes de cada paciente","Hipoclorito 0.5% o alcohol 70%"),
              ("Pisos","Limpieza con desinfectante al inicio y final de jornada","Hipoclorito 0.5% diluido en agua"),
              ("Mesas de mayo o auxiliares","Entre cada procedimiento","Alcohol 70%"),
              ("Tensiómetro, oxímetro (dispositivos)","Entre cada paciente: limpiar con paño con alcohol 70%","Alcohol 70%")]
    tb3=doc.add_table(rows=len(desinfec)+1,cols=3); tb3.style='Table Grid'
    hrow(tb3,["Superficie/equipo","Frecuencia","Agente desinfectante"])
    for i,row in enumerate(desinfec):
        for j,v in enumerate(row): tb3.rows[i+1].cells[j].text=v
        for c in tb3.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"4.3 Preparación de Hipoclorito de Sodio",2)
    B(doc,"El hipoclorito de sodio comercial al 5.25% se diluye para obtener las concentraciones de trabajo:")
    tb4=doc.add_table(rows=4,cols=4); tb4.style='Table Grid'
    hrow(tb4,["Concentración deseada","Uso","Dilución con hipoclorito 5.25%","Vigencia"])
    dilucs=[("0.5% (5000 ppm)","Desinfección de superficies con materia orgánica visible","1 parte de hipoclorito + 9 partes de agua","24 horas"),
            ("0.1% (1000 ppm)","Desinfección de superficies sin materia orgánica visible","1 parte de hipoclorito + 49 partes de agua","24 horas"),
            ("0.05% (500 ppm)","Limpieza general de pisos, baños","1 parte de hipoclorito + 99 partes de agua","24 horas")]
    for i,row in enumerate(dilucs):
        for j,v in enumerate(row): tb4.rows[i+1].cells[j].text=v
        for c in tb4.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()
    B(doc,"IMPORTANTE: El hipoclorito pierde actividad con la luz y el calor. Preparar fresco cada día. No mezclar con otros desinfectantes (especialmente amoniaco - produce gases tóxicos). Usar guantes al manipular.")

    H(doc,"5. MANEJO DE RESIDUOS BIOSANITARIOS")
    B(doc,"En cumplimiento del Decreto 351 de 2014, el consultorio implementa el siguiente plan de gestión de residuos:")
    H(doc,"5.1 Separación en la Fuente",2)
    sep=["Residuos ordinarios (bolsa NEGRA): papel, empaques limpios, residuos no contaminados",
         "Residuos biosanitarios (bolsa ROJA): gasas contaminadas, guantes usados, algodón con sangre, material de vendaje, empaque de material estéril usado",
         "Residuos cortopunzantes (guardián ROJO): agujas, lancetas, jeringas con aguja integrada (NO doblar ni re-encapsular agujas nunca)",
         "Residuos de medicamentos (bolsa ROJA especial): ampollas vacías, frascos de medicamentos"]
    for item in sep: BL(doc,item)
    H(doc,"5.2 Almacenamiento y Disposición",2)
    dis=["Los guardianes se sellan cuando están llenos en sus 3/4 partes",
         "Los residuos biosanitarios se entregan a la empresa gestora de residuos peligrosos: [NOMBRE EMPRESA], Contrato N° [N°]",
         "Frecuencia de recolección: [FRECUENCIA PACTADA]",
         "Los residuos ordinarios se llevan a la canasta de basuras del edificio/sector",
         "Los registros de entrega de residuos peligrosos se conservan en la carpeta de habilitación"]
    for item in dis: BL(doc,item)

    H(doc,"6. BIOSEGURIDAD ESPECÍFICA PARA PROCEDIMIENTOS ESTÉTICOS")
    B(doc,"Los procedimientos estéticos inyectables (toxina botulínica, rellenos, mesoterapia, PRP) implican la ruptura de la barrera cutánea y requieren medidas adicionales de bioseguridad:")
    bio_est=["Usar SIEMPRE guantes estériles para los procedimientos de inyección",
             "Preparar el material en una superficie limpia y desinfectada",
             "Usar material de un solo uso para cada paciente: agujas, jeringas, cánulas",
             "NUNCA compartir frascos de medicamentos abiertos entre pacientes sin uso de adaptadores estériles",
             "La toxina botulínica una vez reconstituida es para un solo paciente o para usar el mismo día",
             "Desechar INMEDIATAMENTE la aguja en el guardián después de cada inyección (no reencapsular)",
             "Si se produce una salpicadura de sangre sobre la piel intacta: lavar con agua y jabón",
             "Si se produce un pinchazo o corte con material contaminado: lavado abundante, reporte a ARL, protocolo de accidente biológico"]
    for item in bio_est: BL(doc,item)

    H(doc,"7. PROTOCOLO DE ACCIDENTE BIOLÓGICO")
    B(doc,"En caso de exposición a sangre u otros fluidos corporales (pinchazo, corte, salpicadura en mucosas):")
    accidente=[("1. Lavado inmediato","Piel: agua y jabón mínimo 5 minutos. Mucosas: solución salina o agua abundante. NO exprimir la herida."),
               ("2. Reporte inmediato","Notificar a la ARL en las primeras 2 horas. Llevar la historia del paciente fuente (si aplica y con confidencialidad)."),
               ("3. Evaluación médica","Ir al centro de atención de la ARL para evaluación y toma de muestras basales del trabajador y del paciente fuente."),
               ("4. Profilaxis post-exposición","Para VIH si el paciente fuente es VIH+ o desconocido en población de alto riesgo: iniciar PEP en las primeras 72 horas. Para VHB: administrar inmunoglobulina si no hay inmunidad y vacunar."),
               ("5. Seguimiento serológico","Basal, 6 semanas, 3 meses, 6 meses para VIH y VHB."),
               ("6. Registro","Diligenciar FURAT (Formulario Único de Reporte de Accidente de Trabajo) con la ARL.")]
    tb5=doc.add_table(rows=len(accidente)+1,cols=2); tb5.style='Table Grid'
    hrow(tb5,["Paso","Acción"])
    for i,(p2,a) in enumerate(accidente):
        tb5.rows[i+1].cells[0].text=p2; tb5.rows[i+1].cells[1].text=a
        for c in tb5.rows[i+1].cells:
            for pp in c.paragraphs:
                for r in pp.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"8. INDICADORES")
    tb6=doc.add_table(rows=4,cols=4); tb6.style='Table Grid'
    hrow(tb6,["Indicador","Fórmula","Meta","Frecuencia"])
    for i,row in enumerate([
        ("% cumplimiento protocolo lavado de manos","(Observaciones conformes/Total observaciones)×100","≥90%","Trimestral"),
        ("% manejo correcto de residuos cortopunzantes","(Residuos en guardián/Total cortopunzantes generados)×100","100%","Mensual"),
        ("N° accidentes biológicos reportados y gestionados","Conteo absoluto (reportados vs ocurridos)","Reportar 100%","Mensual")]):
        for j,v in enumerate(row): tb6.rows[i+1].cells[j].text=v
        for c in tb6.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph(); firmas(doc)
    path=os.path.join(BASE5,"PRO-PP-004_Proceso_Prevencion_Infecciones.docx")
    doc.save(path); print(f"✓ Creado: {path}")

# ===== CONSENTIMIENTOS INFORMADOS =====
def crear_ci(code, titulo, procedimiento, descripcion_proc, beneficios, riesgos, alternativas, instrucciones_post):
    doc=Document(); margins(doc); hf(doc,code,f"Consentimiento Informado - {titulo}")
    portada(doc,code,f"CONSENTIMIENTO INFORMADO\n{titulo}")

    H(doc,"DATOS DEL CONSULTORIO Y LA MÉDICA")
    B(doc,"Consultorio: [NOMBRE DEL CONSULTORIO]\nDirección: [DIRECCIÓN DEL CONSULTORIO], [CIUDAD]\nTeléfono: [TELÉFONO]\nMédica Responsable: Dra. [NOMBRE DE LA MÉDICA]\nTarjeta Profesional N°: [NÚMERO]\nRegistro RETHUS: [NÚMERO]")
    doc.add_paragraph()

    H(doc,"IDENTIFICACIÓN DEL PACIENTE")
    tb=doc.add_table(rows=4,cols=4); tb.style='Table Grid'
    datos=[("Nombre completo:","","Fecha de nacimiento:",""),
           ("N° Documento:","","Tipo de documento:","□ CC  □ CE  □ TI  □ Pasaporte"),
           ("Dirección:","","Teléfono:",""),
           ("Correo electrónico:","","Acompañante (si aplica):","")]
    for i,row in enumerate(datos):
        for j,v in enumerate(row):
            tb.rows[i].cells[j].text=v
            if j%2==0 and v: tb.rows[i].cells[j].paragraphs[0].runs[0].font.bold=True
    doc.add_paragraph()

    H(doc,f"1. DESCRIPCIÓN DEL PROCEDIMIENTO: {titulo.upper()}")
    B(doc,descripcion_proc)

    H(doc,"2. NATURALEZA DEL PROCEDIMIENTO")
    B(doc,f"El procedimiento de {procedimiento} es un ACTO MÉDICO que solo puede ser realizado por un médico legalmente habilitado. La Dra. [NOMBRE DE LA MÉDICA] cuenta con la formación y las certificaciones necesarias para la realización de este procedimiento. Las certificaciones están disponibles en la carpeta de habilitación del consultorio.")

    H(doc,"3. BENEFICIOS ESPERADOS")
    for b in beneficios: BL(doc,b)

    H(doc,"4. RIESGOS, COMPLICACIONES Y EFECTOS ADVERSOS POSIBLES")
    B(doc,"Aunque se tomarán todas las precauciones necesarias para garantizar la seguridad del procedimiento, todo procedimiento médico conlleva riesgos inherentes. Los riesgos específicos de este procedimiento incluyen:")
    H(doc,"4.1 Riesgos Frecuentes (>1% de los casos):",2)
    for r in riesgos[0]: BL(doc,r)
    H(doc,"4.2 Riesgos Infrecuentes pero Importantes (<1% de los casos):",2)
    for r in riesgos[1]: BL(doc,r)
    H(doc,"4.3 Riesgos Raros pero Graves (<0.1% de los casos):",2)
    for r in riesgos[2]: BL(doc,r)

    H(doc,"5. ALTERNATIVAS AL PROCEDIMIENTO")
    for a in alternativas: BL(doc,a)

    H(doc,"6. INSTRUCCIONES POST-PROCEDIMIENTO")
    for i in instrucciones_post: BL(doc,i)

    H(doc,"7. INFORMACIÓN SOBRE FOTOGRAFÍAS CLÍNICAS")
    B(doc,"Para el seguimiento del resultado, se tomarán fotografías clínicas antes y después del procedimiento. Estas fotografías son parte de la historia clínica y son confidenciales. No serán publicadas en medios ni redes sociales sin su autorización explícita por escrito.")
    p=doc.add_paragraph()
    r1=p.add_run("Autorizo la toma de fotografías clínicas con fines de seguimiento médico: ")
    r1.font.size=Pt(11)
    r2=p.add_run("□ SÍ     □ NO")
    r2.font.bold=True; r2.font.size=Pt(11)

    H(doc,"8. DECLARACIÓN DEL PACIENTE")
    B(doc,f"Yo, _____________________________________________, identificado(a) con ____________ N° _____________________, declaro que:")
    declaraciones=["He leído y comprendido la información contenida en este documento",
                   "He tenido la oportunidad de hacer preguntas y todas mis dudas han sido resueltas satisfactoriamente por la Dra. [NOMBRE DE LA MÉDICA]",
                   f"Entiendo en qué consiste el procedimiento de {procedimiento}, cuáles son sus beneficios, riesgos y alternativas",
                   "Acepto voluntariamente que se me realice el procedimiento descrito, sin coacción de ningún tipo",
                   "Entiendo que puedo retirar este consentimiento en cualquier momento antes de que el procedimiento se inicie",
                   "He sido informado(a) del costo del procedimiento y estoy de acuerdo con las condiciones económicas"]
    for d in declaraciones: BL(doc,d)

    H(doc,"9. REVOCACIÓN DEL CONSENTIMIENTO")
    B(doc,"Tengo el derecho de revocar este consentimiento en cualquier momento antes del inicio del procedimiento, sin necesidad de dar explicaciones y sin que ello afecte la calidad de la atención que recibiré en el futuro. Si decido revocar mi consentimiento, lo comunicaré a la médica verbalmente o por escrito.")

    H(doc,"10. FIRMAS")
    doc.add_paragraph()
    tb2=doc.add_table(rows=3,cols=3); tb2.style='Table Grid'
    hrow(tb2,["PACIENTE","TESTIGO (si aplica)","MÉDICA RESPONSABLE"])
    tb2.rows[1].cells[0].text="Nombre: _________________________\n\nFirma: _________________________\n\nFecha: _________________________\n\nHora: _________________________"
    tb2.rows[1].cells[1].text="Nombre: _________________________\n\nFirma: _________________________\n\nFecha: _________________________\n\nParentesco: _____________________"
    tb2.rows[1].cells[2].text="Dra. [NOMBRE DE LA MÉDICA]\nMédica General\nTP N° [NÚMERO]\n\nFirma: _________________________\n\nFecha: _________________________"
    for j in range(3):
        tb2.rows[1].cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in tb2.rows[1].cells[j].paragraphs[0].runs: r.font.size=Pt(10)
    tb2.rows[2].cells[0].text="HUELLA DACTILAR\n(pulgar derecho)\n\n\n\n_________________________"
    tb2.rows[2].cells[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    B(doc,"PARA USO DE LA MÉDICA:", bold=True)
    B(doc,"Fecha del procedimiento: _________________________\nProcedimiento realizado: _________________________\nProducto(s) utilizado(s): _________________________ Lote: _____________ Venc.: _____________\nObservaciones: _________________________\n\nFirma de la médica: _________________________")

    path=os.path.join(BASE5,f"{code}_Consentimiento_Informado_{titulo.replace(' ','_')}.docx")
    doc.save(path); print(f"✓ Creado: {path}")

# ===== FOR-PP-005: Reporte Evento Adverso =====
def crear_for_pp005():
    doc=Document(); margins(doc); hf(doc,"FOR-PP-005","Formato Reporte de Evento Adverso")
    portada(doc,"FOR-PP-005","FORMATO DE REPORTE DE EVENTO ADVERSO E INCIDENTE")

    H(doc,"SECCIÓN 1: DATOS GENERALES DEL REPORTE")
    tb=doc.add_table(rows=5,cols=4); tb.style='Table Grid'
    datos=[("Fecha del evento:","","Hora del evento:",""),
           ("Tipo de reporte:","□ Evento adverso  □ Incidente (near miss)  □ Evento centinela","Servicio donde ocurrió:","□ Consulta general  □ Procedimiento estético"),
           ("Fue detectado por:","□ Médica  □ Paciente  □ Acompañante  □ Personal admin.","¿Llegó al paciente?","□ Sí  □ No"),
           ("N° del reporte:","","Versión del formato:","1.0"),
           ("Elaborado por:","[NOMBRE DE LA MÉDICA]","Cargo:","Médica General - Propietaria")]
    for i,row in enumerate(datos):
        for j,v in enumerate(row):
            tb.rows[i].cells[j].text=v
            if j%2==0 and v and not v.startswith("□"): tb.rows[i].cells[j].paragraphs[0].runs[0].font.bold=True
    doc.add_paragraph()

    H(doc,"SECCIÓN 2: IDENTIFICACIÓN DEL PACIENTE INVOLUCRADO")
    tb2=doc.add_table(rows=3,cols=4); tb2.style='Table Grid'
    datos2=[("Nombre del paciente:","","N° documento:",""),
            ("Edad:","","Sexo:","□ M  □ F  □ Otro"),
            ("Diagnóstico principal:","","¿Primer evento para este paciente?","□ Sí  □ No")]
    for i,row in enumerate(datos2):
        for j,v in enumerate(row): tb2.rows[i].cells[j].text=v
        if 0 < len(tb2.rows[i].cells[0].paragraphs[0].runs):
            tb2.rows[i].cells[0].paragraphs[0].runs[0].font.bold=True
    doc.add_paragraph()

    H(doc,"SECCIÓN 3: DESCRIPCIÓN DEL EVENTO")
    B(doc,"Describir detalladamente qué ocurrió, en qué circunstancias, qué lo precedió y cómo se resolvió:")
    for i in range(8):
        doc.add_paragraph("___________________________________________________________________________")
    doc.add_paragraph()

    H(doc,"SECCIÓN 4: CLASIFICACIÓN DEL EVENTO")
    tb3=doc.add_table(rows=7,cols=2); tb3.style='Table Grid'
    hrow(tb3,["Criterio de clasificación","Selección"])
    clases=[("Tipo de evento",
             "□ Relacionado con medicamentos\n□ Relacionado con procedimiento médico/técnica\n□ Relacionado con dispositivo médico\n□ Relacionado con diagnóstico\n□ Relacionado con comunicación\n□ Relacionado con identificación del paciente\n□ Otro: ___________________________"),
            ("Gravedad del daño al paciente",
             "□ Sin daño (incidente, no llegó al paciente)\n□ Daño leve (molestias menores, resolución espontánea)\n□ Daño moderado (requirió tratamiento médico adicional)\n□ Daño grave (hospitalización, discapacidad temporal)\n□ Daño catastrófico (discapacidad permanente, muerte)"),
            ("Evitabilidad","□ Evitable (se pudo haber prevenido)\n□ Probablemente evitable\n□ Probablemente no evitable\n□ No evitable"),
            ("Causalidad","□ Definitivamente relacionado con la atención\n□ Probablemente relacionado con la atención\n□ Posiblemente relacionado con la atención\n□ No relacionado con la atención"),
            ("Reporte externo requerido","□ No requiere reporte externo\n□ Reportar al INVIMA (farmacovigilancia/tecnovigilancia)\n□ Reportar a SDS/entidad territorial\n□ Reportar a asegurador"),
            ("¿Se informó al paciente?","□ Sí, fecha: _______________ por: _______________\n□ No (justifique): ___________________________")]
    for i,(crit,sel) in enumerate(clases):
        tb3.rows[i+1].cells[0].text=crit; tb3.rows[i+1].cells[1].text=sel
        for c in tb3.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()

    H(doc,"SECCIÓN 5: ANÁLISIS DE CAUSAS (CAUSA RAÍZ)")
    B(doc,"Identificar los factores contribuyentes y la causa raíz del evento:")
    causas=[("Factores del paciente","□ Estado clínico del paciente\n□ Comunicación deficiente del paciente\n□ Incumplimiento de indicaciones"),
            ("Factores del talento humano","□ Conocimiento insuficiente\n□ Habilidad/destreza insuficiente\n□ Fatiga o estrés\n□ Falta de supervisión"),
            ("Factores de los equipos/dispositivos","□ Falla del equipo\n□ Dispositivo defectuoso\n□ Uso inadecuado del equipo"),
            ("Factores del proceso/protocolo","□ Protocolo inexistente\n□ Protocolo deficiente\n□ No seguimiento del protocolo"),
            ("Factores del entorno/infraestructura","□ Condiciones inadecuadas del consultorio\n□ Interrupciones durante el procedimiento\n□ Iluminación insuficiente")]
    tb4=doc.add_table(rows=len(causas)+1,cols=2); tb4.style='Table Grid'
    hrow(tb4,["Categoría de causa","Factores identificados"])
    for i,(cat,fact) in enumerate(causas):
        tb4.rows[i+1].cells[0].text=cat; tb4.rows[i+1].cells[1].text=fact
        for c in tb4.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()
    B(doc,"CAUSA RAÍZ IDENTIFICADA:", bold=True)
    for i in range(3):
        doc.add_paragraph("___________________________________________________________________________")

    H(doc,"SECCIÓN 6: PLAN DE MEJORAMIENTO")
    tb5=doc.add_table(rows=5,cols=4); tb5.style='Table Grid'
    hrow(tb5,["Acción de mejora","Responsable","Fecha compromiso","Estado"])
    for i in range(1,5):
        tb5.rows[i].cells[0].text=""
        tb5.rows[i].cells[3].text="□ Pendiente  □ En curso  □ Completada"
    doc.add_paragraph()

    H(doc,"SECCIÓN 7: SEGUIMIENTO Y CIERRE")
    tb6=doc.add_table(rows=3,cols=2); tb6.style='Table Grid'
    hrow(tb6,["Aspecto","Respuesta"])
    tb6.rows[1].cells[0].text="¿Se implementaron las acciones de mejora?"; tb6.rows[1].cells[1].text="□ Sí  □ No  □ Parcialmente\nObservaciones: ___________________________"
    tb6.rows[2].cells[0].text="¿Se evaluó la efectividad de las acciones?"; tb6.rows[2].cells[1].text="□ Sí  □ No\nResultado: ___________________________"
    doc.add_paragraph()

    H(doc,"SECCIÓN 8: FIRMAS")
    tb7=doc.add_table(rows=2,cols=2); tb7.style='Table Grid'
    hrow(tb7,["REPORTÓ","ANALIZÓ Y APROBÓ"])
    tb7.rows[1].cells[0].text="Nombre: _________________________\nFirma: _________________________\nFecha: _________________________"
    tb7.rows[1].cells[1].text="Nombre: [NOMBRE DE LA MÉDICA]\nFirma: _________________________\nFecha: _________________________"

    path=os.path.join(BASE5,"FOR-PP-005_Formato_Reporte_Evento_Adverso.docx")
    doc.save(path); print(f"✓ Creado: {path}")

if __name__=="__main__":
    print("Generando Carpeta 5 - Parte C...")
    crear_pro_pp003()
    crear_pro_pp004()

    # Consentimiento consulta general
    crear_ci("FOR-PP-001","Consulta de Medicina General","consulta de medicina general",
        "La consulta de medicina general es el acto médico mediante el cual la médica realiza la evaluación de su estado de salud a través de la anamnesis (historia de la enfermedad), el examen físico y el análisis de los resultados de exámenes complementarios, con el fin de establecer un diagnóstico y un plan de manejo personalizado.",
        ["Diagnóstico certero de su condición de salud",
         "Tratamiento adecuado para su condición específica",
         "Prescripción de medicamentos cuando sea necesario",
         "Orientación sobre cambios en el estilo de vida",
         "Solicitud de exámenes complementarios pertinentes",
         "Remisión a especialistas cuando sea requerida"],
        [["Posibles molestias durante el examen físico","Reacción al látex de los guantes (poco frecuente)","Hematoma en el sitio de punción si se toma muestra de sangre"],
         ["Reacción a medicamentos prescritos","Diagnóstico inicial que requiera modificación con nuevos datos o exámenes"],
         ["Eventos adversos graves derivados de medicamentos prescritos (raros, pero posibles)"]],
        ["No recibir tratamiento médico (su condición podría progresar sin manejo)",
         "Buscar atención en otro nivel de complejidad",
         "Tratamientos alternativos o complementarios (preguntar a su médica)"],
        ["Tomar los medicamentos según las indicaciones exactas de la médica",
         "Asistir a la cita de control en la fecha indicada",
         "Consultar urgencias si presenta signos de alarma mencionados por la médica",
         "Informar a la médica si tiene alguna reacción adversa a los medicamentos"])

    # Consentimiento toxina botulínica
    crear_ci("FOR-PP-002","Toxina Botulínica Tipo A","aplicación de toxina botulínica tipo A",
        "La toxina botulínica tipo A es una neurotoxina purificada de origen bacteriano (Clostridium botulinum) que, en las dosis médicas utilizadas, produce una relajación temporal y reversible de los músculos en los que se aplica. Se utiliza para reducir arrugas de expresión (dinámicas), tratar la hiperhidrosis (sudoración excesiva) y otras indicaciones médicas. El efecto es TEMPORAL (dura entre 3 y 6 meses) y el músculo recupera su función gradualmente.",
        ["Reducción visible de las arrugas de expresión tratadas",
         "Aspecto más descansado y rejuvenecido del rostro",
         "Mejoría de la hiperhidrosis (si es la indicación)",
         "Efecto natural cuando se realiza con la técnica y dosis correctas",
         "Procedimiento mínimamente invasivo, sin tiempo de recuperación significativo"],
        [["Eritema (enrojecimiento) transitorio en los sitios de inyección - dura minutos a horas",
          "Hematomas o equimosis en los puntos de inyección - resuelven en 1-2 semanas",
          "Cefalea transitoria el día del procedimiento",
          "Sensación de tensión o pesadez en la zona tratada los primeros días",
          "Edema (hinchazón) leve en el área tratada - dura 24-48 horas"],
         ["Ptosis palpebral (caída del párpado) - puede ocurrir si la toxina difunde al músculo elevador del párpado. Reversible en 4-8 semanas. Se puede tratar con colirio.",
          "Ptosis de cejas - puede ocurrir si se inyecta demasiado cerca de la ceja",
          "Asimetría en el resultado - puede requerir retoque a los 14 días",
          "Efecto insuficiente o ausente (resistencia a la toxina) - poco frecuente"],
         ["Reacción alérgica grave (anafilaxia) - extremadamente rara. El consultorio cuenta con kit de emergencias.",
          "Propagación sistémica de la toxina (descrita con dosis muy altas, no con dosis estéticas) - extremadamente rara",
          "Infección en el sitio de inyección - prevenida con técnica aséptica rigurosa"]],
        ["No realizar el procedimiento: las arrugas permanecerán sin cambios",
         "Tratamientos tópicos con retinoides, ácido hialurónico, vitamina C - mejoría parcial",
         "Procedimientos físicos como radiofrecuencia o ultrasonido - resultados diferentes y más graduales",
         "Procedimientos quirúrgicos - mayor invasividad y costo"],
        ["No masajear ni frotar la zona tratada por 24 horas",
         "Mantener posición vertical (no acostarse) por 4 horas después del procedimiento",
         "No realizar ejercicio intenso el día del procedimiento",
         "No exposición a calor intenso (sauna, vapor, sol directo) por 24 horas",
         "Puede mover normalmente los músculos de la cara (fruncir, sonreír) - esto no difunde la toxina",
         "El efecto comienza a notarse entre los 3 y 7 días, con efecto máximo a los 14 días",
         "Asistir al control a los 14 días para evaluar el resultado y realizar retoques si es necesario"])

    # Consentimiento rellenos dérmicos
    crear_ci("FOR-PP-003","Rellenos Dérmicos con Ácido Hialurónico","aplicación de rellenos dérmicos con ácido hialurónico",
        "El ácido hialurónico es un polisacárido presente naturalmente en los tejidos corporales, especialmente en la piel. En medicina estética, se utiliza en forma de gel reticulado (estabilizado) inyectable, clasificado como dispositivo médico de clase III, para restaurar volumen, corregir arrugas profundas, aumentar labios, proyectar pómulos o hidratación profunda de la piel. El ácido hialurónico se degrada naturalmente por el organismo, por lo que su efecto es TEMPORAL (dura entre 9 y 18 meses según la zona y el tipo de relleno).",
        ["Restauración del volumen perdido por el envejecimiento",
         "Corrección de arrugas estáticas profundas",
         "Aumento y definición de labios",
         "Proyección y contorno de pómulos y mandíbula",
         "Hidratación profunda y luminosidad (con HA no reticulado)",
         "Resultados inmediatos, con mejora progresiva"],
        [["Eritema, edema y sensibilidad en la zona tratada - normal los primeros 3-7 días",
          "Hematomas y equimosis - frecuentes, resuelven en 1-2 semanas",
          "Irregularidades o nódulos palpables - especialmente en labios - generalmente se resuelven espontáneamente o con masaje",
          "Asimetría en el resultado - puede requerir ajuste"],
         ["Efecto Tyndall (coloración azulada visible si el relleno se coloca muy superficialmente)",
          "Granuloma (formación de nódulo tardío) - poco frecuente, puede requerir tratamiento",
          "Infección post-procedimiento - prevenida con técnica aséptica rigurosa"],
         ["COMPLICACIÓN GRAVE: Oclusión vascular (obstrucción de un vaso sanguíneo por el relleno) - RARA pero potencialmente grave. Puede causar necrosis de la piel o, en casos extremadamente raros, ceguera si el relleno alcanza una arteria que irriga el ojo. EL CONSULTORIO CUENTA CON HIALURONIDASA DISPONIBLE PARA EL MANEJO DE EMERGENCIA VASCULAR. Ante cualquier dolor intenso o cambio de color en la piel durante o después del procedimiento, consulte INMEDIATAMENTE.",
          "Reacción alérgica o anafiláctica - extremadamente rara con ácido hialurónico no animal"]],
        ["No realizar el procedimiento",
         "Toxina botulínica (para arrugas dinámicas)",
         "Radiofrecuencia o ultrasonido microfocalizado (Ultherapy) - para laxitud",
         "Procedimientos quirúrgicos (bichectomía, lifting facial) - más invasivos"],
        ["NO presionar ni masajear la zona tratada por 24 horas (salvo indicación médica)",
         "Aplicar hielo en las primeras 24 horas para reducir edema y hematomas (gasa entre el hielo y la piel)",
         "No exposición a calor (sauna, sol directo) por 48 horas",
         "No actividad física intensa por 24 horas",
         "El edema es normal los primeros días - el resultado final se aprecia a las 2 semanas",
         "SIGNOS DE ALARMA que requieren consulta URGENTE: dolor intenso y progresivo, zona de la piel que se pone blanca o violácea, visión borrosa o pérdida de visión"])

    # Consentimiento peeling químico
    crear_ci("FOR-PP-004","Peeling Químico Superficial","peeling químico superficial",
        "El peeling químico superficial es la aplicación controlada de un agente químico (ácido glicólico, ácido tricloroacético TCA en concentraciones bajas, o ácido mandélico) sobre la piel, que produce una exfoliación controlada y estimula la regeneración cutánea. Los peelings superficiales actúan solo en la capa más externa de la piel (epidermis), lo que los hace seguros con mínimo tiempo de recuperación comparado con peelings más profundos. El resultado se mejora con sesiones periódicas.",
        ["Mejora de la textura y luminosidad de la piel",
         "Reducción de manchas e hiperpigmentaciones",
         "Control del acné y reducción de poros",
         "Estimulación del colágeno y efecto rejuvenecedor",
         "Mejoría gradual que se optimiza con sesiones periódicas"],
        [["Eritema (enrojecimiento) - normal durante y después del procedimiento",
          "Sensación de calor, ardor o picazón durante la aplicación - normal",
          "Descamación fina de la piel en los días siguientes (1-5 días) - parte del proceso esperado",
          "Edema leve - especialmente en zona periorbicular"],
         ["Hiperpigmentación post-inflamatoria (manchas oscuras) - más frecuente en fototipos altos (IV-VI). Se previene con fotoprotección estricta y preparación de la piel.",
          "Hipopigmentación (manchas claras) - más rara",
          "Persistencia o intensificación del acné los primeros días (purging - normal transitorio)"],
         ["Quemadura química si hay mala técnica o uso de concentraciones inadecuadas - prevenida con la técnica correcta y verificación de la concentración",
          "Cicatrices (queloides) en personas predispuestas - preguntamos sobre antecedentes antes del procedimiento",
          "Reactivación de herpes oral si hay antecedente - informar antes del procedimiento para profilaxis antiviral"]],
        ["No realizar el procedimiento - textura e hiperpigmentaciones persistirán",
         "Tratamientos tópicos: retinoides, vitamina C, hidroquinona, niacinamida",
         "Tratamientos con luz (IPL, láser) - requieren dispositivos específicos",
         "Procedimientos físicos: microdermoabrasión, dermapen"],
        ["NO frotar ni pelar manualmente la piel que se descama - dejar que se desprenda sola",
         "Mantener la piel bien hidratada con crema humectante suave (sin ácidos ni retinoides)",
         "Usar fotoprotector SPF 50+ TODOS LOS DÍAS por mínimo 4 semanas post-peeling",
         "NO exposición solar directa por 2 semanas",
         "NO usar retinoides, AHA/BHA, vitamina C, ni otros activos exfoliantes hasta resolución completa",
         "NO hacer ejercicio intenso el día del procedimiento",
         "Control en 2-4 semanas para evaluar resultado"])

    crear_for_pp005()
    print("✅ Parte C completada.")
