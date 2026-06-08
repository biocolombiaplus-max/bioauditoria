
# -*- coding: utf-8 -*-
"""Genera MAN-TH-001 Manual de Talento Humano"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = "/home/user/bioauditoria/documentos_habilitacion/CARPETA_1_TALENTO_HUMANO/MAN-TH-001_Manual_Talento_Humano.docx"

AZUL_OSCURO = RGBColor(0, 51, 102)
AZUL_MEDIO = RGBColor(21, 101, 192)
BLANCO = RGBColor(255, 255, 255)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_header_footer(doc, code, title):
    section = doc.sections[0]
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = f"{code} | {title}"; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.runs[0]; run.font.size = Pt(9); run.font.color.rgb = AZUL_OSCURO; run.font.name = 'Calibri'
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)

def add_page_number(paragraph):
    paragraph.clear()
    run = paragraph.add_run("Versión 2.0 | Junio 2025    Página ")
    run.font.size = Pt(9); run.font.name = 'Calibri'
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run2 = paragraph.add_run(); run2.font.size = Pt(9); run2.font.name = 'Calibri'
    run2._r.append(fldChar1); run2._r.append(instrText); run2._r.append(fldChar2)

def h1(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text); run.bold = True; run.font.size = Pt(14); run.font.name = 'Calibri'; run.font.color.rgb = AZUL_OSCURO

def h2(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text); run.bold = True; run.font.size = Pt(12); run.font.name = 'Calibri'; run.font.color.rgb = AZUL_MEDIO

def h3(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text); run.bold = True; run.font.size = Pt(11); run.font.name = 'Calibri'; run.font.color.rgb = AZUL_OSCURO

def body(doc, text):
    p = doc.add_paragraph(text); p.paragraph_format.space_after = Pt(6)
    for run in p.runs: run.font.size = Pt(11); run.font.name = 'Calibri'

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text); run.font.size = Pt(11); run.font.name = 'Calibri'

def table_header_row(table, headers, bg='003366'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]; cell.text = ''; set_cell_bg(cell, bg)
        run = cell.paragraphs[0].add_run(h); run.bold = True; run.font.color.rgb = BLANCO
        run.font.size = Pt(10); run.font.name = 'Calibri'; cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table_row(table, values, shaded=False):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]; cell.text = str(val)
        if shaded: set_cell_bg(cell, 'F2F2F2')
        for para in cell.paragraphs:
            for run in para.runs: run.font.size = Pt(10); run.font.name = 'Calibri'

doc = Document()
style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
add_header_footer(doc, "MAN-TH-001", "Manual de Gestión del Talento Humano")

# PORTADA
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MANUAL DE GESTIÓN DEL TALENTO HUMANO EN SALUD")
run.font.size = Pt(18); run.font.name = 'Calibri'; run.font.color.rgb = AZUL_OSCURO; run.bold = True
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[NOMBRE DEL CONSULTORIO]")
run.font.size = Pt(14); run.font.name = 'Calibri'; run.font.color.rgb = AZUL_MEDIO; run.bold = True
doc.add_paragraph()
t = doc.add_table(rows=5, cols=2); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
pd = [("Código:", "MAN-TH-001"), ("Versión:", "2.0"), ("Fecha:", "Junio 2025"),
      ("Elaboró/Aprobó:", "Dra. [NOMBRE DE LA MÉDICA] - Directora Médica"), ("NIT:", "[NIT DEL CONSULTORIO]")]
for i, (k, v) in enumerate(pd):
    t.rows[i].cells[0].text = k; t.rows[i].cells[1].text = v
    set_cell_bg(t.rows[i].cells[0], 'E8EEF4')
    for cell in t.rows[i].cells:
        for para in cell.paragraphs:
            for run in para.runs: run.font.size = Pt(11); run.font.name = 'Calibri'

doc.add_page_break()

h1(doc, "CONTROL DE VERSIONES")
t_ver = doc.add_table(rows=3, cols=5); t_ver.style = 'Table Grid'
table_header_row(t_ver, ["Versión", "Fecha", "Descripción", "Elaboró", "Aprobó"])
add_table_row(t_ver, ["1.0", "Enero 2024", "Versión inicial del manual de talento humano.", "Dra. [NOMBRE DE LA MÉDICA]", "Dra. [NOMBRE DE LA MÉDICA]"])
add_table_row(t_ver, ["2.0", "Junio 2025", "Ampliación de secciones de educación continuada, certificaciones para estética médica y proceso de evaluación de competencias.", "Dra. [NOMBRE DE LA MÉDICA]", "Dra. [NOMBRE DE LA MÉDICA]"], shaded=True)
doc.add_paragraph()

# 1. INTRODUCCIÓN
h1(doc, "1. INTRODUCCIÓN")
body(doc, "El talento humano constituye el activo más valioso de cualquier organización de salud. En el contexto de la prestación de servicios de salud de primer nivel, la calidad, la idoneidad, la motivación y el bienestar del recurso humano determinan de manera directa la calidad de la atención que reciben los pacientes. Un médico bien seleccionado, adecuadamente formado, evaluado de manera regular y que trabaja en condiciones de bienestar, es la mejor garantía de una atención médica segura, oportuna, pertinente y humana.")
body(doc, "En Colombia, la gestión del talento humano en salud está regulada por la Ley 1164 de 2007 (Ley de Talento Humano en Salud), que establece los principios, los requisitos de habilitación y las condiciones de ejercicio de las profesiones y ocupaciones del área de la salud. Esta ley crea el Registro Único Nacional del Talento Humano en Salud (RETHUS), administrado por el Ministerio de Salud, donde deben estar inscritos todos los profesionales y trabajadores del área de la salud que ejercen en Colombia. La Resolución 3100 de 2019 establece los requisitos de talento humano mínimos para la habilitación de cada tipo de servicio de salud.")
body(doc, "El presente Manual de Gestión del Talento Humano del [NOMBRE DEL CONSULTORIO] tiene como propósito establecer las políticas, procesos y procedimientos para la gestión integral del talento humano desde la planeación y reclutamiento hasta la desvinculación, pasando por la selección, contratación, inducción, evaluación de competencias, educación continuada y bienestar laboral. Este manual aplica a todo el personal que labora en el consultorio, con especial énfasis en el personal asistencial (médico) por su impacto directo en la calidad y seguridad de la atención.")
body(doc, "La Dra. [NOMBRE DE LA MÉDICA], como Directora Médica y propietaria del [NOMBRE DEL CONSULTORIO], es la responsable última de la gestión del talento humano, de garantizar que todo el personal que presta servicios en el consultorio cumple con los requisitos legales y técnicos exigidos, y de crear un ambiente de trabajo que favorezca el bienestar, el desarrollo profesional y la motivación del equipo.")

# 2. POLÍTICA DE TALENTO HUMANO
h1(doc, "2. POLÍTICA DE GESTIÓN DEL TALENTO HUMANO")
body(doc, "El [NOMBRE DEL CONSULTORIO], bajo la dirección de la Dra. [NOMBRE DE LA MÉDICA], está comprometido con la selección, vinculación y desarrollo de un talento humano competente, ético y comprometido con la misión de brindar atención médica integral, segura y de excelente calidad. Reconocemos que nuestros colaboradores son el fundamento de la calidad de nuestra atención y el primer recurso para el cumplimiento de nuestra misión institucional.")
body(doc, "En consecuencia, nos comprometemos a: PRIMERO, seleccionar el personal de acuerdo con los perfiles de cargo establecidos, verificando rigurosamente la idoneidad académica y profesional, los antecedentes y la ausencia de sanciones disciplinarias. SEGUNDO, garantizar que todo el personal asistencial cuente con las habilitaciones, certificaciones y registros vigentes que exige la normatividad colombiana para el ejercicio de sus funciones. TERCERO, brindar una inducción estructurada y completa a todo el personal nuevo, que le permita conocer la organización, sus valores, sus procesos y sus responsabilidades antes de asumir funciones con plena autonomía. CUARTO, evaluar las competencias del personal de manera periódica y objetiva, identificando fortalezas y áreas de mejora, y estableciendo planes de desarrollo acordes con los hallazgos. QUINTO, promover la educación médica continua y la actualización permanente del personal asistencial en las áreas clínicas relevantes para los servicios que presta el consultorio. SEXTO, garantizar condiciones de trabajo seguras, dignas y motivadoras, con especial atención a la prevención de riesgos laborales y al bienestar físico y mental del personal. SÉPTIMO, gestionar el ausentismo de manera proactiva, identificando sus causas y adoptando medidas preventivas. OCTAVO, respetar los derechos laborales del personal y cumplir con todas las obligaciones legales derivadas de la vinculación laboral. NOVENO, realizar los procesos de desvinculación de manera justa, transparente y respetuosa de la dignidad de la persona.")
body(doc, "Esta política se revisa anualmente y se ajusta a los cambios normativos, organizacionales y del entorno. Todos los miembros del equipo del consultorio conocen y acogen esta política como parte de su compromiso con la organización.")

# 3. RECLUTAMIENTO Y SELECCIÓN
h1(doc, "3. PROCESO DE RECLUTAMIENTO, SELECCIÓN Y CONTRATACIÓN")
h2(doc, "3.1 Planeación de la Necesidad de Personal")
body(doc, "La necesidad de vincular nuevo personal puede surgir por: crecimiento en el volumen de servicios que supera la capacidad instalada actual, expansión del portafolio de servicios que requiere nuevas competencias, remplazo de personal desvinculado, o necesidad de apoyo administrativo o de enfermería para mejorar el flujo del proceso. En cualquier caso, antes de iniciar el proceso de selección, la Dra. [NOMBRE DE LA MÉDICA] realiza: análisis de la necesidad real (tiempo completo, medio tiempo, por prestación de servicios), definición del perfil de cargo con los requisitos mínimos obligatorios (educación, certificaciones, experiencia, competencias), y análisis de los recursos financieros disponibles para la vinculación.")

h2(doc, "3.2 Etapas del Proceso de Reclutamiento")
t_rec = doc.add_table(rows=1, cols=4); t_rec.style = 'Table Grid'
table_header_row(t_rec, ["Etapa", "Actividades", "Tiempo Estándar", "Responsable"])
rec_steps = [
    ("1. Publicación de la vacante", "Elaboración del aviso de convocatoria con perfil completo. Publicación en portales de empleo (LinkedIn, Computrabajo, página web de la Universidad si aplica), grupos gremiales de médicos, redes de contactos profesionales.", "3-5 días", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("2. Recepción de hojas de vida", "Recepción de HV por correo electrónico o en físico. Organización y registro de los candidatos.", "Período de postulación: 1-2 semanas", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("3. Preselección de HV", "Revisión de HV contra el perfil del cargo. Verificación de requisitos mínimos obligatorios. Selección de candidatos a entrevistar (3-5 candidatos).", "2-3 días", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("4. Verificación de documentos pre-entrevista", "Verificación en línea de: tarjeta profesional activa (Tribunal Ético de Medicina), registro en RETHUS, ausencia de sanciones disciplinarias activas.", "1 día", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("5. Entrevista de selección", "Entrevista estructurada con preguntas de conocimientos clínicos, casos clínicos, situaciones éticas, motivación y expectativas del candidato.", "45-60 min por candidato", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("6. Evaluación de competencias técnicas", "Aplicación del instrumento de evaluación de competencias técnicas FOR-TH-001 mediante casos clínicos orales o escritos.", "30-45 min", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("7. Verificación de referencias laborales", "Contacto con al menos 2 referencias laborales del candidato. Verificación de desempeño, conducta y motivo de retiro.", "1-2 días", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("8. Exámenes médicos de ingreso", "Examen médico de ingreso: hemograma, serología para hepatitis B y C (o certificado de vacunación hepatitis B), prueba de embarazo si aplica, radiografía de tórax.", "1-2 días (según agenda del laboratorio)", "Médico de medicina laboral o laboratorio clínico"),
    ("9. Selección y oferta laboral", "Selección del candidato idóneo. Comunicación de la oferta laboral con condiciones (honorarios, jornada, funciones, inicio).", "1 día", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("10. Formalización de la vinculación", "Firma del contrato o contrato de prestación de servicios. Afiliación al Sistema de Seguridad Social si aplica. Elaboración de expediente de talento humano.", "1-2 días", "Dra. [NOMBRE DE LA MÉDICA]"),
]
for i, row_data in enumerate(rec_steps):
    add_table_row(t_rec, list(row_data), shaded=(i % 2 == 1))
doc.add_paragraph()

h2(doc, "3.3 Documentos del Expediente de Talento Humano")
body(doc, "El expediente de talento humano de cada persona vinculada al consultorio debe contener los siguientes documentos organizados en carpeta o formato digital seguro:")
bullet(doc, "Hoja de vida (CV actualizado y firmado).")
bullet(doc, "Fotocopia del documento de identidad (ambas caras).")
bullet(doc, "Fotocopia del diploma de grado de médico y de especializaciones relevantes.")
bullet(doc, "Fotocopia de la tarjeta profesional vigente del Tribunal Ético de Medicina.")
bullet(doc, "Certificado de inscripción en el RETHUS (impresión de la consulta en línea, actualizada).")
bullet(doc, "Certificados de estudios de posgrado, diplomados, cursos y actualizaciones pertinentes.")
bullet(doc, "Certificado de BLS/RCP vigente.")
bullet(doc, "Certificados de cursos requeridos para los servicios que va a prestar.")
bullet(doc, "Certificados de los últimos empleos o contratos relevantes.")
bullet(doc, "Resultados de exámenes médicos de ingreso.")
bullet(doc, "Formato de verificación de antecedentes (FOR-TH-002).")
bullet(doc, "Contrato de trabajo o contrato de prestación de servicios firmado.")
bullet(doc, "Formato de inducción diligenciado y firmado (FOR-TH-003).")
bullet(doc, "Evaluaciones de competencias y desempeño periódicas.")
bullet(doc, "Registro de actividades de educación continuada asistidas.")

# 4. INDUCCIÓN
h1(doc, "4. PROGRAMA DE INDUCCIÓN")
body(doc, "Todo el personal que se vincule al [NOMBRE DEL CONSULTORIO] recibirá un proceso de inducción estructurado de tres (3) días antes de asumir funciones de manera autónoma. El objetivo de la inducción es familiarizar al nuevo colaborador con la misión, visión, valores y cultura organizacional del consultorio, los procesos, protocolos y procedimientos operativos relevantes para su cargo, los sistemas de información utilizados, los requisitos del Sistema de Gestión de Calidad, las políticas de seguridad del paciente, y los aspectos administrativos de la vinculación.")
body(doc, "El proceso de inducción es responsabilidad de la Dra. [NOMBRE DE LA MÉDICA] y se realizará con base en el siguiente programa:")

t_ind = doc.add_table(rows=1, cols=4); t_ind.style = 'Table Grid'
table_header_row(t_ind, ["Día", "Horario", "Tema", "Responsable"])
ind_data = [
    ("DÍA 1", "8:00 - 9:00", "Bienvenida y presentación del consultorio: historia, misión, visión, valores, portafolio de servicios, organigrama.", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("DÍA 1", "9:00 - 10:30", "Marco normativo del consultorio: Resolución 3100/2019, SOGCS, Decreto 1011/2006, Ley 23/1981 (Ética Médica). Derechos y deberes de los pacientes.", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("DÍA 1", "10:30 - 12:00", "Sistema de Gestión de Calidad: documentos del SGC, mapa de procesos, indicadores de calidad, PAMEC.", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("DÍA 1", "13:00 - 14:30", "Política y programa de Seguridad del Paciente: eventos adversos, incidentes, caídas, identificación de pacientes, higiene de manos, comunicación efectiva.", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("DÍA 1", "14:30 - 16:00", "Historia Clínica: Resolución 1995/1999, componentes obligatorios, diligenciamiento correcto, manejo de errores, confidencialidad.", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("DÍA 1", "16:00 - 17:00", "Recorrido físico por el consultorio: conocimiento de todas las áreas, equipos, insumos, sistemas de información.", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("DÍA 2", "8:00 - 10:00", "Proceso de consulta de medicina general (PRO-PP-001): fases, tiempos, responsabilidades, indicadores de calidad.", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("DÍA 2", "10:00 - 12:00", "Proceso de procedimientos estéticos no invasivos (PRO-PP-002): marco legal, evaluación previa, técnicas, protocolos de emergencia.", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("DÍA 2", "13:00 - 14:30", "Proceso de referencia y contrarreferencia: criterios, red de referencia, documentación.", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("DÍA 2", "14:30 - 16:00", "Gestión de residuos hospitalarios (PRO-PP-006): clasificación, segregación, almacenamiento, gestión de residuos peligrosos.", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("DÍA 2", "16:00 - 17:00", "Protocolo de accidente biológico: qué hacer ante una punción accidental, exposición a sangre, reporte.", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("DÍA 3", "8:00 - 10:00", "Sistemas de información: manejo del sistema de agenda, historia clínica electrónica o física, consulta en ADRES, consulta en RETHUS.", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("DÍA 3", "10:00 - 12:00", "Atención al usuario: PQRSF, derechos y deberes de los pacientes, encuesta de satisfacción, manejo de situaciones difíciles.", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("DÍA 3", "13:00 - 15:00", "Práctica supervisada: acompañamiento de la Dra. [NOMBRE DE LA MÉDICA] en consultas o procedimientos (como observador).", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("DÍA 3", "15:00 - 16:30", "Evaluación de inducción: cuestionario de 20 preguntas sobre los temas tratados. Aprobación: ≥80% de respuestas correctas.", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("DÍA 3", "16:30 - 17:00", "Firma de formato de inducción FOR-TH-003 y expedición de credencial de autorización para inicio de funciones.", "Dra. [NOMBRE DE LA MÉDICA]"),
]
for i, row_data in enumerate(ind_data):
    add_table_row(t_ind, list(row_data), shaded=(i % 2 == 1))
doc.add_paragraph()

# 5. EVALUACIÓN DE COMPETENCIAS MÉDICAS
h1(doc, "5. EVALUACIÓN DE COMPETENCIAS DEL PERSONAL MÉDICO")
h2(doc, "5.1 Instrumento de Evaluación de Competencias Técnicas del Médico General")
body(doc, "La evaluación de competencias técnicas se aplica al ingreso y anualmente a todo el personal médico. Cada ítem se califica en escala de 1 a 5, donde: 1=Deficiente (no cumple), 2=Insuficiente (cumple con grandes dificultades), 3=Aceptable (cumple con algunas dificultades), 4=Satisfactorio (cumple con mínimas observaciones), 5=Excelente (cumple sin observaciones, puede ser referente). La meta para cada ítem es ≥4. El resultado global se calcula como el promedio de todos los ítems.")

t_comp = doc.add_table(rows=1, cols=4); t_comp.style = 'Table Grid'
table_header_row(t_comp, ["N°", "Competencia Evaluada", "Método de Evaluación", "Calificación (1-5)"])
comp_items = [
    ("1", "Elaboración de anamnesis completa (motivo de consulta, enfermedad actual, antecedentes, revisión por sistemas)", "Caso clínico estandarizado o supervisión directa"),
    ("2", "Técnica de examen físico sistemático y completo (de cabeza a pies)", "Observación directa con lista de verificación"),
    ("3", "Formulación diagnóstica con código CIE-10 correcto y diagnóstico diferencial razonado", "Caso clínico con presentación de diagnóstico"),
    ("4", "Elaboración de plan de manejo farmacológico y no farmacológico según GPC", "Revisión de historia clínica o caso clínico"),
    ("5", "Prescripción médica que cumple Resolución 1478/2006 (todos los campos obligatorios)", "Revisión de fórmulas médicas elaboradas"),
    ("6", "Diligenciamiento completo de la historia clínica según Resolución 1995/1999", "Auditoría de historias clínicas"),
    ("7", "Identificación de criterios de referencia y gestión oportuna del proceso de referencia", "Caso clínico o revisión de referencias previas"),
    ("8", "Manejo de hipertensión arterial según GPC colombiana", "Caso clínico"),
    ("9", "Manejo de diabetes mellitus tipo 2 según GPC colombiana", "Caso clínico"),
    ("10", "Manejo de infección de vías respiratorias superiores y neumonía según GPC", "Caso clínico"),
    ("11", "Manejo de infección de vías urinarias según GPC colombiana", "Caso clínico"),
    ("12", "Identificación y manejo inicial de urgencias menores en primer nivel", "Caso clínico de urgencia"),
    ("13", "Técnica correcta de reanimación cardiopulmonar básica (BLS)", "Simulación en maniquí o evaluación de certificado vigente"),
    ("14", "Notificación de enfermedades de declaración obligatoria (SIVIGILA)", "Pregunta teórica y revisión de casos reportados"),
    ("15", "Conocimiento y aplicación de los derechos de los pacientes", "Casos situacionales"),
    ("16", "Gestión del consentimiento informado: proceso real de obtención y documentación", "Observación directa o revisión de CI en HC"),
    ("17", "Interpretación básica de exámenes de laboratorio de primer nivel (hemograma, química sanguínea básica, uroanálisis)", "Casos clínicos con resultados de laboratorio"),
    ("18", "Interpretación básica de radiografía de tórax (cardiomegalia, condensación, derrame, neumotórax)", "Lectura de radiografías seleccionadas"),
    ("19", "Reporte de eventos adversos e incidentes de seguridad del paciente", "Pregunta sobre proceso y revisión de reportes previos"),
    ("20", "Comunicación efectiva con el paciente: empatía, lenguaje comprensible, verificación de comprensión", "Observación directa de consulta o ECOE"),
    ("21", "Manejo de la historia clínica de procedimientos estéticos: campos específicos, fotografías, registro de productos", "Revisión de HC estéticas (si aplica)"),
    ("22", "Conocimiento de contraindicaciones y complicaciones de procedimientos estéticos (si aplica al cargo)", "Pregunta teórica o caso clínico"),
    ("23", "Protocolo de emergencia vascular por relleno dérmico: pasos, dosis de hialuronidasa (si aplica)", "Simulación del protocolo o pregunta teórica"),
    ("24", "Manejo de residuos biológicos y peligrosos: clasificación y disposición correcta", "Observación directa"),
    ("25", "Actualización en educación médica continua: participación en actividades de actualización relevantes en el último año", "Revisión de certificados de asistencia"),
]
for i, (num, comp, metodo) in enumerate(comp_items):
    add_table_row(t_comp, [num, comp, metodo, "___"], shaded=(i % 2 == 1))
doc.add_paragraph()

h2(doc, "5.2 Evaluación de Competencias Adicionales para Medicina Estética")
body(doc, "Para el personal médico que realiza procedimientos estéticos, adicionalmente se evalúan las siguientes competencias específicas:")
t_comp_est = doc.add_table(rows=1, cols=3); t_comp_est.style = 'Table Grid'
table_header_row(t_comp_est, ["Competencia Estética", "Método de Evaluación", "Calificación (1-5)"])
comp_est = [
    ("Conocimiento anatómico de las zonas de alto riesgo vascular facial (glabela, nariz, temporal, labios, surco nasolabial)", "Pregunta teórica con imagen anatómica"),
    ("Selección correcta del tipo y marca de relleno para cada indicación", "Caso clínico de indicación de relleno"),
    ("Técnica de reconstitución de toxina botulínica (ratio dilución, manipulación de frío)", "Demostración práctica supervisada"),
    ("Técnica de inyección de toxina botulínica en frente, glabela y patas de gallo", "Observación directa o simulación"),
    ("Técnica de inyección de ácido hialurónico en surcos nasolabiales y/o labios", "Observación directa o simulación"),
    ("Reconocimiento de signos de oclusión vascular (blanqueamiento, dolor, cambio de color)", "Caso clínico fotográfico"),
    ("Protocolo completo de emergencia vascular: pasos secuenciales, dosis de hialuronidasa", "Simulación del protocolo sin guía escrita"),
    ("Técnica de aplicación de peeling químico superficial: preparación, aplicación, neutralización", "Observación directa o demostración"),
    ("Evaluación psicológica básica del paciente estético: identificación de trastorno dismórfico corporal", "Caso clínico"),
    ("Fotografía clínica estandarizada: tomas obligatorias, condiciones de luz y fondo", "Evaluación de fotografías tomadas"),
]
for i, (comp, metodo) in enumerate(comp_est):
    add_table_row(t_comp_est, [comp, metodo, "___"], shaded=(i % 2 == 1))
doc.add_paragraph()

# 6. CERTIFICACIONES OBLIGATORIAS
h1(doc, "6. CERTIFICACIONES OBLIGATORIAS DEL PERSONAL MÉDICO")
h2(doc, "6.1 Certificaciones Obligatorias para Médico General (Todos los Servicios)")
t_cert = doc.add_table(rows=1, cols=5); t_cert.style = 'Table Grid'
table_header_row(t_cert, ["Certificación", "Entidad Emisora Válida", "Vigencia", "Cuándo Renovar", "Consecuencia de No Tenerla"])
cert_gen = [
    ("Tarjeta Profesional de Médico vigente", "Tribunal Ético de Medicina de Colombia / Secretaría Seccional de Salud", "Permanente (puede suspenderse por sanciones)", "Verificar anualmente que no tenga sanciones activas en RETHUS", "No puede ejercer la medicina en Colombia. Delito de ejercicio ilegal de la medicina."),
    ("Registro en RETHUS", "Ministerio de Salud y Protección Social", "Permanente (verificar activo)", "Actualizar cuando cambien los datos o se obtengan nuevas certificaciones", "No cumplimiento del requisito de habilitación de la Res. 3100/2019."),
    ("BLS (Soporte Vital Básico) / RCP - American Heart Association o Cruz Roja", "AHA, Cruz Roja Colombiana, universidades y sociedades médicas avaladas", "2 años", "Renovar 3 meses antes del vencimiento", "No cumplimiento del estándar de habilitación de Res. 3100/2019 para médico en consulta de primer nivel."),
    ("Curso de Atención Inicial de Urgencias en Primer Nivel o AIAU", "Universidades con facultad de medicina, sociedades médicas colombianas", "5 años (o según la institución)", "Verificar vigencia. Renovar antes del vencimiento.", "Debilita la capacidad del médico para manejar urgencias menores de manera segura."),
    ("Vacunación completa: Hepatitis B (3 dosis), Tétanos (dTd cada 10 años), Influenza (anual)", "IPS de medicina laboral o EPS del médico", "Según esquema de vacunación", "Verificar en cada actualización del expediente de TH", "Riesgo ocupacional de enfermedad profesional. Incumplimiento de normas de bioseguridad."),
]
for i, row_data in enumerate(cert_gen):
    add_table_row(t_cert, list(row_data), shaded=(i % 2 == 1))
doc.add_paragraph()

h2(doc, "6.2 Certificaciones Específicas para Procedimientos Estéticos")
body(doc, "El médico que realiza procedimientos estéticos no invasivos en el [NOMBRE DEL CONSULTORIO] debe acreditar formación certificada en cada uno de los procedimientos que realiza. La formación en medicina estética no está incluida en el pregrado de medicina, por lo que es responsabilidad del médico obtener la certificación formal antes de realizar cualquier procedimiento. La siguiente tabla detalla los requerimientos:")
t_cert_est = doc.add_table(rows=1, cols=5); t_cert_est.style = 'Table Grid'
table_header_row(t_cert_est, ["Procedimiento", "Tipo de Formación Mínima Requerida", "Instituciones Válidas en Colombia", "Periodicidad de Actualización", "Evidencia en Expediente"])
cert_est_data = [
    ("Toxina Botulínica Tipo A", "Diplomado o curso avalado de mínimo 40 horas teórico-prácticas con demostración de competencia en la técnica de inyección facial.", "Universidades con residencia de dermatología o cirugía plástica, Sociedad Colombiana de Cirugía Plástica (SCCP), Sociedad Colombiana de Dermatología (SCD), Instituto Allergan Training, Instituto de Estudios Avanzados en Estética.", "Actualización cada 2-3 años o ante cambios en técnica o productos.", "Certificado del curso con intensidad horaria, entidad avalante y nombre del médico instructor."),
    ("Rellenos Dérmicos con Ácido Hialurónico", "Diplomado o curso avalado de mínimo 40 horas teórico-prácticas. Debe incluir formación específica en anatomía vascular facial de riesgo y protocolo de emergencia vascular con hialuronidasa.", "SCCP, SCD, Instituto Galderma Training, cursos avalados internacionales (ISAPS, IMCAS).", "Actualización anual sobre nuevas técnicas y protocolo de emergencias.", "Certificado del curso. Constancia de entrenamiento en manejo de oclusión vascular."),
    ("Peelings Químicos Superficiales", "Curso avalado de mínimo 20 horas teórico-prácticas en agentes de peeling, preparación de la piel y manejo de complicaciones.", "SCD, universidades con dermatología, institutos de medicina estética avalados.", "Actualización cada 3-5 años.", "Certificado del curso."),
    ("Mesoterapia Facial", "Curso avalado de mínimo 20 horas teórico-prácticas en técnicas de mesoterapia y productos utilizados.", "Asociación Médica Colombiana de Mesoterapia, SCD, institutos especializados.", "Actualización cada 3 años.", "Certificado del curso."),
    ("Plasma Rico en Plaquetas (PRP)", "Curso avalado de mínimo 16 horas teórico-prácticas en preparación y aplicación de PRP.", "Laboratorios fabricantes de kits de PRP (Regen Lab, Harvest, BioCell), SCD, cursos de medicina regenerativa.", "Actualización cada 3 años o ante cambios en el protocolo del kit utilizado.", "Certificado del curso y manual del kit de PRP utilizado."),
]
for i, row_data in enumerate(cert_est_data):
    add_table_row(t_cert_est, list(row_data), shaded=(i % 2 == 1))
doc.add_paragraph()

# 7. EDUCACIÓN CONTINUADA
h1(doc, "7. PLAN DE EDUCACIÓN MÉDICA CONTINUA")
body(doc, "La educación médica continua (EMC) es una obligación ética y deontológica del médico (Ley 23 de 1981) y un requisito del Sistema de Gestión de Calidad para mantener las competencias actualizadas. El [NOMBRE DEL CONSULTORIO] establece un plan anual de EMC con un mínimo de 20 horas de actualización por año para el personal médico.")
t_emc = doc.add_table(rows=1, cols=6); t_emc.style = 'Table Grid'
table_header_row(t_emc, ["Mes", "Tema", "Modalidad", "Intensidad Horaria", "Institución/Fuente", "Pertinencia"])
emc_data = [
    ("Enero", "Actualización en hipertensión arterial: nuevas guías de manejo 2024-2025. Metas de PA, uso de nuevas moléculas.", "Virtual (webinar)", "3 horas", "Federación Médica Colombiana / Sociedad Colombiana de Cardiología", "Alta: HTA es la primera causa de consulta en primer nivel"),
    ("Febrero", "Novedades en manejo de diabetes mellitus tipo 2: iSGLT2, GLP-1, nuevas GPC colombianas.", "Virtual o presencial", "4 horas", "Asociación Colombiana de Diabetes / MINSALUD", "Alta: DM2 es la segunda enfermedad crónica más prevalente en consulta"),
    ("Marzo", "Taller práctico de reanimación cardiopulmonar BLS (AHA): actualización o primer certificado", "Presencial (simulación)", "8 horas", "American Heart Association / Cruz Roja / Universidad local", "Crítica: requisito de habilitación"),
    ("Abril", "Anatomía facial de riesgo para procedimientos estéticos inyectables. Actualización de protocolo de emergencia vascular.", "Virtual con casos fotográficos", "3 horas", "Sociedad Colombiana de Cirugía Plástica / ISAPS", "Crítica para la seguridad en procedimientos estéticos"),
    ("Mayo", "Novedades en formulación dermatológica para primer nivel: acné, dermatitis atópica, psoriasis leve.", "Virtual (webinar)", "2 horas", "Sociedad Colombiana de Dermatología", "Alta: las enfermedades dermatológicas representan el 15-20% de la consulta de MG"),
    ("Junio", "Actualización en salud mental en primer nivel: depresión, ansiedad, trastorno de adaptación. Herramientas diagnósticas y terapéuticas.", "Virtual o presencial", "4 horas", "Asociación Colombiana de Psiquiatría / MINSALUD (GPC depresión)", "Alta: la carga de trastornos mentales en primer nivel es subestimada"),
    ("Julio", "Farmacovigilancia y uso racional de antibióticos en primer nivel. Resistencia antimicrobiana.", "Virtual (curso corto)", "3 horas", "INVIMA / OPS / Universidad Nacional", "Alta: la prescripción irracional de antibióticos es un problema crítico de salud pública"),
    ("Agosto", "Novedades en toxina botulínica tipo A: nuevas indicaciones, técnicas de inyección avanzada, manejo de complicaciones.", "Presencial (taller)", "6 horas", "Instituto de formación en estética médica / Allergan Training / Galderma Academy", "Alta para el servicio de medicina estética"),
    ("Septiembre", "Manejo de dolor crónico y dolor musculoesquelético en primer nivel. Escala EVA, opioides menores.", "Virtual (webinar)", "2 horas", "Sociedad Colombiana de Anestesiología / ASMEDAS", "Media-alta: el dolor musculoesquelético crónico es una causa frecuente de consulta"),
    ("Octubre", "Actualización en medicina preventiva: tamizaje de cáncer de cuello uterino, mama, colon, hipertensión, diabetes. Esquemas de vacunación del adulto.", "Virtual o autoaprendizaje con material MINSALUD", "3 horas", "Ministerio de Salud y Protección Social / INS", "Alta: la medicina preventiva es un componente central del modelo de APS"),
    ("Noviembre", "Derechos de los pacientes y ética médica: casos colombianos recientes, jurisprudencia de la Corte Constitucional en salud.", "Virtual (conferencia)", "2 horas", "Tribunal Ético de Medicina / Facultades de Medicina", "Alta: el conocimiento ético es una obligación deontológica permanente"),
    ("Diciembre", "Revisión del año en habilitación de servicios de salud: cambios normativos 2025, resoluciones del MINSALUD, novedades del SOGCS.", "Autoaprendizaje con material oficial", "2 horas", "MINSALUD / Secretaría Seccional de Salud / [NOMBRE DEL CONSULTORIO]", "Alta: mantenerse actualizado en normatividad de habilitación"),
]
for i, row_data in enumerate(emc_data):
    add_table_row(t_emc, list(row_data), shaded=(i % 2 == 1))
doc.add_paragraph()
body(doc, "El registro de asistencia y aprobación de cada actividad de EMC se archiva en el expediente de talento humano del profesional. El cumplimiento del plan de EMC se evalúa en la evaluación anual de desempeño. Si un profesional no cumple con el mínimo de horas establecido, se establece un plan de recuperación para el trimestre siguiente.")

# 8. BIENESTAR LABORAL
h1(doc, "8. BIENESTAR LABORAL")
body(doc, "El consultorio [NOMBRE DEL CONSULTORIO] reconoce que el bienestar del personal de salud es un factor determinante de la calidad de la atención. Un médico con bienestar físico y mental, trabajo significativo y condiciones laborales adecuadas brinda una atención de mayor calidad y con mayor seguridad para el paciente. El programa de bienestar laboral del consultorio incluye las siguientes actividades trimestrales:")
t_bienestar = doc.add_table(rows=1, cols=4); t_bienestar.style = 'Table Grid'
table_header_row(t_bienestar, ["Trimestre", "Actividad de Bienestar", "Objetivo", "Responsable"])
bien_data = [
    ("Primer Trimestre (Ene-Mar)", "Evaluación de riesgos psicosociales laborales (aplicación del instrumento de valoración de riesgos psicosociales del MINSALUD). Retroalimentación individual de resultados.", "Identificar factores de estrés laboral, agotamiento emocional (burnout) o insatisfacción laboral para intervención temprana.", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("Segundo Trimestre (Abr-Jun)", "Jornada de autocuidado: actividad física (yoga, pilates, caminata), taller de manejo del estrés, nutrición saludable para el profesional de salud. Espacio de conversación sobre cargas laborales.", "Promover el autocuidado del personal de salud, prevenir el burnout y fortalecer los vínculos del equipo.", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("Tercer Trimestre (Jul-Sep)", "Revisión de condiciones ergonómicas del puesto de trabajo (postura al escribir, altura de la silla y escritorio, luz del consultorio). Ajustes si se requieren.", "Prevenir enfermedades profesionales de origen musculoesquelético (lumbalgia, síndrome del túnel del carpo, tendinitis).", "Dra. [NOMBRE DE LA MÉDICA]"),
    ("Cuarto Trimestre (Oct-Dic)", "Evaluación de satisfacción laboral: encuesta anónima sobre satisfacción con el trabajo, el ambiente laboral, el trato, las condiciones y las oportunidades de desarrollo. Retroalimentación de resultados y plan de acción para el año siguiente.", "Medir el clima laboral, identificar oportunidades de mejora y planificar acciones para el año siguiente.", "Dra. [NOMBRE DE LA MÉDICA]"),
]
for i, row_data in enumerate(bien_data):
    add_table_row(t_bienestar, list(row_data), shaded=(i % 2 == 1))
doc.add_paragraph()

# 9. AUSENTISMO
h1(doc, "9. GESTIÓN DEL AUSENTISMO LABORAL")
body(doc, "El ausentismo laboral en salud tiene consecuencias directas sobre la continuidad de la atención a los pacientes y la carga laboral del personal presente. El [NOMBRE DEL CONSULTORIO] establece las siguientes políticas para la gestión del ausentismo:")
h2(doc, "9.1 Tipos de Ausentismo")
t_ausent = doc.add_table(rows=1, cols=3); t_ausent.style = 'Table Grid'
table_header_row(t_ausent, ["Tipo de Ausentismo", "Definición", "Proceso de Gestión"])
ausen_data = [
    ("Ausentismo por enfermedad", "Ausencia del trabajo por incapacidad médica otorgada por médico tratante.", "Notificación al consultorio tan pronto se conozca la incapacidad. Presentación del original de la incapacidad al reintegro. Si la incapacidad es de más de 2 días, el consultorio debe notificar a los pacientes citados y reprogramar las citas con la menor demora posible."),
    ("Ausentismo por accidente de trabajo", "Ausencia por accidente ocurrido durante el ejercicio de las funciones laborales (incluyendo accidente biológico).", "Reporte inmediato a la ARL. Atención médica de urgencia si aplica. Seguimiento del proceso de calificación por la ARL. Reintegro con concepto médico de aptitud laboral."),
    ("Permisos laborales programados", "Ausencia programada por citas médicas, compromisos académicos, diligencias personales urgentes.", "Solicitud escrita con mínimo 48 horas de anticipación (excepto urgencias). Compensación del tiempo si el contrato lo prevé. Reprogramación de pacientes afectados."),
    ("Ausentismo no justificado", "Ausencia sin notificación previa o sin justificación válida.", "Conversación con el profesional para identificar causas. Si persiste: llamado de atención escrito, inicio de proceso disciplinario según el contrato vigente."),
    ("Ausentismo por capacitación o actividades académicas", "Ausencia programada para asistir a congresos, cursos o actividades de formación.", "El tiempo de capacitación relevante para el servicio puede ser reconocido por el consultorio (negociación caso a caso). Siempre con notificación previa y reprogramación de pacientes."),
]
for i, row_data in enumerate(ausen_data):
    add_table_row(t_ausent, list(row_data), shaded=(i % 2 == 1))
doc.add_paragraph()

# 10. DESVINCULACIÓN
h1(doc, "10. PROCESO DE DESVINCULACIÓN")
body(doc, "La desvinculación del personal, independientemente de su causa (renuncia voluntaria, terminación de contrato, terminación por justa causa, no renovación de contrato, pensión), debe realizarse de manera ordenada, justa y que garantice la continuidad de la atención a los pacientes y la conservación de la información institucional.")
bullet(doc, "Notificación formal: el colaborador debe notificar por escrito su renuncia con el tiempo de anticipación establecido en el contrato (usualmente 30 días para contratos de trabajo o 15 días para prestación de servicios).")
bullet(doc, "Entrega del cargo: el colaborador en proceso de desvinculación debe realizar la entrega formal del cargo, que incluye: historias clínicas físicas bajo su custodia, equipos y materiales a cargo, contraseñas de sistemas de información, documentos institucionales, formatos y sellos.")
bullet(doc, "Paz y salvo: el consultorio emite el paz y salvo al colaborador una vez se ha verificado que no hay pendientes de entrega.")
bullet(doc, "Liquidación de prestaciones: para contratos de trabajo, se realiza la liquidación de prestaciones sociales (cesantías, intereses, prima, vacaciones) de acuerdo con la normatividad laboral colombiana, dentro de los 15 días siguientes a la terminación del contrato.")
bullet(doc, "Entrevista de salida: se realiza una entrevista de salida con el colaborador que se desvincula voluntariamente, para identificar causas de la renuncia, oportunidades de mejora del clima laboral y del contrato de trabajo, y obtener retroalimentación sobre el funcionamiento del consultorio.")
bullet(doc, "Actualización del expediente: el expediente de talento humano se actualiza con el acta de desvinculación, la carta de liquidación y la evaluación de desempeño final.")
bullet(doc, "Confidencialidad post-vinculación: el colaborador desvinculado mantiene la obligación de confidencialidad sobre la información de los pacientes y los datos institucionales del consultorio, con carácter indefinido.")

# 11. BIBLIOGRAFÍA
h1(doc, "11. BIBLIOGRAFÍA")
refs_th = [
    "Congreso de la República de Colombia. Ley 1164 de 2007: Talento Humano en Salud. Bogotá: Congreso de la República; 2007.",
    "Ministerio de Salud y Protección Social de Colombia. Resolución 3100 de 2019: Habilitación de servicios de salud. Bogotá: MINSALUD; 2019.",
    "Ministerio de Salud y Protección Social de Colombia. Decreto 1011 de 2006: SOGCS. Bogotá: MINSALUD; 2006.",
    "Congreso de la República de Colombia. Ley 23 de 1981: Código de Ética Médica. Bogotá: Congreso de la República; 1981.",
    "Ministerio de Salud y Protección Social de Colombia. Rethus - Registro Único Nacional del Talento Humano en Salud. Disponible en: https://www.rethus.gov.co",
    "Organización Panamericana de la Salud (OPS). Manual para la Gestión del Talento Humano en Salud. Washington D.C.: OPS; 2018.",
    "Ministerio de Trabajo de Colombia. Código Sustantivo del Trabajo actualizado. Bogotá: Ministerio de Trabajo; 2023.",
    "Ministerio de Salud y Protección Social de Colombia. Batería de Instrumentos para la Evaluación de Factores de Riesgo Psicosocial. Bogotá: MINSALUD; 2010.",
]
for i, ref in enumerate(refs_th):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1); p.paragraph_format.first_line_indent = Cm(-1)
    run = p.add_run(f"{i+1}. {ref}"); run.font.size = Pt(10); run.font.name = 'Calibri'

doc.save(PATH)
print(f"DOC3 guardado: {PATH}")
