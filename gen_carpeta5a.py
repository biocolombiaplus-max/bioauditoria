#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Carpeta 5 - Parte A: MAN-PP-001, PRO-PP-001"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

BASE5="/home/user/bioauditoria/documentos_habilitacion/CARPETA_5_PROCESOS_PRIORITARIOS"

def margins(doc):
    for s in doc.sections:
        s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5); s.left_margin=Cm(3); s.right_margin=Cm(2.5)

def hf(doc,code,title):
    for s in doc.sections:
        h=s.header; h.paragraphs[0].clear()
        r=h.paragraphs[0].add_run(f"CONSULTORIO [NOMBRE DEL CONSULTORIO]  |  {code}  |  {title}")
        r.font.size=Pt(8); r.font.color.rgb=RGBColor(64,64,64)
        h.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        f=s.footer; f.paragraphs[0].clear()
        r2=f.paragraphs[0].add_run(f"{code} - V1.0 - {datetime.date.today().strftime('%d/%m/%Y')} | Pág. ")
        r2.font.size=Pt(8)
        fld=OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'),'begin')
        f.paragraphs[0].runs[-1]._r.append(fld)
        it=OxmlElement('w:instrText'); it.text='PAGE'
        f.paragraphs[0].runs[-1]._r.append(it)
        fld2=OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'),'end')
        f.paragraphs[0].runs[-1]._r.append(fld2)
        f.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

def portada(doc,code,title):
    doc.add_paragraph(); doc.add_paragraph()
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=t.add_run("CONSULTORIO MÉDICO\n[NOMBRE DEL CONSULTORIO]")
    r.font.size=Pt(18); r.font.bold=True; r.font.color.rgb=RGBColor(31,73,125)
    doc.add_paragraph()
    t2=doc.add_paragraph(); t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=t2.add_run(title); r2.font.size=Pt(15); r2.font.bold=True
    doc.add_paragraph()
    tb=doc.add_table(rows=6,cols=2); tb.style='Table Grid'; tb.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(l,v) in enumerate([("Código:",code),("Versión:","1.0"),("Fecha:",datetime.date.today().strftime('%d/%m/%Y')),
           ("Elaboró:","[NOMBRE DE LA MÉDICA], Médica General"),
           ("Revisó:","[NOMBRE DE LA MÉDICA]"),("Aprobó:","[NOMBRE DE LA MÉDICA]")]):
        tb.rows[i].cells[0].text=l; tb.rows[i].cells[1].text=v
        tb.rows[i].cells[0].paragraphs[0].runs[0].font.bold=True
    doc.add_page_break()
    cv=doc.add_paragraph(); r3=cv.add_run("CONTROL DE VERSIONES")
    r3.font.bold=True; r3.font.size=Pt(13); r3.font.color.rgb=RGBColor(31,73,125)
    cv.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tb2=doc.add_table(rows=3,cols=5); tb2.style='Table Grid'
    hrow(tb2,["Versión","Fecha","Cambio","Elaboró","Aprobó"])
    tb2.rows[1].cells[0].text="1.0"; tb2.rows[1].cells[1].text=datetime.date.today().strftime('%d/%m/%Y')
    tb2.rows[1].cells[2].text="Versión inicial"; tb2.rows[1].cells[3].text="[NOMBRE DE LA MÉDICA]"
    tb2.rows[1].cells[4].text="[NOMBRE DE LA MÉDICA]"
    doc.add_page_break()

def hrow(table,headers,bg="1F497D"):
    row=table.rows[0]
    for i,h in enumerate(headers):
        cell=row.cells[i]; cell.text=""
        run=cell.paragraphs[0].add_run(h)
        run.font.bold=True; run.font.color.rgb=RGBColor(255,255,255); run.font.size=Pt(10)
        cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),bg)
        shd.set(qn('w:color'),'auto'); shd.set(qn('w:val'),'clear'); tcPr.append(shd)

def H(doc,text,lv=1):
    p=doc.add_heading(text,level=lv)
    colors={1:RGBColor(31,73,125),2:RGBColor(46,116,181),3:RGBColor(68,114,196)}
    sizes={1:14,2:12,3:11}
    for run in p.runs:
        run.font.size=Pt(sizes[lv]); run.font.bold=True; run.font.color.rgb=colors[lv]
    return p

def B(doc,text,bold=False):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=p.add_run(text); r.font.size=Pt(11); r.font.bold=bold; r.font.name='Calibri'; return p

def BL(doc,text):
    p=doc.add_paragraph(style='List Bullet'); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=p.add_run(text); r.font.size=Pt(11); r.font.name='Calibri'; return p

def firmas(doc):
    doc.add_paragraph()
    tb=doc.add_table(rows=2,cols=3); tb.style='Table Grid'
    hrow(tb,["ELABORÓ","REVISÓ","APROBÓ"])
    for j in range(3):
        tb.rows[1].cells[j].text="[NOMBRE DE LA MÉDICA]\nMédica General - Propietaria\n\nFirma: _________________________\n\nFecha: "+datetime.date.today().strftime('%d/%m/%Y')
        tb.rows[1].cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

# ===== MAN-PP-001: Manual de Atención al Paciente =====
def crear_man_pp():
    doc=Document(); margins(doc); hf(doc,"MAN-PP-001","Manual de Atención al Paciente")
    portada(doc,"MAN-PP-001","MANUAL DE ATENCIÓN AL PACIENTE")

    H(doc,"1. OBJETIVO")
    B(doc,"Establecer los lineamientos, políticas, flujos y estándares para la atención integral, humanizada y segura de los pacientes que acuden al Consultorio Médico [NOMBRE DEL CONSULTORIO], garantizando el respeto a los derechos de los usuarios, la calidad científica y técnica de los servicios, y el cumplimiento de la normatividad colombiana vigente en materia de atención en salud.")

    H(doc,"2. ALCANCE")
    B(doc,"Aplica a todos los pacientes que acuden al Consultorio [NOMBRE DEL CONSULTORIO] para consultas de medicina general y para procedimientos estéticos no invasivos, en [DIRECCIÓN DEL CONSULTORIO], [CIUDAD]. Cubre el ciclo completo de atención: desde el agendamiento de la cita hasta el seguimiento post-consulta.")

    H(doc,"3. MARCO LEGAL")
    normas=[
        ("Resolución 13437 de 1991","Adopta el Decálogo de los Derechos de los Pacientes en Colombia. Base fundamental para la humanización de la atención."),
        ("Ley 100 de 1993","Sistema General de Seguridad Social en Salud. Establece los derechos de los usuarios del sistema de salud."),
        ("Resolución 1995 de 1999","Establece normas para el manejo de la historia clínica. Componente fundamental de la atención al paciente."),
        ("Decreto 1011 de 2006","SOGCS. Establece el sistema de garantía de calidad, incluyendo el trato humanizado y los derechos de los pacientes."),
        ("Resolución 0256 de 2016","Sistema de Información para la Calidad. Indicadores de atención al paciente."),
        ("Ley 1751 de 2015","Ley Estatutaria de Salud. Establece la salud como derecho fundamental y los estándares de atención."),
        ("Ley 1581 de 2012","Protección de datos personales (Habeas Data). Aplica al manejo de información de los pacientes."),
        ("Ley 23 de 1981","Ética médica. Establece los principios de la relación médico-paciente."),
        ("Resolución 2003 de 2014","Habilitación de servicios. Estándares de procesos prioritarios de atención."),
        ("Resolución 3100 de 2019","Estándares actualizados de habilitación para procesos prioritarios de atención."),
    ]
    tb=doc.add_table(rows=len(normas)+1,cols=2); tb.style='Table Grid'
    hrow(tb,["Norma","Descripción"])
    for i,(n,d) in enumerate(normas):
        tb.rows[i+1].cells[0].text=n; tb.rows[i+1].cells[1].text=d
        for c in tb.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()

    H(doc,"4. DERECHOS DE LOS PACIENTES")
    B(doc,"En cumplimiento de la Resolución 13437 de 1991 y la Ley 1751 de 2015, el Consultorio [NOMBRE DEL CONSULTORIO] garantiza a todos sus pacientes los siguientes derechos:")
    derechos=[
        ("Derecho a recibir atención digna","Todo paciente tiene derecho a recibir atención de salud digna, con calidad, oportunidad y continuidad, sin discriminación por razones de raza, sexo, edad, religión, condición socioeconómica o cualquier otra."),
        ("Derecho a ser informado","El paciente tiene derecho a recibir información clara, veraz y completa sobre su estado de salud, el diagnóstico, el plan de tratamiento, los riesgos y las alternativas disponibles, en lenguaje comprensible."),
        ("Derecho a la confidencialidad","La información contenida en la historia clínica y los datos de salud del paciente son confidenciales y no pueden ser divulgados sin su autorización, salvo las excepciones establecidas por ley."),
        ("Derecho al consentimiento informado","El paciente tiene derecho a aceptar o rechazar cualquier procedimiento diagnóstico o terapéutico después de haber sido informado de manera suficiente sobre sus características, beneficios, riesgos y alternativas."),
        ("Derecho a la libre elección","El paciente tiene derecho a elegir libremente a su médico y a la institución de salud que lo atenderá."),
        ("Derecho a recibir orientación","El paciente tiene derecho a recibir orientación sobre cómo acceder a los demás niveles de atención cuando sea necesario."),
        ("Derecho a presentar quejas y reclamos","El paciente tiene derecho a presentar quejas, reclamos y sugerencias relacionadas con la atención recibida y a recibir respuesta oportuna."),
        ("Derecho a morir dignamente","Cuando sea el caso, el paciente tiene derecho a que se respete su voluntad de morir con dignidad, sin aplicación de medidas desproporcionadas."),
        ("Derecho a la segunda opinión","El paciente tiene derecho a solicitar la opinión de otro médico sobre su diagnóstico o tratamiento sin que ello afecte negativamente su atención."),
        ("Derecho a la privacidad","El paciente tiene derecho a que se respete su privacidad e intimidad durante el examen físico y los procedimientos médicos."),
    ]
    for num,(der,desc) in enumerate(derechos,1):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r1=p.add_run(f"{num}. {der}: "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(desc); r2.font.size=Pt(11)

    H(doc,"5. DEBERES DE LOS PACIENTES")
    B(doc,"En concordancia con los derechos, los pacientes del Consultorio tienen los siguientes deberes:")
    deberes=["Cuidar su salud y la de su familia, adoptando estilos de vida saludables",
             "Proporcionar información veraz y completa sobre su estado de salud, síntomas, antecedentes y medicamentos que toma",
             "Cumplir con las indicaciones y el plan de tratamiento acordado con el médico",
             "Asistir puntualmente a las citas programadas o cancelarlas con anticipación si no puede asistir",
             "Tratar con respeto y dignidad al personal del consultorio",
             "Respetar las normas del consultorio y las indicaciones del personal de salud",
             "Usar adecuadamente y de manera racional los servicios de salud",
             "Cumplir con las obligaciones económicas pactadas con el consultorio",
             "No suministrar información falsa para acceder a beneficios no correspondidos"]
    for item in deberes: BL(doc,item)

    H(doc,"6. FLUJO DE ATENCIÓN DEL PACIENTE")
    H(doc,"6.1 Ciclo Completo de Atención",2)
    flujo=[("1. Agendamiento de cita","El paciente solicita cita a través de los medios disponibles: teléfono [TELÉFONO], WhatsApp [NÚMERO] o presencialmente. Se registran los datos básicos: nombre, documento, motivo de consulta. Se asigna fecha y hora según disponibilidad.","Administrativo / Médica"),
           ("2. Llegada y recepción","El paciente llega al consultorio. Se verifica la identidad con documento de identificación. Se confirma el tipo de consulta (medicina general o estética). Se registra la hora de llegada.","Administrativo / Médica"),
           ("3. Pre-consulta","Se registran datos vitales si aplica (tensión, temperatura, oximetría, talla, peso). Se verifica si el paciente trae resultados de exámenes o remisiones previas. Se verifica si es paciente nuevo (diligenciar datos en sistema/historia) o seguimiento.","Médica / Auxiliar"),
           ("4. Atención médica","Se realiza la consulta médica o el procedimiento estético según el proceso correspondiente (PRO-PP-001 o PRO-PP-002). Se elabora la historia clínica completa.","Médica"),
           ("5. Post-consulta","Se entrega la fórmula médica, incapacidad u órdenes de exámenes si aplica. Se explican las indicaciones al paciente. Se agenda cita de control si es necesario. Se cobra el servicio.","Médica / Administrativo"),
           ("6. Seguimiento","Para casos que lo requieran: llamada o mensaje de seguimiento 24-48 horas después. Para procedimientos estéticos: cita de seguimiento a los 15 días.","Médica"),
           ("7. Satisfacción","Aplicación de encuesta de satisfacción (periodicidad mensual). Gestión de quejas y reclamos.","Médica")]
    tb2=doc.add_table(rows=len(flujo)+1,cols=3); tb2.style='Table Grid'
    hrow(tb2,["Paso","Descripción","Responsable"])
    for i,(paso,desc,resp) in enumerate(flujo):
        tb2.rows[i+1].cells[0].text=paso; tb2.rows[i+1].cells[1].text=desc; tb2.rows[i+1].cells[2].text=resp
        for c in tb2.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"7. TIEMPOS DE ESPERA Y OPORTUNIDAD")
    B(doc,"El consultorio establece los siguientes tiempos estándar de atención:")
    tb3=doc.add_table(rows=6,cols=3); tb3.style='Table Grid'
    hrow(tb3,["Tipo de atención","Tiempo máximo de espera","Tiempo estándar de consulta"])
    tiempos=[("Consulta de medicina general (programada)","≤ 30 minutos desde la hora de la cita","20-30 minutos"),
             ("Consulta de medicina general (urgencia menor)","≤ 60 minutos desde la llegada","20-30 minutos"),
             ("Consulta de valoración estética","≤ 30 minutos desde la hora de la cita","30-45 minutos"),
             ("Procedimiento estético (con valoración previa)","≤ 15 minutos desde la hora pactada","Variable según procedimiento"),
             ("Entrega de resultados de laboratorio","≤ 20 minutos una vez el médico los revise","10-15 minutos")]
    for i,row in enumerate(tiempos):
        for j,v in enumerate(row): tb3.rows[i+1].cells[j].text=v
        for c in tb3.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"8. HUMANIZACIÓN DE LA ATENCIÓN")
    H(doc,"8.1 Principios de Humanización",2)
    B(doc,"El Consultorio [NOMBRE DEL CONSULTORIO] se fundamenta en el trato humanizado como eje central de su práctica médica. Los siguientes principios guían la atención:")
    principios_hum=[
        ("Dignidad:","Cada paciente es una persona única y valiosa, que merece ser tratada con respeto independientemente de su condición económica, social, cultural o de salud."),
        ("Empatía:","La médica propietaria se pone en el lugar del paciente, comprende sus miedos, expectativas y necesidades, respondiendo de manera sensible a sus preocupaciones."),
        ("Comunicación efectiva:","Se utiliza un lenguaje claro, sin tecnicismos excesivos, para explicar el diagnóstico y el plan de tratamiento. Se verifica que el paciente haya comprendido la información."),
        ("Escucha activa:","Se dedica tiempo suficiente para escuchar al paciente, sin interrupciones innecesarias, reconociendo que la historia del paciente es el elemento más valioso del diagnóstico."),
        ("Respeto a la autonomía:","Se respeta el derecho del paciente a tomar decisiones informadas sobre su salud, incluso cuando no coinciden con la recomendación médica."),
        ("Confidencialidad:","Se protege estrictamente la privacidad del paciente y la confidencialidad de su información de salud."),
        ("Continuidad:","Se garantiza la continuidad de la atención, el seguimiento de los casos y la referencia oportuna cuando se requiere atención de mayor complejidad."),
    ]
    for term,desc in principios_hum:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r1=p.add_run("• "+term+" "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(desc); r2.font.size=Pt(11)

    H(doc,"9. GESTIÓN DE QUEJAS Y RECLAMOS")
    H(doc,"9.1 Proceso de Gestión",2)
    B(doc,"El consultorio tiene un proceso establecido para la recepción, análisis y respuesta a quejas, reclamos y sugerencias de los pacientes:")
    quejas=[("Recepción","El paciente puede presentar su queja o reclamo de forma verbal (durante o después de la consulta), por escrito (formato disponible en el consultorio) o por medios digitales (correo electrónico, WhatsApp).","Inmediata"),
            ("Registro","Se registra la queja con: fecha, nombre del paciente (si autoriza), descripción detallada del problema, nombre de quien recibe.","Inmediata"),
            ("Análisis","La médica propietaria analiza la queja, determina la causa raíz y establece la respuesta y las medidas correctivas.","≤ 3 días hábiles"),
            ("Respuesta","Se comunica al paciente la respuesta y las acciones tomadas por el canal más apropiado.","≤ 5 días hábiles"),
            ("Seguimiento","Se verifica que las medidas correctivas se implementaron y que el paciente quedó satisfecho con la respuesta.","≤ 15 días"),
            ("Mejora continua","Las quejas recurrentes se analizan en la revisión mensual de calidad para implementar mejoras sistémicas.","Mensual")]
    tb4=doc.add_table(rows=len(quejas)+1,cols=4); tb4.style='Table Grid'
    hrow(tb4,["Etapa","Descripción","Responsable","Tiempo"])
    for i,(e,d,t) in enumerate(quejas):
        tb4.rows[i+1].cells[0].text=e; tb4.rows[i+1].cells[1].text=d
        tb4.rows[i+1].cells[2].text="[NOMBRE DE LA MÉDICA]"; tb4.rows[i+1].cells[3].text=t
        for c in tb4.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"10. PRIVACIDAD Y DIGNIDAD DEL PACIENTE")
    privacidad=["El consultorio médico debe tener privacidad visual y relativa privacidad auditiva durante la consulta",
                "El examen físico se realiza con el mínimo de exposición necesaria para el diagnóstico, con biombo o cortina si es necesario",
                "Se solicita al paciente su autorización para el registro fotográfico (procedimientos estéticos)",
                "La información del paciente no se discute en lugares públicos del consultorio",
                "El personal externo (proveedores, técnicos) no debe tener acceso al área clínica durante la atención de pacientes",
                "Las historias clínicas se almacenan de manera segura con acceso restringido (MAN-HC-001)",
                "En el tratamiento de datos personales se cumplen los principios de la Ley 1581 de 2012"]
    for item in privacidad: BL(doc,item)

    H(doc,"11. ATENCIÓN A POBLACIONES ESPECIALES")
    H(doc,"11.1 Niños y Adolescentes",2)
    B(doc,"Para la atención de menores de edad: el consentimiento informado es otorgado por el padre, madre o representante legal. Para adolescentes entre 14-17 años, se puede escuchar directamente al menor y considerar su asentimiento. Se garantiza privacidad en la consulta del adolescente salvo situaciones de riesgo.")
    H(doc,"11.2 Adultos Mayores",2)
    B(doc,"Para la atención de adultos mayores: se utiliza un tono de voz adecuado y se evita el lenguaje condescendiente. Se asegura que el paciente comprende las indicaciones (puede requerirse acompañante). Se tienen en cuenta las interacciones medicamentosas frecuentes en polimedicados.")
    H(doc,"11.3 Personas con Discapacidad",2)
    B(doc,"Se garantiza accesibilidad al consultorio según la Ley 361 de 1997. Se adapta la comunicación a las necesidades del paciente (hipoacusia, baja visión, etc.). Se solicita intérprete o apoyo cuando sea necesario.")

    H(doc,"12. INDICADORES")
    tb5=doc.add_table(rows=6,cols=4); tb5.style='Table Grid'
    hrow(tb5,["Indicador","Fórmula","Meta","Frecuencia"])
    inds=[("Satisfacción de usuarios","(Pacientes satisfechos/Total encuestados)×100","≥90%","Mensual"),
          ("Oportunidad en la asignación de citas","% de citas asignadas en ≤3 días hábiles","≥90%","Mensual"),
          ("Tiempo de espera dentro del consultorio","Promedio minutos entre hora cita y hora atención","≤30 min","Mensual"),
          ("Tasa de quejas y reclamos","N° quejas/100 consultas","Reducción progresiva","Mensual"),
          ("% quejas con respuesta en ≤5 días","(Quejas respondidas a tiempo/Total quejas)×100","100%","Mensual")]
    for i,row in enumerate(inds):
        for j,v in enumerate(row): tb5.rows[i+1].cells[j].text=v
        for c in tb5.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph(); firmas(doc)
    path=os.path.join(BASE5,"MAN-PP-001_Manual_Atencion_Paciente.docx")
    doc.save(path); print(f"✓ Creado: {path}")

# ===== PRO-PP-001: Proceso Consulta Medicina General =====
def crear_pro_pp001():
    doc=Document(); margins(doc); hf(doc,"PRO-PP-001","Proceso de Consulta de Medicina General")
    portada(doc,"PRO-PP-001","PROCESO DE CONSULTA DE MEDICINA GENERAL")

    H(doc,"1. OBJETIVO")
    B(doc,"Definir el proceso estandarizado y detallado de la consulta de medicina general del Consultorio [NOMBRE DEL CONSULTORIO], desde la recepción del paciente hasta el seguimiento post-consulta, garantizando una atención médica de calidad, segura, humanizada y conforme a las guías de práctica clínica colombianas vigentes y los estándares de habilitación de la Resolución 3100 de 2019.")

    H(doc,"2. ALCANCE")
    B(doc,"Aplica a todas las consultas de medicina general realizadas en el Consultorio [NOMBRE DEL CONSULTORIO] por la Dra. [NOMBRE DE LA MÉDICA]. Cubre los servicios de medicina general: atención de enfermedades agudas, enfermedades crónicas, actividades preventivas, expedición de certificados e incapacidades, y remisión a especialistas.")

    H(doc,"3. MARCO LEGAL")
    for n in ["Resolución 3100 de 2019 - Habilitación del servicio de consulta de medicina general",
              "Resolución 1995 de 1999 - Historia clínica",
              "Resolución 1478 de 2006 - Prescripción de medicamentos",
              "Ley 23 de 1981 - Ética médica",
              "Guías de Práctica Clínica del Ministerio de Salud vigentes",
              "Resolución 0256 de 2016 - Indicadores de calidad",
              "Decreto 1011 de 2006 - SOGCS"]: BL(doc,n)

    H(doc,"4. DESCRIPCIÓN PASO A PASO DEL PROCESO")
    H(doc,"4.1 PASO 1: Preparación del Consultorio",2)
    B(doc,"Antes de iniciar la jornada de atención, la médica propietaria verifica:")
    prep=["Lavado de manos o uso de alcohol glicerinado al inicio de la jornada",
          "Verificación del funcionamiento de todos los equipos: tensiómetro, oxímetro, glucómetro, termómetro, otoscopio, oftalmoscopio",
          "Disponibilidad del botiquín de urgencias completo con medicamentos en fecha",
          "Disponibilidad de insumos: guantes, gasas, agujas, jeringas, papel para camilla, desinfectantes",
          "Revisión de la agenda del día: número de pacientes, tipos de consulta, casos especiales",
          "Apertura del sistema de información (software médico o sistema físico de historias clínicas)"]
    for item in prep: BL(doc,item)

    H(doc,"4.2 PASO 2: Recepción e Identificación del Paciente",2)
    B(doc,"La identificación correcta del paciente es una medida fundamental de seguridad (Resolución 0256/2016):")
    ident=["Solicitar al paciente su documento de identidad: CC, TI, CE, pasaporte",
           "Verificar que el nombre y número de documento correspondan",
           "Para pacientes nuevos: diligenciar los datos de identificación completos en la historia clínica",
           "Para pacientes en seguimiento: buscar y verificar la historia clínica existente",
           "Registrar la hora de llegada y el motivo de consulta",
           "Verificar si el paciente tiene alergias conocidas (especialmente a medicamentos) y registrar en la historia",
           "Tomar signos vitales pre-consulta si hay auxiliar de enfermería, o la médica los toma durante la consulta"]
    for item in ident: BL(doc,item)

    H(doc,"4.3 PASO 3: Anamnesis (Historia de la Enfermedad Actual)",2)
    B(doc,"La anamnesis es el elemento más valioso del diagnóstico médico. Se realiza siguiendo el método clínico:")
    anamnesis_items=["Saludo y presentación de la médica (si es paciente nuevo)",
                     "Establecimiento del motivo de consulta en palabras del paciente",
                     "Historia de la enfermedad actual: inicio, evolución, localización, calidad, irradiación, intensidad (escala 1-10 para dolor), factores que mejoran/empeoran, síntomas asociados, tratamientos previos",
                     "Antecedentes personales patológicos: enfermedades crónicas (HTA, DM, hipotiroidismo, etc.), cirugías, hospitalizaciones previas",
                     "Antecedentes farmacológicos: medicamentos actuales (nombre, dosis, frecuencia), alergias medicamentosas",
                     "Antecedentes familiares: enfermedades hereditarias relevantes para el motivo de consulta",
                     "Antecedentes gineco-obstétricos (para mujeres): FUR, ciclos menstruales, embarazos, anticonceptivos, última citología, mamografía",
                     "Hábitos y estilos de vida: tabaquismo, alcohol, actividad física, alimentación, ocupación",
                     "Revisión por sistemas: interrogatorio sistemático de los sistemas orgánicos relevantes"]
    for item in anamnesis_items: BL(doc,item)

    H(doc,"4.4 PASO 4: Examen Físico",2)
    B(doc,"El examen físico se realiza de manera sistemática, con respeto a la dignidad del paciente:")
    examen=[("Signos vitales","Tensión arterial (brazo derecho, con el paciente sentado 5 minutos antes), frecuencia cardíaca, frecuencia respiratoria, temperatura, oximetría de pulso, talla, peso, IMC"),
            ("Examen general","Aspecto general, estado de conciencia, orientación, hidratación, coloración de piel y mucosas, facies"),
            ("Piel y faneras","Cambios en piel (manchas, lesiones, erupciones), uñas, cabello - especialmente relevante en valoración estética"),
            ("Cabeza y cuello","Cráneo, ojos (pupilas, reflejos, agudeza visual básica), oídos (otoscopia), nariz, boca/faringe (amígdalas, mucosa, dientes), cuello (ganglios, tiroides, pulsos carotídeos, ingurgitación yugular)"),
            ("Tórax","Inspección (simetría, tipo respiratorio), palpación (expansión, frémito), percusión (matidez, timpanismo), auscultación (murmullo vesicular, ruidos cardíacos: ritmo, tono, soplos)"),
            ("Abdomen","Inspección (cicatrices, distensión), auscultación (ruidos intestinales), percusión (timpanismo, matidez hepática/esplénica), palpación superficial y profunda (masa, dolor, visceromegalias), signos de irritación peritoneal"),
            ("Extremidades","Pulsos periféricos, edemas (con godet), varices, movilidad articular básica, reflejos osteotendinosos (si indicado)"),
            ("Neurológico básico","Estado mental (orientación, memoria reciente), pares craneales básicos (si indicado), coordinación, marcha (si indicado)"),
            ("Examen ginecológico","Solo cuando está indicado por el motivo de consulta y con consentimiento explícito de la paciente")]
    for area,desc in examen:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r1=p.add_run("• "+area+": "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(desc); r2.font.size=Pt(11)

    H(doc,"4.5 PASO 5: Diagnóstico e Impresión Diagnóstica",2)
    B(doc,"Con base en la anamnesis y el examen físico, la médica establece el diagnóstico:")
    diag=["Formular la impresión diagnóstica con códigos CIE-10 (Clasificación Internacional de Enfermedades, décima revisión)",
          "Establecer el diagnóstico diferencial cuando corresponda",
          "Determinar si se requieren exámenes complementarios para confirmar el diagnóstico o para descartar diagnósticos diferenciales",
          "Identificar las comorbilidades activas del paciente que impactan el diagnóstico principal",
          "Registrar todo en la historia clínica de manera clara y legible"]
    for item in diag: BL(doc,item)

    H(doc,"4.6 PASO 6: Plan de Manejo",2)
    B(doc,"El plan de manejo es la síntesis de las decisiones terapéuticas para el paciente:")
    H(doc,"4.6.1 Prescripción de Medicamentos (Resolución 1478 de 2006)",3)
    prescripcion=["La receta médica debe contener: lugar y fecha de expedición, nombre del paciente y documento de identidad, nombre del medicamento (DCI o genérico preferiblemente), presentación farmacéutica, dosis, vía de administración, frecuencia, duración del tratamiento",
                  "Firma y sello de la médica con tarjeta profesional",
                  "Para medicamentos controlados: usar receta especial según normativa de la Resolución 1478/2006",
                  "Explicar al paciente la forma de tomar los medicamentos, efectos secundarios frecuentes y signos de alarma",
                  "Verificar alergias antes de prescribir"]
    for item in prescripcion: BL(doc,item)

    H(doc,"4.6.2 Órdenes de Exámenes",3)
    B(doc,"Las órdenes de exámenes deben indicar: nombre del paciente, documento, tipo de examen solicitado, diagnóstico presuntivo (código CIE-10), indicación clínica, datos de la médica (nombre, TP), firma y sello. Se explica al paciente las indicaciones de preparación para los exámenes.")

    H(doc,"4.6.3 Incapacidades Médicas",3)
    B(doc,"La incapacidad médica debe ser expedida solo cuando hay justificación clínica real. Debe contener: nombre del paciente, documento, diagnóstico (CIE-10), número de días de incapacidad, fecha de inicio y fin, actividades que puede o no puede realizar, datos del médico, firma y sello. La incapacidad no puede extenderse más allá de lo clínicamente justificado.")

    H(doc,"4.6.4 Remisión a Especialistas (PRO-RCR-001)",3)
    B(doc,"Se remite al paciente cuando el motivo de consulta requiere atención especializada o de mayor complejidad. La remisión se realiza siguiendo el proceso de referencia establecido en MAN-RCR-001.")

    H(doc,"4.7 PASO 7: Educación al Paciente",2)
    edu=["Explicar en lenguaje sencillo el diagnóstico y lo que significa para el paciente",
         "Instruir sobre el régimen de tratamiento: medicamentos, medidas no farmacológicas, cambios en el estilo de vida",
         "Informar sobre los signos de alarma que requieren consulta de urgencias inmediata",
         "Promover hábitos de vida saludables relevantes para la condición del paciente",
         "Resolver las dudas del paciente antes de terminar la consulta",
         "Verificar que el paciente comprendió las instrucciones (técnica de retroinformación: pedir al paciente que repita las instrucciones clave)"]
    for item in edu: BL(doc,item)

    H(doc,"4.8 PASO 8: Cierre de la Consulta y Registro",2)
    cierre=["Completar el registro de la historia clínica con todos los elementos requeridos por la Resolución 1995/1999",
            "Registrar: fecha y hora de la consulta, motivo de consulta, anamnesis relevante, hallazgos del examen físico, diagnóstico(s) con código CIE-10, plan de manejo, medicamentos prescritos, exámenes solicitados, remisiones, próxima cita",
            "Firma de la médica en la historia clínica",
            "Cobro del servicio y entrega de recibo/factura",
            "Agendamiento de la cita de control si es necesario",
            "Acompañar al paciente hasta la salida con amabilidad"]
    for item in cierre: BL(doc,item)

    H(doc,"5. ACTIVIDADES PREVENTIVAS EN CONSULTA (DYPE)")
    B(doc,"En cada consulta, la médica aprovecha la oportunidad para realizar actividades de Detección Temprana y Protección Específica (DYPE) según la edad y el género del paciente:")
    dype=[("Adultos ≥18 años","Control de presión arterial (mínimo anual), toma de talla y peso (IMC), glicemia en ayunas (mínimo cada 3 años en >35 años), perfil lipídico"),
          ("Mujeres 25-65 años","Recordar vigencia de citología cervicovaginal, autoexamen de mama, mamografía (≥50 años cada 2 años)"),
          ("Adultos ≥40 años","Tamizaje de cáncer colorrectal (sangre oculta en heces, colonoscopia según riesgo), glaucoma"),
          ("Vacunación","Verificar vigencia del esquema PAI: tétanos (refuerzo cada 10 años), hepatitis B, influenza anual")]
    for edad,acc in dype:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r1=p.add_run("• "+edad+": "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(acc); r2.font.size=Pt(11)

    H(doc,"6. INDICADORES DEL PROCESO")
    tb=doc.add_table(rows=6,cols=4); tb.style='Table Grid'
    hrow(tb,["Indicador","Fórmula","Meta","Frecuencia"])
    inds=[("Calidad de historia clínica (HC completas)","(HC con todos los campos requeridos/Total HC auditadas)×100","≥95%","Mensual"),
          ("Satisfacción del paciente en consulta","(Pacientes satisfechos/Total encuestados)×100","≥90%","Mensual"),
          ("% consultas con diagnóstico CIE-10 registrado","(Consultas con CIE-10/Total consultas)×100","100%","Mensual"),
          ("% pacientes con plan de manejo documentado","(Pacientes con plan/Total consultas)×100","100%","Mensual"),
          ("Tasa de eventos adversos en consulta","N° eventos adversos/100 consultas","<1%","Mensual")]
    for i,row in enumerate(inds):
        for j,v in enumerate(row): tb.rows[i+1].cells[j].text=v
        for c in tb.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph(); firmas(doc)
    path=os.path.join(BASE5,"PRO-PP-001_Proceso_Consulta_Medicina_General.docx")
    doc.save(path); print(f"✓ Creado: {path}")

if __name__=="__main__":
    print("Generando Carpeta 5 - Parte A...")
    crear_man_pp(); crear_pro_pp001()
    print("✅ Parte A completada.")
