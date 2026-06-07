#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
GENERADOR DE DOCUMENTOS - CARPETA 1: TALENTO HUMANO
Sistema de Habilitación - Resolución 3100 de 2019
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

BASE_DIR = "/home/user/bioauditoria/documentos_habilitacion/CARPETA_1_TALENTO_HUMANO"

def set_page_margins(doc, top=2.5, bottom=2.5, left=3, right=2.5):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)

def add_header_footer(doc, code, title):
    for section in doc.sections:
        header = section.header
        header.paragraphs[0].clear()
        run = header.paragraphs[0].add_run(f"CONSULTORIO MÉDICO [NOMBRE DEL CONSULTORIO]  |  {code}  |  {title}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        footer = section.footer
        footer.paragraphs[0].clear()
        run2 = footer.paragraphs[0].add_run(f"{code} - Versión 1.0 - {datetime.date.today().strftime('%d/%m/%Y')}    |    Página ")
        run2.font.size = Pt(8)
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), 'begin')
        footer.paragraphs[0].runs[-1]._r.append(fld)
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        footer.paragraphs[0].runs[-1]._r.append(instrText)
        fld2 = OxmlElement('w:fldChar')
        fld2.set(qn('w:fldCharType'), 'end')
        footer.paragraphs[0].runs[-1]._r.append(fld2)
        footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def style_heading(paragraph, level=1):
    if level == 1:
        paragraph.style = 'Heading 1'
        for run in paragraph.runs:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        paragraph.style = 'Heading 2'
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        paragraph.style = 'Heading 3'
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

def add_title(doc, text, size=14, bold=True, color=RGBColor(0x1F, 0x49, 0x7D), center=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return p

def add_body(doc, text, size=11, bold=False, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Calibri'
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    return p

def add_table_header_row(table, headers, bg_color="1F497D"):
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), bg_color)
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

def add_portada(doc, codigo, titulo, version="1.0"):
    doc.add_paragraph()
    doc.add_paragraph()
    add_title(doc, "CONSULTORIO MÉDICO", 20)
    add_title(doc, "[NOMBRE DEL CONSULTORIO]", 18)
    doc.add_paragraph()
    doc.add_paragraph()
    add_title(doc, titulo, 16)
    doc.add_paragraph()

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    datos = [
        ("Código:", codigo),
        ("Versión:", version),
        ("Fecha de elaboración:", datetime.date.today().strftime('%d/%m/%Y')),
        ("Elaborado por:", "[NOMBRE DE LA MÉDICA], Médica General"),
        ("Revisado por:", "[NOMBRE DE LA MÉDICA], Médica General - Propietaria"),
        ("Aprobado por:", "[NOMBRE DE LA MÉDICA], Médica General - Propietaria"),
    ]
    for i, (label, val) in enumerate(datos):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = val
        for cell in table.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)
    doc.add_page_break()

def add_control_versiones(doc):
    h = doc.add_paragraph()
    run = h.add_run("CONTROL DE VERSIONES")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=4, cols=5)
    table.style = 'Table Grid'
    headers = ["Versión", "Fecha", "Descripción del cambio", "Elaboró", "Aprobó"]
    add_table_header_row(table, headers)

    data = [
        ["1.0", datetime.date.today().strftime('%d/%m/%Y'), "Versión inicial del documento", "[NOMBRE DE LA MÉDICA]", "[NOMBRE DE LA MÉDICA]"],
        ["", "", "", "", ""],
        ["", "", "", "", ""],
    ]
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            table.rows[i+1].cells[j].text = val
            for para in table.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_page_break()

# ============================================================
# DOCUMENTO 1: MAN-TH-001 Manual de Talento Humano
# ============================================================
def crear_manual_talento_humano():
    doc = Document()
    set_page_margins(doc)
    add_header_footer(doc, "MAN-TH-001", "Manual de Talento Humano")

    add_portada(doc, "MAN-TH-001", "MANUAL DE GESTIÓN DEL TALENTO HUMANO")
    add_control_versiones(doc)

    # TABLA DE CONTENIDO (manual)
    add_title(doc, "TABLA DE CONTENIDO", 13)
    toc_items = [
        "1. Objetivo", "2. Alcance", "3. Marco Legal", "4. Definiciones",
        "5. Política de Talento Humano", "6. Estructura Organizacional",
        "7. Proceso de Selección y Vinculación", "8. Inducción y Reinducción",
        "9. Evaluación de Competencias", "10. Educación Continuada",
        "11. Gestión de Turnos y Disponibilidad", "12. Bienestar Laboral",
        "13. Manejo de Situaciones Especiales", "14. Responsables",
        "15. Indicadores", "16. Anexos"
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.runs[0].font.size = Pt(11)
    doc.add_page_break()

    # SECCIÓN 1: OBJETIVO
    h = doc.add_heading("1. OBJETIVO", level=1)
    style_heading(h, 1)
    add_body(doc, "Establecer los lineamientos, políticas, procedimientos y criterios para la gestión integral del talento humano del Consultorio Médico [NOMBRE DEL CONSULTORIO], garantizando que el personal cuente con las competencias, idoneidad y condiciones necesarias para brindar una atención médica de calidad, segura y humanizada a los pacientes, cumpliendo con todos los requisitos establecidos en la Resolución 3100 de 2019 del Ministerio de Salud y Protección Social de Colombia y demás normatividad vigente aplicable.")
    add_body(doc, "Este manual busca asegurar que el recurso humano del consultorio sea seleccionado, vinculado, capacitado y evaluado bajo criterios técnicos y legales que garanticen la prestación segura y efectiva de los servicios de salud, incluyendo los servicios de medicina general y los procedimientos estéticos no invasivos dentro del marco legal permitido para médicos generales.")

    # SECCIÓN 2: ALCANCE
    h2 = doc.add_heading("2. ALCANCE", level=1)
    style_heading(h2, 1)
    add_body(doc, "El presente manual aplica a:")
    items_alcance = [
        "La médica general propietaria del consultorio: [NOMBRE DE LA MÉDICA]",
        "Todo el personal de apoyo administrativo y/o asistencial que preste sus servicios en el consultorio, de manera directa o indirecta",
        "El proceso de selección, vinculación, inducción, capacitación y evaluación de todo el talento humano",
        "Los servicios habilitados: Consulta de Medicina General y Procedimientos Estéticos No Invasivos",
        "El sitio de atención: [DIRECCIÓN DEL CONSULTORIO], [CIUDAD]"
    ]
    for item in items_alcance:
        add_bullet(doc, item)

    # SECCIÓN 3: MARCO LEGAL
    h3 = doc.add_heading("3. MARCO LEGAL", level=1)
    style_heading(h3, 1)
    add_body(doc, "El presente manual se fundamenta en la siguiente normatividad colombiana vigente:")

    leyes = [
        ("Ley 100 de 1993", "Por la cual se crea el Sistema General de Seguridad Social en Salud en Colombia. Establece las bases para la organización y prestación de servicios de salud, incluyendo los criterios de calidad y los requisitos para los prestadores de servicios de salud."),
        ("Ley 23 de 1981", "Por la cual se dictan normas en materia de ética médica. Establece los deberes y obligaciones del médico en el ejercicio de su profesión, los principios de la relación médico-paciente y los estándares de conducta profesional."),
        ("Ley 14 de 1962", "Por la cual se dictan normas relativas al ejercicio de la medicina y cirugía. Regula el ejercicio profesional de la medicina en Colombia, estableciendo los requisitos de habilitación profesional."),
        ("Decreto 3380 de 1981", "Por el cual se reglamenta la Ley 23 de 1981. Define con precisión los conceptos de acto médico, prescripción médica y los límites de la práctica médica."),
        ("Decreto 1011 de 2006", "Por el cual se establece el Sistema Obligatorio de Garantía de Calidad de la Atención de Salud del Sistema General de Seguridad Social en Salud (SOGCS). Define los componentes del sistema de garantía de calidad, incluyendo el Sistema Único de Habilitación."),
        ("Resolución 3100 de 2019", "Por la cual se definen los procedimientos y condiciones de inscripción de los prestadores de servicios de salud y de habilitación de los servicios de salud. Establece los estándares mínimos de habilitación, incluyendo los requisitos de talento humano para cada servicio."),
        ("Resolución 2003 de 2014", "Por la cual se definen los procedimientos y condiciones de inscripción de los prestadores de servicios de salud. Normativa predecesora que aún sirve de referencia para muchos estándares de habilitación."),
        ("Decreto 1072 de 2015", "Decreto Único Reglamentario del Sector Trabajo. Establece las normas generales de las relaciones laborales en Colombia, aplicables al personal del consultorio."),
        ("Ley 1562 de 2012", "Por la cual se modifica el sistema de riesgos laborales. Establece las obligaciones del empleador en materia de seguridad y salud en el trabajo."),
        ("Resolución 0256 de 2016", "Por la cual se dictan disposiciones en relación con el Sistema de Información para la Calidad. Define los indicadores de calidad de los servicios de salud que deben ser medidos y reportados."),
    ]

    table = doc.add_table(rows=len(leyes)+1, cols=2)
    table.style = 'Table Grid'
    add_table_header_row(table, ["Norma", "Descripción"], "1F497D")
    for i, (norma, desc) in enumerate(leyes):
        table.rows[i+1].cells[0].text = norma
        table.rows[i+1].cells[1].text = desc
        for cell in table.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # SECCIÓN 4: DEFINICIONES
    h4 = doc.add_heading("4. DEFINICIONES Y TERMINOLOGÍA", level=1)
    style_heading(h4, 1)

    definiciones = [
        ("Talento Humano en Salud:", "Conjunto de personas que ejercen una ocupación relacionada con la prestación de servicios de salud, ya sea en el campo asistencial, administrativo, de investigación o de formación."),
        ("Habilitación:", "Proceso mediante el cual el Estado verifica el cumplimiento de las condiciones mínimas de capacidad tecnológica y científica, de suficiencia patrimonial y financiera y de capacidad técnico-administrativa para la entrada y permanencia en el mercado de la prestación de servicios de salud."),
        ("Competencia:", "Capacidad de aplicar conocimientos, habilidades y actitudes en el desempeño de una función productiva, de acuerdo con criterios de desempeño y evidencias requeridas."),
        ("Médico General:", "Profesional graduado en medicina con título universitario reconocido por el Estado colombiano, con Tarjeta Profesional vigente expedida por el Tribunal Ético de Medicina."),
        ("Inducción:", "Proceso mediante el cual se suministra al trabajador nuevo o reincorporado, información necesaria sobre la institución, sus funciones, responsabilidades y el entorno laboral."),
        ("Reinducción:", "Proceso de actualización y refrescamiento sobre las políticas, procedimientos y cambios organizacionales que se realiza periódicamente a todo el personal."),
        ("Educación Continuada:", "Proceso sistemático y permanente de actualización de conocimientos, habilidades y actitudes del personal de salud, orientado a mejorar la calidad de la atención."),
        ("Evaluación de Desempeño:", "Proceso sistemático y periódico de estimación cuantitativa y cualitativa del grado de eficacia con que las personas llevan a cabo las actividades y responsabilidades de los puestos que desempeñan."),
        ("RETHUS:", "Registro Único Nacional del Talento Humano en Salud. Sistema de información que registra y certifica los datos de formación y experiencia del talento humano en salud en Colombia."),
        ("Tarjeta Profesional:", "Documento expedido por el Tribunal Ético de Medicina que habilita al médico para ejercer legalmente la profesión en Colombia."),
    ]

    for term, defi in definiciones:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r1 = p.add_run(term + " ")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(defi)
        r2.font.size = Pt(11)

    doc.add_paragraph()

    # SECCIÓN 5: POLÍTICA DE TALENTO HUMANO
    h5 = doc.add_heading("5. POLÍTICA DE TALENTO HUMANO", level=1)
    style_heading(h5, 1)
    add_body(doc, "El Consultorio Médico [NOMBRE DEL CONSULTORIO], bajo la dirección de la Dra. [NOMBRE DE LA MÉDICA], se compromete a mantener un talento humano altamente calificado, ético y comprometido con la excelencia en la prestación de servicios de salud. Esta política se fundamenta en los siguientes principios:")

    h5a = doc.add_heading("5.1 Principios Rectores de la Política de Talento Humano", level=2)
    style_heading(h5a, 2)

    principios = [
        ("Idoneidad y Competencia:", "Todo el personal que preste servicios en el consultorio debe contar con la formación, títulos, certificaciones y competencias necesarias para el desempeño de sus funciones, en cumplimiento de la Resolución 3100 de 2019."),
        ("Ética Profesional:", "El ejercicio de las funciones debe estar guiado por los principios éticos establecidos en la Ley 23 de 1981 y el Código de Ética Médica, garantizando siempre el bienestar del paciente."),
        ("Actualización Permanente:", "El consultorio promueve y facilita la educación continuada del personal, reconociendo que el conocimiento médico evoluciona constantemente y que la actualización es un deber ético y legal."),
        ("Bienestar Laboral:", "Se garantizan condiciones de trabajo dignas, seguras y que promuevan el bienestar físico, mental y social del personal, en cumplimiento de la normatividad de riesgos laborales."),
        ("Mejora Continua:", "Se establecen mecanismos periódicos de evaluación del desempeño y retroalimentación constructiva, orientados a la mejora continua de las competencias del personal."),
        ("Respeto y Dignidad:", "Las relaciones laborales se rigen por el respeto mutuo, la dignidad humana y la no discriminación por ningún motivo."),
    ]

    for term, desc in principios:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.5)
        r1 = p.add_run("• " + term + " ")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(desc)
        r2.font.size = Pt(11)

    # SECCIÓN 6: ESTRUCTURA ORGANIZACIONAL
    h6 = doc.add_heading("6. ESTRUCTURA ORGANIZACIONAL", level=1)
    style_heading(h6, 1)
    add_body(doc, "El Consultorio Médico [NOMBRE DEL CONSULTORIO] cuenta con la siguiente estructura organizacional para la prestación de sus servicios:")

    h6a = doc.add_heading("6.1 Cargos y Funciones", level=2)
    style_heading(h6a, 2)

    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Table Grid'
    add_table_header_row(table2, ["Cargo", "Perfil Requerido", "Funciones Principales"])
    cargos = [
        ("Médica General - Propietaria\n[NOMBRE DE LA MÉDICA]",
         "- Título de Médica y Cirujana\n- Tarjeta Profesional vigente\n- Registro RETHUS activo\n- Formación en procedimientos estéticos no invasivos",
         "- Atención de consultas médicas\n- Realización de procedimientos estéticos no invasivos\n- Gestión del consultorio\n- Supervisión de la calidad"),
        ("Personal Administrativo\n(si aplica)",
         "- Bachiller o técnico en administración\n- Conocimientos básicos en salud\n- Manejo de sistemas informáticos",
         "- Agendamiento de citas\n- Facturación\n- Atención al usuario\n- Gestión documental"),
        ("Personal de Apoyo\n(si aplica)",
         "- Formación en auxiliar de servicios generales\n- Conocimiento en bioseguridad básica",
         "- Aseo y desinfección\n- Apoyo logístico"),
        ("Proveedor externo de mantenimiento\n(si aplica)",
         "- Certificaciones técnicas específicas\n- Contratos vigentes",
         "- Mantenimiento de equipos biomédicos\n- Mantenimiento de instalaciones"),
    ]
    for i, (cargo, perfil, funcs) in enumerate(cargos):
        table2.rows[i+1].cells[0].text = cargo
        table2.rows[i+1].cells[1].text = perfil
        table2.rows[i+1].cells[2].text = funcs
        for cell in table2.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # SECCIÓN 7: PROCESO DE SELECCIÓN Y VINCULACIÓN
    h7 = doc.add_heading("7. PROCESO DE SELECCIÓN Y VINCULACIÓN", level=1)
    style_heading(h7, 1)

    h7a = doc.add_heading("7.1 Requisitos Obligatorios de Habilitación", level=2)
    style_heading(h7a, 2)
    add_body(doc, "De acuerdo con la Resolución 3100 de 2019, para el servicio de Consulta de Medicina General, el talento humano debe cumplir con los siguientes requisitos mínimos:")

    req_items = [
        "Título de Médico y Cirujano otorgado por institución de educación superior legalmente reconocida en Colombia, o título extranjero debidamente revalidado según las normas vigentes.",
        "Tarjeta Profesional vigente expedida por el Tribunal Ético de Medicina, sin sanciones disciplinarias vigentes.",
        "Inscripción activa y vigente en el Registro Único Nacional del Talento Humano en Salud (RETHUS) del Ministerio de Salud y Protección Social.",
        "Certificado de antecedentes disciplinarios emitido por el Tribunal Ético de Medicina, sin sanciones vigentes.",
        "Certificado de antecedentes judiciales (Policía Nacional) sin anotaciones que impidan el ejercicio profesional.",
        "Certificado de antecedentes fiscales (Contraloría General de la República) sin anotaciones.",
        "Certificado de antecedentes en el Registro de Autores, Participes y Determinadores de Procesos Disciplinarios (SIRI) de la Procuraduría General de la Nación.",
        "Certificados de formación y capacitación en los procedimientos específicos que realizará (para procedimientos estéticos no invasivos, certificados de los cursos específicos realizados).",
        "Tarjeta de vacunación al día (mínimo Hepatitis B, Tétanos, Influenza).",
        "Exámenes médicos de ingreso con énfasis osteomusicular y de riesgo biológico.",
    ]
    for item in req_items:
        add_bullet(doc, item)

    h7b = doc.add_heading("7.2 Documentos Adicionales Requeridos", level=2)
    style_heading(h7b, 2)
    add_body(doc, "Adicionalmente a los requisitos de habilitación, se requieren los siguientes documentos:")
    docs_adicionales = [
        "Hoja de vida actualizada con soportes de estudios y experiencia",
        "Fotocopia del documento de identidad",
        "Soportes de estudios de pregrado y posgrado (si aplica)",
        "Certificados laborales de experiencia previa",
        "Certificados de cursos, diplomados y especializaciones relevantes",
        "Para procedimientos estéticos no invasivos: certificados de cursos específicos en toxina botulínica, rellenos dérmicos, peelings químicos, mesoterapia y/o PRP emitidos por instituciones reconocidas",
        "RUT actualizado",
        "Número de cuenta bancaria para pagos (si aplica)",
        "Afiliación a EPS, fondo de pensiones y ARL vigente",
    ]
    for item in docs_adicionales:
        add_bullet(doc, item)

    add_body(doc, "NOTA: Para la Dra. [NOMBRE DE LA MÉDICA] como médica propietaria, los siguientes certificados de formación en procedimientos estéticos deben estar archivados en la carpeta de habilitación:")
    add_bullet(doc, "[CERTIFICADO DE CURSO EN TOXINA BOTULÍNICA TIPO A - institución y fecha]")
    add_bullet(doc, "[CERTIFICADO DE CURSO EN RELLENOS DÉRMICOS CON ÁCIDO HIALURÓNICO - institución y fecha]")
    add_bullet(doc, "[CERTIFICADO DE CURSO EN PEELINGS QUÍMICOS - institución y fecha]")
    add_bullet(doc, "[CERTIFICADO DE CURSO EN MESOTERAPIA - institución y fecha]")
    add_bullet(doc, "[CERTIFICADO DE CURSO EN PLASMA RICO EN PLAQUETAS - institución y fecha]")

    h7c = doc.add_heading("7.3 Etapas del Proceso de Selección", level=2)
    style_heading(h7c, 2)

    etapas = [
        ("ETAPA 1", "Definición del perfil requerido", "Se establece el perfil del cargo con los requisitos técnicos, legales y personales necesarios para el puesto. Se verifica que los requisitos cumplan con la normatividad de habilitación vigente."),
        ("ETAPA 2", "Publicación de la vacante", "Se publica la vacante a través de los medios disponibles (redes profesionales, bolsas de empleo médico, referencias de colegas). Se especifican claramente los requisitos obligatorios."),
        ("ETAPA 3", "Revisión de hojas de vida", "Se revisan las hojas de vida recibidas verificando el cumplimiento de los requisitos mínimos. Se verifica la autenticidad de los títulos y certificaciones a través de RETHUS y los sistemas oficiales."),
        ("ETAPA 4", "Entrevista", "Se realiza entrevista técnica y personal para evaluar conocimientos, competencias y actitudes del candidato. Para personal médico, se incluye evaluación de casos clínicos."),
        ("ETAPA 5", "Verificación de documentos", "Se verifican todos los documentos obligatorios: tarjeta profesional en RETHUS, antecedentes disciplinarios, judiciales, fiscales y penales."),
        ("ETAPA 6", "Exámenes preocupacionales", "Se realizan los exámenes médicos de ingreso requeridos, con énfasis en riesgo biológico y condiciones para el cargo específico."),
        ("ETAPA 7", "Vinculación y firma de contrato", "Se formaliza la vinculación mediante contrato laboral o de prestación de servicios según corresponda, en cumplimiento del Código Sustantivo del Trabajo."),
    ]

    table3 = doc.add_table(rows=len(etapas)+1, cols=3)
    table3.style = 'Table Grid'
    add_table_header_row(table3, ["Etapa", "Denominación", "Descripción"])
    for i, (etapa, nombre, desc) in enumerate(etapas):
        table3.rows[i+1].cells[0].text = etapa
        table3.rows[i+1].cells[1].text = nombre
        table3.rows[i+1].cells[2].text = desc
        for cell in table3.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # SECCIÓN 8: INDUCCIÓN Y REINDUCCIÓN
    h8 = doc.add_heading("8. PROCESO DE INDUCCIÓN Y REINDUCCIÓN", level=1)
    style_heading(h8, 1)

    h8a = doc.add_heading("8.1 Programa de Inducción", level=2)
    style_heading(h8a, 2)
    add_body(doc, "Todo el personal que se vincule al consultorio, independientemente del tipo de contrato, debe recibir un proceso de inducción antes de iniciar sus actividades asistenciales o administrativas. El programa de inducción tiene una duración mínima de 8 horas y cubre los siguientes temas:")

    temas_induccion = [
        ("Módulo 1 - Institucional (2 horas):", [
            "Historia y misión del consultorio",
            "Valores institucionales y cultura organizacional",
            "Estructura organizacional y canales de comunicación",
            "Políticas institucionales generales",
            "Reglamento interno de trabajo",
        ]),
        ("Módulo 2 - Normativo y de Calidad (2 horas):", [
            "Sistema de Garantía de Calidad en Salud (Decreto 1011/2006)",
            "Resolución 3100 de 2019 - Estándares de habilitación",
            "Política de calidad y seguridad del paciente",
            "Sistema de gestión documental del consultorio",
            "Indicadores de calidad",
        ]),
        ("Módulo 3 - Seguridad del Paciente (2 horas):", [
            "Política de seguridad del paciente (Resolución 0256/2016)",
            "Identificación segura del paciente",
            "Prevención y reporte de eventos adversos",
            "Protocolo de manejo de emergencias",
            "Derechos y deberes de los pacientes",
        ]),
        ("Módulo 4 - Bioseguridad y Manejo de Residuos (2 horas):", [
            "Precauciones estándar y universales",
            "Protocolo de lavado de manos (OMS)",
            "Uso de elementos de protección personal (EPP)",
            "Manejo de residuos biosanitarios (Decreto 351/2014)",
            "Riesgo biológico y medidas de prevención",
        ]),
    ]

    for modulo, items in temas_induccion:
        p = doc.add_paragraph()
        r = p.add_run(modulo)
        r.font.bold = True
        r.font.size = Pt(11)
        for item in items:
            add_bullet(doc, item)

    h8b = doc.add_heading("8.2 Programa de Reinducción", level=2)
    style_heading(h8b, 2)
    add_body(doc, "La reinducción se realizará con una frecuencia mínima anual y adicionalmente cuando se presenten los siguientes eventos:")
    reinduccion_eventos = [
        "Cambios significativos en los procesos o procedimientos del consultorio",
        "Actualización o cambio de normatividad aplicable",
        "Después de períodos de ausencia del trabajador superiores a 30 días continuos",
        "Cuando se detecten fallas o no conformidades relacionadas con el desconocimiento de procedimientos",
        "Ingreso de nuevos servicios o tecnologías al consultorio",
        "Resultados desfavorables en la evaluación de desempeño relacionados con conocimiento de procesos",
    ]
    for item in reinduccion_eventos:
        add_bullet(doc, item)

    add_body(doc, "El registro de participación en los procesos de inducción y reinducción debe quedar en el expediente del trabajador y en el registro de asistencia correspondiente, firmado por el trabajador y por la médica propietaria.")

    # SECCIÓN 9: EVALUACIÓN DE COMPETENCIAS
    h9 = doc.add_heading("9. EVALUACIÓN DE COMPETENCIAS Y DESEMPEÑO", level=1)
    style_heading(h9, 1)

    add_body(doc, "La evaluación del desempeño es un proceso sistemático y periódico que permite valorar el nivel de cumplimiento de las funciones, competencias y objetivos del personal del consultorio. Este proceso tiene una periodicidad anual, con seguimientos semestrales.")

    h9a = doc.add_heading("9.1 Dimensiones de la Evaluación", level=2)
    style_heading(h9a, 2)

    dimensiones = [
        ("Competencias Técnicas (40%):", "Evalúa los conocimientos y habilidades específicas para el desempeño del cargo: conocimientos médicos actualizados, manejo de procedimientos, uso de equipos, gestión de historia clínica, cumplimiento de protocolos."),
        ("Competencias en Seguridad del Paciente (20%):", "Evalúa el cumplimiento de las políticas de seguridad: identificación correcta del paciente, reporte de eventos adversos, cumplimiento de listas de chequeo, lavado de manos."),
        ("Competencias Humanísticas y de Comunicación (20%):", "Evalúa la calidad de la relación con el paciente: comunicación efectiva, respeto, empatía, manejo de información confidencial, atención a quejas y sugerencias."),
        ("Competencias Administrativas y de Calidad (10%):", "Evalúa el cumplimiento de los procesos administrativos: diligenciamiento correcto de historias clínicas, gestión documental, reporte oportuno de indicadores."),
        ("Cumplimiento de Normas y Reglamento (10%):", "Evalúa el cumplimiento de las normas institucionales: puntualidad, presentación personal, uso de EPP, participación en capacitaciones."),
    ]

    for dim, desc in dimensiones:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r1 = p.add_run("• " + dim + " ")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(desc)
        r2.font.size = Pt(11)

    h9b = doc.add_heading("9.2 Escala de Calificación", level=2)
    style_heading(h9b, 2)

    table_cal = doc.add_table(rows=6, cols=3)
    table_cal.style = 'Table Grid'
    add_table_header_row(table_cal, ["Calificación", "Rango", "Descripción"])
    cal_data = [
        ["Sobresaliente", "4.5 - 5.0", "Supera consistentemente las expectativas del cargo en todas las dimensiones evaluadas"],
        ["Satisfactorio", "4.0 - 4.4", "Cumple con las expectativas del cargo en la mayoría de las dimensiones evaluadas"],
        ["Aceptable", "3.0 - 3.9", "Cumple básicamente con las expectativas pero presenta áreas de mejora identificadas"],
        ["En mejoramiento", "2.0 - 2.9", "No cumple plenamente con las expectativas. Requiere plan de mejoramiento"],
        ["Insatisfactorio", "Menor a 2.0", "No cumple con las expectativas del cargo. Se activan medidas correctivas"],
    ]
    for i, row in enumerate(cal_data):
        for j, val in enumerate(row):
            table_cal.rows[i+1].cells[j].text = val
            for para in table_cal.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # SECCIÓN 10: EDUCACIÓN CONTINUADA
    h10 = doc.add_heading("10. PROGRAMA DE EDUCACIÓN CONTINUADA", level=1)
    style_heading(h10, 1)

    add_body(doc, "El consultorio reconoce que la educación médica continua es un deber ético y legal del médico general, establecido en la Ley 23 de 1981 y los principios del Sistema de Garantía de Calidad. Se establece un programa anual de educación continuada que contempla:")

    h10a = doc.add_heading("10.1 Plan Anual de Capacitación", level=2)
    style_heading(h10a, 2)

    capacitaciones = [
        ("Anual (mínimo)", "Actualización en guías de práctica clínica colombianas vigentes", "Médica General", "Interna/Externa", "8 horas"),
        ("Anual (mínimo)", "Soporte básico de vida (BLS/RCP) - recertificación", "Todo el personal", "Externa", "8 horas"),
        ("Semestral", "Actualización en bioseguridad y manejo de residuos", "Todo el personal", "Interna", "4 horas"),
        ("Semestral", "Seguridad del paciente y reporte de eventos adversos", "Todo el personal", "Interna", "4 horas"),
        ("Anual", "Actualización en procedimientos estéticos no invasivos", "Médica General", "Externa", "[HORAS]"),
        ("Anual", "Actualización normativa en habilitación de servicios de salud", "Médica General", "Externa/Interna", "4 horas"),
        ("Anual", "Manejo de emergencias en consultorio", "Médica General", "Externa", "8 horas"),
        ("Según necesidad", "Cursos específicos identificados en evaluación de desempeño", "Según evaluación", "Externa", "Variable"),
    ]

    table_cap = doc.add_table(rows=len(capacitaciones)+1, cols=5)
    table_cap.style = 'Table Grid'
    add_table_header_row(table_cap, ["Frecuencia", "Tema", "Dirigido a", "Modalidad", "Intensidad"])
    for i, row in enumerate(capacitaciones):
        for j, val in enumerate(row):
            table_cap.rows[i+1].cells[j].text = val
            for para in table_cap.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    add_body(doc, "NOTA IMPORTANTE: Para mantener la habilitación de los servicios de procedimientos estéticos no invasivos, la médica propietaria debe acreditar actualización periódica en los procedimientos que realiza. Se recomienda participar al menos en un congreso o curso de actualización en medicina estética por año, con los correspondientes certificados archivados en la carpeta de habilitación.")

    # SECCIÓN 11: GESTIÓN DE TURNOS
    h11 = doc.add_heading("11. GESTIÓN DE TURNOS Y DISPONIBILIDAD", level=1)
    style_heading(h11, 1)

    add_body(doc, "El consultorio opera bajo el siguiente esquema de disponibilidad, garantizando la continuidad de la atención y el cumplimiento de los tiempos de respuesta establecidos:")

    h11a = doc.add_heading("11.1 Horario de Atención", level=2)
    style_heading(h11a, 2)

    add_body(doc, "Horario habitual de atención del consultorio:")
    add_bullet(doc, "Días hábiles: Lunes a Viernes de [HORA DE APERTURA] a [HORA DE CIERRE]")
    add_bullet(doc, "Sábados: [HORARIO DE SÁBADO o 'No aplica']")
    add_bullet(doc, "El horario exacto debe ser registrado en el formulario de inscripción al REPS y estar publicado visiblemente en el consultorio")

    h11b = doc.add_heading("11.2 Disponibilidad para Urgencias", level=2)
    style_heading(h11b, 2)
    add_body(doc, "El consultorio NO está habilitado para atención de urgencias. En caso de que un paciente presente una complicación o situación de urgencia durante la atención:")
    urgencias = [
        "Se activa el protocolo de manejo de emergencias del consultorio (PRO-PP-003)",
        "Se llama al número de emergencias 123 para traslado si es necesario",
        "Se estabiliza al paciente con los recursos disponibles en el consultorio mientras llega el apoyo",
        "Se remite al paciente a la red de urgencias correspondiente según su aseguradora",
        "Se documenta el evento en el formato de reporte de eventos adversos (FOR-PP-005)",
    ]
    for item in urgencias:
        add_bullet(doc, item)

    # SECCIÓN 12: BIENESTAR LABORAL
    h12 = doc.add_heading("12. BIENESTAR LABORAL", level=1)
    style_heading(h12, 1)

    add_body(doc, "El consultorio se compromete con el bienestar integral de su talento humano, reconociendo que el bienestar del personal es fundamental para la calidad de la atención a los pacientes.")

    h12a = doc.add_heading("12.1 Seguridad y Salud en el Trabajo", level=2)
    style_heading(h12a, 2)
    add_body(doc, "En cumplimiento de la Ley 1562 de 2012 y el Decreto 1072 de 2015, el consultorio implementa las siguientes medidas de seguridad y salud en el trabajo:")
    sst = [
        "Afiliación de todo el personal a la ARL correspondiente antes del inicio de actividades",
        "Identificación y evaluación de los riesgos ocupacionales específicos del cargo (riesgo biológico, ergonómico, psicosocial)",
        "Provisión de elementos de protección personal (EPP) suficientes y adecuados: guantes, tapabocas, gafas de protección, bata",
        "Vacunación del personal asistencial: Hepatitis B (esquema completo), Tétanos, Influenza (anual)",
        "Protocolo de accidente de trabajo con riesgo biológico: lavado inmediato, reporte a ARL, seguimiento",
        "Exámenes médicos periódicos de monitoreo según los riesgos identificados",
        "Condiciones ergonómicas adecuadas en el puesto de trabajo",
    ]
    for item in sst:
        add_bullet(doc, item)

    h12b = doc.add_heading("12.2 Salud Mental y Prevención del Burnout", level=2)
    style_heading(h12b, 2)
    add_body(doc, "El ejercicio de la medicina general, especialmente en consultorios unipersonales, puede generar altos niveles de estrés y riesgo de síndrome de burnout. El consultorio promueve las siguientes estrategias de bienestar mental:")
    bienestar_mental = [
        "Establecer límites claros entre el tiempo de trabajo y el tiempo personal",
        "Organizar la agenda de manera que permita pausas activas durante la jornada laboral",
        "Participar en redes de apoyo entre pares médicos (sociedades científicas, grupos de colegas)",
        "Buscar supervisión clínica o interconsulta con colegas cuando se presenten casos complejos o emocionalmente difíciles",
        "Acceder a los programas de bienestar ofrecidos por la ARL a la que está afiliada",
        "Reconocer los signos de agotamiento profesional y buscar apoyo oportuno",
    ]
    for item in bienestar_mental:
        add_bullet(doc, item)

    # SECCIÓN 13: MANEJO DE SITUACIONES ESPECIALES
    h13 = doc.add_heading("13. MANEJO DE SITUACIONES ESPECIALES DE TALENTO HUMANO", level=1)
    style_heading(h13, 1)

    h13a = doc.add_heading("13.1 Accidente de Trabajo con Riesgo Biológico", level=2)
    style_heading(h13a, 2)
    accidente_bio = [
        "Lavado inmediato y abundante de la zona afectada con agua y jabón (mínimo 5 minutos para piel, solución salina para mucosas)",
        "Notificación inmediata a la ARL dentro de las primeras 2 horas",
        "Evaluación médica en el centro de atención designado por la ARL",
        "Inicio del protocolo de profilaxis post-exposición si está indicado (para VIH, VHB)",
        "Seguimiento serológico según protocolo (basal, 6 semanas, 3 meses, 6 meses)",
        "Diligenciamiento del formulario de reporte de accidente de trabajo (FURAT)",
        "Análisis de la causa del accidente y medidas correctivas preventivas",
    ]
    for item in accidente_bio:
        add_bullet(doc, item)

    h13b = doc.add_heading("13.2 Ausencia o Incapacidad del Personal", level=2)
    style_heading(h13b, 2)
    add_body(doc, "En caso de incapacidad o ausencia justificada de la médica propietaria, se procederá de la siguiente manera:")
    ausencia = [
        "Notificación oportuna a los pacientes con citas programadas (mínimo con 24 horas de anticipación cuando sea posible)",
        "Reprogramación de las citas para cuando la médica pueda atenderlas o remisión a otro prestador si el caso lo requiere con urgencia",
        "Para ausencias prolongadas superiores a 15 días, evaluar la posibilidad de vincular un médico sustituto temporal que cumpla con todos los requisitos de habilitación",
        "Actualización del horario de atención en el REPS si la ausencia modifica los horarios de manera permanente",
        "El médico sustituto debe ser inscrito en el REPS antes de iniciar actividades asistenciales",
    ]
    for item in ausencia:
        add_bullet(doc, item)

    # SECCIÓN 14: RESPONSABLES
    h14 = doc.add_heading("14. RESPONSABLES", level=1)
    style_heading(h14, 1)

    table_resp = doc.add_table(rows=5, cols=3)
    table_resp.style = 'Table Grid'
    add_table_header_row(table_resp, ["Cargo", "Responsabilidad", "Periodicidad"])
    resp_data = [
        ("Médica General - Propietaria\n[NOMBRE DE LA MÉDICA]",
         "Responsable principal de la gestión del talento humano. Aprueba vinculaciones, evalúa desempeño, aprueba el plan de capacitación y supervisa el cumplimiento de este manual.",
         "Permanente"),
        ("Médica General - Propietaria\n[NOMBRE DE LA MÉDICA]",
         "Mantiene actualizados sus propios documentos de habilitación (tarjeta profesional, RETHUS, certificaciones).",
         "Según vencimiento de cada documento"),
        ("Personal Administrativo\n(si aplica)",
         "Apoya la gestión documental del personal, agenda capacitaciones, lleva el registro de documentos.",
         "Permanente"),
        ("Responsable de Gestión de Calidad\n[NOMBRE DE LA MÉDICA]",
         "Verifica el cumplimiento de los indicadores de talento humano y reporta al SOGCS según corresponda.",
         "Mensual/Anual"),
    ]
    for i, (cargo, resp, per) in enumerate(resp_data):
        table_resp.rows[i+1].cells[0].text = cargo
        table_resp.rows[i+1].cells[1].text = resp
        table_resp.rows[i+1].cells[2].text = per
        for cell in table_resp.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # SECCIÓN 15: INDICADORES
    h15 = doc.add_heading("15. INDICADORES DE SEGUIMIENTO", level=1)
    style_heading(h15, 1)

    table_ind = doc.add_table(rows=7, cols=5)
    table_ind.style = 'Table Grid'
    add_table_header_row(table_ind, ["Indicador", "Fórmula", "Meta", "Frecuencia", "Responsable"])
    ind_data = [
        ("% cumplimiento requisitos de habilitación del talento humano",
         "(N° de requisitos cumplidos / N° total de requisitos exigidos) × 100",
         "100%", "Semestral", "[NOMBRE DE LA MÉDICA]"),
        ("% cumplimiento del plan de capacitación anual",
         "(N° de capacitaciones realizadas / N° de capacitaciones programadas) × 100",
         "≥ 80%", "Anual", "[NOMBRE DE LA MÉDICA]"),
        ("% de personal con evaluación de desempeño realizada",
         "(N° de evaluaciones realizadas / N° de trabajadores) × 100",
         "100%", "Anual", "[NOMBRE DE LA MÉDICA]"),
        ("% documentos de talento humano vigentes",
         "(N° de documentos vigentes / N° total de documentos requeridos) × 100",
         "100%", "Semestral", "[NOMBRE DE LA MÉDICA]"),
        ("N° de accidentes de trabajo",
         "Número absoluto de accidentes de trabajo reportados",
         "0", "Mensual", "[NOMBRE DE LA MÉDICA]"),
        ("% cumplimiento del programa de vacunación",
         "(N° de vacunas aplicadas / N° de vacunas programadas) × 100",
         "100%", "Anual", "[NOMBRE DE LA MÉDICA]"),
    ]
    for i, row in enumerate(ind_data):
        for j, val in enumerate(row):
            table_ind.rows[i+1].cells[j].text = val
            for para in table_ind.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # SECCIÓN 16: FIRMAS
    h16 = doc.add_heading("16. APROBACIÓN Y FIRMAS", level=1)
    style_heading(h16, 1)
    doc.add_paragraph()

    table_firmas = doc.add_table(rows=3, cols=3)
    table_firmas.style = 'Table Grid'
    add_table_header_row(table_firmas, ["ELABORÓ", "REVISÓ", "APROBÓ"])

    cargos_firma = [
        ("[NOMBRE DE LA MÉDICA]\nMédica General\nPropietaria del Consultorio",
         "[NOMBRE DE LA MÉDICA]\nMédica General\nPropietaria del Consultorio",
         "[NOMBRE DE LA MÉDICA]\nMédica General\nPropietaria del Consultorio"),
        ("Firma: _________________________\n\nFecha: " + datetime.date.today().strftime('%d/%m/%Y'),
         "Firma: _________________________\n\nFecha: " + datetime.date.today().strftime('%d/%m/%Y'),
         "Firma: _________________________\n\nFecha: " + datetime.date.today().strftime('%d/%m/%Y')),
    ]
    for i, row in enumerate(cargos_firma):
        for j, val in enumerate(row):
            table_firmas.rows[i+1].cells[j].text = val
            for para in table_firmas.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    path = os.path.join(BASE_DIR, "MAN-TH-001_Manual_Talento_Humano.docx")
    doc.save(path)
    print(f"✓ Creado: {path}")

# ============================================================
# DOCUMENTO 2: PER-TH-001 Perfil del Cargo
# ============================================================
def crear_perfil_cargo():
    doc = Document()
    set_page_margins(doc)
    add_header_footer(doc, "PER-TH-001", "Perfil del Cargo Médico General")
    add_portada(doc, "PER-TH-001", "PERFIL DEL CARGO: MÉDICO GENERAL PROPIETARIO")
    add_control_versiones(doc)

    h1 = doc.add_heading("1. OBJETIVO", level=1)
    style_heading(h1, 1)
    add_body(doc, "Definir el perfil integral del cargo de Médico General Propietario del Consultorio [NOMBRE DEL CONSULTORIO], estableciendo con precisión los requisitos de formación, experiencia, competencias técnicas y personales, responsabilidades y criterios de desempeño que garanticen la prestación de servicios de salud de calidad, incluyendo la atención de consultas de medicina general y la realización de procedimientos estéticos no invasivos dentro del marco legal establecido para médicos generales en Colombia.")

    h2 = doc.add_heading("2. ALCANCE", level=1)
    style_heading(h2, 1)
    add_body(doc, "El presente documento aplica específicamente al cargo de Médico General Propietario del Consultorio [NOMBRE DEL CONSULTORIO], ejercido por la Dra. [NOMBRE DE LA MÉDICA]. Los criterios aquí establecidos deben cumplirse de manera permanente durante toda la vigencia de la habilitación del servicio.")

    h3 = doc.add_heading("3. MARCO LEGAL", level=1)
    style_heading(h3, 1)
    add_body(doc, "El presente perfil de cargo se fundamenta en la siguiente normatividad:")

    normas = [
        ("Ley 23 de 1981", "Normas de ética médica. Establece los deberes del médico y el marco del ejercicio profesional."),
        ("Ley 14 de 1962", "Regula el ejercicio de la medicina y cirugía en Colombia. Establece los requisitos para el ejercicio legal de la profesión."),
        ("Decreto 3380 de 1981", "Reglamenta la Ley 23/1981. Define el acto médico y sus implicaciones legales."),
        ("Resolución 3100 de 2019", "Define los procedimientos de habilitación. Establece los requisitos de talento humano por servicio, incluyendo medicina general."),
        ("Ley 711 de 2001", "Reglamenta el ejercicio de la cosmetología. Establece la diferencia entre actos de cosmetología y actos médicos, siendo los procedimientos con medicamentos o que impliquen riesgo (como toxina botulínica, rellenos, peelings profundos, mesoterapia, PRP) de exclusiva competencia médica."),
        ("Resolución 2263 de 2004", "Establece los requisitos para la apertura y funcionamiento de centros de estética y similares. Delimita el alcance de los procedimientos estéticos NO médicos."),
        ("Decreto 780 de 2016", "Decreto Único Reglamentario del Sector Salud. Consolida la normatividad aplicable a la prestación de servicios de salud."),
        ("Resolución 1441 de 2013", "Predecesora de la 3100/2019 en estándares de habilitación. Algunos conceptos aún vigentes."),
        ("Ley 1164 de 2007", "Disposiciones en materia del Talento Humano en Salud. Regula la formación, ejercicio y gestión del talento humano en salud en Colombia."),
        ("Decreto 4725 de 2005", "Regula el régimen de registros sanitarios de dispositivos médicos. Aplicable a los dispositivos médicos utilizados en los procedimientos estéticos."),
    ]

    table = doc.add_table(rows=len(normas)+1, cols=2)
    table.style = 'Table Grid'
    add_table_header_row(table, ["Norma", "Relevancia para el Cargo"])
    for i, (n, d) in enumerate(normas):
        table.rows[i+1].cells[0].text = n
        table.rows[i+1].cells[1].text = d
        for cell in table.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

    h4 = doc.add_heading("4. IDENTIFICACIÓN DEL CARGO", level=1)
    style_heading(h4, 1)

    table_id = doc.add_table(rows=8, cols=2)
    table_id.style = 'Table Grid'
    id_data = [
        ("Nombre del cargo:", "Médica General - Propietaria"),
        ("Titular actual:", "[NOMBRE DE LA MÉDICA]"),
        ("Tipo de vinculación:", "Propietaria / [TIPO DE CONTRATO si aplica]"),
        ("Jornada laboral:", "[HORAS SEMANALES] horas semanales"),
        ("Lugar de trabajo:", "[DIRECCIÓN DEL CONSULTORIO], [CIUDAD]"),
        ("Dependencia jerárquica:", "No aplica (propietaria)"),
        ("Cargo que supervisa:", "Personal administrativo y de apoyo (si aplica)"),
        ("Fecha de elaboración del perfil:", datetime.date.today().strftime('%d/%m/%Y')),
    ]
    for i, (label, val) in enumerate(id_data):
        table_id.rows[i].cells[0].text = label
        table_id.rows[i].cells[1].text = val
        table_id.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in table_id.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)
    doc.add_paragraph()

    h5 = doc.add_heading("5. PROPÓSITO PRINCIPAL DEL CARGO", level=1)
    style_heading(h5, 1)
    add_body(doc, "Brindar atención médica integral de primer nivel a los pacientes del Consultorio [NOMBRE DEL CONSULTORIO], realizando consultas de medicina general y procedimientos estéticos no invasivos dentro del marco legal establecido para médicos generales en Colombia, garantizando la más alta calidad científica, ética y humanística en cada acto médico, contribuyendo al mejoramiento del estado de salud y la satisfacción de los usuarios.")

    h6 = doc.add_heading("6. REQUISITOS DEL CARGO", level=1)
    style_heading(h6, 1)

    h6a = doc.add_heading("6.1 Requisitos de Formación Académica", level=2)
    style_heading(h6a, 2)

    h6a1 = doc.add_heading("6.1.1 Formación Básica Obligatoria", level=3)
    style_heading(h6a1, 3)
    add_body(doc, "ESENCIAL:", bold=True)
    form_basica = [
        "Título de Médico y Cirujano otorgado por universidad legalmente reconocida en Colombia o título extranjero debidamente revalidado ante el Ministerio de Educación y reconocido por el Ministerio de Salud",
        "Tarjeta Profesional vigente expedida por el Tribunal Ético de Medicina (TEM) - sin sanciones disciplinarias activas",
        "Registro activo y vigente en el Registro Único Nacional del Talento Humano en Salud (RETHUS)",
    ]
    for item in form_basica:
        add_bullet(doc, item)

    h6a2 = doc.add_heading("6.1.2 Formación Específica para Procedimientos Estéticos No Invasivos", level=3)
    style_heading(h6a2, 3)
    add_body(doc, "IMPORTANTE - MARCO LEGAL: Los procedimientos estéticos no invasivos descritos en este perfil son de EXCLUSIVA COMPETENCIA MÉDICA. La Ley 711 de 2001 y la Resolución 2263 de 2004 establecen claramente que los procedimientos que involucren la aplicación de medicamentos (toxina botulínica, ácido hialurónico), sustancias químicas con potencial de daño (peelings ácidos), técnicas que involucren la piel en su estructura o procedimientos de medicina regenerativa (PRP, mesoterapia), son actos médicos que SOLO pueden ser realizados por médicos legalmente habilitados. Su realización por personal no médico constituye ejercicio ilegal de la medicina.")

    add_body(doc, "Para la realización de estos procedimientos, la médica debe contar con:")
    form_estetica = [
        "[CERTIFICADO] Curso de Toxina Botulínica Tipo A - Institución: [NOMBRE INSTITUCIÓN] - Fecha: [FECHA] - Intensidad: [HORAS]",
        "[CERTIFICADO] Curso de Rellenos Dérmicos con Ácido Hialurónico - Institución: [NOMBRE INSTITUCIÓN] - Fecha: [FECHA] - Intensidad: [HORAS]",
        "[CERTIFICADO] Curso de Peelings Químicos (TCA y Glicólico) - Institución: [NOMBRE INSTITUCIÓN] - Fecha: [FECHA] - Intensidad: [HORAS]",
        "[CERTIFICADO] Curso de Mesoterapia Facial y Corporal - Institución: [NOMBRE INSTITUCIÓN] - Fecha: [FECHA] - Intensidad: [HORAS]",
        "[CERTIFICADO] Curso de Plasma Rico en Plaquetas (PRP) - Institución: [NOMBRE INSTITUCIÓN] - Fecha: [FECHA] - Intensidad: [HORAS]",
        "Certificados expedidos por instituciones médicas reconocidas (sociedades científicas, hospitales universitarios, instituciones de educación superior con programas de medicina)",
    ]
    for item in form_estetica:
        add_bullet(doc, item)

    h6b = doc.add_heading("6.2 Experiencia Requerida", level=2)
    style_heading(h6b, 2)
    exp_items = [
        "Experiencia mínima de [X AÑOS/MESES] en atención de consulta de medicina general",
        "Experiencia práctica en procedimientos estéticos no invasivos supervisada durante los cursos de formación específica",
        "Experiencia comprobada en el manejo de urgencias básicas y reanimación cardiopulmonar",
    ]
    for item in exp_items:
        add_bullet(doc, item)

    h6c = doc.add_heading("6.3 Conocimientos Técnicos Específicos", level=2)
    style_heading(h6c, 2)

    h6c1 = doc.add_heading("6.3.1 Medicina General", level=3)
    style_heading(h6c1, 3)
    conocimientos_gen = [
        "Semiología médica completa: anamnesis, examen físico por sistemas",
        "Diagnóstico diferencial en patologías frecuentes de consulta externa de primer nivel",
        "Manejo de enfermedades crónicas no transmisibles: hipertensión, diabetes, dislipidemia",
        "Manejo de patologías infecciosas frecuentes: respiratorias, gastrointestinales, urinarias",
        "Prescripción racional de medicamentos (Resolución 1478/2006)",
        "Interpretación de exámenes de laboratorio e imágenes diagnósticas básicas",
        "Manejo de urgencias médicas básicas y reanimación cardiopulmonar (BLS/ACLS)",
        "Elaboración de historia clínica completa y legible (Resolución 1995/1999)",
        "Expedición de incapacidades y certificados médicos",
        "Remisión y contraremisión de pacientes",
        "Guías de práctica clínica colombianas vigentes (GPC del MinSalud)",
        "Programa de Detección Temprana y Protección Específica - DYPE",
    ]
    for item in conocimientos_gen:
        add_bullet(doc, item)

    h6c2 = doc.add_heading("6.3.2 Procedimientos Estéticos No Invasivos", level=3)
    style_heading(h6c2, 3)
    conocimientos_estetica = [
        "Anatomía facial detallada: capas de la piel, músculos de la expresión, planos de inyección, vascularización facial, zonas de riesgo",
        "Fisiología de la piel y el envejecimiento cutáneo",
        "Farmacología de la toxina botulínica tipo A: mecanismo de acción, dosis, dilución, conservación, efectos adversos",
        "Farmacología del ácido hialurónico: tipos (reticulado/no reticulado), características, indicaciones, técnicas de inyección",
        "Clasificación de arrugas: dinámicas vs estáticas, escalas de valoración (Glogau, Fitzpatrick)",
        "Técnicas de inyección: lineal, retrotraza, depósito, abanico, multipunto",
        "Indicaciones y contraindicaciones absolutas y relativas de cada procedimiento",
        "Manejo de complicaciones: hematoma, equimosis, ptosis, necrosis, oclusión vascular (protocolo hialuronidasa)",
        "Clasificación de peelings: superficiales, medios y profundos - solo superficiales para médicos generales",
        "Agentes químicos para peeling superficial: ácido glicólico (hasta 70%), ácido tricloroacético (TCA ≤20%), ácido mandélico",
        "Fototipo cutáneo: clasificación de Fitzpatrick, relevancia para procedimientos estéticos",
        "Mesoterapia: cócteles autorizados, técnicas de infiltración, indicaciones faciales y corporales",
        "Plasma Rico en Plaquetas (PRP): procesamiento, activación, técnicas de aplicación",
        "Consentimiento informado en medicina estética: elementos esenciales, documentación",
        "Regulación INVIMA de productos biológicos y medicamentos de uso estético",
    ]
    for item in conocimientos_estetica:
        add_bullet(doc, item)

    h6d = doc.add_heading("6.4 Habilidades y Competencias", level=2)
    style_heading(h6d, 2)

    h6d1 = doc.add_heading("6.4.1 Competencias Técnicas", level=3)
    style_heading(h6d1, 3)
    comp_tec = [
        "Destreza manual para la realización de procedimientos de inyección y técnicas estéticas",
        "Capacidad de diagnóstico clínico preciso y rápido",
        "Manejo de equipos médicos del consultorio",
        "Elaboración de historias clínicas completas y de calidad",
        "Interpretación y aplicación de guías de práctica clínica",
        "Prescripción racional de medicamentos",
        "Manejo de urgencias y emergencias médicas",
    ]
    for item in comp_tec:
        add_bullet(doc, item)

    h6d2 = doc.add_heading("6.4.2 Competencias Blandas y Relacionales", level=3)
    style_heading(h6d2, 3)
    comp_blandas = [
        "Comunicación efectiva con los pacientes: capacidad de explicar diagnósticos y procedimientos en lenguaje comprensible",
        "Empatía y escucha activa: capacidad de entender las necesidades y expectativas del paciente",
        "Ética profesional: toma de decisiones basada en principios éticos y en el mejor interés del paciente",
        "Manejo de situaciones difíciles: capacidad de comunicar malas noticias, manejar pacientes difíciles",
        "Trabajo autónomo y toma de decisiones: capacidad de actuar con independencia y criterio propio",
        "Orientación al aprendizaje continuo: disposición para la actualización permanente",
        "Gestión del tiempo: organización eficiente de la agenda de atención",
        "Manejo del estrés: capacidad de mantener la calma y la eficiencia en situaciones de presión",
    ]
    for item in comp_blandas:
        add_bullet(doc, item)

    h7 = doc.add_heading("7. FUNCIONES Y RESPONSABILIDADES DEL CARGO", level=1)
    style_heading(h7, 1)

    h7a = doc.add_heading("7.1 Funciones Asistenciales - Medicina General", level=2)
    style_heading(h7a, 2)
    func_general = [
        "Realizar la consulta de medicina general completa: anamnesis, examen físico, diagnóstico e impresión diagnóstica, plan de manejo.",
        "Solicitar, interpretar y hacer seguimiento de los exámenes complementarios necesarios (laboratorios, imágenes diagnósticas, otros).",
        "Prescribir los medicamentos necesarios en la receta médica oficial, cumpliendo la Resolución 1478 de 2006.",
        "Expedir incapacidades médicas cuando estén clínicamente justificadas.",
        "Emitir certificados médicos veraces y oportunos.",
        "Remitir a los pacientes que requieran atención especializada, siguiendo el sistema de referencia y contrarreferencia (MAN-RCR-001).",
        "Realizar actividades de promoción de la salud y prevención de la enfermedad en cada consulta.",
        "Atender adecuadamente las urgencias que se presenten en el consultorio y gestionar la remisión cuando sea necesario.",
        "Participar en programas de salud pública cuando sea requerido por las autoridades sanitarias.",
    ]
    for item in func_general:
        add_bullet(doc, item)

    h7b = doc.add_heading("7.2 Funciones Asistenciales - Procedimientos Estéticos No Invasivos", level=2)
    style_heading(h7b, 2)
    add_body(doc, "AVISO LEGAL: Las siguientes funciones corresponden a actos médicos de exclusiva competencia del médico legalmente habilitado. Ninguna de estas actividades puede ser delegada a personal no médico (esteticistas, cosmetólogos, enfermeras, técnicos) bajo ninguna circunstancia, so pena de incurrir en ejercicio ilegal de la medicina y responsabilidad penal, civil y disciplinaria.")

    func_estetica = [
        "Realizar la valoración estética médica integral: anamnesis estética, examen físico del área a tratar, fotografías clínicas estandarizadas.",
        "Establecer el diagnóstico estético médico y formular el plan de tratamiento individualizado.",
        "Aplicar toxina botulínica tipo A (INVIMA) en indicaciones aprobadas: glabela, frente, patas de gallo, hiperhidrosis, bruxismo, entre otras.",
        "Aplicar rellenos dérmicos con ácido hialurónico (dispositivo médico con registro INVIMA) en labios, surcos, mejillas, según indicación clínica.",
        "Realizar peelings químicos superficiales: ácido glicólico (hasta 70%), TCA (hasta 20%), ácido mandélico.",
        "Realizar mesoterapia facial y corporal con cócteles de uso médico y sustancias de aplicación intradérmica.",
        "Realizar tratamientos con Plasma Rico en Plaquetas (PRP) autólogo para rejuvenecimiento facial y caída de cabello.",
        "Obtener el consentimiento informado específico para cada procedimiento estético ANTES de realizarlo.",
        "Documentar adecuadamente cada procedimiento en la historia clínica estética (FOR-HC-002).",
        "Realizar el seguimiento post-procedimiento y el manejo de las complicaciones que se presenten.",
    ]
    for item in func_estetica:
        add_bullet(doc, item)

    h7c = doc.add_heading("7.3 Funciones Administrativas y de Gestión de Calidad", level=2)
    style_heading(h7c, 2)
    func_admin = [
        "Mantener actualizados todos los documentos requeridos para la habilitación del servicio (tarjeta profesional, RETHUS, certificados de formación, póliza de responsabilidad civil, etc.).",
        "Supervisar el cumplimiento de todos los manuales, procedimientos y protocolos del sistema de gestión de calidad del consultorio.",
        "Verificar que los equipos biomédicos cuenten con el mantenimiento y calibración al día.",
        "Asegurar el almacenamiento adecuado de los medicamentos y dispositivos médicos utilizados.",
        "Reportar los eventos adversos y complicaciones en el formato establecido (FOR-PP-005).",
        "Mantener la confidencialidad de la historia clínica (Ley 1581/2012).",
        "Implementar y hacer seguimiento a los indicadores de calidad del consultorio.",
        "Gestionar las quejas y reclamos de los pacientes de manera oportuna y constructiva.",
        "Cumplir con las obligaciones de reporte al SIVIGILA cuando sea pertinente.",
        "Renovar la habilitación del consultorio según las condiciones y plazos establecidos.",
    ]
    for item in func_admin:
        add_bullet(doc, item)

    h8 = doc.add_heading("8. INDICADORES DE DESEMPEÑO DEL CARGO", level=1)
    style_heading(h8, 1)

    table_kpi = doc.add_table(rows=8, cols=4)
    table_kpi.style = 'Table Grid'
    add_table_header_row(table_kpi, ["Indicador", "Fórmula de cálculo", "Meta", "Frecuencia"])
    kpi_data = [
        ("Satisfacción del paciente", "(N° pacientes satisfechos / N° pacientes encuestados) × 100", "≥ 90%", "Mensual"),
        ("Calidad de historia clínica", "(N° HC con todos los componentes / N° HC auditadas) × 100", "≥ 95%", "Mensual"),
        ("Cumplimiento de protocolos", "(N° de atenciones con protocolo / N° total atenciones) × 100", "100%", "Mensual"),
        ("Eventos adversos reportados", "N° de eventos adversos reportados en el período", "100% reportados", "Mensual"),
        ("Cumplimiento plan de capacitación", "(Capacitaciones realizadas / Capacitaciones programadas) × 100", "≥ 80%", "Anual"),
        ("Documentos habilitación vigentes", "(Docs vigentes / Docs requeridos) × 100", "100%", "Semestral"),
        ("Consentimientos informados obtenidos", "(Procedimientos con CI / Total procedimientos) × 100", "100%", "Mensual"),
    ]
    for i, row in enumerate(kpi_data):
        for j, val in enumerate(row):
            table_kpi.rows[i+1].cells[j].text = val
            for para in table_kpi.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # Firmas
    h_f = doc.add_heading("9. APROBACIÓN Y FIRMAS", level=1)
    style_heading(h_f, 1)
    doc.add_paragraph()
    table_firmas = doc.add_table(rows=3, cols=3)
    table_firmas.style = 'Table Grid'
    add_table_header_row(table_firmas, ["ELABORÓ", "REVISÓ", "APROBÓ"])
    firmas_data = [
        ("[NOMBRE DE LA MÉDICA]\nMédica General\nPropietaria",
         "[NOMBRE DE LA MÉDICA]\nMédica General\nPropietaria",
         "[NOMBRE DE LA MÉDICA]\nMédica General\nPropietaria"),
        ("Firma: _________________________\nFecha: " + datetime.date.today().strftime('%d/%m/%Y'),
         "Firma: _________________________\nFecha: " + datetime.date.today().strftime('%d/%m/%Y'),
         "Firma: _________________________\nFecha: " + datetime.date.today().strftime('%d/%m/%Y')),
    ]
    for i, row in enumerate(firmas_data):
        for j, val in enumerate(row):
            table_firmas.rows[i+1].cells[j].text = val
            for para in table_firmas.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    path = os.path.join(BASE_DIR, "PER-TH-001_Perfil_Cargo_Medico_General.docx")
    doc.save(path)
    print(f"✓ Creado: {path}")

# ============================================================
# DOCUMENTO 3: PRO-TH-001 Proceso Selección y Contratación
# ============================================================
def crear_proceso_seleccion():
    doc = Document()
    set_page_margins(doc)
    add_header_footer(doc, "PRO-TH-001", "Proceso de Selección y Contratación")
    add_portada(doc, "PRO-TH-001", "PROCESO DE SELECCIÓN Y CONTRATACIÓN DEL TALENTO HUMANO")
    add_control_versiones(doc)

    h1 = doc.add_heading("1. OBJETIVO", level=1)
    style_heading(h1, 1)
    add_body(doc, "Establecer el proceso estandarizado para la selección, verificación, contratación e inducción del personal que preste sus servicios en el Consultorio Médico [NOMBRE DEL CONSULTORIO], asegurando que cumplan con todos los requisitos legales y de habilitación establecidos por la Resolución 3100 de 2019 y demás normatividad vigente, garantizando la idoneidad y competencia del talento humano para la prestación segura y de calidad de los servicios de salud.")

    h2 = doc.add_heading("2. ALCANCE", level=1)
    style_heading(h2, 1)
    add_body(doc, "Este proceso aplica a la selección y contratación de todo el personal que preste servicios en el Consultorio [NOMBRE DEL CONSULTORIO], incluyendo: personal asistencial médico, personal de apoyo administrativo, personal de servicios generales y personal externo contratado (proveedores de mantenimiento, etc.). Para la médica propietaria, aplica en la verificación periódica de la vigencia de sus documentos de habilitación.")

    h3 = doc.add_heading("3. MARCO LEGAL", level=1)
    style_heading(h3, 1)
    leyes_sel = [
        "Código Sustantivo del Trabajo (CST) - contratación laboral",
        "Decreto 1072 de 2015 - Decreto Único Reglamentario del Trabajo",
        "Resolución 3100 de 2019 - Requisitos de talento humano por servicio de salud",
        "Ley 1164 de 2007 - Talento Humano en Salud",
        "Ley 23 de 1981 - Ética médica",
        "Ley 100 de 1993 - SGSSS - obligaciones de afiliación al sistema de seguridad social",
        "Ley 1562 de 2012 - Riesgos laborales",
        "Decreto 1011 de 2006 - SOGCS",
    ]
    for item in leyes_sel:
        add_bullet(doc, item)

    h4 = doc.add_heading("4. DEFINICIONES", level=1)
    style_heading(h4, 1)
    defs = [
        ("Proceso de selección:", "Conjunto de actividades orientadas a identificar y elegir el candidato más idóneo para ocupar un cargo, con base en el perfil establecido y los requisitos legales aplicables."),
        ("Verificación de idoneidad:", "Proceso de comprobación de la autenticidad y vigencia de los títulos, registros, certificaciones y antecedentes del candidato, a través de las fuentes oficiales correspondientes."),
        ("RETHUS:", "Registro Único Nacional del Talento Humano en Salud. Sistema oficial del Ministerio de Salud donde se verifica la habilitación del personal de salud."),
        ("Inducción:", "Proceso de orientación al nuevo trabajador sobre la organización, sus procesos, políticas y el entorno laboral."),
        ("Prueba de período de prueba:", "Período inicial de la relación laboral (máximo 2 meses para contratos a término indefinido) en el cual ambas partes evalúan la conveniencia del vínculo."),
    ]
    for term, defi in defs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r1 = p.add_run(term + " ")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(defi)
        r2.font.size = Pt(11)

    h5 = doc.add_heading("5. DESCRIPCIÓN DEL PROCESO", level=1)
    style_heading(h5, 1)

    h5a = doc.add_heading("5.1 Flujograma del Proceso", level=2)
    style_heading(h5a, 2)
    add_body(doc, "El proceso de selección y contratación sigue la siguiente secuencia:")

    flujo = [
        ("PASO 1", "Identificación de la necesidad", "La médica propietaria identifica la necesidad de vincular personal al consultorio (nueva contratación, sustitución, ampliación de capacidad). Se define el tipo de cargo y el perfil requerido."),
        ("PASO 2", "Elaboración/revisión del perfil", "Se revisa o elabora el perfil del cargo requerido, asegurando que cumpla con los requisitos de habilitación de la Resolución 3100/2019. Documento de referencia: PER-TH-001."),
        ("PASO 3", "Convocatoria", "Se publican los requisitos del cargo a través de los canales disponibles: redes profesionales (LinkedIn médico, bolsas de empleo de universidades médicas), referencias de colegas, anuncios en colegios médicos."),
        ("PASO 4", "Recepción y preselección de hojas de vida", "Se reciben las hojas de vida y se realiza una preselección verificando el cumplimiento de los requisitos obligatorios (título, tarjeta profesional, certificaciones requeridas). Se descartan los que no cumplen el perfil mínimo."),
        ("PASO 5", "Verificación inicial de documentos", "Se verifica en RETHUS la habilitación del candidato. Se consulta el Tribunal Ético de Medicina para antecedentes disciplinarios. Se consultan los registros de antecedentes judiciales y penales."),
        ("PASO 6", "Entrevista técnica y personal", "Se realiza la entrevista con la médica propietaria. La entrevista incluye: evaluación de casos clínicos, preguntas sobre competencias técnicas, evaluación de competencias personales y valores."),
        ("PASO 7", "Solicitud y verificación completa de documentos", "Se solicita la carpeta completa de documentos al candidato seleccionado. Se verifican todos los documentos físicos y su autenticidad. Se revisan los antecedentes completos."),
        ("PASO 8", "Exámenes preocupacionales", "Se realizan los exámenes médicos de ingreso con el médico ocupacional designado o con la ARL. Los exámenes incluyen: hemoleucograma, VDRL, HBsAg, serología para VIH (con consentimiento), examen físico con énfasis osteomusicular."),
        ("PASO 9", "Decisión de vinculación", "Con base en toda la información recopilada, la médica propietaria toma la decisión de vinculación o rechazo. La decisión se fundamenta en criterios técnicos, legales y de idoneidad."),
        ("PASO 10", "Formalización del contrato", "Se elabora el contrato laboral o de prestación de servicios según corresponda, en cumplimiento del Código Sustantivo del Trabajo. Se firman los documentos de vinculación: contrato, confidencialidad, políticas internas."),
        ("PASO 11", "Afiliación al sistema de seguridad social", "Se realiza la afiliación a EPS, fondo de pensiones y ARL antes del primer día de trabajo. Se registra al trabajador en la planilla PILA."),
        ("PASO 12", "Inducción", "Se realiza el proceso de inducción según el programa establecido en el Manual de Talento Humano (MAN-TH-001), con registro de asistencia."),
        ("PASO 13", "Seguimiento en período de prueba", "Durante el período de prueba se hace seguimiento semanal a la adaptación del trabajador. Al finalizar el período, se realiza una evaluación formal."),
    ]

    table_flujo = doc.add_table(rows=len(flujo)+1, cols=3)
    table_flujo.style = 'Table Grid'
    add_table_header_row(table_flujo, ["Paso", "Actividad", "Descripción"])
    for i, (paso, act, desc) in enumerate(flujo):
        table_flujo.rows[i+1].cells[0].text = paso
        table_flujo.rows[i+1].cells[1].text = act
        table_flujo.rows[i+1].cells[2].text = desc
        for cell in table_flujo.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    h6 = doc.add_heading("6. VERIFICACIÓN PERIÓDICA DE DOCUMENTOS", level=1)
    style_heading(h6, 1)
    add_body(doc, "Los documentos de habilitación del personal tienen fechas de vencimiento que deben ser monitoreadas permanentemente. Se establece el siguiente calendario de verificación:")

    docs_verificacion = [
        ("Tarjeta Profesional Médica", "Según indicación del TEM (verificar vigencia)", "Anual o según vencimiento", "RETHUS, Tribunal Ético de Medicina"),
        ("Certificado Antecedentes Disciplinarios (TEM)", "No mayor a 3 meses en el momento de la habilitación", "Renovar cuando sea requerido por autoridades", "Tribunal Ético de Medicina"),
        ("Certificado Antecedentes Judiciales", "No mayor a 3 meses en el momento de la habilitación", "Renovar cuando sea requerido", "Policía Nacional"),
        ("Certificado Antecedentes Fiscales (Contraloría)", "Sin fecha de vencimiento formal, pero se actualiza anualmente", "Anual", "Contraloría General de la República"),
        ("Certificado SIRI (Procuraduría)", "Sin fecha de vencimiento formal", "Anual o cuando sea requerido", "Procuraduría General"),
        ("Registro RETHUS", "Vigencia permanente (verificar actividad)", "Anual", "RETHUS MinSalud"),
        ("Vacunas", "Según esquema: Hepatitis B (cada 10 años o nunca si completó 3 dosis), Tétanos (cada 10 años), Influenza (anual)", "Según esquema", "Institución vacunadora"),
        ("Certificados de cursos de estética", "Sin vencimiento, pero se recomienda actualizar cada 2 años", "Cada 2 años (recomendado)", "Institución capacitadora"),
    ]

    table_docs = doc.add_table(rows=len(docs_verificacion)+1, cols=4)
    table_docs.style = 'Table Grid'
    add_table_header_row(table_docs, ["Documento", "Vigencia", "Frecuencia de verificación", "Fuente de verificación"])
    for i, row in enumerate(docs_verificacion):
        for j, val in enumerate(row):
            table_docs.rows[i+1].cells[j].text = val
            for para in table_docs.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    h7 = doc.add_heading("7. RESPONSABLES DEL PROCESO", level=1)
    style_heading(h7, 1)
    responsables = [
        ("Médica General - Propietaria\n[NOMBRE DE LA MÉDICA]", "Responsable de todo el proceso de selección y contratación. Toma la decisión final de vinculación. Aprueba el perfil del cargo. Realiza la entrevista técnica."),
        ("Personal Administrativo (si aplica)", "Apoya la recepción de hojas de vida, coordina las entrevistas, gestiona la documentación y apoya los trámites de afiliación al sistema de seguridad social."),
    ]
    table_resp2 = doc.add_table(rows=len(responsables)+1, cols=2)
    table_resp2.style = 'Table Grid'
    add_table_header_row(table_resp2, ["Cargo", "Responsabilidades en el proceso"])
    for i, (c, r) in enumerate(responsables):
        table_resp2.rows[i+1].cells[0].text = c
        table_resp2.rows[i+1].cells[1].text = r
        for cell in table_resp2.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

    h8 = doc.add_heading("8. DOCUMENTOS RELACIONADOS", level=1)
    style_heading(h8, 1)
    docs_rel = [
        "MAN-TH-001 - Manual de Gestión del Talento Humano",
        "PER-TH-001 - Perfil del Cargo: Médico General Propietario",
        "FOR-TH-001 - Formato de Evaluación de Desempeño",
        "MAN-GC-001 - Manual de Calidad",
    ]
    for item in docs_rel:
        add_bullet(doc, item)

    h9 = doc.add_heading("9. INDICADORES DEL PROCESO", level=1)
    style_heading(h9, 1)
    indicadores_sel = [
        ("% procesos de selección completados en los pasos establecidos", "(Procesos con todos los pasos / Total procesos) × 100", "100%"),
        ("% candidatos vinculados que cumplen perfil completo", "(Candidatos con perfil completo / Total vinculados) × 100", "100%"),
        ("Tiempo promedio del proceso de selección", "Días hábiles desde identificación de necesidad hasta vinculación", "≤ 30 días hábiles"),
    ]
    table_ind2 = doc.add_table(rows=len(indicadores_sel)+1, cols=3)
    table_ind2.style = 'Table Grid'
    add_table_header_row(table_ind2, ["Indicador", "Fórmula", "Meta"])
    for i, row in enumerate(indicadores_sel):
        for j, val in enumerate(row):
            table_ind2.rows[i+1].cells[j].text = val
            for para in table_ind2.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

    # Firmas
    table_firmas = doc.add_table(rows=3, cols=3)
    table_firmas.style = 'Table Grid'
    add_table_header_row(table_firmas, ["ELABORÓ", "REVISÓ", "APROBÓ"])
    for i in range(1, 3):
        for j in range(3):
            table_firmas.rows[i].cells[j].text = "[NOMBRE DE LA MÉDICA]\nMédica General - Propietaria\n\nFirma: _________________________\n\nFecha: " + datetime.date.today().strftime('%d/%m/%Y')
            table_firmas.rows[i].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in table_firmas.rows[i].cells[j].paragraphs[0].runs:
                run.font.size = Pt(10)

    path = os.path.join(BASE_DIR, "PRO-TH-001_Proceso_Seleccion_Contratacion.docx")
    doc.save(path)
    print(f"✓ Creado: {path}")

# ============================================================
# DOCUMENTO 4: FOR-TH-001 Formato Evaluación de Desempeño
# ============================================================
def crear_formato_evaluacion():
    doc = Document()
    set_page_margins(doc)
    add_header_footer(doc, "FOR-TH-001", "Formato de Evaluación de Desempeño")
    add_portada(doc, "FOR-TH-001", "FORMATO DE EVALUACIÓN DE DESEMPEÑO DEL PERSONAL")
    add_control_versiones(doc)

    h1 = doc.add_heading("INSTRUCCIONES DE DILIGENCIAMIENTO", level=1)
    style_heading(h1, 1)
    add_body(doc, "El presente formato debe ser diligenciado por la médica propietaria con una periodicidad anual para todo el personal del consultorio. Para la médica propietaria, sirve como autoevaluación de sus competencias y desempeño. El proceso de evaluación debe ser participativo: primero el evaluado realiza su autoevaluación, luego el evaluador realiza la evaluación, y finalmente se hace una reunión de retroalimentación.")
    add_body(doc, "ESCALA DE CALIFICACIÓN: 1=Insatisfactorio | 2=En mejoramiento | 3=Aceptable | 4=Satisfactorio | 5=Sobresaliente")

    h2 = doc.add_heading("SECCIÓN 1: DATOS DE IDENTIFICACIÓN", level=1)
    style_heading(h2, 1)

    table_id = doc.add_table(rows=7, cols=4)
    table_id.style = 'Table Grid'
    id_labels = [
        ("Nombre del evaluado:", "", "Cargo:", ""),
        ("Período de evaluación:", "", "Fecha de evaluación:", ""),
        ("Nombre del evaluador:", "", "Cargo del evaluador:", ""),
        ("Fecha de vinculación:", "", "Tipo de vinculación:", ""),
        ("Área de trabajo:", "", "Horario:", ""),
        ("¿Es evaluación de ingreso?", "□ Sí  □ No", "¿Es evaluación de período de prueba?", "□ Sí  □ No"),
        ("Número de evaluación:", "", "Versión del formato:", "1.0"),
    ]
    for i, row in enumerate(id_labels):
        for j in range(4):
            table_id.rows[i].cells[j].text = row[j]
            if j % 2 == 0:
                for run in table_id.rows[i].cells[j].paragraphs[0].runs:
                    run.font.bold = True
            for run in table_id.rows[i].cells[j].paragraphs[0].runs:
                run.font.size = Pt(10)

    doc.add_paragraph()

    h3 = doc.add_heading("SECCIÓN 2: EVALUACIÓN DE COMPETENCIAS TÉCNICAS (40% del total)", level=1)
    style_heading(h3, 1)
    add_body(doc, "Calificar cada ítem en la escala del 1 al 5. Los espacios de 'Autoevaluación' los completa el evaluado y los de 'Evaluación' los completa el evaluador.")

    comp_tecnicas = [
        ("2.1", "Conocimiento médico actualizado y aplicación de guías de práctica clínica"),
        ("2.2", "Calidad del proceso diagnóstico (anamnesis, examen físico, diagnóstico diferencial)"),
        ("2.3", "Pertinencia y racionalidad de la prescripción médica"),
        ("2.4", "Calidad y completitud de la historia clínica"),
        ("2.5", "Cumplimiento de los protocolos de atención establecidos"),
        ("2.6", "Manejo adecuado de urgencias y emergencias"),
        ("2.7", "Destreza técnica en los procedimientos realizados"),
        ("2.8", "Conocimiento y aplicación de normas de bioseguridad"),
        ("2.9", "Uso correcto de equipos e instrumentos médicos"),
        ("2.10", "Para procedimientos estéticos: técnica adecuada, manejo de complicaciones"),
    ]

    table_ct = doc.add_table(rows=len(comp_tecnicas)+2, cols=5)
    table_ct.style = 'Table Grid'
    add_table_header_row(table_ct, ["N°", "Competencia", "Autoevaluación\n(1-5)", "Evaluación\n(1-5)", "Observaciones"])
    for i, (num, comp) in enumerate(comp_tecnicas):
        table_ct.rows[i+1].cells[0].text = num
        table_ct.rows[i+1].cells[1].text = comp
        table_ct.rows[i+1].cells[2].text = ""
        table_ct.rows[i+1].cells[3].text = ""
        table_ct.rows[i+1].cells[4].text = ""
        for cell in table_ct.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    table_ct.rows[-1].cells[0].text = ""
    table_ct.rows[-1].cells[1].text = "PROMEDIO COMPETENCIAS TÉCNICAS"
    table_ct.rows[-1].cells[2].text = ""
    table_ct.rows[-1].cells[3].text = ""
    table_ct.rows[-1].cells[4].text = "Peso: 40%"
    for cell in table_ct.rows[-1].cells:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)

    doc.add_paragraph()

    h4 = doc.add_heading("SECCIÓN 3: EVALUACIÓN EN SEGURIDAD DEL PACIENTE (20% del total)", level=1)
    style_heading(h4, 1)

    comp_sp = [
        ("3.1", "Identificación correcta del paciente antes de cada procedimiento"),
        ("3.2", "Obtención de consentimiento informado cuando es requerido"),
        ("3.3", "Reporte oportuno de eventos adversos e incidentes"),
        ("3.4", "Cumplimiento del protocolo de lavado de manos"),
        ("3.5", "Uso correcto de elementos de protección personal (EPP)"),
        ("3.6", "Cumplimiento de las listas de verificación (checklists) establecidas"),
        ("3.7", "Comunicación efectiva sobre el plan de manejo al paciente"),
    ]

    table_sp = doc.add_table(rows=len(comp_sp)+2, cols=5)
    table_sp.style = 'Table Grid'
    add_table_header_row(table_sp, ["N°", "Competencia", "Autoevaluación\n(1-5)", "Evaluación\n(1-5)", "Observaciones"])
    for i, (num, comp) in enumerate(comp_sp):
        table_sp.rows[i+1].cells[0].text = num
        table_sp.rows[i+1].cells[1].text = comp
        for cell in table_sp.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    table_sp.rows[-1].cells[1].text = "PROMEDIO SEGURIDAD DEL PACIENTE"
    table_sp.rows[-1].cells[4].text = "Peso: 20%"
    for cell in table_sp.rows[-1].cells:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)

    doc.add_paragraph()

    h5 = doc.add_heading("SECCIÓN 4: COMPETENCIAS HUMANÍSTICAS Y DE COMUNICACIÓN (20%)", level=1)
    style_heading(h5, 1)

    comp_hum = [
        ("4.1", "Trato respetuoso, empático y digno con todos los pacientes"),
        ("4.2", "Comunicación clara y comprensible con el paciente sobre su diagnóstico y plan"),
        ("4.3", "Manejo adecuado de la información confidencial del paciente"),
        ("4.4", "Atención a las necesidades psicosociales del paciente"),
        ("4.5", "Manejo constructivo de quejas y reclamos"),
        ("4.6", "Respeto a la autonomía del paciente en la toma de decisiones"),
    ]

    table_hum = doc.add_table(rows=len(comp_hum)+2, cols=5)
    table_hum.style = 'Table Grid'
    add_table_header_row(table_hum, ["N°", "Competencia", "Autoevaluación\n(1-5)", "Evaluación\n(1-5)", "Observaciones"])
    for i, (num, comp) in enumerate(comp_hum):
        table_hum.rows[i+1].cells[0].text = num
        table_hum.rows[i+1].cells[1].text = comp
        for cell in table_hum.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    table_hum.rows[-1].cells[1].text = "PROMEDIO COMPETENCIAS HUMANÍSTICAS"
    table_hum.rows[-1].cells[4].text = "Peso: 20%"
    for cell in table_hum.rows[-1].cells:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)

    doc.add_paragraph()

    h6 = doc.add_heading("SECCIÓN 5: COMPETENCIAS ADMINISTRATIVAS Y NORMATIVAS (20%)", level=1)
    style_heading(h6, 1)

    comp_adm = [
        ("5.1", "Cumplimiento de la jornada laboral y puntualidad"),
        ("5.2", "Organización y gestión eficiente del tiempo de consulta"),
        ("5.3", "Cumplimiento de los indicadores de calidad establecidos"),
        ("5.4", "Participación activa en capacitaciones y actividades de educación continuada"),
        ("5.5", "Mantenimiento vigente de todos los documentos de habilitación"),
        ("5.6", "Cumplimiento del reglamento interno y políticas del consultorio"),
        ("5.7", "Presentación personal y uso del uniforme/bata según normas"),
    ]

    table_adm = doc.add_table(rows=len(comp_adm)+2, cols=5)
    table_adm.style = 'Table Grid'
    add_table_header_row(table_adm, ["N°", "Competencia", "Autoevaluación\n(1-5)", "Evaluación\n(1-5)", "Observaciones"])
    for i, (num, comp) in enumerate(comp_adm):
        table_adm.rows[i+1].cells[0].text = num
        table_adm.rows[i+1].cells[1].text = comp
        for cell in table_adm.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    table_adm.rows[-1].cells[1].text = "PROMEDIO COMPETENCIAS ADMINISTRATIVAS"
    table_adm.rows[-1].cells[4].text = "Peso: 20%"
    for cell in table_adm.rows[-1].cells:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)

    doc.add_paragraph()

    h7 = doc.add_heading("SECCIÓN 6: CALIFICACIÓN FINAL", level=1)
    style_heading(h7, 1)

    table_final = doc.add_table(rows=8, cols=4)
    table_final.style = 'Table Grid'
    add_table_header_row(table_final, ["Sección", "Promedio Autoevaluación", "Promedio Evaluación", "Puntaje Ponderado"])
    secciones = [
        ("Competencias Técnicas (40%)", "", "", ""),
        ("Seguridad del Paciente (20%)", "", "", ""),
        ("Competencias Humanísticas (20%)", "", "", ""),
        ("Competencias Administrativas (20%)", "", "", ""),
        ("CALIFICACIÓN FINAL PONDERADA", "", "", ""),
        ("NIVEL DE DESEMPEÑO", "", "", ""),
        ("¿Requiere plan de mejoramiento?", "□ Sí  □ No", "", ""),
    ]
    for i, row in enumerate(secciones):
        for j, val in enumerate(row):
            table_final.rows[i+1].cells[j].text = val
            for run in table_final.rows[i+1].cells[j].paragraphs[0].runs:
                run.font.size = Pt(10)
    for cell in table_final.rows[5].cells:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

    doc.add_paragraph()

    h8 = doc.add_heading("SECCIÓN 7: FORTALEZAS IDENTIFICADAS", level=1)
    style_heading(h8, 1)
    for i in range(5):
        doc.add_paragraph(f"{i+1}. ________________________________________________________________________")

    h9 = doc.add_heading("SECCIÓN 8: OPORTUNIDADES DE MEJORA", level=1)
    style_heading(h9, 1)
    for i in range(5):
        doc.add_paragraph(f"{i+1}. ________________________________________________________________________")

    h10 = doc.add_heading("SECCIÓN 9: PLAN DE MEJORAMIENTO (si aplica)", level=1)
    style_heading(h10, 1)

    table_plan = doc.add_table(rows=5, cols=4)
    table_plan.style = 'Table Grid'
    add_table_header_row(table_plan, ["Oportunidad de mejora", "Acción a implementar", "Responsable", "Fecha límite"])
    for i in range(1, 5):
        for j in range(4):
            table_plan.rows[i].cells[j].text = ""

    doc.add_paragraph()

    h11 = doc.add_heading("SECCIÓN 10: FIRMAS Y COMPROMISOS", level=1)
    style_heading(h11, 1)

    add_body(doc, "El evaluado declara que ha participado activamente en el proceso de evaluación, que ha sido informado de los resultados y que se compromete a implementar las acciones del plan de mejoramiento.")
    doc.add_paragraph()

    table_firmas = doc.add_table(rows=3, cols=2)
    table_firmas.style = 'Table Grid'
    add_table_header_row(table_firmas, ["EVALUADO(A)", "EVALUADOR(A)"])
    firmas = [
        ("Nombre: _________________________\nCargo: _________________________\nFirma: _________________________\nFecha: _________________________",
         "Nombre: [NOMBRE DE LA MÉDICA]\nCargo: Médica General - Propietaria\nFirma: _________________________\nFecha: _________________________"),
        ("Comentarios del evaluado:\n\n___________________________________\n___________________________________\n___________________________________",
         "Comentarios del evaluador:\n\n___________________________________\n___________________________________\n___________________________________"),
    ]
    for i, (e1, e2) in enumerate(firmas):
        table_firmas.rows[i+1].cells[0].text = e1
        table_firmas.rows[i+1].cells[1].text = e2
        for cell in table_firmas.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    path = os.path.join(BASE_DIR, "FOR-TH-001_Formato_Evaluacion_Desempeno.docx")
    doc.save(path)
    print(f"✓ Creado: {path}")

if __name__ == "__main__":
    print("Generando documentos de CARPETA 1 - TALENTO HUMANO...")
    crear_manual_talento_humano()
    crear_perfil_cargo()
    crear_proceso_seleccion()
    crear_formato_evaluacion()
    print("\n✅ Carpeta 1 completada.")
