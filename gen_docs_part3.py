#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera PRO-GC-003 PQR y los 3 documentos de CARPETA 9"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

DARK_BLUE = RGBColor(0, 51, 102)
OUT_DIR_8 = "/home/user/bioauditoria/documentos_habilitacion/CARPETA_8_GESTION_CALIDAD"
OUT_DIR_9 = "/home/user/bioauditoria/documentos_habilitacion/CARPETA_9_HABILITACION_NORMATIVIDAD"
os.makedirs(OUT_DIR_9, exist_ok=True)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14) if level==1 else (Pt(12) if level==2 else Pt(11))
    run.font.color.rgb = DARK_BLUE; run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name='Calibri'; run.font.size=Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table_header(table, headers):
    row = table.rows[0]
    for i, h in enumerate(headers):
        if i >= len(row.cells): break
        cell = row.cells[i]; cell.text = h
        set_cell_bg(cell, '003366')
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold=True; run.font.color.rgb=RGBColor(255,255,255)
                run.font.size=Pt(10); run.font.name='Calibri'
            para.alignment=WD_ALIGN_PARAGRAPH.CENTER

def add_portada(doc, titulo, codigo):
    doc.add_paragraph(); doc.add_paragraph()
    for txt,sz,bold in [("[NOMBRE DEL CONSULTORIO]",18,True),
                         ("Consultorio Médico General y Procedimientos Estéticos No Invasivos",13,False),
                         (titulo,16,True)]:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(txt); r.bold=bold; r.font.size=Pt(sz); r.font.color.rgb=DARK_BLUE; r.font.name='Calibri'
        if not bold: doc.add_paragraph()
    doc.add_paragraph()
    t=doc.add_table(rows=5,cols=2); t.style='Table Grid'
    for i,(k,v) in enumerate([("Código:",codigo),("Versión:","1.0"),("Fecha:","Junio 2025"),
                                ("Responsable:","[NOMBRE DE LA MÉDICA]"),("Dirección:","[DIRECCIÓN DEL CONSULTORIO]")]):
        t.rows[i].cells[0].text=k; t.rows[i].cells[1].text=v
        set_cell_bg(t.rows[i].cells[0],'BDD7EE')
        for cell in t.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(11)
    doc.add_page_break()

def add_cv(doc):
    add_heading(doc,"CONTROL DE VERSIONES",1)
    t=doc.add_table(rows=2,cols=5); t.style='Table Grid'
    add_table_header(t,["Versión","Fecha","Descripción","Elaboró","Aprobó"])
    vals=["1.0","Junio 2025","Creación del documento","[NOMBRE DE LA MÉDICA]","[NOMBRE DE LA MÉDICA]"]
    for j,v in enumerate(vals):
        t.rows[1].cells[j].text=v
        for para in t.rows[1].cells[j].paragraphs:
            for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

def setup_doc(titulo, codigo):
    doc=Document()
    doc.styles['Normal'].font.name='Calibri'; doc.styles['Normal'].font.size=Pt(11)
    sec=doc.sections[0]
    sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5); sec.left_margin=Cm(3); sec.right_margin=Cm(2.5)
    hp=sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
    hp.clear(); hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=hp.add_run(f"{titulo} | {codigo}"); r.font.size=Pt(9); r.font.name='Calibri'; r.font.color.rgb=DARK_BLUE
    fp=sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    fp.clear(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=fp.add_run("Versión 1.0 | Junio 2025 | Página "); r2.font.size=Pt(9); r2.font.name='Calibri'
    fld1=OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'),'begin')
    instr=OxmlElement('w:instrText'); instr.text='PAGE'
    fld2=OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'),'end')
    r3=fp.add_run(); r3._r.append(fld1); r3._r.append(instr); r3._r.append(fld2)
    r3.font.size=Pt(9); r3.font.name='Calibri'
    return doc

def bullet(doc, text):
    pb=doc.add_paragraph(text,style='List Bullet')
    for run in pb.runs: run.font.name='Calibri'; run.font.size=Pt(11)

# =====================================================================
# DOCUMENTO 5: PRO-GC-003 PQR
# =====================================================================
def gen_pqr():
    titulo="PROCESO DE ATENCIÓN DE PETICIONES, QUEJAS, RECLAMOS Y SUGERENCIAS (PQRS)"
    codigo="PRO-GC-003"
    doc=setup_doc(titulo,codigo)
    add_portada(doc,titulo,codigo); add_cv(doc)
    add_heading(doc,"TABLA DE CONTENIDO",1)
    for item in ["1. Objetivo","2. Alcance","3. Marco Legal","4. Definiciones",
                 "5. Canales de Recepción","6. Clasificación de PQRS",
                 "7. Tiempos de Respuesta","8. Flujograma de Atención",
                 "9. Formato de Registro","10. Análisis de Causas",
                 "11. Acciones de Mejora","12. Indicadores de Satisfacción",
                 "13. Responsables","14. Registros"]:
        add_body(doc, item)
    doc.add_page_break()

    add_heading(doc,"1. OBJETIVO",1)
    add_body(doc,"Establecer el proceso para recibir, clasificar, atender y dar respuesta oportuna a las Peticiones, Quejas, Reclamos y Sugerencias (PQRS) presentadas por los usuarios del consultorio [NOMBRE DEL CONSULTORIO], garantizando el respeto por sus derechos, la mejora continua de los servicios y el cumplimiento de la normatividad vigente en materia de atención al usuario en salud.")

    add_heading(doc,"2. ALCANCE",1)
    add_body(doc,"Aplica a todas las PQRS recibidas de usuarios, familiares, acompañantes, entidades o cualquier persona relacionada con la prestación de servicios de salud del consultorio [NOMBRE DEL CONSULTORIO], a través de cualquier canal de recepción habilitado.")

    add_heading(doc,"3. MARCO LEGAL",1)
    normas=[("Ley 1122 de 2007","Por la cual se hacen algunas modificaciones en el Sistema General de Seguridad Social en Salud. Establece el derecho de los usuarios a presentar peticiones, quejas y reclamos."),
            ("Resolución 1552 de 2013","Por la cual se definen los indicadores de monitoreo del Sistema de Información para la Calidad, incluyendo la satisfacción del usuario."),
            ("Ley 1437 de 2011","Código de Procedimiento Administrativo y de lo Contencioso Administrativo. Regula el derecho de petición."),
            ("Ley 23 de 1981","Código de Ética Médica. Deberes del médico con los pacientes."),
            ("Resolución 13437 de 1991","Derechos de los pacientes. Lista de derechos que deben conocer los usuarios."),
            ("Ley 1755 de 2015","Por medio de la cual se regula el Derecho Fundamental de Petición.")]
    t=doc.add_table(rows=len(normas)+1,cols=2); t.style='Table Grid'
    add_table_header(t,["Norma","Descripción"])
    for i,(n,d) in enumerate(normas):
        t.rows[i+1].cells[0].text=n; t.rows[i+1].cells[1].text=d
        set_cell_bg(t.rows[i+1].cells[0],'EBF3FB')
        for cell in t.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc,"4. DEFINICIONES",1)
    defs=[("Petición","Solicitud de información, documentos o aclaración sobre los servicios prestados."),
          ("Queja","Expresión de insatisfacción del usuario sobre la atención recibida, sin pretensión de resarcimiento."),
          ("Reclamo","Solicitud de corrección o compensación por parte del usuario que considera se han vulnerado sus derechos."),
          ("Sugerencia","Propuesta del usuario orientada a mejorar la calidad de los servicios prestados."),
          ("PQRS","Acrónimo de Petición, Queja, Reclamo y Sugerencia. Sistema de atención al usuario."),
          ("Tiempo de respuesta","Plazo máximo establecido para dar respuesta a una PQRS desde su recepción."),
          ("Causa raíz","Factor fundamental que origina un problema y cuya eliminación evita la recurrencia.")]
    for term,defi in defs:
        p=doc.add_paragraph()
        r=p.add_run(f"{term}: "); r.bold=True; r.font.name='Calibri'; r.font.size=Pt(11)
        r2=p.add_run(defi); r2.font.name='Calibri'; r2.font.size=Pt(11)

    add_heading(doc,"5. CANALES DE RECEPCIÓN",1)
    add_body(doc,"El consultorio [NOMBRE DEL CONSULTORIO] habilita los siguientes canales para la recepción de PQRS:")
    t2=doc.add_table(rows=6,cols=4); t2.style='Table Grid'
    add_table_header(t2,["Canal","Descripción","Disponibilidad","Responsable de Recepción"])
    canales=[("Presencial","Buzón de sugerencias en recepción y atención directa al usuario","Horario de atención","Personal de recepción / [NOMBRE DE LA MÉDICA]"),
             ("Telefónico","Número de contacto del consultorio: [TELÉFONO]","Horario de atención","Personal de recepción"),
             ("Correo electrónico","[CORREO DEL CONSULTORIO]","24/7 (respuesta en horario laboral)","[NOMBRE DE LA MÉDICA]"),
             ("Redes sociales / WhatsApp","Mensajería a través de los canales oficiales del consultorio","24/7 (respuesta en horario laboral)","[NOMBRE DE LA MÉDICA]"),
             ("Formulario físico","Formato de PQRS disponible en recepción","Horario de atención","Personal de recepción")]
    for i,row in enumerate(canales):
        for j,val in enumerate(row):
            t2.rows[i+1].cells[j].text=val
            for para in t2.rows[i+1].cells[j].paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
        if i%2==0:
            for j in range(4): set_cell_bg(t2.rows[i+1].cells[j],'EBF3FB')
    doc.add_paragraph()

    add_heading(doc,"6. CLASIFICACIÓN DE PQRS",1)
    t3=doc.add_table(rows=5,cols=4); t3.style='Table Grid'
    add_table_header(t3,["Tipo","Descripción","Ejemplos","Tiempo de Respuesta"])
    clasif=[("PETICIÓN","Solicitud de información o documentos","Solicitud de copia de historia clínica, certificado de atención, información sobre tarifas","10 días hábiles"),
            ("QUEJA","Expresión de insatisfacción sin pretensión de resarcimiento","Mal trato por parte del personal, instalaciones inadecuadas, tiempo de espera excesivo","15 días hábiles"),
            ("RECLAMO","Solicitud de corrección por vulneración de derechos","Cobro excesivo, error en diagnóstico, negligencia percibida","15 días hábiles"),
            ("SUGERENCIA","Propuesta de mejora","Ideas para mejorar el servicio, equipamiento o atención","No requiere respuesta formal, se agradece y evalúa")]
    for i,row in enumerate(clasif):
        for j,val in enumerate(row):
            t3.rows[i+1].cells[j].text=val
            for para in t3.rows[i+1].cells[j].paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc,"7. TIEMPOS DE RESPUESTA",1)
    add_body(doc,"Los tiempos de respuesta se establecen según el tipo de PQRS y la normatividad vigente:")
    t4=doc.add_table(rows=5,cols=4); t4.style='Table Grid'
    add_table_header(t4,["Tipo de PQRS","Tiempo Máximo de Respuesta","Base Legal","Observaciones"])
    tiempos=[("Petición de información general","10 días hábiles","Ley 1755/2015","Contados desde la radicación"),
             ("Queja","15 días hábiles","Ley 1122/2007","Con posibilidad de prórroga de 15 días más"),
             ("Reclamo","15 días hábiles","Ley 1122/2007","Con posibilidad de prórroga de 15 días más"),
             ("Petición de documentos (HC)","5 días hábiles","Res. 1995/1999","Copia de historia clínica")]
    for i,row in enumerate(tiempos):
        for j,val in enumerate(row):
            t4.rows[i+1].cells[j].text=val
            for para in t4.rows[i+1].cells[j].paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
        if i%2==0:
            for j in range(4): set_cell_bg(t4.rows[i+1].cells[j],'EBF3FB')
    doc.add_paragraph()

    add_heading(doc,"8. FLUJOGRAMA DE ATENCIÓN DE PQRS",1)
    add_body(doc,"El proceso de atención de PQRS sigue el siguiente flujograma:")
    pasos_flujo=[
        ("1. RECEPCIÓN","El usuario presenta la PQRS a través de cualquiera de los canales habilitados. El personal de recepción o [NOMBRE DE LA MÉDICA] recibe y registra la PQRS en el formato de registro (FOR-GC-003-01)."),
        ("2. RADICACIÓN","Se asigna número de radicado, fecha y hora de recepción. Se informa al usuario sobre el tiempo de respuesta y el número de radicado para seguimiento."),
        ("3. CLASIFICACIÓN","Se clasifica la PQRS según el tipo (petición, queja, reclamo, sugerencia). Se determina el responsable de dar respuesta."),
        ("4. ANÁLISIS","[NOMBRE DE LA MÉDICA] analiza la PQRS, revisa la historia clínica si aplica, entrevista al personal involucrado y busca la causa raíz del problema."),
        ("5. FORMULACIÓN DE RESPUESTA","Se elabora la respuesta fundamentada en la normatividad y en los hallazgos del análisis. Se define si hay acciones correctivas."),
        ("6. RESPUESTA AL USUARIO","Se comunica la respuesta al usuario dentro del tiempo establecido, por el canal que eligió o por escrito. Se informa sobre las acciones de mejora tomadas."),
        ("7. CIERRE Y SEGUIMIENTO","Se cierra el caso en el registro. Si se generaron acciones correctivas, se hace seguimiento a su implementación. Se consolidan las PQRS para análisis de tendencias."),
    ]
    t5=doc.add_table(rows=len(pasos_flujo)+1,cols=2); t5.style='Table Grid'
    add_table_header(t5,["Etapa","Descripción del Proceso"])
    for i,(etapa,desc) in enumerate(pasos_flujo):
        t5.rows[i+1].cells[0].text=etapa; t5.rows[i+1].cells[1].text=desc
        set_cell_bg(t5.rows[i+1].cells[0],'BDD7EE' if i%2==0 else 'EBF3FB')
        for cell in t5.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(9)
    doc.add_paragraph()

    add_heading(doc,"9. FORMATO DE REGISTRO DE PQRS",1)
    add_body(doc,"Cada PQRS debe registrarse en el formato FOR-GC-003-01, que contiene los siguientes campos:")
    t6=doc.add_table(rows=13,cols=2); t6.style='Table Grid'
    add_table_header(t6,["Campo","Descripción / Instrucción de Diligenciamiento"])
    campos=[("N° de Radicado","Número consecutivo asignado automáticamente o manualmente"),
            ("Fecha y hora de recepción","DD/MM/AAAA HH:MM"),
            ("Canal de recepción","Presencial / Telefónico / Correo / WhatsApp / Formulario"),
            ("Tipo de PQRS","Petición / Queja / Reclamo / Sugerencia"),
            ("Datos del usuario","Nombre completo, tipo y número de documento, teléfono, correo electrónico"),
            ("Descripción detallada","Narración completa de la PQRS en palabras del usuario"),
            ("Servicio o área involucrada","Área o proceso relacionado con la PQRS"),
            ("Responsable asignado","Nombre de quien atenderá la PQRS"),
            ("Fecha límite de respuesta","Calculada según tipo de PQRS"),
            ("Análisis y causa raíz","Hallazgos del análisis"),
            ("Acciones de mejora tomadas","Descripción de las acciones correctivas o preventivas"),
            ("Respuesta al usuario","Resumen de la respuesta enviada al usuario")]
    for i,(c,d) in enumerate(campos):
        t6.rows[i+1].cells[0].text=c; t6.rows[i+1].cells[1].text=d
        set_cell_bg(t6.rows[i+1].cells[0],'EBF3FB')
        for cell in t6.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc,"10. ANÁLISIS DE CAUSAS",1)
    add_heading(doc,"10.1 Técnica de los 5 ¿Por Qué?",2)
    add_body(doc,"La técnica de los 5 ¿Por Qué? consiste en preguntarse sucesivamente '¿por qué?' ante cada respuesta obtenida, hasta llegar a la causa raíz del problema. Ejemplo aplicado:")
    t7=doc.add_table(rows=6,cols=2); t7.style='Table Grid'
    add_table_header(t7,["Pregunta","Respuesta (Ejemplo)"])
    cinco_pq=[("¿Por qué el usuario se quejó?","Por tiempo de espera excesivo para ser atendido"),
              ("¿Por qué hubo tiempo de espera excesivo?","Porque las citas no estaban bien distribuidas en la agenda"),
              ("¿Por qué las citas no estaban bien distribuidas?","Porque no se tiene en cuenta el tiempo real de cada tipo de consulta"),
              ("¿Por qué no se tiene en cuenta el tiempo real?","Porque no hay tiempos estándar definidos por tipo de consulta"),
              ("¿Por qué no hay tiempos estándar definidos?","Porque nunca se ha realizado un análisis del tiempo de atención (CAUSA RAÍZ)")]
    for i,(p,r) in enumerate(cinco_pq):
        t7.rows[i+1].cells[0].text=p; t7.rows[i+1].cells[1].text=r
        set_cell_bg(t7.rows[i+1].cells[0],'BDD7EE')
        for cell in t7.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc,"10.2 Espina de Pescado (Ishikawa)",2)
    add_body(doc,"El diagrama de causa-efecto (espina de pescado) identifica las causas desde 6 categorías (6M): Método, Mano de obra, Maquinaria/Equipo, Materiales, Medición y Medio ambiente. Para cada queja o reclamo de mayor impacto, se aplica este análisis:")
    t8=doc.add_table(rows=7,cols=2); t8.style='Table Grid'
    add_table_header(t8,["Categoría (6M)","Posibles Causas a Investigar"])
    seis_m=[("Método","Procesos no estandarizados, ausencia de protocolos, procedimientos inadecuados"),
            ("Mano de obra","Falta de capacitación, desmotivación, comunicación deficiente con el usuario"),
            ("Maquinaria/Equipo","Equipos defectuosos, falta de mantenimiento, tecnología inadecuada"),
            ("Materiales","Insumos de baja calidad, falta de materiales, productos no apropiados"),
            ("Medición","Falta de monitoreo, indicadores no definidos, ausencia de seguimiento"),
            ("Medio ambiente","Infraestructura inadecuada, ruido, falta de privacidad, condiciones incómodas")]
    for i,(cat,causas) in enumerate(seis_m):
        t8.rows[i+1].cells[0].text=cat; t8.rows[i+1].cells[1].text=causas
        set_cell_bg(t8.rows[i+1].cells[0],'EBF3FB')
        for cell in t8.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc,"11. ACCIONES DE MEJORA",1)
    add_body(doc,"Cada PQRS que identifique un problema en la calidad del servicio debe generar acciones de mejora. Estas se incorporan al plan de mejoramiento del PAMEC.")
    t9=doc.add_table(rows=5,cols=6); t9.style='Table Grid'
    add_table_header(t9,["N°","PQRS Relacionada","Causa Raíz Identificada","Acción de Mejora","Responsable","Fecha Compromiso"])
    for i in range(4):
        for j in range(6):
            t9.rows[i+1].cells[j].text=""
            for para in t9.rows[i+1].cells[j].paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc,"12. INDICADORES DE SATISFACCIÓN",1)
    t10=doc.add_table(rows=6,cols=5); t10.style='Table Grid'
    add_table_header(t10,["Indicador","Fórmula","Meta","Periodicidad","Fuente"])
    ind=[("Índice de satisfacción global","Usuarios satisfechos / Total encuestados x 100","≥ 90%","Trimestral","Encuesta de satisfacción"),
         ("Tasa de PQRS","N° PQRS / Total consultas x 100","< 2%","Mensual","Registro de PQRS"),
         ("Oportunidad en respuesta de PQRS","PQRS respondidas a tiempo / Total PQRS x 100","100%","Mensual","Registro de PQRS"),
         ("Resolución efectiva de PQRS","PQRS resueltas satisfactoriamente / Total PQRS x 100","≥ 90%","Trimestral","Encuesta de seguimiento"),
         ("Tasa de reincidencia","PQRS por el mismo motivo / Total PQRS x 100","< 5%","Trimestral","Registro de PQRS")]
    for i,row in enumerate(ind):
        for j,val in enumerate(row):
            t10.rows[i+1].cells[j].text=val
            for para in t10.rows[i+1].cells[j].paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(9)
        if i%2==0:
            for j in range(5): set_cell_bg(t10.rows[i+1].cells[j],'EBF3FB')

    path=os.path.join(OUT_DIR_8,"PRO-GC-003_Proceso_Atencion_Quejas_Reclamos.docx")
    doc.save(path); print(f"Guardado: {path}")

# =====================================================================
# DOCUMENTO 6: MAN-HAB-001 Manual de Habilitación
# =====================================================================
def gen_manual_habilitacion():
    titulo="MANUAL DE HABILITACIÓN DE SERVICIOS DE SALUD"
    codigo="MAN-HAB-001"
    doc=setup_doc(titulo,codigo)
    add_portada(doc,titulo,codigo); add_cv(doc)
    add_heading(doc,"TABLA DE CONTENIDO",1)
    for item in ["1. Objetivo","2. Alcance","3. Marco Legal",
                 "4. Sistema Obligatorio de Garantía de Calidad (SOGCS)",
                 "5. Proceso de Habilitación en Colombia",
                 "6. Estándares de Habilitación Resolución 3100/2019 - Medicina General",
                 "7. Lista de Verificación de Estándares",
                 "8. Formulario de Inscripción en el REPS",
                 "9. Renovación de Habilitación",
                 "10. Visita de Verificación: Preparación y Qué Esperar",
                 "11. Sanciones por Incumplimiento","12. Responsables"]:
        add_body(doc,item)
    doc.add_page_break()

    add_heading(doc,"1. OBJETIVO",1)
    add_body(doc,"Orientar al consultorio [NOMBRE DEL CONSULTORIO] en el proceso de habilitación de servicios de salud, brindando información clara y completa sobre los requisitos normativos, los estándares a cumplir, el proceso de inscripción en el Registro Especial de Prestadores de Servicios de Salud (REPS) y los aspectos clave para superar la visita de verificación de la Secretaría de Salud.")

    add_heading(doc,"2. ALCANCE",1)
    add_body(doc,"El presente manual aplica para el proceso de habilitación del servicio de consulta externa de medicina general y procedimientos estéticos no invasivos del consultorio [NOMBRE DEL CONSULTORIO], ubicado en [DIRECCIÓN DEL CONSULTORIO], ante la Secretaría de Salud del municipio/departamento correspondiente.")

    add_heading(doc,"3. MARCO LEGAL",1)
    normas=[("Ley 100 de 1993","Crea el Sistema General de Seguridad Social en Salud."),
            ("Decreto 1011 de 2006","Establece el Sistema Obligatorio de Garantía de Calidad de la Atención de Salud (SOGCS)."),
            ("Resolución 1043 de 2006","Condiciones que deben cumplir los Prestadores de Servicios de Salud para habilitar sus servicios."),
            ("Resolución 2003 de 2014","Define procedimientos y condiciones de inscripción de prestadores y habilitación de servicios."),
            ("Resolución 3100 de 2019","Actualiza y define los procedimientos y condiciones de inscripción de prestadores de servicios de salud y habilitación de servicios de salud. NORMA VIGENTE."),
            ("Decreto 780 de 2016","Decreto Único Reglamentario del Sector Salud y Protección Social.")]
    t=doc.add_table(rows=len(normas)+1,cols=2); t.style='Table Grid'
    add_table_header(t,["Norma","Descripción"])
    for i,(n,d) in enumerate(normas):
        t.rows[i+1].cells[0].text=n; t.rows[i+1].cells[1].text=d
        set_cell_bg(t.rows[i+1].cells[0],'EBF3FB')
        for cell in t.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc,"4. SISTEMA OBLIGATORIO DE GARANTÍA DE CALIDAD (SOGCS)",1)
    add_body(doc,"El SOGCS es el conjunto de instituciones, normas, requisitos, mecanismos y procesos deliberados y sistemáticos del sector salud para generar, mantener y mejorar la calidad de los servicios de salud. Está compuesto por cuatro componentes:")
    t2=doc.add_table(rows=5,cols=3); t2.style='Table Grid'
    add_table_header(t2,["Componente","Descripción","Entidad Responsable"])
    sogcs=[("1. Habilitación","Sistema de registro de los prestadores y sus servicios que cumplan con condiciones mínimas de capacidad tecnológica, científica y suficiencia patrimonial y financiera.","Secretarías de Salud (Departamental/Municipal)"),
           ("2. Auditoría para el Mejoramiento (PAMEC)","Mecanismo sistemático y continuo de evaluación y mejoramiento de la calidad observada respecto de la calidad esperada.","Prestadores de Servicios de Salud"),
           ("3. Sistema de Información para la Calidad","Sistema de información que permite identificar, recolectar, procesar y analizar información sobre la calidad de los servicios de salud.","Ministerio de Salud / Prestadores"),
           ("4. Acreditación en Salud","Proceso voluntario y periódico de autoevaluación interna y revisión externa de los procesos y resultados que garantizan y mejoran la calidad.","ICONTEC (Entidad Acreditadora)")]
    for i,row in enumerate(sogcs):
        for j,val in enumerate(row):
            t2.rows[i+1].cells[j].text=val
            for para in t2.rows[i+1].cells[j].paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
        if i%2==0:
            for j in range(3): set_cell_bg(t2.rows[i+1].cells[j],'EBF3FB')
    doc.add_paragraph()

    add_heading(doc,"5. PROCESO DE HABILITACIÓN EN COLOMBIA - PASO A PASO",1)
    pasos_hab=[("PASO 1: Autoevaluación","Realizar una autoevaluación exhaustiva de las condiciones del consultorio comparadas con los estándares de la Resolución 3100 de 2019. Identificar las brechas y desarrollar un plan de acción para subsanarlas antes de presentarse ante la Secretaría."),
               ("PASO 2: Preparación de documentos","Organizar toda la documentación requerida: hojas de vida del personal, certificados de formación, contratos, inventario de equipos, fichas técnicas, manuales, protocolos, pólizas de responsabilidad civil, etc."),
               ("PASO 3: Inscripción en el REPS","Acceder al portal del Registro Especial de Prestadores de Servicios de Salud (REPS) en: www.sispro.gov.co. Diligenciar el formulario de inscripción con los datos del prestador y los servicios a habilitar."),
               ("PASO 4: Radicación ante la Secretaría de Salud","Radicar la solicitud de habilitación ante la Secretaría de Salud del municipio/departamento. Presentar los documentos requeridos y el pago de la tarifa si aplica."),
               ("PASO 5: Visita de verificación","La Secretaría de Salud programa una visita de verificación al consultorio. Durante la visita, los funcionarios verificarán el cumplimiento de los estándares de habilitación in situ."),
               ("PASO 6: Concepto de la visita","Tras la visita, la Secretaría emite un concepto favorable o desfavorable. Si hay hallazgos, se concede un plazo para subsanarlos."),
               ("PASO 7: Habilitación y REPS","Con concepto favorable, el prestador queda habilitado e inscrito en el REPS. El certificado de habilitación está disponible en el portal REPS."),
               ("PASO 8: Mantenimiento","El prestador debe mantener las condiciones de habilitación durante toda la vigencia. Debe reportar cambios que modifiquen las condiciones habilitadas.")]
    t3=doc.add_table(rows=len(pasos_hab)+1,cols=2); t3.style='Table Grid'
    add_table_header(t3,["Paso","Descripción"])
    for i,(p,d) in enumerate(pasos_hab):
        t3.rows[i+1].cells[0].text=p; t3.rows[i+1].cells[1].text=d
        set_cell_bg(t3.rows[i+1].cells[0],'BDD7EE')
        for cell in t3.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc,"6. ESTÁNDARES DE HABILITACIÓN RESOLUCIÓN 3100/2019 - MEDICINA GENERAL",1)
    add_body(doc,"La Resolución 3100 de 2019 establece los siguientes grupos de estándares para la habilitación de servicios de salud:")
    estandares=[
        ("1. Talento Humano","Requisitos relacionados con el personal que presta los servicios de salud: formación, certificaciones, títulos, contratos y perfiles de cargo.",["Médico con título profesional y tarjeta profesional vigente","Certificados de formación complementaria actualizados","Contratos o convenios formalizados","Reglamento de trabajo o manual de funciones"]),
        ("2. Infraestructura",  "Condiciones físicas de las instalaciones donde se prestan los servicios de salud.",["Áreas específicas para cada servicio claramente delimitadas","Dimensiones mínimas según el tipo de servicio","Ventilación e iluminación adecuadas","Accesibilidad para personas con discapacidad (Ley 361/1997)","Baño para usuarios","Área de lavado de manos en el consultorio"]),
        ("3. Dotación","Equipos biomédicos, instrumental e insumos necesarios para la prestación del servicio.",["Equipo de cómputo para historia clínica electrónica o papelería","Tensiómetro calibrado","Estetoscopio","Termómetro","Báscula","Equipo para procedimientos estéticos habilitados"]),
        ("4. Medicamentos y Dispositivos Médicos","Condiciones para el almacenamiento y manejo de medicamentos e insumos.",["Medicamentos esenciales para la atención de urgencias básicas (si aplica)","Epinefrina disponible para manejo de anafilaxia","Condiciones de almacenamiento adecuadas (temperatura, luz, humedad)"]),
        ("5. Procesos Prioritarios","Documentación y cumplimiento de los procesos asistenciales críticos.",["Protocolos y guías de práctica clínica","Protocolo de manejo de residuos hospitalarios","Protocolo de lavado de manos","Protocolo de esterilización y desinfección","Consentimientos informados","Plan de emergencias"]),
        ("6. Historia Clínica y Registros","Requisitos para el diligenciamiento, conservación y confidencialidad de la historia clínica.",["Historia clínica completa por cada usuario atendido","Consentimiento informado firmado por el paciente","Archivo de historias clínicas organizado y seguro","Confidencialidad garantizada"]),
        ("7. Interdependencia de Servicios","Capacidad del prestador para articularse con otros niveles de atención cuando sea necesario.",["Mecanismo formal de referencia y contrarreferencia","Convenio o acuerdo con servicio de urgencias","Directorio de instituciones de referencia"]),
        ("8. Referencia y Contrarreferencia","Procesos para la remisión de pacientes a otros niveles de atención.",["Protocolo de remisión de pacientes","Formato de remisión","Registro de pacientes remitidos"]),
    ]
    for nombre,desc,criterios in estandares:
        add_heading(doc,f"6.{estandares.index((nombre,desc,criterios))+1} Estándar: {nombre}",2)
        add_body(doc,desc)
        for c in criterios: bullet(doc,c)
        doc.add_paragraph()

    add_heading(doc,"8. FORMULARIO DE INSCRIPCIÓN EN EL REPS",1)
    add_body(doc,"El Registro Especial de Prestadores de Servicios de Salud (REPS) es el sistema de información a través del cual los prestadores de servicios de salud se inscriben para habilitar sus servicios. El formulario de inscripción contiene los siguientes módulos:")
    modulos=[("Módulo 1: Datos del Prestador","Nombre o razón social, NIT o cédula, representante legal, dirección, municipio, teléfono, correo electrónico."),
             ("Módulo 2: Servicios a Habilitar","Código y nombre de cada servicio según el Manual de Inscripción de Prestadores (Res. 3100/2019), complejidad, sede de prestación."),
             ("Módulo 3: Talento Humano","Nombre, documento, título profesional, tarjeta profesional de cada profesional de salud."),
             ("Módulo 4: Capacidad Instalada","Número de consultorios, camas (si aplica), equipos biomédicos principales."),
             ("Módulo 5: Declaración","Declaración juramentada del cumplimiento de los estándares de habilitación.")]
    t4=doc.add_table(rows=len(modulos)+1,cols=2); t4.style='Table Grid'
    add_table_header(t4,["Módulo","Información Requerida"])
    for i,(m,info) in enumerate(modulos):
        t4.rows[i+1].cells[0].text=m; t4.rows[i+1].cells[1].text=info
        set_cell_bg(t4.rows[i+1].cells[0],'BDD7EE')
        for cell in t4.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc,"9. RENOVACIÓN DE HABILITACIÓN",1)
    add_body(doc,"La habilitación de servicios de salud no tiene una periodicidad fija de renovación en la Resolución 3100 de 2019. Sin embargo, el prestador debe:")
    for item in ["Actualizar la información en el REPS cuando haya cambios en los servicios, el personal, la infraestructura o la dotación.",
                 "Mantener las condiciones de habilitación en todo momento durante la vigencia.",
                 "Informar a la Secretaría de Salud sobre la suspensión o clausura de servicios.",
                 "Actualizar anualmente la declaración de cumplimiento de estándares en el REPS.",
                 "Verificar que las certificaciones del personal se mantengan vigentes.",
                 "Mantener al día los contratos de mantenimiento preventivo de equipos.",
                 "Conservar el archivo documental actualizado y disponible para inspección."]:
        bullet(doc,item)

    add_heading(doc,"10. VISITA DE VERIFICACIÓN: PREPARACIÓN Y QUÉ ESPERAR",1)
    add_heading(doc,"10.1 Preparación para la Visita",2)
    prep=[("Documentación","Tener organizada toda la documentación en carpetas separadas por estándar (Talento Humano, Infraestructura, Dotación, etc.)."),
          ("Instalaciones","Las instalaciones deben estar limpias, organizadas y señalizadas. Los equipos deben estar funcionando y con certificados de calibración al día."),
          ("Personal","Todo el personal debe conocer los protocolos y procedimientos. Deben poder explicar su rol y sus responsabilidades."),
          ("Equipos y dotación","Los equipos deben estar limpios, con etiquetas de identificación, y con sus hojas de vida actualizadas."),
          ("Residuos","El área de almacenamiento de residuos debe estar ordenada y los recipientes correctamente identificados."),
          ("Historia clínica","Tener disponibles ejemplos de historias clínicas completas y los formatos de consentimiento informado.")]
    t5=doc.add_table(rows=len(prep)+1,cols=2); t5.style='Table Grid'
    add_table_header(t5,["Aspecto a Preparar","Acciones"])
    for i,(a,acc) in enumerate(prep):
        t5.rows[i+1].cells[0].text=a; t5.rows[i+1].cells[1].text=acc
        set_cell_bg(t5.rows[i+1].cells[0],'EBF3FB')
        for cell in t5.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc,"10.2 Durante la Visita",2)
    for item in ["Recibir amablemente a los visitadores de la Secretaría de Salud.",
                 "Presentar la documentación solicitada de forma organizada.",
                 "Acompañar a los visitadores durante el recorrido por las instalaciones.",
                 "Responder con honestidad y precisión las preguntas formuladas.",
                 "Tomar nota de todos los hallazgos y observaciones realizados.",
                 "No discutir ni confrontar; si hay discrepancias, presentarlas por escrito en la respuesta.",
                 "Solicitar copia del acta de visita al finalizar."]:
        bullet(doc,item)

    add_heading(doc,"11. SANCIONES POR INCUMPLIMIENTO",1)
    add_body(doc,"De conformidad con la Ley 9 de 1979, el Decreto 1011 de 2006 y la Resolución 3100 de 2019, las sanciones por incumplimiento de las normas de habilitación pueden incluir:")
    t6=doc.add_table(rows=5,cols=3); t6.style='Table Grid'
    add_table_header(t6,["Tipo de Sanción","Descripción","Norma"])
    sanciones=[("Amonestación escrita","Primera medida ante hallazgos menores. Se concede plazo para corrección.","Ley 9/1979, Res. 3100/2019"),
               ("Multa","Sanción económica proporcional a la gravedad del incumplimiento.","Ley 9/1979"),
               ("Suspensión del servicio","Cierre temporal del servicio hasta que se subsanen las condiciones que generaron el riesgo.","Res. 3100/2019"),
               ("Clausura definitiva","Cierre definitivo del establecimiento en casos de incumplimiento grave o reiterativo.","Ley 9/1979")]
    for i,row in enumerate(sanciones):
        for j,val in enumerate(row):
            t6.rows[i+1].cells[j].text=val
            for para in t6.rows[i+1].cells[j].paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
        if i%2==0:
            for j in range(3): set_cell_bg(t6.rows[i+1].cells[j],'FCE4D6')

    path=os.path.join(OUT_DIR_9,"MAN-HAB-001_Manual_Habilitacion_Servicios_Salud.docx")
    doc.save(path); print(f"Guardado: {path}")

# =====================================================================
# DOCUMENTO 7: CHK-HAB-001 Lista de Verificación
# =====================================================================
def gen_lista_verificacion():
    titulo="LISTA DE VERIFICACIÓN DE ESTÁNDARES RESOLUCIÓN 3100 DE 2019"
    codigo="CHK-HAB-001"
    doc=setup_doc(titulo,codigo)
    add_portada(doc,titulo,codigo); add_cv(doc)
    add_heading(doc,"INSTRUCCIONES DE USO",1)
    add_body(doc,"Esta lista de verificación está diseñada para realizar la autoevaluación del cumplimiento de los estándares de habilitación establecidos en la Resolución 3100 de 2019 para el servicio de consulta externa de medicina general y procedimientos estéticos no invasivos. Diligencie cada criterio marcando con X en la columna correspondiente y registre la evidencia disponible.")
    add_body(doc,"CONVENCIONES: C = Cumple | NC = No Cumple | EP = En Proceso de Cumplimiento")
    doc.add_paragraph()

    def add_estandar_table(doc, estandar_titulo, criterios):
        add_heading(doc, estandar_titulo, 2)
        ncols = 7
        t=doc.add_table(rows=len(criterios)+1, cols=ncols)
        t.style='Table Grid'
        add_table_header(t,["N°","Criterio de Verificación","C","NC","EP","Evidencia Requerida","Responsable"])
        for i,row in enumerate(criterios):
            for j,val in enumerate(row):
                if j < ncols:
                    t.rows[i+1].cells[j].text=val
                    for para in t.rows[i+1].cells[j].paragraphs:
                        for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(8)
            if i%2==0:
                for j in range(ncols): set_cell_bg(t.rows[i+1].cells[j],'F2F7FB')
        doc.add_paragraph()

    # Estándar 1: Talento Humano
    add_heading(doc,"ESTÁNDAR 1: TALENTO HUMANO",1)
    criterios_th=[
        ("TH-01","El médico que presta el servicio cuenta con título profesional de médico o médico-cirujano","","","","Diploma, acta de grado","[NOMBRE DE LA MÉDICA]"),
        ("TH-02","El médico cuenta con tarjeta profesional vigente expedida por el Tribunal Ético de Medicina","","","","Tarjeta profesional","[NOMBRE DE LA MÉDICA]"),
        ("TH-03","El médico cuenta con certificados de formación complementaria actualizados (cursos de actualización, posgrados si aplica)","","","","Certificados vigentes","[NOMBRE DE LA MÉDICA]"),
        ("TH-04","Existe un perfil de cargo o manual de funciones para cada cargo","","","","Manual de funciones","[NOMBRE DE LA MÉDICA]"),
        ("TH-05","El personal auxiliar o de apoyo cuenta con los certificados o títulos exigidos para su cargo","","","","Diplomas/certificados del personal","[NOMBRE DE LA MÉDICA]"),
        ("TH-06","El personal auxiliar tiene el carnet de vacunación al día (Hepatitis B, Tétanos)","","","","Carnet de vacunación","[NOMBRE DE LA MÉDICA]"),
        ("TH-07","Existe contrato o vinculación laboral formalizada para todo el personal","","","","Contratos laborales","[NOMBRE DE LA MÉDICA]"),
        ("TH-08","El médico cuenta con ACLS (Soporte Vital Cardiovascular Avanzado) vigente si realiza procedimientos","","","","Certificado ACLS","[NOMBRE DE LA MÉDICA]"),
    ]
    add_estandar_table(doc,"1.1 Consulta Externa Medicina General",criterios_th)

    # Estándar 2: Infraestructura
    add_heading(doc,"ESTÁNDAR 2: INFRAESTRUCTURA",1)
    criterios_inf=[
        ("INF-01","El consultorio cuenta con sala de espera con sillas suficientes para los usuarios","","","","Observación directa","[NOMBRE DE LA MÉDICA]"),
        ("INF-02","El área del consultorio médico es de mínimo 9 m² (independiente de la sala de espera)","","","","Planos o medición","[NOMBRE DE LA MÉDICA]"),
        ("INF-03","El consultorio cuenta con lavamanos de uso exclusivo para el médico","","","","Observación directa","[NOMBRE DE LA MÉDICA]"),
        ("INF-04","Existe baño para usuarios separado del área clínica","","","","Observación directa","[NOMBRE DE LA MÉDICA]"),
        ("INF-05","Las instalaciones cuentan con ventilación natural o mecánica adecuada","","","","Observación directa","[NOMBRE DE LA MÉDICA]"),
        ("INF-06","La iluminación es suficiente para realizar los procedimientos clínicos","","","","Observación directa","[NOMBRE DE LA MÉDICA]"),
        ("INF-07","El establecimiento cuenta con señalización de seguridad y evacuación","","","","Observación directa / fotos","[NOMBRE DE LA MÉDICA]"),
        ("INF-08","Existen medidas de accesibilidad para personas con discapacidad (rampas, barandas)","","","","Observación directa","[NOMBRE DE LA MÉDICA]"),
        ("INF-09","Las paredes, pisos y techos son de material lavable y en buen estado","","","","Observación directa","[NOMBRE DE LA MÉDICA]"),
        ("INF-10","Existe área diferenciada para procedimientos estéticos (si aplica, separada del consultorio general)","","","","Observación directa/planos","[NOMBRE DE LA MÉDICA]"),
        ("INF-11","El consultorio cuenta con extintor vigente y en lugar accesible","","","","Extintor con etiqueta de revisión","[NOMBRE DE LA MÉDICA]"),
        ("INF-12","Existe área de almacenamiento de residuos hospitalarios claramente identificada","","","","Observación directa","[NOMBRE DE LA MÉDICA]"),
    ]
    add_estandar_table(doc,"2.1 Condiciones de Infraestructura Física",criterios_inf)

    # Estándar 3: Dotación
    add_heading(doc,"ESTÁNDAR 3: DOTACIÓN",1)
    criterios_dot=[
        ("DOT-01","Tensiómetro (esfigmomanómetro) calibrado y con certificado vigente","","","","Certificado de calibración","[NOMBRE DE LA MÉDICA]"),
        ("DOT-02","Estetoscopio en buen estado de funcionamiento","","","","Observación directa","[NOMBRE DE LA MÉDICA]"),
        ("DOT-03","Termómetro","","","","Observación directa","[NOMBRE DE LA MÉDICA]"),
        ("DOT-04","Báscula con tallímetro o tallímetro independiente","","","","Observación directa, certificado calibración","[NOMBRE DE LA MÉDICA]"),
        ("DOT-05","Camilla de examen en buen estado con sábana o cubierta desechable","","","","Observación directa","[NOMBRE DE LA MÉDICA]"),
        ("DOT-06","Linterna o fuente de luz para examen clínico","","","","Observación directa","[NOMBRE DE LA MÉDICA]"),
        ("DOT-07","Equipo de cómputo para historia clínica o papelería para HC manual","","","","Observación directa","[NOMBRE DE LA MÉDICA]"),
        ("DOT-08","Todos los equipos tienen hoja de vida con registros de mantenimiento","","","","Hojas de vida de equipos","[NOMBRE DE LA MÉDICA]"),
        ("DOT-09","Existe cronograma de mantenimiento preventivo de equipos","","","","Cronograma de mantenimiento","[NOMBRE DE LA MÉDICA]"),
        ("DOT-10","Para procedimientos estéticos: equipos específicos habilitados con fichas técnicas","","","","Fichas técnicas, manual del fabricante","[NOMBRE DE LA MÉDICA]"),
        ("DOT-11","Existe dotación para manejo de urgencias básicas (epinefrina, ambu, oxígeno si aplica)","","","","Inventario de botiquín de urgencias","[NOMBRE DE LA MÉDICA]"),
    ]
    add_estandar_table(doc,"3.1 Dotación y Equipos Médicos",criterios_dot)

    # Estándar 4: Medicamentos
    add_heading(doc,"ESTÁNDAR 4: MEDICAMENTOS Y DISPOSITIVOS MÉDICOS",1)
    criterios_med=[
        ("MED-01","Los medicamentos almacenados están dentro de su fecha de vencimiento","","","","Revisión física","[NOMBRE DE LA MÉDICA]"),
        ("MED-02","Los medicamentos están almacenados en condiciones adecuadas (temperatura, luz, humedad)","","","","Observación directa, termómetro ambiental","[NOMBRE DE LA MÉDICA]"),
        ("MED-03","Existe inventario actualizado de medicamentos","","","","Inventario de medicamentos","[NOMBRE DE LA MÉDICA]"),
        ("MED-04","Los medicamentos termolábiles (toxina botulínica, etc.) tienen cadena de frío garantizada","","","","Registro de temperaturas del refrigerador","[NOMBRE DE LA MÉDICA]"),
        ("MED-05","Los dispositivos médicos (agujas, jeringas, etc.) son de uso único y estériles","","","","Observación directa, fechas de vencimiento","[NOMBRE DE LA MÉDICA]"),
        ("MED-06","Existe epinefrina disponible para manejo de reacciones anafilácticas","","","","Inventario de urgencias","[NOMBRE DE LA MÉDICA]"),
        ("MED-07","Los medicamentos vencidos se gestionan como residuos peligrosos","","","","Registros de gestión de residuos","[NOMBRE DE LA MÉDICA]"),
    ]
    add_estandar_table(doc,"4.1 Medicamentos y Dispositivos Médicos",criterios_med)

    # Estándar 5: Procesos Prioritarios
    add_heading(doc,"ESTÁNDAR 5: PROCESOS PRIORITARIOS",1)
    criterios_pp=[
        ("PP-01","Existe protocolo de lavado de manos (higiene de manos) documentado y disponible","","","","Protocolo documentado","[NOMBRE DE LA MÉDICA]"),
        ("PP-02","El personal conoce y aplica la técnica correcta de lavado de manos","","","","Observación directa","[NOMBRE DE LA MÉDICA]"),
        ("PP-03","Existe protocolo de esterilización y desinfección de equipos e instrumental","","","","Protocolo documentado","[NOMBRE DE LA MÉDICA]"),
        ("PP-04","Existen guías de práctica clínica para las patologías más frecuentes","","","","Guías disponibles y accesibles","[NOMBRE DE LA MÉDICA]"),
        ("PP-05","Existe plan de gestión integral de residuos hospitalarios (PGIRHS) documentado","","","","PGIRHS documentado","[NOMBRE DE LA MÉDICA]"),
        ("PP-06","Se realiza separación en la fuente de residuos con los colores establecidos","","","","Observación directa, recipientes disponibles","[NOMBRE DE LA MÉDICA]"),
        ("PP-07","Existe plan de emergencias y contingencias documentado","","","","Plan de emergencias documentado","[NOMBRE DE LA MÉDICA]"),
        ("PP-08","Existe protocolo para el reporte y análisis de eventos adversos","","","","Protocolo documentado, formato de reporte","[NOMBRE DE LA MÉDICA]"),
        ("PP-09","Existe formato de consentimiento informado para los procedimientos realizados","","","","Formatos de consentimiento informado","[NOMBRE DE LA MÉDICA]"),
        ("PP-10","Existe PAMEC documentado e implementado","","","","PAMEC documentado","[NOMBRE DE LA MÉDICA]"),
        ("PP-11","Para procedimientos estéticos: existe protocolo específico para cada procedimiento","","","","Protocolos de procedimientos estéticos","[NOMBRE DE LA MÉDICA]"),
        ("PP-12","El personal está capacitado en los procedimientos que realiza","","","","Certificados de capacitación","[NOMBRE DE LA MÉDICA]"),
    ]
    add_estandar_table(doc,"5.1 Procesos Prioritarios y Seguros",criterios_pp)

    # Estándar 6: Historia Clínica
    add_heading(doc,"ESTÁNDAR 6: HISTORIA CLÍNICA Y REGISTROS",1)
    criterios_hc=[
        ("HC-01","Se diligencia historia clínica completa por cada usuario atendido","","","","Revisión de muestra de HC","[NOMBRE DE LA MÉDICA]"),
        ("HC-02","La historia clínica incluye: anamnesis, examen físico, diagnóstico, tratamiento, evolución","","","","Revisión de HC","[NOMBRE DE LA MÉDICA]"),
        ("HC-03","El consentimiento informado está firmado por el paciente antes del procedimiento","","","","Revisión de HC con procedimientos","[NOMBRE DE LA MÉDICA]"),
        ("HC-04","Las historias clínicas están archivadas de forma segura y confidencial","","","","Observación del archivo","[NOMBRE DE LA MÉDICA]"),
        ("HC-05","El archivo de HC garantiza la conservación por mínimo 20 años desde la última atención","","","","Sistema de archivo documentado","[NOMBRE DE LA MÉDICA]"),
        ("HC-06","Las HC están protegidas contra acceso no autorizado","","","","Mecanismos de seguridad","[NOMBRE DE LA MÉDICA]"),
        ("HC-07","Si la HC es electrónica, existe sistema de backup y respaldo","","","","Sistema de backup documentado","[NOMBRE DE LA MÉDICA]"),
    ]
    add_estandar_table(doc,"6.1 Historia Clínica",criterios_hc)

    # Estándar 7 y 8
    add_heading(doc,"ESTÁNDAR 7: INTERDEPENDENCIA DE SERVICIOS",1)
    criterios_inter=[
        ("INT-01","Existe mecanismo formal para la remisión de pacientes a otros niveles de atención","","","","Protocolo de remisión","[NOMBRE DE LA MÉDICA]"),
        ("INT-02","Existe directorio actualizado de instituciones de referencia (urgencias, hospitalización, laboratorio)","","","","Directorio disponible","[NOMBRE DE LA MÉDICA]"),
        ("INT-03","Existe convenio o acuerdo con al menos un servicio de urgencias para la atención de emergencias","","","","Convenio o carta de intención","[NOMBRE DE LA MÉDICA]"),
        ("INT-04","El consultorio cuenta con formato de remisión que incluye información clínica del paciente","","","","Formato de remisión","[NOMBRE DE LA MÉDICA]"),
    ]
    add_estandar_table(doc,"7.1 Interdependencia de Servicios",criterios_inter)

    # Procedimientos estéticos adicionales
    add_heading(doc,"CRITERIOS ADICIONALES PARA PROCEDIMIENTOS ESTÉTICOS NO INVASIVOS",1)
    criterios_est=[
        ("EST-01","El médico cuenta con formación específica en los procedimientos estéticos que realiza","","","","Diplomas/certificados de formación en estética","[NOMBRE DE LA MÉDICA]"),
        ("EST-02","Existe protocolo específico para toxina botulínica (si se aplica)","","","","Protocolo documentado","[NOMBRE DE LA MÉDICA]"),
        ("EST-03","Existe protocolo específico para ácido hialurónico (si se aplica)","","","","Protocolo documentado","[NOMBRE DE LA MÉDICA]"),
        ("EST-04","Existe protocolo para mesoterapia (si se aplica)","","","","Protocolo documentado","[NOMBRE DE LA MÉDICA]"),
        ("EST-05","Los productos utilizados en los procedimientos estéticos cuentan con registro INVIMA vigente","","","","Etiquetas o consulta INVIMA","[NOMBRE DE LA MÉDICA]"),
        ("EST-06","Existe protocolo de atención de reacciones adversas a los productos utilizados","","","","Protocolo documentado","[NOMBRE DE LA MÉDICA]"),
        ("EST-07","Se dispone de hialuronidasa para revertir complicaciones de ácido hialurónico (si aplica)","","","","Inventario","[NOMBRE DE LA MÉDICA]"),
        ("EST-08","Existe consentimiento informado específico para cada tipo de procedimiento estético","","","","Formatos de CI por procedimiento","[NOMBRE DE LA MÉDICA]"),
        ("EST-09","Se realiza fotografía clínica antes y después de los procedimientos (con autorización del paciente)","","","","Sistema de archivo fotográfico","[NOMBRE DE LA MÉDICA]"),
    ]
    add_estandar_table(doc,"Criterios Específicos para Procedimientos Estéticos",criterios_est)

    add_heading(doc,"RESUMEN DE AUTOEVALUACIÓN",1)
    t_res=doc.add_table(rows=10,cols=4); t_res.style='Table Grid'
    add_table_header(t_res,["Estándar","Total Criterios","Criterios Cumplidos","% Cumplimiento"])
    estandares_res=[("1. Talento Humano","8","",""),("2. Infraestructura","12","",""),
                    ("3. Dotación","11","",""),("4. Medicamentos","7","",""),
                    ("5. Procesos Prioritarios","12","",""),("6. Historia Clínica","7","",""),
                    ("7. Interdependencia","4","",""),("8. Procedimientos Estéticos","9","",""),
                    ("TOTAL","70","","")]
    for i,row in enumerate(estandares_res):
        for j,val in enumerate(row):
            t_res.rows[i+1].cells[j].text=val
            for para in t_res.rows[i+1].cells[j].paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
        if i==len(estandares_res)-1:
            for j in range(4): set_cell_bg(t_res.rows[i+1].cells[j],'003366')
            for cell in t_res.rows[i+1].cells:
                for para in cell.paragraphs:
                    for run in para.runs: run.bold=True; run.font.color.rgb=RGBColor(255,255,255)

    path=os.path.join(OUT_DIR_9,"CHK-HAB-001_Lista_Verificacion_Estandares_Resolucion_3100.docx")
    doc.save(path); print(f"Guardado: {path}")

# =====================================================================
# DOCUMENTO 8: FOR-HAB-001 Acta de Autoevaluación
# =====================================================================
def gen_acta_autoevaluacion():
    titulo="ACTA DE VISITA DE AUTOEVALUACIÓN PARA HABILITACIÓN"
    codigo="FOR-HAB-001"
    doc=setup_doc(titulo,codigo)
    add_portada(doc,titulo,codigo); add_cv(doc)

    add_heading(doc,"I. DATOS DEL PRESTADOR",1)
    t1=doc.add_table(rows=10,cols=2); t1.style='Table Grid'
    campos_prestador=[("Nombre del Consultorio/Prestador:","[NOMBRE DEL CONSULTORIO]"),
                      ("NIT o Cédula:","[NIT/CÉDULA]"),
                      ("Representante Legal / Director Médico:","[NOMBRE DE LA MÉDICA]"),
                      ("Tarjeta Profesional:","[NÚMERO DE TARJETA PROFESIONAL]"),
                      ("Dirección:","[DIRECCIÓN DEL CONSULTORIO]"),
                      ("Municipio/Departamento:","[MUNICIPIO], [DEPARTAMENTO]"),
                      ("Teléfono:","[TELÉFONO]"),
                      ("Correo electrónico:","[CORREO ELECTRÓNICO]"),
                      ("Fecha de la autoevaluación:","[DD/MM/AAAA]"),
                      ("Evaluado por:","[NOMBRE DE LA MÉDICA] - Directora del Consultorio")]
    for i,(k,v) in enumerate(campos_prestador):
        t1.rows[i].cells[0].text=k; t1.rows[i].cells[1].text=v
        set_cell_bg(t1.rows[i].cells[0],'BDD7EE')
        for cell in t1.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(11)
    doc.add_paragraph()

    add_heading(doc,"II. SERVICIOS HABILITADOS / A HABILITAR",1)
    t2=doc.add_table(rows=4,cols=5); t2.style='Table Grid'
    add_table_header(t2,["Código Servicio","Nombre del Servicio","Complejidad","Sede","Estado"])
    servicios=[("191","Consulta externa de medicina general","Baja","[DIRECCIÓN DEL CONSULTORIO]","A habilitar"),
               ("Procedimientos","Procedimientos estéticos no invasivos","Baja","[DIRECCIÓN DEL CONSULTORIO]","A habilitar"),
               ("","[Otros servicios si aplica]","","","")]
    for i,row in enumerate(servicios):
        for j,val in enumerate(row):
            t2.rows[i+1].cells[j].text=val
            for para in t2.rows[i+1].cells[j].paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc,"III. TABLA DE HALLAZGOS POR ESTÁNDAR",1)
    add_body(doc,"La siguiente tabla resume los hallazgos identificados durante la autoevaluación. Los ítems marcados como 'No Cumple' deben incluirse en el plan de acción inmediata.")

    t3=doc.add_table(rows=20,cols=6); t3.style='Table Grid'
    add_table_header(t3,["N°","Estándar","Criterio","Hallazgo / Observación","Estado\n(C/NC/EP)","Urgencia\n(Alta/Med/Baja)"])
    hallazgos_ejemplo=[
        ("1","Talento Humano","Tarjeta profesional","Tarjeta profesional vigente","C",""),
        ("2","Talento Humano","Vacunación personal","Verificar carnet de hepatitis B del personal auxiliar","EP","Alta"),
        ("3","Infraestructura","Señalización","Falta señal de evacuación en el pasillo","NC","Alta"),
        ("4","Infraestructura","Accesibilidad","No hay rampa de acceso para sillas de ruedas","NC","Media"),
        ("5","Dotación","Calibración de equipos","Tensiómetro sin certificado de calibración vigente","NC","Alta"),
        ("6","Dotación","Hoja de vida equipos","Algunos equipos sin hoja de vida actualizada","NC","Media"),
        ("7","Medicamentos","Cadena de frío","Registro de temperaturas del refrigerador no está al día","NC","Alta"),
        ("8","Procesos Prioritarios","PGIRHS","PGIRHS existe pero no está completamente implementado","EP","Alta"),
        ("9","Procesos Prioritarios","Consentimientos","Faltan consentimientos para procedimientos estéticos","NC","Alta"),
        ("10","Historia Clínica","Archivo HC","Archivo de HC sin mecanismo de restricción de acceso","NC","Alta"),
        ("11","Interdependencia","Directorio de referencia","Directorio de IPS de referencia no está actualizado","NC","Media"),
        ("12","Proc. Estéticos","Registro INVIMA productos","Verificar registro INVIMA de productos utilizados","EP","Alta"),
        ("","","","","",""),
        ("","","","","",""),
        ("","","","","",""),
        ("","","","","",""),
        ("","","","","",""),
        ("","","","","",""),
        ("","","","","",""),
    ]
    for i,row in enumerate(hallazgos_ejemplo):
        for j,val in enumerate(row):
            t3.rows[i+1].cells[j].text=val
            for para in t3.rows[i+1].cells[j].paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(9)
        if row[4]=="NC":
            set_cell_bg(t3.rows[i+1].cells[4],'FF0000')
            for para in t3.rows[i+1].cells[4].paragraphs:
                for run in para.runs: run.bold=True; run.font.color.rgb=RGBColor(255,255,255)
        elif row[4]=="EP":
            set_cell_bg(t3.rows[i+1].cells[4],'FFC000')
        elif row[4]=="C":
            set_cell_bg(t3.rows[i+1].cells[4],'70AD47')
    doc.add_paragraph()

    add_heading(doc,"IV. RESUMEN DE CUMPLIMIENTO",1)
    t4=doc.add_table(rows=9,cols=4); t4.style='Table Grid'
    add_table_header(t4,["Estándar","Total Criterios","Cumple","No Cumple / En Proceso"])
    estandares_res=[("1. Talento Humano","8","",""),("2. Infraestructura","12","",""),
                    ("3. Dotación","11","",""),("4. Medicamentos","7","",""),
                    ("5. Procesos Prioritarios","12","",""),("6. Historia Clínica","7","",""),
                    ("7. Interdependencia","4","",""),("TOTAL","61","","")]
    for i,row in enumerate(estandares_res):
        for j,val in enumerate(row):
            t4.rows[i+1].cells[j].text=val
            for para in t4.rows[i+1].cells[j].paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(10)
        if i==len(estandares_res)-1:
            for j in range(4): set_cell_bg(t4.rows[i+1].cells[j],'BDD7EE')
            for cell in t4.rows[i+1].cells:
                for para in cell.paragraphs:
                    for run in para.runs: run.bold=True
    doc.add_paragraph()

    add_heading(doc,"V. PLAN DE ACCIÓN INMEDIATA",1)
    add_body(doc,"Los hallazgos de No Cumple con urgencia Alta deben resolverse antes de solicitar la visita de habilitación a la Secretaría de Salud:")
    t5=doc.add_table(rows=10,cols=6); t5.style='Table Grid'
    add_table_header(t5,["N°","Hallazgo","Acción Correctiva","Responsable","Fecha\nCompr.","Recursos\nNecesarios"])
    acciones=[
        ("1","Señalización de evacuación","Adquirir e instalar señales fotoluminiscentes","[NOMBRE DE LA MÉDICA]","","Señales NTC 1867"),
        ("2","Calibración de tensiómetro","Llevar a calibrar con empresa certificada","[NOMBRE DE LA MÉDICA]","","Empresa de metrología"),
        ("3","Registro de temperaturas","Implementar formato diario de registro de temperatura","[NOMBRE DE LA MÉDICA]","","Formato FOR-MED-002"),
        ("4","Consentimientos estéticos","Elaborar y aprobar CI para cada procedimiento","[NOMBRE DE LA MÉDICA]","","Asesoría jurídica si necesario"),
        ("5","Archivo HC con acceso restringido","Instalar cerradura al archivador de HC","[NOMBRE DE LA MÉDICA]","","Material de ferretería"),
        ("6","Directorio de referencia","Actualizar directorio con IPS de la zona","[NOMBRE DE LA MÉDICA]","","Directorio de IPS actualizadas"),
        ("7","Registro INVIMA productos","Verificar en www.invima.gov.co cada producto","[NOMBRE DE LA MÉDICA]","","Internet, portal INVIMA"),
    ]
    for i,row in enumerate(acciones):
        for j,val in enumerate(row):
            t5.rows[i+1].cells[j].text=val
            for para in t5.rows[i+1].cells[j].paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(9)
        if i%2==0:
            for j in range(6): set_cell_bg(t5.rows[i+1].cells[j],'EBF3FB')
    for i in range(len(acciones)+1, 10):
        for j in range(6):
            t5.rows[i].cells[j].text=""
            for para in t5.rows[i].cells[j].paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(9)
    doc.add_paragraph()

    add_heading(doc,"VI. COMPROMISOS Y FECHAS",1)
    add_body(doc,"El consultorio [NOMBRE DEL CONSULTORIO] se compromete a implementar todas las acciones correctivas identificadas en el plan de acción inmediata antes del [FECHA PROPUESTA] con el fin de garantizar el cumplimiento de todos los estándares de la Resolución 3100 de 2019 previo a la solicitud de visita de habilitación ante la Secretaría de Salud.")
    doc.add_paragraph()
    add_body(doc,"Compromisos específicos:")
    compromisos=["Subsanar todos los hallazgos de urgencia Alta en un plazo máximo de 30 días.",
                 "Subsanar todos los hallazgos de urgencia Media en un plazo máximo de 60 días.",
                 "Realizar una nueva autoevaluación después de implementadas las acciones para verificar el cumplimiento.",
                 "Solicitar la visita de habilitación solo cuando el porcentaje de cumplimiento sea ≥ 95%."]
    for c in compromisos: bullet(doc,c)
    doc.add_paragraph()

    add_heading(doc,"VII. FIRMAS",1)
    doc.add_paragraph()
    doc.add_paragraph()
    t6=doc.add_table(rows=4,cols=3); t6.style='Table Grid'
    add_table_header(t6,["Rol","Nombre","Firma / Fecha"])
    firmas=[("Directora Médica / Responsable de la Habilitación","[NOMBRE DE LA MÉDICA]","Firma: _____________ Fecha: _________"),
            ("Testigo (si aplica)","[NOMBRE DEL TESTIGO]","Firma: _____________ Fecha: _________"),
            ("Revisor externo (si aplica)","[NOMBRE DEL REVISOR]","Firma: _____________ Fecha: _________")]
    for i,row in enumerate(firmas):
        for j,val in enumerate(row):
            t6.rows[i+1].cells[j].text=val
            for para in t6.rows[i+1].cells[j].paragraphs:
                for run in para.runs: run.font.name='Calibri'; run.font.size=Pt(11)
    doc.add_paragraph()
    add_body(doc,"Nota: Esta acta de autoevaluación es un documento interno del consultorio y no reemplaza la visita oficial de habilitación realizada por la Secretaría de Salud. Su propósito es preparar al prestador para el proceso oficial de habilitación.")

    path=os.path.join(OUT_DIR_9,"FOR-HAB-001_Acta_Visita_Autoevaluacion.docx")
    doc.save(path); print(f"Guardado: {path}")

if __name__=="__main__":
    gen_pqr()
    gen_manual_habilitacion()
    gen_lista_verificacion()
    gen_acta_autoevaluacion()
    print("Documentos 5, 6, 7 y 8 completados.")
