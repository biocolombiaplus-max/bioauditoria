#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

BASE3="/home/user/bioauditoria/documentos_habilitacion/CARPETA_3_DOTACION"

def margins(doc):
    for s in doc.sections:
        s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5); s.left_margin=Cm(3); s.right_margin=Cm(2.5)

def hf(doc,code,title):
    for s in doc.sections:
        h=s.header; h.paragraphs[0].clear()
        r=h.paragraphs[0].add_run(f"CONSULTORIO [NOMBRE DEL CONSULTORIO]  |  {code}  |  {title}")
        r.font.size=Pt(8); r.font.color.rgb=RGBColor(64,64,64)
        h.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        f=s.footer; f.paragraphs[0].clear()
        r2=f.paragraphs[0].add_run(f"{code} - Versión 1.0 - {datetime.date.today().strftime('%d/%m/%Y')}  | Pág. ")
        r2.font.size=Pt(8)
        fld=OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'),'begin')
        f.paragraphs[0].runs[-1]._r.append(fld)
        it=OxmlElement('w:instrText'); it.text='PAGE'
        f.paragraphs[0].runs[-1]._r.append(it)
        fld2=OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'),'end')
        f.paragraphs[0].runs[-1]._r.append(fld2)
        f.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

def portada(doc,code,title):
    doc.add_paragraph(); doc.add_paragraph()
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=t.add_run("CONSULTORIO MÉDICO\n[NOMBRE DEL CONSULTORIO]")
    r.font.size=Pt(18); r.font.bold=True; r.font.color.rgb=RGBColor(31,73,125)
    doc.add_paragraph()
    t2=doc.add_paragraph(); t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=t2.add_run(title); r2.font.size=Pt(15); r2.font.bold=True
    doc.add_paragraph()
    tb=doc.add_table(rows=6,cols=2); tb.style='Table Grid'; tb.alignment=WD_TABLE_ALIGNMENT.CENTER
    datos=[("Código:",code),("Versión:","1.0"),("Fecha:",datetime.date.today().strftime('%d/%m/%Y')),
           ("Elaboró:","[NOMBRE DE LA MÉDICA], Médica General"),
           ("Revisó:","[NOMBRE DE LA MÉDICA], Propietaria"),
           ("Aprobó:","[NOMBRE DE LA MÉDICA], Propietaria")]
    for i,(l,v) in enumerate(datos):
        tb.rows[i].cells[0].text=l; tb.rows[i].cells[1].text=v
        tb.rows[i].cells[0].paragraphs[0].runs[0].font.bold=True
    doc.add_page_break()
    cv=doc.add_paragraph(); r3=cv.add_run("CONTROL DE VERSIONES")
    r3.font.bold=True; r3.font.size=Pt(13); r3.font.color.rgb=RGBColor(31,73,125)
    cv.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tb2=doc.add_table(rows=3,cols=5); tb2.style='Table Grid'
    hrow(tb2,["Versión","Fecha","Cambio","Elaboró","Aprobó"])
    tb2.rows[1].cells[0].text="1.0"; tb2.rows[1].cells[1].text=datetime.date.today().strftime('%d/%m/%Y')
    tb2.rows[1].cells[2].text="Versión inicial"; tb2.rows[1].cells[3].text="[NOMBRE DE LA MÉDICA]"
    tb2.rows[1].cells[4].text="[NOMBRE DE LA MÉDICA]"
    doc.add_page_break()

def hrow(table,headers,bg="1F497D"):
    row=table.rows[0]
    for i,h in enumerate(headers):
        cell=row.cells[i]; cell.text=""
        run=cell.paragraphs[0].add_run(h)
        run.font.bold=True; run.font.color.rgb=RGBColor(255,255,255); run.font.size=Pt(10)
        cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),bg)
        shd.set(qn('w:color'),'auto'); shd.set(qn('w:val'),'clear'); tcPr.append(shd)

def H(doc,text,lv=1):
    p=doc.add_heading(text,level=lv)
    colors={1:RGBColor(31,73,125),2:RGBColor(46,116,181),3:RGBColor(68,114,196)}
    sizes={1:14,2:12,3:11}
    for run in p.runs:
        run.font.size=Pt(sizes[lv]); run.font.bold=True; run.font.color.rgb=colors[lv]
    return p

def B(doc,text,bold=False):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=p.add_run(text); r.font.size=Pt(11); r.font.bold=bold; r.font.name='Calibri'
    return p

def BL(doc,text):
    p=doc.add_paragraph(style='List Bullet'); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=p.add_run(text); r.font.size=Pt(11); r.font.name='Calibri'; return p

def firmas(doc):
    doc.add_paragraph()
    tb=doc.add_table(rows=2,cols=3); tb.style='Table Grid'
    hrow(tb,["ELABORÓ","REVISÓ","APROBÓ"])
    for j in range(3):
        tb.rows[1].cells[j].text="[NOMBRE DE LA MÉDICA]\nMédica General - Propietaria\n\nFirma: _________________________\n\nFecha: "+datetime.date.today().strftime('%d/%m/%Y')
        tb.rows[1].cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

# =========== MAN-DOT-001 ===========
def crear_man_dot():
    doc=Document(); margins(doc); hf(doc,"MAN-DOT-001","Manual de Gestión de Equipos Biomédicos")
    portada(doc,"MAN-DOT-001","MANUAL DE GESTIÓN DE EQUIPOS BIOMÉDICOS Y DOTACIÓN")

    H(doc,"1. OBJETIVO")
    B(doc,"Establecer los lineamientos para la gestión integral de los equipos biomédicos, dispositivos médicos e instrumental clínico del Consultorio Médico [NOMBRE DEL CONSULTORIO], incluyendo su adquisición, inventario, mantenimiento, calibración, tecnovigilancia y disposición final, garantizando el cumplimiento de la Resolución 3100 de 2019 y el Decreto 4725 de 2005, y asegurando que todos los equipos y dispositivos utilizados en la atención de los pacientes sean seguros, funcionales y con registro INVIMA vigente.")

    H(doc,"2. ALCANCE")
    B(doc,"Aplica a todos los equipos médicos, dispositivos médicos, instrumental clínico y elementos de dotación del Consultorio [NOMBRE DEL CONSULTORIO], incluyendo los equipos utilizados en la consulta de medicina general y los equipos y materiales utilizados en los procedimientos estéticos no invasivos (toxina botulínica tipo A, ácido hialurónico, materiales para peelings, mesoterapia y PRP).")

    H(doc,"3. MARCO LEGAL")
    normas=[
        ("Resolución 3100 de 2019","Estándares de dotación (equipos e instrumental) para cada servicio de salud habilitado."),
        ("Decreto 4725 de 2005","Régimen de registros sanitarios, permiso de comercialización y vigilancia sanitaria de dispositivos médicos. Define los requisitos de registro INVIMA para equipos y dispositivos médicos."),
        ("Resolución 4816 de 2008","Programa Nacional de Tecnovigilancia. Establece la obligatoriedad de reportar eventos adversos e incidentes relacionados con dispositivos médicos al INVIMA."),
        ("Decreto 780 de 2016","Decreto Único del Sector Salud. Consolida regulación sobre dispositivos médicos y tecnologías en salud."),
        ("Resolución 2003 de 2014","Estándares de habilitación anteriores a 3100/2019, aún útiles como referencia."),
        ("ISO 13485","Estándar internacional de sistemas de gestión de calidad para dispositivos médicos (referencia)."),
        ("NTC-ISO 9001","Sistema de gestión de calidad (referencia para gestión de equipos)."),
    ]
    tb=doc.add_table(rows=len(normas)+1,cols=2); tb.style='Table Grid'
    hrow(tb,["Norma","Descripción"])
    for i,(n,d) in enumerate(normas):
        tb.rows[i+1].cells[0].text=n; tb.rows[i+1].cells[1].text=d
        for c in tb.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()

    H(doc,"4. DEFINICIONES")
    defs=[("Equipo biomédico:","Dispositivo médico que requiere calibración, mantenimiento, reparación, capacitación del usuario y gestión de obsolescencia. Incluye tensiómetros, glucómetros, electrocardiógrafos, etc."),
          ("Dispositivo médico:","Cualquier instrumento, aparato, implemento, máquina, software u otro artículo destinado a ser usado en seres humanos con fines de diagnóstico, prevención, monitoreo o tratamiento."),
          ("Registro sanitario INVIMA:","Autorización expedida por el INVIMA que permite la fabricación, importación, procesamiento, envase, empaque, transporte, almacenamiento, expendio y comercialización de un dispositivo médico en Colombia."),
          ("Tecnovigilancia:","Conjunto de actividades que tienen por objeto la identificación y la evaluación de los efectos adversos serios e incidentes que puedan ser atribuidos a dispositivos médicos."),
          ("Calibración:","Conjunto de operaciones que establecen la relación entre los valores indicados por un instrumento de medición y los valores conocidos o convenidos de un patrón de referencia."),
          ("Mantenimiento preventivo:","Actividades programadas para garantizar el funcionamiento óptimo del equipo y prevenir fallas."),
          ("Mantenimiento correctivo:","Reparación de equipos que han presentado falla o deterioro en su funcionamiento."),
          ("Vida útil:","Período durante el cual un equipo puede operar de manera segura y con la precisión requerida."),]
    for term,defi in defs:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r1=p.add_run(term+" "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(defi); r2.font.size=Pt(11)

    H(doc,"5. INVENTARIO DE EQUIPOS REQUERIDOS PARA MEDICINA GENERAL")
    H(doc,"5.1 Equipos Básicos de Consulta Médica",2)
    B(doc,"Según la Resolución 3100 de 2019, el servicio de Consulta de Medicina General debe contar con los siguientes equipos e instrumental:")
    equipos_gen=[
        ("Tensiómetro (esfigmomanómetro)","1 unidad mínimo","Aneroide o digital, con varios tamaños de brazalete","Cada 6 meses o según fabricante","[MARCA/MODELO]\nS/N: [NÚMERO SERIAL]\nRegistro INVIMA: [NÚMERO]"),
        ("Fonendoscopio (estetoscopio)","1 unidad","Biauricular, con campana y membrana","No requiere calibración, inspección mensual","[MARCA/MODELO]"),
        ("Termómetro","1 unidad","Clínico, digital o de mercurio (preferir digital)","Anual","[MARCA/MODELO]"),
        ("Oxímetro de pulso","1 unidad","Pulsioxímetro digital, con alarmas","Anual o según fabricante","[MARCA/MODELO]\nS/N: [NÚMERO SERIAL]"),
        ("Glucómetro + tiras reactivas","1 unidad","Con certificación vigente, tiras no vencidas","Semestral con control de calidad","[MARCA/MODELO]\nLote tiras: [LOTE]"),
        ("Linterna de exploración","1 unidad","Con batería cargada o recargable","Mensual (verificar funcionamiento)","[MARCA/MODELO]"),
        ("Martillo de reflejos","1 unidad","De goma, mango largo","Inspección mensual","[MARCA/MODELO]"),
        ("Otoscopio","1 unidad","Con espéculos desechables de varios tamaños","Mensual (batería, luz)","[MARCA/MODELO]"),
        ("Oftalmoscopio","1 unidad o combinado con otoscopio","Con fuente de luz adecuada","Mensual","[MARCA/MODELO]"),
        ("Cinta métrica","1 unidad","Flexible, lavable","Inspección mensual","[MARCA/MODELO]"),
        ("Tallímetro","1 unidad","Con precisión de 1 mm","Anual","[MARCA/MODELO]"),
        ("Báscula","1 unidad","Con tallímetro integrado (deseable), precisión 100g","Semestral o anual (calibración)","[MARCA/MODELO]\nS/N: [NÚMERO SERIAL]"),
        ("Camilla de examen","1 unidad","Con cobertura lavable, regulable, freno","Inspección mensual","[MARCA/MODELO]"),
        ("Negatoscopio","1 unidad (si se usa radiología análoga) o monitor si es digital","Funcionamiento adecuado de la luz","Semestral","[MARCA/MODELO]"),
        ("Equipo de urgencias básico","1 set","Incluye: ambú, cánulas Guedel (varios tamaños), mascarillas, O2 portátil","Mensual (verificar completo y en fecha)","Ver kit completo en PRO-PP-003"),
    ]
    tb2=doc.add_table(rows=len(equipos_gen)+1,cols=5); tb2.style='Table Grid'
    hrow(tb2,["Equipo","Cantidad","Especificaciones","Mantenimiento/Calibración","Datos del equipo actual"])
    for i,row in enumerate(equipos_gen):
        for j,v in enumerate(row):
            tb2.rows[i+1].cells[j].text=v
            for c in tb2.rows[i+1].cells[j].paragraphs:
                for r in c.runs: r.font.size=Pt(8)
    doc.add_paragraph()

    H(doc,"5.2 Equipos e Insumos para Procedimientos Estéticos No Invasivos",2)
    B(doc,"AVISO IMPORTANTE: Los siguientes equipos y materiales son para procedimientos que constituyen ACTOS MÉDICOS de competencia exclusiva de la médica general. Los productos inyectables (toxina botulínica, ácido hialurónico, cócteles de mesoterapia) deben contar con registro sanitario INVIMA vigente como medicamento o dispositivo médico según corresponda.")

    equipos_est=[
        ("Refrigerador médico o nevera con control de temperatura","1 unidad","Para almacenamiento de toxina botulínica a 2-8°C y otros biológicos. Con termómetro de control.","Mensual (temperatura), Anual (técnico)","[MARCA/MODELO]\nS/N: [N°]\nRango: 2-8°C"),
        ("Centrífuga (para PRP)","1 unidad","Centrífuga de laboratorio con protocolo para PRP: mínimo 1500-3000 RPM, con temporizador","Semestral (calibración RPM)","[MARCA/MODELO]\nS/N: [N°]"),
        ("Lámpara de luz fría / magnificación","1 unidad","Para examen detallado de piel. Lupa o luz de Wood para valoración dermatológica","Mensual (funcionamiento)","[MARCA/MODELO]"),
        ("Cámara fotográfica clínica","1 unidad","Para documentación fotográfica pre y post procedimiento. Puede ser smartphone con calidad suficiente + fondo neutro","N/A","[MARCA/MODELO]"),
        ("Set de agujas para inyección","Stock permanente","Agujas 30G y 32G para toxina; 27G y 25G para rellenos; agujas de mesoterapia 4mm y 13mm","Verificación mensual de stock y fechas de vencimiento","[PROVEEDOR]"),
        ("Cánulas romas para rellenos","Stock permanente","Cánulas 25G, 23G de diferentes longitudes para aplicación de ácido hialurónico","Verificación mensual","[PROVEEDOR]"),
        ("Jeringas de insulina 1ml","Stock permanente","Para dilución y aplicación de toxina botulínica","Verificación mensual","[PROVEEDOR]"),
        ("Kit de emergencias estéticas","1 set completo","Incluye OBLIGATORIAMENTE: hialuronidasa (para emergencia por oclusión vascular con HA), solución salina, adrenalina, corticoide inyectable, antihistamínico","Mensual (verificar completo, vencimientos)","Ver detalle en PRO-PP-002"),
        ("Sistema de drenaje o aspiración (para PRP)","1 set","Tubos vacutainer con citrato, tubos para PRP, agujas de extracción","Verificación mensual de stock","[PROVEEDOR]"),
        ("Equipo de crioterapia / cold packs","1 unidad","Para anestesia local previa a procedimientos. Bolsas de frío o dispositivo de crioanestesia","Mensual","[MARCA/MODELO]"),
    ]
    tb3=doc.add_table(rows=len(equipos_est)+1,cols=5); tb3.style='Table Grid'
    hrow(tb3,["Equipo/Insumo","Cantidad","Especificaciones","Mantenimiento","Datos actuales"])
    for i,row in enumerate(equipos_est):
        for j,v in enumerate(row):
            tb3.rows[i+1].cells[j].text=v
            for c in tb3.rows[i+1].cells[j].paragraphs:
                for r in c.runs: r.font.size=Pt(8)
    doc.add_paragraph()

    H(doc,"6. PROGRAMA DE MANTENIMIENTO DE EQUIPOS BIOMÉDICOS")
    H(doc,"6.1 Mantenimiento Preventivo",2)
    B(doc,"El mantenimiento preventivo de los equipos biomédicos se realiza según el siguiente cronograma:")
    mant_data=[
        ("Tensiómetro","Verificación de calibración con patrón de referencia. Inspección de manguito, válvulas, tubería","Semestral","Técnico biomédico certificado","Empresa de mantenimiento biomédico"),
        ("Báscula","Verificación con pesas patrón certificadas. Ajuste si hay desviación","Anual","Técnico biomédico","ONAC-acreditado si posible"),
        ("Oxímetro de pulso","Verificación con simulador de SPO2. Limpieza del sensor","Anual","Técnico biomédico","Servicio técnico autorizado"),
        ("Glucómetro","Control de calidad con solución control del fabricante","Mensual / con cada lote de tiras","Médica","Interno"),
        ("Refrigerador médico","Verificación de temperatura (monitoreo diario), limpieza, revisión de sellos y compresores","Mensual monitoreo / Semestral técnico","Técnico refrigeración","Servicio técnico autorizado"),
        ("Centrífuga (PRP)","Verificación de RPM con tacómetro, limpieza de rotor, revisión de escobillas","Semestral","Técnico biomédico","Empresa especializada"),
        ("Oxígeno portátil","Verificación de presión del cilindro, estado del regulador y mascarillas","Mensual","Médica / proveedor O2","Empresa proveedora de gases medicinales"),
        ("Equipo de urgencias","Verificación de completo del kit, fechas de vencimiento de medicamentos, funcionamiento del ambú","Mensual","Médica","Interno"),
    ]
    tb4=doc.add_table(rows=len(mant_data)+1,cols=5); tb4.style='Table Grid'
    hrow(tb4,["Equipo","Actividad","Frecuencia","Ejecuta","Proveedor"])
    for i,row in enumerate(mant_data):
        for j,v in enumerate(row):
            tb4.rows[i+1].cells[j].text=v
            for c in tb4.rows[i+1].cells[j].paragraphs:
                for r in c.runs: r.font.size=Pt(8)
    doc.add_paragraph()

    H(doc,"7. TECNOVIGILANCIA")
    B(doc,"En cumplimiento de la Resolución 4816 de 2008, el consultorio implementa el programa de tecnovigilancia para el reporte de eventos adversos e incidentes relacionados con los dispositivos médicos utilizados en la atención.")
    H(doc,"7.1 ¿Qué debe reportarse?",2)
    reportar=[
        "Cualquier mal funcionamiento, deterioro o inexactitud del etiquetado o instrucciones de uso de un dispositivo médico que haya causado o pudiera causar la muerte o deterioro grave del estado de salud del paciente",
        "Razones técnicas que, de volver a producirse, pudieran ocasionar la muerte o el deterioro grave del estado de salud del paciente",
        "Problemas de calidad, efectos adversos o reacciones inesperadas derivadas del uso de los dispositivos médicos con registro INVIMA",
        "Fallas de equipos durante procedimientos que puedan afectar al paciente",
        "Reacciones adversas al uso de dispositivos médicos para estética: rellenos con HA, toxina botulínica (bulto en el lugar de inyección diferente a lo esperado, migración, etc.)",
    ]
    for item in reportar: BL(doc,item)
    H(doc,"7.2 ¿Cómo reportar?",2)
    B(doc,"El reporte de tecnovigilancia se realiza a través del portal del INVIMA (www.invima.gov.co) en el módulo de Reportes de Tecnovigilancia. El reporte debe incluir: identificación del producto (nombre, lote, registro INVIMA), descripción del evento, paciente afectado (datos de identificación reservados), consecuencias y manejo del evento.")

    H(doc,"8. GESTIÓN DE DISPOSITIVOS MÉDICOS PARA PROCEDIMIENTOS ESTÉTICOS")
    H(doc,"8.1 Requisitos Regulatorios INVIMA",2)
    B(doc,"Todos los productos utilizados en los procedimientos estéticos del consultorio deben contar con registro sanitario INVIMA vigente. Se establece la siguiente verificación:")
    productos_est=[
        ("Toxina Botulínica Tipo A","Medicamento biológico","Registro como medicamento (no como cosmético). Cadena de frío 2-8°C. Solo distribuido por importadores autorizados por el INVIMA.","[MARCA COMERCIAL]\nRegistro INVIMA: [NÚMERO]\nProveedor: [NOMBRE]"),
        ("Ácido Hialurónico (relleno dérmico)","Dispositivo médico Clase III","Registro INVIMA como dispositivo médico de clase III (alto riesgo). Verificar que no sea de uso cosmético sino médico.","[MARCA COMERCIAL]\nRegistro INVIMA: [NÚMERO]\nProveedor: [NOMBRE]"),
        ("Ácido Glicólico para peeling","Medicamento / cosmético médico","Verificar concentración y registro correspondiente. TCA debe ser de grado médico.","[MARCA COMERCIAL]\nConcentración: [%]\nProveedor: [NOMBRE]"),
        ("TCA (Ácido Tricloroacético) ≤20%","Reactivo médico","Solo en concentraciones ≤20% para peeling superficial. Adquirir en laboratorio farmacéutico o distribuidor médico certificado.","Concentración: ≤20%\nProveedor farmacéutico: [NOMBRE]"),
        ("Cócteles de mesoterapia","Medicamentos inyectables","Los principios activos utilizados deben ser medicamentos con registro INVIMA. No usar productos sin registro ni de origen incierto.","[PRINCIPIOS ACTIVOS USADOS]\nProveedor: [NOMBRE]"),
        ("Tubos para PRP","Dispositivo médico","Tubos con citrato o ACD certificados para uso médico en Colombia.","[MARCA COMERCIAL]\nRegistro INVIMA: [NÚMERO]"),
        ("Hialuronidasa (kit de emergencia)","Medicamento","Disponible en todo momento cuando se aplique ácido hialurónico. Para manejo de oclusión vascular.","[MARCA COMERCIAL]\nRegistro INVIMA: [NÚMERO]\nVencimiento: [FECHA]"),
    ]
    tb5=doc.add_table(rows=len(productos_est)+1,cols=4); tb5.style='Table Grid'
    hrow(tb5,["Producto","Clasificación regulatoria","Requisitos de verificación","Datos del producto actual"])
    for i,row in enumerate(productos_est):
        for j,v in enumerate(row):
            tb5.rows[i+1].cells[j].text=v
            for c in tb5.rows[i+1].cells[j].paragraphs:
                for r in c.runs: r.font.size=Pt(8)
    doc.add_paragraph()

    H(doc,"9. INDICADORES")
    tb6=doc.add_table(rows=5,cols=4); tb6.style='Table Grid'
    hrow(tb6,["Indicador","Fórmula","Meta","Frecuencia"])
    inds=[("% equipos con mantenimiento al día","(Equipos con mant. vigente/Total equipos)×100","100%","Semestral"),
          ("% dispositivos médicos con registro INVIMA vigente","(Productos con INVIMA vigente/Total productos)×100","100%","Trimestral"),
          ("N° eventos adversos relacionados con dispositivos médicos","Conteo absoluto reportados a INVIMA","Reportar 100%","Mensual"),
          ("% equipos calibrados según cronograma","(Equipos calibrados/Equipos que requieren calibración)×100","100%","Anual")]
    for i,row in enumerate(inds):
        for j,v in enumerate(row): tb6.rows[i+1].cells[j].text=v
        for c in tb6.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph(); firmas(doc)
    path=os.path.join(BASE3,"MAN-DOT-001_Manual_Gestion_Equipos_Biomedicos.docx")
    doc.save(path); print(f"✓ Creado: {path}")

# =========== PRO-DOT-001 ===========
def crear_pro_dot():
    doc=Document(); margins(doc); hf(doc,"PRO-DOT-001","Proceso de Mantenimiento de Equipos")
    portada(doc,"PRO-DOT-001","PROCESO DE MANTENIMIENTO DE EQUIPOS BIOMÉDICOS")

    H(doc,"1. OBJETIVO")
    B(doc,"Definir el proceso estandarizado para el mantenimiento preventivo y correctivo de los equipos biomédicos del Consultorio [NOMBRE DEL CONSULTORIO], garantizando que los equipos funcionen correctamente, sean seguros para los pacientes y el personal, y cumplan con los estándares de precisión requeridos para la atención médica de calidad.")

    H(doc,"2. ALCANCE")
    B(doc,"Aplica a todos los equipos biomédicos e instrumental clínico del consultorio: equipos de diagnóstico, equipos de monitoreo, equipos para procedimientos estéticos, equipos de urgencias y todos los elementos de dotación que requieran mantenimiento periódico.")

    H(doc,"3. DESCRIPCIÓN DEL PROCESO")
    pasos=[("1","Identificación del equipo","Se identifica el equipo que requiere mantenimiento (preventivo según cronograma o correctivo por falla detectada). Se ubica en el inventario (FOR-DOT-001)."),
           ("2","Verificación del tipo de mantenimiento","Se determina si es mantenimiento preventivo (programado según cronograma MAN-DOT-001) o correctivo (por falla o deterioro detectado)."),
           ("3","Contacto con el proveedor de mantenimiento","Se contacta al proveedor de mantenimiento biomédico certificado correspondiente al tipo de equipo."),
           ("4","Ejecución del mantenimiento","El técnico realiza el mantenimiento o calibración. Se verifica que el equipo funcione correctamente al finalizar."),
           ("5","Registro del mantenimiento","Se diligencia el registro de mantenimiento: fecha, actividad, técnico, resultados, próximo mantenimiento. Se actualiza el FOR-DOT-001."),
           ("6","Verificación de funcionamiento post-mantenimiento","La médica propietaria verifica que el equipo funcione correctamente antes de volver a usarlo en pacientes."),
           ("7","Gestión de equipos en falla","Si el equipo no puede ser reparado o está fuera de la vida útil, se retira del servicio y se gestiona su disposición final según la normatividad de RAEE."),]
    tb=doc.add_table(rows=len(pasos)+1,cols=3); tb.style='Table Grid'
    hrow(tb,["Paso","Actividad","Descripción"])
    for i,(n,a,d) in enumerate(pasos):
        tb.rows[i+1].cells[0].text=n; tb.rows[i+1].cells[1].text=a; tb.rows[i+1].cells[2].text=d
        for c in tb.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"4. PROCEDIMIENTO PARA BAJA DE EQUIPOS")
    B(doc,"Cuando un equipo biomédico llega al final de su vida útil, presenta fallas irreparables o es sustituido por uno más moderno, se procede así:")
    baja=["Identificar el equipo a dar de baja, con sus datos de inventario",
          "Solicitar concepto técnico escrito del proveedor de mantenimiento que justifique la baja",
          "Retirar el equipo del inventario activo de equipos del consultorio",
          "Gestionar la disposición final según corresponda: residuos de aparatos eléctricos y electrónicos (RAEE), donación, venta a empresa de mantenimiento, etc.",
          "Registrar la baja en el formato de inventario FOR-DOT-001 con fecha y justificación",
          "Si el equipo es reemplazado, registrar el nuevo equipo en el inventario con todos sus datos"]
    for item in baja: BL(doc,item)

    H(doc,"5. CONTROL DE CALIDAD DE GLUCÓMETRO")
    B(doc,"El glucómetro requiere un procedimiento especial de control de calidad que se realiza internamente:")
    gluco=["Realizar el control de calidad con la solución control del fabricante según las instrucciones: nivel bajo y nivel alto",
           "El resultado debe estar dentro del rango esperado indicado en el inserto de las tiras reactivas",
           "Si el resultado está fuera del rango: verificar la solución control (no vencida), verificar las tiras (no vencidas, almacenadas correctamente), limpiar el sensor del glucómetro",
           "Si persiste el problema: contactar el servicio técnico del fabricante",
           "Registrar los controles de calidad en el libro de registro del glucómetro (fecha, lote de tiras, resultado del control, nombre de quien realiza)"]
    for item in gluco: BL(doc,item)

    H(doc,"6. TEMPERATURA DEL REFRIGERADOR MÉDICO")
    B(doc,"El refrigerador para almacenamiento de toxina botulínica requiere monitoreo diario de temperatura:")
    temp=["Verificar y registrar la temperatura del refrigerador TODOS los días hábiles, dos veces al día (mañana y tarde)",
          "Rango aceptable: 2-8°C. Temperatura objetivo: 4-5°C",
          "Si la temperatura está fuera de rango: verificar puerta bien cerrada, volumen de carga, funcionamiento del compresor",
          "Si la temperatura supera 8°C por más de 4 horas: contactar al proveedor de mantenimiento de refrigeración",
          "Si la toxina botulínica fue expuesta a temperaturas fuera del rango recomendado por el fabricante, consultar el inserto antes de usar y contactar al distribuidor",
          "Registrar cualquier incidencia de temperatura en el libro de control de temperatura"]
    for item in temp: BL(doc,item)

    H(doc,"7. INDICADORES")
    tb2=doc.add_table(rows=4,cols=4); tb2.style='Table Grid'
    hrow(tb2,["Indicador","Fórmula","Meta","Frecuencia"])
    for i,row in enumerate([
        ("% cumplimiento mantenimientos programados","Realizados/Programados×100","≥90%","Semestral"),
        ("Tiempo promedio resolución mantenimiento correctivo","Días hábiles hasta resolución (promedio)","≤5 días","Mensual"),
        ("% equipos con controles de calidad al día","Equipos con CQ/Total que requieren CQ×100","100%","Mensual")]):
        for j,v in enumerate(row): tb2.rows[i+1].cells[j].text=v
        for c in tb2.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph(); firmas(doc)
    path=os.path.join(BASE3,"PRO-DOT-001_Proceso_Mantenimiento_Equipos.docx")
    doc.save(path); print(f"✓ Creado: {path}")

# =========== FOR-DOT-001 ===========
def crear_for_dot():
    doc=Document(); margins(doc); hf(doc,"FOR-DOT-001","Inventario de Equipos y Dispositivos")
    portada(doc,"FOR-DOT-001","INVENTARIO DE EQUIPOS Y DISPOSITIVOS MÉDICOS")

    H(doc,"SECCIÓN 1: DATOS GENERALES DEL INVENTARIO")
    tb=doc.add_table(rows=3,cols=4); tb.style='Table Grid'
    datos=[("Fecha de actualización:","","Responsable:",""),
           ("Consultorio:","[NOMBRE DEL CONSULTORIO]","Ciudad:","[CIUDAD]"),
           ("Versión del formato:","1.0","N° de equipos registrados:","")]
    for i,row in enumerate(datos):
        for j,v in enumerate(row):
            tb.rows[i].cells[j].text=v
            if j%2==0 and v: tb.rows[i].cells[j].paragraphs[0].runs[0].font.bold=True
    doc.add_paragraph()

    H(doc,"SECCIÓN 2: INVENTARIO DE EQUIPOS BIOMÉDICOS - MEDICINA GENERAL",2)
    cols_inv=["N°","Nombre del equipo","Marca/Modelo","N° Serial","Registro INVIMA","Fecha adquisición","Fecha último mant.","Fecha próx. mant.","Estado","Observaciones"]
    equipos_lista=["Tensiómetro","Fonendoscopio","Termómetro digital","Oxímetro de pulso","Glucómetro","Linterna de exploración","Martillo de reflejos","Otoscopio","Oftalmoscopio/otoscopio combinado","Cinta métrica","Tallímetro","Báscula","Camilla de examen","Negatoscopio","Ambú (adulto)","Cánulas de Guedel (set)","Oxígeno portátil + regulador","Desfibrilador DEA (si disponible)"]
    tb2=doc.add_table(rows=len(equipos_lista)+1,cols=len(cols_inv)); tb2.style='Table Grid'
    hrow(tb2,cols_inv)
    for i,eq in enumerate(equipos_lista):
        tb2.rows[i+1].cells[0].text=str(i+1)
        tb2.rows[i+1].cells[1].text=eq
        for c in tb2.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(8)
    doc.add_paragraph()

    H(doc,"SECCIÓN 3: INVENTARIO DE EQUIPOS PARA PROCEDIMIENTOS ESTÉTICOS",2)
    equipos_est=["Refrigerador médico (toxina botulínica)","Termómetro de refrigerador","Centrífuga (PRP)","Lámpara de magnificación","Cámara fotográfica clínica","Cold pack / crioterapia"]
    tb3=doc.add_table(rows=len(equipos_est)+1,cols=len(cols_inv)); tb3.style='Table Grid'
    hrow(tb3,cols_inv)
    for i,eq in enumerate(equipos_est):
        tb3.rows[i+1].cells[0].text=str(i+1); tb3.rows[i+1].cells[1].text=eq
        for c in tb3.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(8)
    doc.add_paragraph()

    H(doc,"SECCIÓN 4: REGISTRO DE MANTENIMIENTOS",2)
    cols_mant=["N°","Equipo","Tipo mant.","Fecha","Actividad realizada","Técnico","Empresa","Resultado","Próximo mant."]
    tb4=doc.add_table(rows=11,cols=len(cols_mant)); tb4.style='Table Grid'
    hrow(tb4,cols_mant)
    for i in range(1,11):
        tb4.rows[i].cells[0].text=str(i)
        for c in tb4.rows[i].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(8)
    doc.add_paragraph()

    H(doc,"SECCIÓN 5: CONTROL DE TEMPERATURA REFRIGERADOR MÉDICO",2)
    cols_temp=["Fecha","Temp. AM (°C)","Temp. PM (°C)","¿Dentro rango 2-8°C?","Observaciones","Firma"]
    tb5=doc.add_table(rows=16,cols=len(cols_temp)); tb5.style='Table Grid'
    hrow(tb5,cols_temp)
    for i in range(1,16):
        tb5.rows[i].cells[3].text="□ Sí  □ No"
        for c in tb5.rows[i].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"SECCIÓN 6: FIRMA DE RESPONSABLE")
    doc.add_paragraph("Médica responsable: [NOMBRE DE LA MÉDICA]\n\nFirma: _________________________\n\nFecha: _________________________")

    path=os.path.join(BASE3,"FOR-DOT-001_Inventario_Equipos_Dispositivos.docx")
    doc.save(path); print(f"✓ Creado: {path}")

if __name__=="__main__":
    print("Generando CARPETA 3 - DOTACIÓN...")
    crear_man_dot(); crear_pro_dot(); crear_for_dot()
    print("✅ Carpeta 3 completada.")
