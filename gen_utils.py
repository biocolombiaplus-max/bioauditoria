"""Utilidades comunes para generación de documentos de habilitación."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

TODAY = datetime.date.today().strftime("%d/%m/%Y")
YEAR = "2025"

# Colores corporativos
COLOR_AZUL = RGBColor(0x1F, 0x49, 0x7D)
COLOR_AZUL_CLARO = RGBColor(0xBD, 0xD7, 0xEE)
COLOR_GRIS = RGBColor(0x70, 0x70, 0x70)
COLOR_NEGRO = RGBColor(0x00, 0x00, 0x00)
COLOR_BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_ROJO = RGBColor(0xC0, 0x00, 0x00)

CONSULTORIO = "[NOMBRE DEL CONSULTORIO MÉDICO]"
MEDICA = "[NOMBRE DE LA MÉDICA]"
DIRECCION = "[DIRECCIÓN DEL CONSULTORIO]"
CIUDAD = "[CIUDAD]"
TELEFONO = "[TELÉFONO]"
REPS = "[REGISTRO ESPECIAL DE PRESTADORES DE SALUD - REPS]"
REG_MEDICA = "[NÚMERO DE REGISTRO MÉDICO / TARJETA PROFESIONAL]"
NIT = "[NIT O CÉDULA]"


def nueva_doc():
    doc = Document()
    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)
    return doc


def set_fuente(run, nombre="Calibri", tamaño=11, negrita=False, italica=False, color=None):
    run.font.name = nombre
    run.font.size = Pt(tamaño)
    run.font.bold = negrita
    run.font.italic = italica
    if color:
        run.font.color.rgb = color


def parrafo_estilo(doc, texto, nivel=0, negrita=False, italica=False,
                   tamaño=11, color=None, alineacion=WD_ALIGN_PARAGRAPH.LEFT,
                   espacio_antes=6, espacio_despues=6):
    p = doc.add_paragraph()
    p.alignment = alineacion
    p.paragraph_format.space_before = Pt(espacio_antes)
    p.paragraph_format.space_after = Pt(espacio_despues)
    if nivel > 0:
        p.paragraph_format.left_indent = Cm(nivel * 0.75)
    run = p.add_run(texto)
    set_fuente(run, tamaño=tamaño, negrita=negrita, italica=italica, color=color)
    return p


def titulo_seccion(doc, texto, nivel=1):
    if nivel == 1:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(texto)
        set_fuente(run, tamaño=13, negrita=True, color=COLOR_AZUL)
        # Línea inferior simulada con sombreado
        p.paragraph_format.left_indent = Cm(0)
        return p
    elif nivel == 2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(texto)
        set_fuente(run, tamaño=12, negrita=True, color=COLOR_AZUL)
        return p
    elif nivel == 3:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(texto)
        set_fuente(run, tamaño=11, negrita=True, color=COLOR_NEGRO)
        return p
    else:
        p = doc.add_paragraph()
        run = p.add_run(texto)
        set_fuente(run, tamaño=11, negrita=True, italica=True)
        return p


def agregar_viñeta(doc, texto, nivel=0, tamaño=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1 + nivel * 0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(texto)
    set_fuente(run, tamaño=tamaño)
    return p


def tabla_simple(doc, encabezados, filas, ancho_cols=None):
    """Crea tabla con encabezados azules y filas alternadas."""
    n_cols = len(encabezados)
    tabla = doc.add_table(rows=1, cols=n_cols)
    tabla.style = 'Table Grid'
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezados
    fila_enc = tabla.rows[0]
    for i, enc in enumerate(encabezados):
        cell = fila_enc.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(enc)
        set_fuente(run, tamaño=10, negrita=True, color=COLOR_BLANCO)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Color de fondo azul
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1F497D')
        cell._tc.get_or_add_tcPr().append(shading)

    # Filas de datos
    for idx, fila in enumerate(filas):
        row = tabla.add_row()
        fill = 'BDD7EE' if idx % 2 == 0 else 'FFFFFF'
        for i, val in enumerate(fila):
            cell = row.cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_fuente(run, tamaño=10)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), fill)
            cell._tc.get_or_add_tcPr().append(shading)

    if ancho_cols:
        for i, ancho in enumerate(ancho_cols):
            for row in tabla.rows:
                row.cells[i].width = Cm(ancho)

    doc.add_paragraph()
    return tabla


def portada(doc, titulo, codigo, version="1.0", proceso="", paginas_aprox=""):
    """Genera portada estándar del documento."""
    doc.add_paragraph()
    doc.add_paragraph()

    # Logo/encabezado institucional simulado
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(CONSULTORIO)
    set_fuente(run, tamaño=18, negrita=True, color=COLOR_AZUL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Consultorio de Medicina General")
    set_fuente(run, tamaño=13, color=COLOR_GRIS)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{DIRECCION} | {CIUDAD}")
    set_fuente(run, tamaño=11, color=COLOR_GRIS)

    doc.add_paragraph()
    doc.add_paragraph()

    # Línea decorativa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 55)
    set_fuente(run, color=COLOR_AZUL, tamaño=12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(titulo)
    set_fuente(run, tamaño=20, negrita=True, color=COLOR_AZUL)

    doc.add_paragraph()
    doc.add_paragraph()

    # Tabla de identificación
    tabla_id = doc.add_table(rows=7, cols=2)
    tabla_id.style = 'Table Grid'
    tabla_id.alignment = WD_TABLE_ALIGNMENT.CENTER

    datos = [
        ("CÓDIGO", codigo),
        ("VERSIÓN", version),
        ("FECHA DE EMISIÓN", TODAY),
        ("PROCESO", proceso if proceso else "Sistema de Gestión de Calidad"),
        ("ELABORADO POR", MEDICA),
        ("REVISADO POR", MEDICA),
        ("APROBADO POR", MEDICA),
    ]

    for i, (clave, valor) in enumerate(datos):
        row = tabla_id.rows[i]
        c0 = row.cells[0]
        c1 = row.cells[1]
        c0.text = ""
        run0 = c0.paragraphs[0].add_run(clave)
        set_fuente(run0, tamaño=11, negrita=True, color=COLOR_BLANCO)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading0 = OxmlElement('w:shd')
        shading0.set(qn('w:val'), 'clear')
        shading0.set(qn('w:color'), 'auto')
        shading0.set(qn('w:fill'), '1F497D')
        c0._tc.get_or_add_tcPr().append(shading0)

        c1.text = ""
        run1 = c1.paragraphs[0].add_run(valor)
        set_fuente(run1, tamaño=11)
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in tabla_id.rows:
        row.cells[0].width = Cm(6)
        row.cells[1].width = Cm(10)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 55)
    set_fuente(run, color=COLOR_AZUL, tamaño=12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"REPS: {REPS}")
    set_fuente(run, tamaño=10, italica=True, color=COLOR_GRIS)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"NIT/Cédula: {NIT}  |  Tel: {TELEFONO}")
    set_fuente(run, tamaño=10, italica=True, color=COLOR_GRIS)

    doc.add_page_break()


def control_versiones(doc):
    titulo_seccion(doc, "CONTROL DE VERSIONES Y CAMBIOS")
    encabezados = ["Versión", "Fecha", "Descripción del Cambio", "Elaboró", "Aprobó"]
    filas = [
        ["1.0", TODAY, "Emisión inicial del documento para proceso de habilitación ante la Secretaría de Salud", MEDICA, MEDICA],
        ["", "", "", "", ""],
        ["", "", "", "", ""],
    ]
    tabla_simple(doc, encabezados, filas, ancho_cols=[1.5, 2.5, 7, 3, 3])
    doc.add_page_break()


def tabla_contenido_texto(doc, items):
    """items = [(título, página_aprox), ...]"""
    titulo_seccion(doc, "TABLA DE CONTENIDO")
    parrafo_estilo(doc, "(Actualizar números de página usando Word: Referencias → Actualizar tabla)",
                   italica=True, tamaño=9, color=COLOR_GRIS)
    for titulo, pag in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        tab = p.paragraph_format.tab_stops
        run = p.add_run(f"{titulo}")
        set_fuente(run, tamaño=10)
        run2 = p.add_run(f"\t{pag}")
        set_fuente(run2, tamaño=10)
    doc.add_page_break()


def seccion_firmas(doc):
    titulo_seccion(doc, "FIRMAS DE APROBACIÓN")
    parrafo_estilo(doc, "El presente documento ha sido elaborado, revisado y aprobado por:")
    doc.add_paragraph()
    tabla = doc.add_table(rows=3, cols=3)
    tabla.style = 'Table Grid'
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    roles = [
        ("ELABORÓ", MEDICA, "Médica General Propietaria"),
        ("REVISÓ", MEDICA, "Médica General Propietaria"),
        ("APROBÓ", MEDICA, "Médica General Propietaria"),
    ]

    # Fila de títulos
    for i, (rol, nombre, cargo) in enumerate(roles):
        cell_titulo = tabla.rows[0].cells[i]
        cell_titulo.text = ""
        run = cell_titulo.paragraphs[0].add_run(rol)
        set_fuente(run, tamaño=11, negrita=True, color=COLOR_BLANCO)
        cell_titulo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1F497D')
        cell_titulo._tc.get_or_add_tcPr().append(shading)

    for i, (rol, nombre, cargo) in enumerate(roles):
        cell_nombre = tabla.rows[1].cells[i]
        cell_nombre.text = ""
        cell_nombre.paragraphs[0].add_run("\n\n_________________________\n")
        run = cell_nombre.paragraphs[0].add_run(nombre)
        set_fuente(run, tamaño=10, negrita=True)
        cell_nombre.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_cargo = tabla.rows[2].cells[i]
        cell_cargo.text = ""
        run = cell_cargo.paragraphs[0].add_run(cargo)
        set_fuente(run, tamaño=10)
        run2 = cell_cargo.paragraphs[0].add_run(f"\nFecha: {TODAY}")
        set_fuente(run2, tamaño=10)
        cell_cargo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in tabla.rows:
        for cell in row.cells:
            cell.width = Cm(5.5)

    doc.add_paragraph()


def marco_legal_general(doc, normas_adicionales=None):
    titulo_seccion(doc, "3. MARCO LEGAL Y NORMATIVO")
    parrafo_estilo(doc,
        "El presente documento se fundamenta en las siguientes disposiciones legales vigentes en la República de Colombia:",
        tamaño=11)

    normas_base = [
        ("Constitución Política de Colombia 1991", "Artículo 49: La atención de la salud y el saneamiento ambiental son servicios públicos a cargo del Estado."),
        ("Ley 100 de 1993", "Por la cual se crea el sistema de seguridad social integral. Establece las bases del Sistema General de Seguridad Social en Salud (SGSSS)."),
        ("Decreto 1011 de 2006", "Por el cual se establece el Sistema Obligatorio de Garantía de Calidad de la Atención de Salud del Sistema General de Seguridad Social en Salud (SOGCS)."),
        ("Resolución 3100 de 2019", "Por la cual se definen los procedimientos y condiciones de inscripción de los prestadores de servicios de salud y de habilitación de los servicios de salud. Deroga la Resolución 2003 de 2014."),
        ("Resolución 2003 de 2014", "Antecedente normativo que definió los procedimientos de habilitación. Vigente en aspectos complementarios."),
        ("Ley 23 de 1981", "Por la cual se dictan normas en materia de ética médica. Código de Ética Médica colombiano."),
        ("Decreto 3380 de 1981", "Por el cual se reglamenta la Ley 23 de 1981 sobre ética médica."),
        ("Resolución 1995 de 1999", "Por la cual se establecen normas para el manejo de la Historia Clínica."),
        ("Ley 1438 de 2011", "Por medio de la cual se reforma el Sistema General de Seguridad Social en Salud y se dictan otras disposiciones. Refuerza la Atención Primaria en Salud."),
        ("Resolución 0256 de 2016", "Por la cual se dictan disposiciones en relación con el Sistema de Información para la Calidad y se establecen los indicadores para el monitoreo de la calidad en salud."),
        ("Decreto 780 de 2016", "Por medio del cual se expide el Decreto Único Reglamentario del Sector Salud y Protección Social."),
    ]

    if normas_adicionales:
        normas_base.extend(normas_adicionales)

    for norma, descripcion in normas_base:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(3)
        run1 = p.add_run(f"• {norma}: ")
        set_fuente(run1, tamaño=10.5, negrita=True)
        run2 = p.add_run(descripcion)
        set_fuente(run2, tamaño=10.5)


def guardar(doc, ruta):
    doc.save(ruta)
    print(f"  ✓ Guardado: {ruta.split('/')[-1]}")
