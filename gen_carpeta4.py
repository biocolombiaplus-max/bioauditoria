#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

BASE4="/home/user/bioauditoria/documentos_habilitacion/CARPETA_4_MEDICAMENTOS"

def margins(doc):
    for s in doc.sections:
        s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5); s.left_margin=Cm(3); s.right_margin=Cm(2.5)

def hf(doc,code,title):
    for s in doc.sections:
        h=s.header; h.paragraphs[0].clear()
        r=h.paragraphs[0].add_run(f"CONSULTORIO [NOMBRE DEL CONSULTORIO]  |  {code}  |  {title}")
        r.font.size=Pt(8); r.font.color.rgb=RGBColor(64,64,64)
        h.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        f=s.footer; f.paragraphs[0].clear()
        r2=f.paragraphs[0].add_run(f"{code} - V1.0 - {datetime.date.today().strftime('%d/%m/%Y')} | Pág. ")
        r2.font.size=Pt(8)
        fld=OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'),'begin')
        f.paragraphs[0].runs[-1]._r.append(fld)
        it=OxmlElement('w:instrText'); it.text='PAGE'
        f.paragraphs[0].runs[-1]._r.append(it)
        fld2=OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'),'end')
        f.paragraphs[0].runs[-1]._r.append(fld2)
        f.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

def portada(doc,code,title):
    doc.add_paragraph(); doc.add_paragraph()
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=t.add_run("CONSULTORIO MÉDICO\n[NOMBRE DEL CONSULTORIO]")
    r.font.size=Pt(18); r.font.bold=True; r.font.color.rgb=RGBColor(31,73,125)
    doc.add_paragraph()
    t2=doc.add_paragraph(); t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=t2.add_run(title); r2.font.size=Pt(15); r2.font.bold=True
    doc.add_paragraph()
    tb=doc.add_table(rows=6,cols=2); tb.style='Table Grid'; tb.alignment=WD_TABLE_ALIGNMENT.CENTER
    datos=[("Código:",code),("Versión:","1.0"),("Fecha:",datetime.date.today().strftime('%d/%m/%Y')),
           ("Elaboró:","[NOMBRE DE LA MÉDICA], Médica General"),
           ("Revisó:","[NOMBRE DE LA MÉDICA]"),("Aprobó:","[NOMBRE DE LA MÉDICA]")]
    for i,(l,v) in enumerate(datos):
        tb.rows[i].cells[0].text=l; tb.rows[i].cells[1].text=v
        tb.rows[i].cells[0].paragraphs[0].runs[0].font.bold=True
    doc.add_page_break()
    cv=doc.add_paragraph(); r3=cv.add_run("CONTROL DE VERSIONES")
    r3.font.bold=True; r3.font.size=Pt(13); r3.font.color.rgb=RGBColor(31,73,125)
    cv.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tb2=doc.add_table(rows=3,cols=5); tb2.style='Table Grid'
    hrow(tb2,["Versión","Fecha","Cambio","Elaboró","Aprobó"])
    tb2.rows[1].cells[0].text="1.0"; tb2.rows[1].cells[1].text=datetime.date.today().strftime('%d/%m/%Y')
    tb2.rows[1].cells[2].text="Versión inicial"; tb2.rows[1].cells[3].text="[NOMBRE DE LA MÉDICA]"
    tb2.rows[1].cells[4].text="[NOMBRE DE LA MÉDICA]"
    doc.add_page_break()

def hrow(table,headers,bg="1F497D"):
    row=table.rows[0]
    for i,h in enumerate(headers):
        cell=row.cells[i]; cell.text=""
        run=cell.paragraphs[0].add_run(h)
        run.font.bold=True; run.font.color.rgb=RGBColor(255,255,255); run.font.size=Pt(10)
        cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),bg)
        shd.set(qn('w:color'),'auto'); shd.set(qn('w:val'),'clear'); tcPr.append(shd)

def H(doc,text,lv=1):
    p=doc.add_heading(text,level=lv)
    colors={1:RGBColor(31,73,125),2:RGBColor(46,116,181),3:RGBColor(68,114,196)}
    sizes={1:14,2:12,3:11}
    for run in p.runs:
        run.font.size=Pt(sizes[lv]); run.font.bold=True; run.font.color.rgb=colors[lv]
    return p

def B(doc,text,bold=False):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=p.add_run(text); r.font.size=Pt(11); r.font.bold=bold; r.font.name='Calibri'; return p

def BL(doc,text):
    p=doc.add_paragraph(style='List Bullet'); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=p.add_run(text); r.font.size=Pt(11); r.font.name='Calibri'; return p

def firmas(doc):
    doc.add_paragraph()
    tb=doc.add_table(rows=2,cols=3); tb.style='Table Grid'
    hrow(tb,["ELABORÓ","REVISÓ","APROBÓ"])
    for j in range(3):
        tb.rows[1].cells[j].text="[NOMBRE DE LA MÉDICA]\nMédica General - Propietaria\n\nFirma: _________________________\n\nFecha: "+datetime.date.today().strftime('%d/%m/%Y')
        tb.rows[1].cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

# =========== MAN-MED-001 ===========
def crear_man_med():
    doc=Document(); margins(doc); hf(doc,"MAN-MED-001","Manual de Gestión de Medicamentos")
    portada(doc,"MAN-MED-001","MANUAL DE GESTIÓN DE MEDICAMENTOS Y DISPOSITIVOS MÉDICOS TERAPÉUTICOS")

    H(doc,"1. OBJETIVO")
    B(doc,"Establecer los lineamientos para la gestión integral de medicamentos y dispositivos médicos terapéuticos del Consultorio [NOMBRE DEL CONSULTORIO], incluyendo su selección, adquisición, almacenamiento, dispensación, control de calidad, farmacovigilancia y disposición final, garantizando el cumplimiento de la normatividad colombiana vigente y la seguridad de los pacientes en el uso de estos productos.")

    H(doc,"2. ALCANCE")
    B(doc,"Aplica a todos los medicamentos y dispositivos médicos terapéuticos disponibles en el consultorio: botiquín de urgencias para la consulta de medicina general, y medicamentos/dispositivos médicos utilizados en los procedimientos estéticos no invasivos (toxina botulínica tipo A, ácido hialurónico, agentes para peelings, cócteles de mesoterapia, etc.).")

    H(doc,"3. MARCO LEGAL")
    normas=[
        ("Ley 212 de 1995","Reglamenta la profesión de químico farmacéutico y establece los requisitos para la gestión de medicamentos."),
        ("Decreto 2200 de 2005","Reglamenta el servicio farmacéutico. Establece las condiciones para la gestión del medicamento en las instituciones de salud."),
        ("Resolución 1478 de 2006","Establece las normas para el control, seguimiento y vigilancia de la importación, exportación, procesamiento, síntesis, fabricación, distribución, dispensación, compra, venta y uso de sustancias que se utilizan como precursores o como sustancias psicoactivas."),
        ("Resolución 2004009455","Adopta el Programa Nacional de Farmacovigilancia en Colombia. Establece la obligatoriedad de reporte de reacciones adversas a medicamentos (RAM) al INVIMA."),
        ("Resolución 4816 de 2008","Programa Nacional de Tecnovigilancia. Para dispositivos médicos terapéuticos."),
        ("Decreto 4725 de 2005","Regula registros sanitarios de dispositivos médicos. Aplica al ácido hialurónico, tubos PRP, etc."),
        ("Resolución 3100 de 2019","Estándares de habilitación relacionados con medicamentos y dispositivos médicos disponibles en el servicio."),
        ("Decreto 780 de 2016","Decreto Único del Sector Salud. Consolida normativa sobre medicamentos y dispositivos."),
        ("Resolución 1403 de 2007","Modelo de Gestión del Servicio Farmacéutico. Adoptado como referencia."),
        ("Ley 9 de 1979","Código Sanitario Nacional. Normas generales sobre alimentos, medicamentos y productos similares."),
    ]
    tb=doc.add_table(rows=len(normas)+1,cols=2); tb.style='Table Grid'
    hrow(tb,["Norma","Descripción"])
    for i,(n,d) in enumerate(normas):
        tb.rows[i+1].cells[0].text=n; tb.rows[i+1].cells[1].text=d
        for c in tb.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()

    H(doc,"4. DEFINICIONES")
    defs=[("Medicamento:","Toda preparación o producto farmacéutico empleado para la prevención, diagnóstico, tratamiento, mitigación y cura de enfermedades, para el alivio de síntomas y la modificación de procesos fisiológicos normales o patológicos."),
          ("Dispositivo médico terapéutico:","Dispositivo médico utilizado con fines terapéuticos directos en el paciente: ácido hialurónico, apósitos activos, etc."),
          ("Registro sanitario INVIMA:","Autorización expedida por el Instituto Nacional de Vigilancia de Medicamentos y Alimentos para la comercialización de un medicamento o dispositivo médico en Colombia."),
          ("Cadena de frío:","Sistema de almacenamiento y distribución que mantiene los biológicos y termoestables entre los rangos de temperatura requeridos desde el fabricante hasta el usuario final."),
          ("Farmacovigilancia:","Ciencia y actividades relacionadas con la detección, evaluación, comprensión y prevención de efectos adversos o cualquier otro problema relacionado con medicamentos."),
          ("Reacción Adversa a Medicamento (RAM):","Cualquier respuesta a un medicamento que sea nociva y no intencionada, y que tenga lugar a dosis normalmente usadas en el ser humano para la profilaxis, el diagnóstico o el tratamiento."),
          ("Botiquín:","Conjunto de medicamentos e insumos médicos destinados a la atención de urgencias médicas básicas."),
          ("Fecha de vencimiento:","Fecha a partir de la cual el fabricante no garantiza la estabilidad y efectividad del medicamento o dispositivo médico."),]
    for term,defi in defs:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r1=p.add_run(term+" "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(defi); r2.font.size=Pt(11)

    H(doc,"5. BOTIQUÍN DE URGENCIAS PARA MEDICINA GENERAL")
    H(doc,"5.1 Composición del Botiquín de Urgencias",2)
    B(doc,"El consultorio debe mantener permanentemente disponible un botiquín de urgencias con los medicamentos e insumos necesarios para el manejo inicial de las situaciones de urgencia más frecuentes que puedan presentarse durante la consulta:")

    botiquin=[
        ("Adrenalina (epinefrina) 1mg/1ml","2 ampollas","Para reacción anafiláctica","Refrigeración no obligatoria; lugar fresco y oscuro","[REGISTRO INVIMA]"),
        ("Hidrocortisona 500mg IV","2 viales","Para reacción alérgica severa, insuficiencia suprarrenal aguda","Temperatura ambiente","[REGISTRO INVIMA]"),
        ("Difenhidramina (antihistamínico) inyectable","2 ampollas","Para reacciones alérgicas moderadas","Temperatura ambiente","[REGISTRO INVIMA]"),
        ("Nitroglicerina sublingual (spray o tabletas)","1 frasco","Para crisis anginosa","Temperatura ambiente, protegido de luz","[REGISTRO INVIMA]"),
        ("Dexametasona 4mg/ml inyectable","2 ampollas","Antiinflamatorio, antialérgico","Temperatura ambiente","[REGISTRO INVIMA]"),
        ("Glucosa 50% en jeringa 50ml","1 unidad","Para hipoglucemia severa","Temperatura ambiente","[REGISTRO INVIMA]"),
        ("Metoclopramida 10mg inyectable","2 ampollas","Para náuseas y vómitos","Temperatura ambiente","[REGISTRO INVIMA]"),
        ("Diazepam 10mg inyectable","2 ampollas","Para crisis convulsivas","Temperatura ambiente","[REGISTRO INVIMA]"),
        ("Solución salina NaCl 0.9% 250ml o 500ml","2 bolsas","Para rehidratación, dilución de medicamentos","Temperatura ambiente","[REGISTRO INVIMA]"),
        ("Atropina 1mg inyectable","2 ampollas","Para bradicardia","Temperatura ambiente","[REGISTRO INVIMA]"),
        ("Salbutamol (ventolin) inhalador","1 unidad","Para broncoespasmo","Temperatura ambiente","[REGISTRO INVIMA]"),
        ("Furosemida 20mg inyectable","2 ampollas","Para edema agudo de pulmón","Temperatura ambiente","[REGISTRO INVIMA]"),
        ("Ácido acetilsalicílico 300mg tabletas","5 tabletas","Para síndrome coronario agudo","Temperatura ambiente","[REGISTRO INVIMA]"),
    ]
    tb2=doc.add_table(rows=len(botiquin)+1,cols=5); tb2.style='Table Grid'
    hrow(tb2,["Medicamento","Cantidad mínima","Indicación de urgencia","Condiciones de almacenamiento","Registro INVIMA"])
    for i,row in enumerate(botiquin):
        for j,v in enumerate(row):
            tb2.rows[i+1].cells[j].text=v
            for c in tb2.rows[i+1].cells[j].paragraphs:
                for r in c.runs: r.font.size=Pt(8)
    doc.add_paragraph()

    H(doc,"5.2 Insumos del Botiquín",2)
    insumos=["Jeringas de 1ml, 3ml, 5ml, 10ml (mínimo 5 de cada una)","Agujas 21G, 23G, 25G (mínimo 5 de cada una)","Catéter IV (venoclisis) N° 18, 20, 22 (mínimo 2 de cada uno)","Llave de 3 vías (2 unidades)","Equipo de venoclisis (2 unidades)","Guantes estériles talla S, M, L (2 pares de cada talla)","Gasas estériles 10x10 cm (1 caja)","Esparadrapo / cinta adhesiva médica","Tensiómetro, fonendoscopio, oxímetro","Ambú adulto (1 unidad)","Cánulas de Guedel N° 4, 5, 6 (1 de cada una)","Mascarilla de alto flujo con reservorio","Mascarilla de venturi","Oxígeno medicinal portátil con regulador","Tabla de soporte cardíaco para RCP"]
    for item in insumos: BL(doc,item)

    H(doc,"6. MEDICAMENTOS PARA PROCEDIMIENTOS ESTÉTICOS NO INVASIVOS")
    H(doc,"6.1 Toxina Botulínica Tipo A",2)
    B(doc,"La toxina botulínica tipo A es un medicamento biológico de uso exclusivo médico. Su gestión en el consultorio debe cumplir con los siguientes lineamientos:")
    B(doc,"CADENA DE FRÍO - REQUISITO CRÍTICO:", bold=True)
    tb3_data=[
        ("Temperatura de almacenamiento","2-8°C. NUNCA congelar. NUNCA exponer a temperatura >8°C por más de 4 horas."),
        ("Registro INVIMA requerido","Verificar que el producto tenga registro INVIMA vigente como medicamento (no cosmético)"),
        ("Proveedor autorizado","Adquirir ÚNICAMENTE de importadores o distribuidores autorizados por el INVIMA para toxina botulínica"),
        ("Fecha de vencimiento","Verificar antes de cada uso. Registrar lote y vencimiento en la historia clínica del paciente"),
        ("Reconstitución","Reconstituir según instrucciones del fabricante con solución salina 0.9% estéril sin preservantes. Usar la concentración correcta (típicamente 2.5-4 UI/0.1ml según técnica y fabricante)"),
        ("Almacenamiento post-reconstitución","Una vez reconstituida: usar en el mismo día (máximo 4 horas si se mantiene refrigerada). Algunos fabricantes permiten hasta 24h a 2-8°C. Verificar inserto."),
        ("Disposición de sobrante","Desnaturalizar con solución de hipoclorito antes de desechar. Desechar en residuos biosanitarios."),
        ("Marcas disponibles en Colombia","Botox® (AbbVie), Dysport® (Ipsen), Xeomin® (Merz) - verificar registro INVIMA vigente de cada marca"),
    ]
    tb3=doc.add_table(rows=len(tb3_data)+1,cols=2); tb3.style='Table Grid'
    hrow(tb3,["Aspecto","Especificación"])
    for i,(k,v) in enumerate(tb3_data):
        tb3.rows[i+1].cells[0].text=k; tb3.rows[i+1].cells[1].text=v
        for c in tb3.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()

    H(doc,"6.2 Ácido Hialurónico (Rellenos Dérmicos)",2)
    B(doc,"El ácido hialurónico para rellenos dérmicos es un dispositivo médico de Clase III (alto riesgo) según el Decreto 4725 de 2005 y la normativa INVIMA:")
    tb4_data=[
        ("Clasificación regulatoria","Dispositivo médico Clase III. Requiere registro INVIMA como dispositivo médico (NO como cosmético)."),
        ("Temperatura almacenamiento","Temperatura ambiente (15-25°C) protegido de la luz solar directa. No refrigerar. No congelar."),
        ("Tipos en Colombia","Reticulado (para voluminización: labios, surcos, pómulos) vs no reticulado (para hidratación dérmica). Varios rangos de viscosidad."),
        ("Proveedor","Adquirir de distribuidores autorizados con documentación del registro INVIMA. Exigir factura con lote y registro."),
        ("Verificación antes del uso","Comprobar: aspecto del gel (sin partículas, sin cambio de color), fecha de vencimiento, integridad del envase, registro INVIMA."),
        ("Hialuronidasa (antídoto obligatorio)","Mantener SIEMPRE hialuronidasa disponible cuando se aplique ácido hialurónico. Protocolo de emergencia por oclusión vascular incluido en PRO-PP-002."),
    ]
    tb4=doc.add_table(rows=len(tb4_data)+1,cols=2); tb4.style='Table Grid'
    hrow(tb4,["Aspecto","Especificación"])
    for i,(k,v) in enumerate(tb4_data):
        tb4.rows[i+1].cells[0].text=k; tb4.rows[i+1].cells[1].text=v
        for c in tb4.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()

    H(doc,"7. ALMACENAMIENTO DE MEDICAMENTOS")
    H(doc,"7.1 Condiciones Generales de Almacenamiento",2)
    cond_alm=["Temperatura controlada: temperatura ambiente (15-25°C) para medicamentos que no requieren cadena de frío. Monitorear con termómetro ambiental.",
              "Protección de la luz: mantener medicamentos fotosensibles en sus envases originales y en lugar oscuro",
              "Humedad controlada: evitar almacenar en lugares húmedos o con cambios bruscos de temperatura (no en baños ni cerca de ventanas)",
              "Identificación: todos los medicamentos deben estar claramente identificados y en sus envases originales",
              "Rotación de stock (FEFO - First Expired First Out): los medicamentos con fecha de vencimiento más próxima se usan primero",
              "Separación: medicamentos para uso en pacientes separados de cualquier otro producto de limpieza o uso personal",
              "Acceso controlado: acceso restringido a personal autorizado (médica propietaria)",
              "Refrigerador exclusivo para medicamentos: no mezclar con alimentos ni otros productos",
              "Registro de entradas y salidas en el formato FOR-MED-001"]
    for item in cond_alm: BL(doc,item)

    H(doc,"7.2 Almacenamiento según Condición Especial",2)
    tb5=doc.add_table(rows=5,cols=3); tb5.style='Table Grid'
    hrow(tb5,["Condición de almacenamiento","Temperatura","Medicamentos del consultorio"])
    rows5=[("Refrigeración (cadena de frío)","2-8°C","Toxina botulínica tipo A, vacunas si aplica, insulina"),
           ("Temperatura ambiente controlada","15-25°C","Ácido hialurónico, mayoría de medicamentos del botiquín, peelings"),
           ("Proteger de la luz","Temperatura ambiente","Difenhidramina, metoclopramida, algunas soluciones IV"),
           ("Controlados (si aplica)","Temperatura ambiente bajo llave","Diazepam (si se tiene) - en caso de tener este tipo de medicamentos")]
    for i,row in enumerate(rows5):
        for j,v in enumerate(row): tb5.rows[i+1].cells[j].text=v
        for c in tb5.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"8. FARMACOVIGILANCIA")
    H(doc,"8.1 Programa de Farmacovigilancia",2)
    B(doc,"En cumplimiento de la Resolución 2004009455, el consultorio implementa el siguiente programa de farmacovigilancia para la detección y reporte de reacciones adversas a medicamentos (RAM):")

    fv_pasos=[("Identificación","El personal identifica cualquier reacción inesperada en el paciente que pueda estar relacionada con un medicamento administrado o prescrito en el consultorio."),
              ("Documentación inicial","Se registra en la historia clínica del paciente: medicamento sospechoso, dosis, vía de administración, fecha de inicio, descripción de la reacción, manejo dado."),
              ("Evaluación de causalidad","Se evalúa la causalidad según el algoritmo de Naranjo: ¿es probable, posible o dudosa la relación medicamento-reacción adversa?"),
              ("Notificación al INVIMA","Las RAM graves (que requirieron hospitalización, causaron discapacidad, fueron mortales o congénitas) deben reportarse al INVIMA a través del Programa Nacional de Farmacovigilancia en www.invima.gov.co"),
              ("Seguimiento del paciente","Hacer seguimiento al paciente hasta la resolución del evento adverso. Documentar evolución."),
              ("Análisis y aprendizaje","Analizar la causa del evento adverso y establecer acciones para prevenir recurrencias: cambio de medicamento, ajuste de dosis, contraindicación en el paciente específico."),]
    tb6=doc.add_table(rows=len(fv_pasos)+1,cols=3); tb6.style='Table Grid'
    hrow(tb6,["Paso","Actividad","Descripción"])
    for i,(a,d) in enumerate(fv_pasos):
        tb6.rows[i+1].cells[0].text=str(i+1); tb6.rows[i+1].cells[1].text=a; tb6.rows[i+1].cells[2].text=d
        for c in tb6.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"8.2 Farmacovigilancia Específica para Productos Estéticos",2)
    B(doc,"Para los productos utilizados en procedimientos estéticos, se deben reportar los siguientes eventos:")
    eventos_est=["Reacciones alérgicas o anafilácticas a la toxina botulínica o al ácido hialurónico",
                 "Efectos adversos inesperados de la toxina botulínica: ptosis severa, disfagia, parálisis muscular no esperada, propagación del efecto",
                 "Complicaciones vasculares por rellenos dérmicos: necrosis, ceguera, ACV",
                 "Infecciones graves post-procedimiento",
                 "Formación de granulomas o nódulos por rellenos dérmicos",
                 "Cualquier evento adverso grave atribuible a un dispositivo médico o medicamento utilizado en estética"]
    for item in eventos_est: BL(doc,item)

    H(doc,"9. CONTROL DE FECHAS DE VENCIMIENTO")
    B(doc,"Se establece el siguiente sistema para el control de fechas de vencimiento de medicamentos y dispositivos médicos:")
    venc=["Revisión mensual: El primer día hábil de cada mes se realiza una revisión de las fechas de vencimiento de todos los medicamentos y dispositivos médicos del consultorio",
          "Rotación FEFO: Los productos con vencimiento más próximo se colocan al frente para ser usados primero",
          "Alerta temprana: Los productos que vencen en los próximos 3 meses se marcan con señal visible (ej. cinta de color rojo)",
          "Retiro de vencidos: Los medicamentos vencidos se retiran inmediatamente del inventario activo y se almacenan separados para disposición final",
          "Disposición final: Los medicamentos vencidos se entregan a la empresa gestora de residuos farmacéuticos o se devuelven al proveedor según el plan de retorno",
          "Registro: Cada revisión de fechas de vencimiento se registra en el formato FOR-MED-001"]
    for item in venc: BL(doc,item)

    H(doc,"10. INDICADORES")
    tb7=doc.add_table(rows=5,cols=4); tb7.style='Table Grid'
    hrow(tb7,["Indicador","Fórmula","Meta","Frecuencia"])
    inds=[("% botiquín completo","(Medicamentos disponibles/Medicamentos requeridos)×100","100%","Mensual"),
          ("% medicamentos sin vencer","(Medicamentos vigentes/Total medicamentos)×100","100%","Mensual"),
          ("N° RAM reportadas al INVIMA","Conteo de reportes realizados","Reportar 100% de RAM graves","Mensual"),
          ("% productos con registro INVIMA vigente","(Productos con INVIMA/Total productos)×100","100%","Trimestral")]
    for i,row in enumerate(inds):
        for j,v in enumerate(row): tb7.rows[i+1].cells[j].text=v
        for c in tb7.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph(); firmas(doc)
    path=os.path.join(BASE4,"MAN-MED-001_Manual_Gestion_Medicamentos.docx")
    doc.save(path); print(f"✓ Creado: {path}")

# =========== PRO-MED-001 ===========
def crear_pro_med():
    doc=Document(); margins(doc); hf(doc,"PRO-MED-001","Proceso de Almacenamiento de Medicamentos")
    portada(doc,"PRO-MED-001","PROCESO DE ALMACENAMIENTO Y CONTROL DE MEDICAMENTOS")

    H(doc,"1. OBJETIVO")
    B(doc,"Definir el proceso estandarizado para el almacenamiento, control y gestión de los medicamentos y dispositivos médicos terapéuticos del Consultorio [NOMBRE DEL CONSULTORIO], garantizando sus condiciones de calidad, seguridad y eficacia desde la recepción hasta su uso en el paciente.")

    H(doc,"2. ALCANCE")
    B(doc,"Aplica a todos los medicamentos e insumos del botiquín de urgencias y a todos los medicamentos y dispositivos médicos utilizados en procedimientos estéticos no invasivos.")

    H(doc,"3. DESCRIPCIÓN DEL PROCESO")

    H(doc,"3.1 Recepción de Medicamentos",2)
    pasos_recep=[("Verificación del pedido","Cotejar lo recibido con la orden de compra: nombre del medicamento, presentación, concentración, cantidad"),
                 ("Verificación del proveedor","Comprobar que el proveedor esté autorizado por el INVIMA para el producto específico"),
                 ("Verificación del producto","Comprobar: registro INVIMA en el empaque, fecha de vencimiento (mínimo 6 meses de vigencia al momento de recepción), lote, condiciones del embalaje, cadena de frío si aplica"),
                 ("Rechazo si no cumple","Rechazar productos sin registro INVIMA, vencidos, con embalaje dañado, sin cadena de frío cuando se requería, o de origen dudoso"),
                 ("Registro","Registrar la recepción en FOR-MED-001 con: fecha, producto, lote, vencimiento, cantidad, proveedor, número de factura"),
                 ("Almacenamiento inmediato","Almacenar inmediatamente en las condiciones correctas según el tipo de producto")]
    for i,(act,desc) in enumerate(pasos_recep):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r1=p.add_run(f"Paso {i+1} - {act}: "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(desc); r2.font.size=Pt(11)

    H(doc,"3.2 Almacenamiento de la Toxina Botulínica - Protocolo Específico",2)
    B(doc,"La toxina botulínica requiere manejo especial por ser un medicamento biológico termoestable:")
    tb=doc.add_table(rows=8,cols=2); tb.style='Table Grid'
    hrow(tb,["Etapa","Procedimiento específico"])
    etapas_tb=[("Al recibir","Verificar que llegó con indicadores de temperatura (si el proveedor los incluye). Colocar INMEDIATAMENTE en el refrigerador a 2-8°C. Registrar temperatura de llegada."),
               ("Almacenamiento","En refrigerador exclusivo a 2-8°C. NO congelar. NO colocar en la puerta del refrigerador (hay más variación de temperatura). Colocar en el centro de la nevera."),
               ("Monitoreo de temperatura","Verificar temperatura dos veces al día. Registrar en el libro de control. Alarma si supera 8°C."),
               ("Antes del uso","Sacar del refrigerador mínimo 15-30 minutos antes del uso (para que llegue a temperatura de confort para el paciente). No más de 1 hora antes de la reconstitución."),
               ("Reconstitución","Con solución salina 0.9% sin preservantes, jeringa de insulina 1ml. Según instrucciones del fabricante. Rotular el frasco con fecha y hora de reconstitución, concentración, nombre del producto."),
               ("Post-reconstitución","Usar idealmente el mismo día. Máximo 4 horas a temperatura ambiente o 24h refrigerada (verificar inserto del fabricante). No recongelar."),
               ("Sobrante","Inactivar con hipoclorito 0.5% antes de desechar. Desechar en guardianes o bolsa roja según el inserto del fabricante.")]
    for i,(e,p2) in enumerate(etapas_tb):
        tb.rows[i+1].cells[0].text=e; tb.rows[i+1].cells[1].text=p2
        for c in tb.rows[i+1].cells:
            for pp in c.paragraphs:
                for r in pp.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"3.3 Control Mensual de Medicamentos",2)
    B(doc,"El primer día hábil de cada mes se realiza la revisión mensual del inventario de medicamentos:")
    control_mensual=["Verificar que todos los medicamentos del botiquín están disponibles y en las cantidades mínimas requeridas",
                     "Revisar fecha de vencimiento de todos los productos: retirar los vencidos, marcar los que vencen en los próximos 3 meses",
                     "Verificar condiciones de almacenamiento: temperatura, humedad, protección de la luz",
                     "Revisar el registro de temperatura del refrigerador: verificar que no hubo incidentes de temperatura",
                     "Verificar el stock de insumos de procedimientos estéticos: agujas, jeringas, cánulas, rellenos, toxina",
                     "Actualizar el formato FOR-MED-001 con los hallazgos y las acciones realizadas",
                     "Realizar los pedidos de reposición necesarios"]
    for item in control_mensual: BL(doc,item)

    H(doc,"4. INDICADORES")
    tb2=doc.add_table(rows=4,cols=4); tb2.style='Table Grid'
    hrow(tb2,["Indicador","Fórmula","Meta","Frecuencia"])
    for i,row in enumerate([
        ("% control mensual medicamentos realizado","Controles realizados/Meses del período×100","100%","Mensual"),
        ("% medicamentos vencidos detectados","Medicamentos vencidos/Total medicamentos×100","0%","Mensual"),
        ("% incidentes de cadena de frío documentados y manejados","Incidentes documentados/Total incidentes×100","100%","Mensual")]):
        for j,v in enumerate(row): tb2.rows[i+1].cells[j].text=v
        for c in tb2.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph(); firmas(doc)
    path=os.path.join(BASE4,"PRO-MED-001_Proceso_Almacenamiento_Medicamentos.docx")
    doc.save(path); print(f"✓ Creado: {path}")

# =========== FOR-MED-001 ===========
def crear_for_med():
    doc=Document(); margins(doc); hf(doc,"FOR-MED-001","Control de Stock de Medicamentos")
    portada(doc,"FOR-MED-001","FORMATO DE CONTROL DE STOCK DE MEDICAMENTOS")

    H(doc,"SECCIÓN 1: DATOS DEL CONTROL")
    tb=doc.add_table(rows=3,cols=4); tb.style='Table Grid'
    datos=[("Mes/Año:","","Responsable:",""),
           ("Fecha de revisión:","","N° de revisión:",""),
           ("¿Todos los medicamentos del botiquín disponibles?","□ Sí  □ No","¿Todos en fecha?","□ Sí  □ No")]
    for i,row in enumerate(datos):
        for j,v in enumerate(row):
            tb.rows[i].cells[j].text=v
            if j%2==0 and v: tb.rows[i].cells[j].paragraphs[0].runs[0].font.bold=True
    doc.add_paragraph()

    H(doc,"SECCIÓN 2: INVENTARIO DE MEDICAMENTOS DEL BOTIQUÍN",2)
    cols=["Medicamento","Presentación","Stock mínimo","Stock actual","Lote","Fecha venc.","Estado","Acción"]
    meds_bot=["Adrenalina 1mg/1ml amp","Ampollas 1ml","2","","","","","",
              "Hidrocortisona 500mg vial","Viales","2","","","","","",
              "Difenhidramina iny","Ampollas","2","","","","","",
              "Nitroglicerina sublingual","Frasco spray","1","","","","","",
              "Dexametasona 4mg/ml","Ampollas","2","","","","","",
              "Glucosa 50% jeringa","Jeringas 50ml","1","","","","","",
              "Metoclopramida 10mg","Ampollas","2","","","","","",
              "Diazepam 10mg iny","Ampollas","2","","","","","",
              "Solución salina 0.9% 250ml","Bolsas","2","","","","","",
              "Atropina 1mg","Ampollas","2","","","","","",
              "Salbutamol inhalador","Inhalador","1","","","","","",
              "Furosemida 20mg","Ampollas","2","","","","","",
              "ASA 300mg tabletas","Tabletas","5","","","","",""]
    tb2=doc.add_table(rows=len(meds_bot)//8+1,cols=len(cols)); tb2.style='Table Grid'
    hrow(tb2,cols)
    meds_list=["Adrenalina 1mg/1ml ampolla","Hidrocortisona 500mg vial","Difenhidramina inyectable",
               "Nitroglicerina sublingual spray","Dexametasona 4mg/ml iny","Glucosa 50% jeringa 50ml",
               "Metoclopramida 10mg iny","Diazepam 10mg iny","Solución salina 0.9% 250ml",
               "Atropina 1mg iny","Salbutamol inhalador","Furosemida 20mg iny","ASA 300mg tabletas"]
    mins=["2","2","2","1","2","1","2","2","2","2","1","2","5"]
    for i,med in enumerate(meds_list):
        if i+1 < len(tb2.rows):
            tb2.rows[i+1].cells[0].text=med; tb2.rows[i+1].cells[2].text=mins[i]
            for c in tb2.rows[i+1].cells:
                for p in c.paragraphs:
                    for r in p.runs: r.font.size=Pt(8)

    # Rearrange: build fresh table
    tb2b=doc.add_table(rows=len(meds_list)+1,cols=len(cols)); tb2b.style='Table Grid'
    hrow(tb2b,cols)
    for i,med in enumerate(meds_list):
        tb2b.rows[i+1].cells[0].text=med; tb2b.rows[i+1].cells[2].text=mins[i]
        tb2b.rows[i+1].cells[7].text="□ OK  □ Pedir  □ Retirar"
        for c in tb2b.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(8)
    doc.add_paragraph()

    H(doc,"SECCIÓN 3: MEDICAMENTOS PARA PROCEDIMIENTOS ESTÉTICOS",2)
    cols3=["Producto","Tipo","Stock actual","Lote","Fecha venc.","Registro INVIMA","Estado"]
    prods_est=["Toxina Botulínica Tipo A (nombre comercial)","Ácido Hialurónico relleno (nombre comercial)","Ácido Hialurónico hidratación (si aplica)","Ácido glicólico para peeling","TCA (ácido tricloroacético) ≤20%","Hialuronidasa (antídoto)","Cócteles de mesoterapia (especificar)","Tubos para PRP"]
    tb3=doc.add_table(rows=len(prods_est)+1,cols=len(cols3)); tb3.style='Table Grid'
    hrow(tb3,cols3)
    for i,prod in enumerate(prods_est):
        tb3.rows[i+1].cells[0].text=prod
        tb3.rows[i+1].cells[6].text="□ OK  □ Pedir  □ Retirar"
        for c in tb3.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(8)
    doc.add_paragraph()

    H(doc,"SECCIÓN 4: REGISTRO DE TEMPERATURA DEL REFRIGERADOR (TOXINA BOTULÍNICA)",2)
    cols4=["Fecha","Temp. AM (°C)","Temp. PM (°C)","¿Dentro de 2-8°C?","Observaciones","Firma"]
    tb4=doc.add_table(rows=32,cols=len(cols4)); tb4.style='Table Grid'
    hrow(tb4,cols4)
    for i in range(1,32):
        tb4.rows[i].cells[3].text="□ Sí  □ No"
        for c in tb4.rows[i].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(8)
    doc.add_paragraph()

    H(doc,"SECCIÓN 5: REGISTRO DE ENTRADAS Y SALIDAS DE MEDICAMENTOS",2)
    cols5=["Fecha","Medicamento/Dispositivo","Lote","Vencimiento","Cantidad entrada","Cantidad salida","Stock resultante","Motivo","Firma"]
    tb5=doc.add_table(rows=16,cols=len(cols5)); tb5.style='Table Grid'
    hrow(tb5,cols5)
    for i in range(1,16):
        for c in tb5.rows[i].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(7)
    doc.add_paragraph()

    doc.add_paragraph("Revisión realizada por: [NOMBRE DE LA MÉDICA]\n\nFirma: _________________________\nFecha: _________________________")

    path=os.path.join(BASE4,"FOR-MED-001_Control_Stock_Medicamentos.docx")
    doc.save(path); print(f"✓ Creado: {path}")

if __name__=="__main__":
    print("Generando CARPETA 4 - MEDICAMENTOS...")
    crear_man_med(); crear_pro_med(); crear_for_med()
    print("✅ Carpeta 4 completada.")
