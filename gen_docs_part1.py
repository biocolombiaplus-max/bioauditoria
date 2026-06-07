#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera PRO-PP-005 PAMEC y PRO-PP-006 Residuos Hospitalarios"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

DARK_BLUE = RGBColor(0, 51, 102)
LIGHT_BLUE = RGBColor(189, 215, 238)
HEADER_GRAY = RGBColor(214, 220, 229)
OUT_DIR_5 = "/home/user/bioauditoria/documentos_habilitacion/CARPETA_5_PROCESOS_PRIORITARIOS"

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table_header(table, headers, bg='003366'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255,255,255)
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_portada(doc, titulo, codigo, descripcion):
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[NOMBRE DEL CONSULTORIO]")
    run.bold = True; run.font.size = Pt(18); run.font.color.rgb = DARK_BLUE; run.font.name = 'Calibri'

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Consultorio Médico General y Procedimientos Estéticos No Invasivos")
    r2.font.size = Pt(13); r2.font.color.rgb = DARK_BLUE; r2.font.name = 'Calibri'

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(titulo)
    r3.bold = True; r3.font.size = Pt(16); r3.font.name = 'Calibri'

    doc.add_paragraph()
    t = doc.add_table(rows=5, cols=2)
    t.style = 'Table Grid'
    datos = [
        ("Código:", codigo),
        ("Versión:", "1.0"),
        ("Fecha:", "Junio 2025"),
        ("Responsable:", "[NOMBRE DE LA MÉDICA]"),
        ("Dirección:", "[DIRECCIÓN DEL CONSULTORIO]"),
    ]
    for i, (k,v) in enumerate(datos):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        set_cell_bg(t.rows[i].cells[0], 'BDD7EE')
        for cell in t.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Calibri'; run.font.size = Pt(11)
    doc.add_page_break()

def add_control_versiones(doc):
    add_heading(doc, "CONTROL DE VERSIONES", 1)
    t = doc.add_table(rows=3, cols=5)
    t.style = 'Table Grid'
    add_table_header(t, ["Versión","Fecha","Descripción del Cambio","Elaboró","Aprobó"])
    data = [
        ("1.0","Junio 2025","Creación del documento","[NOMBRE DE LA MÉDICA]","[NOMBRE DE LA MÉDICA]"),
    ]
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            cell = t.rows[i+1].cells[j]
            cell.text = val
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2.5)
    return doc

def add_header_footer(doc, titulo, codigo):
    from docx.oxml.ns import qn
    sec = doc.sections[0]
    header = sec.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(f"{titulo} | {codigo}")
    run.font.size = Pt(9); run.font.name = 'Calibri'; run.font.color.rgb = DARK_BLUE

    footer = sec.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = fp.add_run("Versión 1.0 | Junio 2025 | Página ")
    run2.font.size = Pt(9); run2.font.name = 'Calibri'
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3 = fp.add_run()
    run3._r.append(fldChar1); run3._r.append(instrText); run3._r.append(fldChar2)
    run3.font.size = Pt(9); run3.font.name = 'Calibri'

# =====================================================================
# DOCUMENTO 1: PRO-PP-005 PAMEC
# =====================================================================
def gen_pamec():
    doc = setup_doc()
    titulo = "PROGRAMA DE AUDITORÍA PARA EL MEJORAMIENTO DE LA CALIDAD (PAMEC)"
    codigo = "PRO-PP-005"
    add_portada(doc, titulo, codigo, "")
    add_header_footer(doc, titulo, codigo)
    add_control_versiones(doc)

    add_heading(doc, "TABLA DE CONTENIDO", 1)
    toc_items = [
        "1. Objetivo", "2. Alcance", "3. Marco Legal", "4. Definiciones",
        "5. Autoevaluación de Calidad", "6. Selección de Prioridades",
        "7. Definición de Calidad Esperada", "8. Medición de Calidad Observada",
        "9. Plan de Mejoramiento", "10. Seguimiento e Indicadores",
        "11. Indicadores Obligatorios", "12. Responsables", "13. Registros"
    ]
    for item in toc_items:
        add_body(doc, item)
    doc.add_page_break()

    add_heading(doc, "1. OBJETIVO", 1)
    add_body(doc, "Implementar el Programa de Auditoría para el Mejoramiento de la Calidad (PAMEC) en el consultorio médico [NOMBRE DEL CONSULTORIO], con el fin de mejorar continuamente la calidad de los servicios de salud prestados, mediante procesos de autoevaluación, identificación de brechas, implementación de planes de mejora y seguimiento sistemático de resultados, en cumplimiento del Sistema Obligatorio de Garantía de Calidad en Salud (SOGCS).")

    add_heading(doc, "2. ALCANCE", 1)
    add_body(doc, "El presente PAMEC aplica a todos los procesos asistenciales y administrativos del consultorio [NOMBRE DEL CONSULTORIO], ubicado en [DIRECCIÓN DEL CONSULTORIO], que presta servicios de consulta externa de medicina general y procedimientos estéticos no invasivos. Involucra a todo el personal asistencial y administrativo del establecimiento.")

    add_heading(doc, "3. MARCO LEGAL", 1)
    normas = [
        ("Decreto 1011 de 2006", "Por el cual se establece el Sistema Obligatorio de Garantía de Calidad de la Atención de Salud del Sistema General de Seguridad Social en Salud (SOGCS)."),
        ("Resolución 1445 de 2006", "Por la cual se definen las funciones de la Entidad Acreditadora y se adoptan otras disposiciones. Define los procesos de auditoría para el mejoramiento de la calidad."),
        ("Resolución 0256 de 2016", "Por la cual se dictan disposiciones en relación con el Sistema de Información para la Calidad y se establecen los indicadores para el monitoreo de la calidad de los servicios de salud."),
        ("Resolución 2003 de 2014", "Por la cual se definen los procedimientos y condiciones de inscripción de los Prestadores de Servicios de Salud y de habilitación de servicios de salud."),
        ("Resolución 3100 de 2019", "Por la cual se definen los procedimientos y condiciones de inscripción de los prestadores de servicios de salud y de habilitación de servicios de salud."),
        ("Ley 100 de 1993", "Por la cual se crea el sistema de seguridad social integral y se dictan otras disposiciones."),
    ]
    t = doc.add_table(rows=len(normas)+1, cols=2)
    t.style = 'Table Grid'
    add_table_header(t, ["Norma", "Descripción"])
    for i, (n, d) in enumerate(normas):
        t.rows[i+1].cells[0].text = n
        t.rows[i+1].cells[1].text = d
        set_cell_bg(t.rows[i+1].cells[0], 'EBF3FB')
        for cell in t.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc, "4. DEFINICIONES", 1)
    defs = [
        ("PAMEC", "Programa de Auditoría para el Mejoramiento de la Calidad de la Atención de Salud. Es el mecanismo sistemático y continuo de evaluación y mejoramiento de la calidad observada respecto de la calidad esperada de la atención de salud."),
        ("Auditoría de calidad", "Proceso sistemático, independiente y documentado para obtener evidencias y evaluarlas de manera objetiva con el fin de determinar la extensión en que se cumplen los criterios de auditoría."),
        ("Calidad esperada", "Nivel de calidad que los prestadores de servicios de salud deben alcanzar, expresado en estándares definidos por el Sistema Obligatorio de Garantía de Calidad."),
        ("Calidad observada", "Nivel real de calidad que se identifica en la prestación de los servicios de salud, medido a través de indicadores."),
        ("Plan de mejoramiento", "Conjunto de acciones correctivas y preventivas orientadas a cerrar las brechas identificadas entre la calidad observada y la calidad esperada."),
        ("Indicador", "Medida cuantitativa que puede usarse como guía para controlar y valorar la calidad de las diferentes actividades."),
        ("Evento adverso", "Resultado clínico adverso que se deriva de la atención en salud y no de la enfermedad subyacente del paciente."),
        ("Oportunidad en la atención", "Posibilidad que tiene el usuario de obtener los servicios que requiere, sin que se presenten retrasos que pongan en riesgo su vida o su salud."),
    ]
    for term, defi in defs:
        p = doc.add_paragraph()
        run = p.add_run(f"{term}: ")
        run.bold = True; run.font.name = 'Calibri'; run.font.size = Pt(11)
        run2 = p.add_run(defi)
        run2.font.name = 'Calibri'; run2.font.size = Pt(11)

    add_heading(doc, "5. AUTOEVALUACIÓN DE CALIDAD", 1)
    add_body(doc, "La autoevaluación es el punto de partida del PAMEC. Consiste en la valoración interna que realiza el prestador sobre el cumplimiento de los estándares de calidad definidos en el SOGCS. Para el consultorio [NOMBRE DEL CONSULTORIO], la autoevaluación se realiza con periodicidad semestral y abarca los siguientes componentes:")

    add_heading(doc, "5.1 Metodología de Autoevaluación", 2)
    add_body(doc, "La autoevaluación se realiza mediante la aplicación de listas de chequeo basadas en los estándares de habilitación (Resolución 3100 de 2019) y los estándares de acreditación. El proceso incluye:")
    pasos = [
        "Revisión documental: verificación de la existencia y actualización de protocolos, guías y manuales.",
        "Observación directa: evaluación de las condiciones físicas, de equipos y de procesos en el punto de atención.",
        "Entrevistas al personal: verificación del conocimiento y aplicación de los procedimientos.",
        "Revisión de historias clínicas: evaluación del cumplimiento de estándares de registro.",
        "Medición de indicadores: cálculo de indicadores de calidad del período evaluado.",
        "Encuestas de satisfacción: aplicación de instrumentos de medición de satisfacción del usuario.",
    ]
    for p_item in pasos:
        pb = doc.add_paragraph(p_item, style='List Bullet')
        for run in pb.runs:
            run.font.name = 'Calibri'; run.font.size = Pt(11)

    add_heading(doc, "5.2 Tabla de Autoevaluación por Componentes", 2)
    headers_ae = ["Componente", "Criterio Evaluado", "Puntaje Máx.", "Puntaje Obtenido", "% Cumplimiento", "Brecha"]
    comp_data = [
        ("Talento Humano", "Perfil, competencias, certificaciones", "100", "", "", ""),
        ("Infraestructura", "Condiciones físicas, señalización", "100", "", "", ""),
        ("Dotación", "Equipos, mantenimiento, calibración", "100", "", "", ""),
        ("Medicamentos", "Almacenamiento, control, cadena de frío", "100", "", "", ""),
        ("Procesos Asistenciales", "Historia clínica, consentimientos", "100", "", "", ""),
        ("Seguridad del Paciente", "Barreras de seguridad, eventos adversos", "100", "", "", ""),
        ("Gestión de Residuos", "PGIRH, separación en fuente", "100", "", "", ""),
        ("Satisfacción del Usuario", "Encuestas, PQR, tiempos de espera", "100", "", "", ""),
    ]
    t2 = doc.add_table(rows=len(comp_data)+1, cols=6)
    t2.style = 'Table Grid'
    add_table_header(t2, headers_ae)
    for i, row in enumerate(comp_data):
        for j, val in enumerate(row):
            t2.rows[i+1].cells[j].text = val
            for para in t2.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
        if i % 2 == 0:
            for j in range(6):
                set_cell_bg(t2.rows[i+1].cells[j], 'F2F7FB')
    doc.add_paragraph()

    add_heading(doc, "6. SELECCIÓN DE PRIORIDADES", 1)
    add_body(doc, "Con base en los resultados de la autoevaluación, se seleccionan las prioridades de mejoramiento considerando los siguientes criterios:")
    criterios = [
        "Magnitud: frecuencia e impacto del problema en la calidad de la atención.",
        "Trascendencia: importancia del problema para los usuarios y para la organización.",
        "Factibilidad: posibilidad real de intervención con los recursos disponibles.",
        "Vulnerabilidad: susceptibilidad del problema a ser mejorado.",
        "Obligatoriedad: requerimientos normativos que deben cumplirse.",
    ]
    for c in criterios:
        pb = doc.add_paragraph(c, style='List Bullet')
        for run in pb.runs:
            run.font.name = 'Calibri'; run.font.size = Pt(11)

    add_heading(doc, "6.1 Matriz de Priorización", 2)
    t3 = doc.add_table(rows=7, cols=7)
    t3.style = 'Table Grid'
    add_table_header(t3, ["Problema Identificado","Magnitud\n(1-5)","Trascend.\n(1-5)","Factib.\n(1-5)","Vulnerab.\n(1-5)","Puntaje\nTotal","Prioridad"])
    problemas = [
        ("Oportunidad en la asignación de citas","","","","","",""),
        ("Eventos adversos en procedimientos","","","","","",""),
        ("Satisfacción del usuario < 90%","","","","","",""),
        ("Cumplimiento de historia clínica","","","","","",""),
        ("Gestión de residuos hospitalarios","","","","","",""),
        ("Competencias del personal","","","","","",""),
    ]
    for i, row in enumerate(problemas):
        for j, val in enumerate(row):
            t3.rows[i+1].cells[j].text = val
            for para in t3.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
    doc.add_paragraph()

    add_heading(doc, "7. DEFINICIÓN DE CALIDAD ESPERADA", 1)
    add_body(doc, "La calidad esperada se expresa en metas concretas y medibles para cada uno de los indicadores priorizados. Para el consultorio [NOMBRE DEL CONSULTORIO] se establecen las siguientes metas de calidad esperada:")

    t4 = doc.add_table(rows=8, cols=5)
    t4.style = 'Table Grid'
    add_table_header(t4, ["Indicador","Fórmula de Cálculo","Meta Esperada","Periodicidad","Fuente"])
    indicadores_ce = [
        ("Oportunidad en consulta médica general","Días entre solicitud y atención / Total citas","≤ 3 días","Mensual","Agenda de citas"),
        ("Satisfacción global del usuario","Usuarios satisfechos / Total encuestados x 100","≥ 90%","Trimestral","Encuesta de satisfacción"),
        ("Tasa de eventos adversos","N° EA / Total consultas x 1000","< 1 por 1000","Mensual","Reporte de EA"),
        ("Cumplimiento de historia clínica","HC completas / Total HC revisadas x 100","≥ 95%","Semestral","Auditoría de HC"),
        ("Oportunidad cirugía/procedimiento electivo","Días entre solicitud y realización","≤ 30 días","Mensual","Registro de procedimientos"),
        ("Tasa de infección en procedimientos","N° infecciones / Total procedimientos x 100","0%","Mensual","Seguimiento postprocedimiento"),
        ("Cumplimiento del PGIRH","Criterios cumplidos / Total criterios x 100","≥ 95%","Semestral","Lista de chequeo PGIRH"),
    ]
    for i, row in enumerate(indicadores_ce):
        for j, val in enumerate(row):
            t4.rows[i+1].cells[j].text = val
            for para in t4.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
        if i % 2 == 0:
            for j in range(5):
                set_cell_bg(t4.rows[i+1].cells[j], 'EBF3FB')
    doc.add_paragraph()

    add_heading(doc, "8. MEDICIÓN DE CALIDAD OBSERVADA", 1)
    add_body(doc, "La medición de calidad observada se realiza a través del monitoreo continuo de indicadores, auditorías de procesos y análisis de eventos adversos. Los resultados se comparan con la calidad esperada para identificar brechas y generar acciones de mejora.")

    add_heading(doc, "8.1 Proceso de Medición", 2)
    add_body(doc, "El proceso de medición incluye las siguientes etapas:")
    etapas = [
        "Recolección de datos: el personal asistencial registra la información en los formatos establecidos.",
        "Consolidación: mensualmente se consolidan los datos en la hoja de cálculo de indicadores.",
        "Cálculo: se aplican las fórmulas de cada indicador.",
        "Análisis: se comparan los resultados con las metas y el período anterior.",
        "Reporte: los resultados se presentan en la reunión de análisis de calidad.",
        "Acciones: se generan acciones correctivas cuando no se cumplen las metas.",
    ]
    for e in etapas:
        pb = doc.add_paragraph(e, style='List Bullet')
        for run in pb.runs:
            run.font.name = 'Calibri'; run.font.size = Pt(11)

    add_heading(doc, "8.2 Registro de Medición de Indicadores", 2)
    t5 = doc.add_table(rows=5, cols=8)
    t5.style = 'Table Grid'
    add_table_header(t5, ["Indicador","Ene","Feb","Mar","Abr","May","Jun","Promedio\nSemestre"])
    ind_med = [
        ("Oportunidad consulta (días)","","","","","","",""),
        ("Satisfacción usuario (%)","","","","","","",""),
        ("Eventos adversos (tasa)","","","","","","",""),
        ("Cumplimiento HC (%)","","","","","","",""),
    ]
    for i, row in enumerate(ind_med):
        for j, val in enumerate(row):
            t5.rows[i+1].cells[j].text = val
            for para in t5.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
    doc.add_paragraph()

    add_heading(doc, "9. PLAN DE MEJORAMIENTO", 1)
    add_body(doc, "El plan de mejoramiento es el conjunto de acciones correctivas y preventivas que se implementan para cerrar las brechas identificadas entre la calidad observada y la calidad esperada. Se estructura con base en el ciclo PHVA (Planear, Hacer, Verificar, Actuar).")

    add_heading(doc, "9.1 Tabla de Acciones Correctivas y Preventivas", 2)
    t6 = doc.add_table(rows=9, cols=8)
    t6.style = 'Table Grid'
    add_table_header(t6, ["N°","Problema/Brecha Identificada","Causa Raíz","Acción de Mejora","Responsable","Fecha\nInicio","Fecha\nFin","Estado"])
    acciones = [
        ("1","Oportunidad en consulta > 3 días","Alta demanda / agenda insuficiente","Ampliar horarios de atención y optimizar agenda","[NOMBRE DE LA MÉDICA]","","","Pendiente"),
        ("2","Satisfacción < 90%","Tiempos de espera prolongados","Implementar sistema de gestión de citas en línea","[NOMBRE DE LA MÉDICA]","","","Pendiente"),
        ("3","Registro incompleto HC","Falta de tiempo / no hay hábito","Capacitación en diligenciamiento y auditoría mensual","[NOMBRE DE LA MÉDICA]","","","Pendiente"),
        ("4","Sin reporte de eventos adversos","Desconocimiento del proceso","Capacitación en reporte y análisis de EA","[NOMBRE DE LA MÉDICA]","","","Pendiente"),
        ("5","PGIRH incompleto","Falta de capacitación en segregación","Jornada de capacitación en gestión de residuos","[NOMBRE DE LA MÉDICA]","","","Pendiente"),
        ("6","Equipos sin mantenimiento documentado","No hay cronograma formal","Elaborar e implementar cronograma de mantenimiento","[NOMBRE DE LA MÉDICA]","","","Pendiente"),
        ("7","Personal sin actualización","No hay plan de capacitación","Elaborar e implementar plan anual de capacitación","[NOMBRE DE LA MÉDICA]","","","Pendiente"),
        ("8","Sin indicadores de procedimientos estéticos","No hay sistema de medición","Diseñar e implementar fichas técnicas de indicadores","[NOMBRE DE LA MÉDICA]","","","Pendiente"),
    ]
    for i, row in enumerate(acciones):
        for j, val in enumerate(row):
            t6.rows[i+1].cells[j].text = val
            for para in t6.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(8)
        if i % 2 == 0:
            for j in range(8):
                set_cell_bg(t6.rows[i+1].cells[j], 'F2F7FB')
    doc.add_paragraph()

    add_heading(doc, "9.2 Ciclo PHVA del Plan de Mejoramiento", 2)
    fases_phva = [
        ("PLANEAR", "Identificar los problemas prioritarios, analizar sus causas raíz, definir las acciones de mejora, establecer metas, responsables y cronograma."),
        ("HACER", "Implementar las acciones de mejora planificadas según el cronograma establecido. Documentar las intervenciones realizadas."),
        ("VERIFICAR", "Medir los resultados de las acciones implementadas. Comparar con las metas establecidas. Evaluar la efectividad de las intervenciones."),
        ("ACTUAR", "Si las acciones fueron efectivas, estandarizarlas y documentarlas. Si no fueron efectivas, reiniciar el ciclo con nuevas acciones."),
    ]
    t7 = doc.add_table(rows=5, cols=2)
    t7.style = 'Table Grid'
    add_table_header(t7, ["Fase PHVA", "Actividades"])
    colores_phva = ['003366', '1F497D', '4472C4', '70AD47']
    for i, (fase, act) in enumerate(fases_phva):
        t7.rows[i+1].cells[0].text = fase
        t7.rows[i+1].cells[1].text = act
        set_cell_bg(t7.rows[i+1].cells[0], 'BDD7EE')
        for cell in t7.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc, "10. SEGUIMIENTO E INDICADORES", 1)
    add_body(doc, "El seguimiento del PAMEC se realiza a través de la medición periódica de indicadores y la verificación del cumplimiento de las acciones de mejora. Se establece una reunión trimestral de análisis de calidad donde se revisan los resultados y se toman decisiones.")

    add_heading(doc, "10.1 Cronograma de Seguimiento", 2)
    t8 = doc.add_table(rows=7, cols=5)
    t8.style = 'Table Grid'
    add_table_header(t8, ["Actividad","Ene-Mar","Abr-Jun","Jul-Sep","Oct-Dic"])
    crono = [
        ("Medición de indicadores mensuales","X","X","X","X"),
        ("Reunión análisis de calidad","X","X","X","X"),
        ("Autoevaluación semestral","","X","","X"),
        ("Auditoría interna de historias clínicas","","X","","X"),
        ("Encuesta de satisfacción","","X","","X"),
        ("Informe de gestión PAMEC","","X","","X"),
    ]
    for i, row in enumerate(crono):
        for j, val in enumerate(row):
            t8.rows[i+1].cells[j].text = val
            for para in t8.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

    add_heading(doc, "11. INDICADORES OBLIGATORIOS PARA MEDICINA GENERAL", 1)
    add_body(doc, "De conformidad con la Resolución 0256 de 2016, los prestadores de servicios de salud deben reportar al sistema de información para la calidad (SIVIGILA/SISPRO) los siguientes indicadores obligatorios:")

    add_heading(doc, "11.1 Indicador de Oportunidad en la Consulta Médica General", 2)
    t9 = doc.add_table(rows=7, cols=2)
    t9.style = 'Table Grid'
    ficha_op = [
        ("Nombre del indicador:", "Oportunidad en la asignación de cita en la consulta médica general"),
        ("Definición:", "Número de días calendario que transcurren entre el día que el usuario solicita una cita para consulta médica general y el día que es atendido"),
        ("Fórmula:", "Sumatoria de días entre solicitud y atención de todos los usuarios / Número total de usuarios atendidos en el período"),
        ("Unidad de medida:", "Días"),
        ("Periodicidad de reporte:", "Trimestral"),
        ("Meta:", "≤ 3 días calendario"),
        ("Fuente de información:", "Libro/sistema de registro de citas"),
    ]
    for i, (k, v) in enumerate(ficha_op):
        t9.rows[i].cells[0].text = k
        t9.rows[i].cells[1].text = v
        set_cell_bg(t9.rows[i].cells[0], 'BDD7EE')
        for cell in t9.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc, "11.2 Indicador de Satisfacción Global del Usuario", 2)
    t10 = doc.add_table(rows=7, cols=2)
    t10.style = 'Table Grid'
    ficha_sat = [
        ("Nombre del indicador:", "Satisfacción global del usuario de consulta externa"),
        ("Definición:", "Proporción de usuarios que manifiestan estar satisfechos o muy satisfechos con la atención recibida en el consultorio"),
        ("Fórmula:", "(Número de usuarios satisfechos o muy satisfechos / Total de usuarios encuestados) x 100"),
        ("Unidad de medida:", "Porcentaje (%)"),
        ("Periodicidad de reporte:", "Trimestral"),
        ("Meta:", "≥ 90%"),
        ("Fuente de información:", "Encuesta de satisfacción del usuario"),
    ]
    for i, (k, v) in enumerate(ficha_sat):
        t10.rows[i].cells[0].text = k
        t10.rows[i].cells[1].text = v
        set_cell_bg(t10.rows[i].cells[0], 'BDD7EE')
        for cell in t10.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc, "11.3 Indicador de Eventos Adversos", 2)
    t11 = doc.add_table(rows=8, cols=2)
    t11.style = 'Table Grid'
    ficha_ea = [
        ("Nombre del indicador:", "Tasa de eventos adversos en consulta externa"),
        ("Definición:", "Número de eventos adversos presentados en la prestación de servicios de consulta externa y procedimientos, por cada 1000 consultas/procedimientos realizados"),
        ("Fórmula:", "(Número de eventos adversos reportados / Total de consultas y procedimientos realizados) x 1000"),
        ("Unidad de medida:", "Tasa por 1000 consultas/procedimientos"),
        ("Periodicidad de reporte:", "Mensual"),
        ("Meta:", "< 1 evento adverso por 1000 consultas/procedimientos"),
        ("Fuente de información:", "Formato de reporte de eventos adversos"),
        ("Nota:", "Incluye eventos adversos en procedimientos de medicina estética no invasiva"),
    ]
    for i, (k, v) in enumerate(ficha_ea):
        t11.rows[i].cells[0].text = k
        t11.rows[i].cells[1].text = v
        set_cell_bg(t11.rows[i].cells[0], 'BDD7EE')
        for cell in t11.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc, "12. RESPONSABLES", 1)
    t12 = doc.add_table(rows=5, cols=3)
    t12.style = 'Table Grid'
    add_table_header(t12, ["Cargo / Responsable", "Responsabilidad en el PAMEC", "Frecuencia"])
    responsables = [
        ("[NOMBRE DE LA MÉDICA]\nMédica General - Directora", "Liderar el PAMEC, aprobar el plan de mejoramiento, presentar informes de calidad, tomar decisiones basadas en indicadores", "Permanente"),
        ("Personal asistencial", "Reportar eventos adversos, diligenciar correctamente la historia clínica, participar en actividades de mejora", "Permanente"),
        ("Personal administrativo", "Aplicar encuestas de satisfacción, registrar PQR, medir oportunidad en la atención", "Mensual"),
    ]
    for i, row in enumerate(responsables):
        for j, val in enumerate(row):
            t12.rows[i+1].cells[j].text = val
            for para in t12.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc, "13. REGISTROS", 1)
    t13 = doc.add_table(rows=8, cols=4)
    t13.style = 'Table Grid'
    add_table_header(t13, ["Nombre del Registro", "Código", "Tiempo de Retención", "Responsable Custodia"])
    registros = [
        ("Hoja de cálculo de indicadores de calidad", "FOR-GC-001", "5 años", "[NOMBRE DE LA MÉDICA]"),
        ("Acta de reunión de análisis de calidad", "FOR-GC-002", "5 años", "[NOMBRE DE LA MÉDICA]"),
        ("Encuesta de satisfacción del usuario", "FOR-GC-003", "2 años", "[NOMBRE DE LA MÉDICA]"),
        ("Formato de reporte de eventos adversos", "FOR-GC-004", "5 años", "[NOMBRE DE LA MÉDICA]"),
        ("Plan de mejoramiento", "FOR-GC-005", "5 años", "[NOMBRE DE LA MÉDICA]"),
        ("Informe semestral del PAMEC", "FOR-GC-006", "5 años", "[NOMBRE DE LA MÉDICA]"),
        ("Lista de chequeo autoevaluación", "FOR-GC-007", "5 años", "[NOMBRE DE LA MÉDICA]"),
    ]
    for i, row in enumerate(registros):
        for j, val in enumerate(row):
            t13.rows[i+1].cells[j].text = val
            for para in t13.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
        if i % 2 == 0:
            for j in range(4):
                set_cell_bg(t13.rows[i+1].cells[j], 'EBF3FB')

    path = os.path.join(OUT_DIR_5, "PRO-PP-005_PAMEC_Programa_Auditoria_Mejoramiento_Calidad.docx")
    doc.save(path)
    print(f"Guardado: {path}")

# =====================================================================
# DOCUMENTO 2: PRO-PP-006 Residuos Hospitalarios
# =====================================================================
def gen_residuos():
    doc = setup_doc()
    titulo = "PLAN DE GESTIÓN INTEGRAL DE RESIDUOS HOSPITALARIOS Y SIMILARES (PGIRHS)"
    codigo = "PRO-PP-006"
    add_portada(doc, titulo, codigo, "")
    add_header_footer(doc, titulo, codigo)
    add_control_versiones(doc)

    add_heading(doc, "TABLA DE CONTENIDO", 1)
    toc = ["1. Objetivo","2. Alcance","3. Marco Legal","4. Definiciones","5. Clasificación de Residuos",
           "6. Caracterización de Residuos","7. Separación en la Fuente","8. Almacenamiento Interno",
           "9. Ruta Sanitaria Interna","10. Recolección y Transporte Externo","11. Registro PGIRHS",
           "12. Indicadores de Generación","13. Capacitación del Personal","14. Responsables","15. Registros"]
    for item in toc:
        add_body(doc, item)
    doc.add_page_break()

    add_heading(doc, "1. OBJETIVO", 1)
    add_body(doc, "Establecer el Plan de Gestión Integral de Residuos Hospitalarios y Similares (PGIRHS) para el consultorio [NOMBRE DEL CONSULTORIO], con el fin de garantizar el manejo adecuado, seguro y ambientalmente responsable de todos los residuos generados en la prestación de servicios de salud, en cumplimiento de la normatividad colombiana vigente, protegiendo la salud del personal, los usuarios y el medio ambiente.")

    add_heading(doc, "2. ALCANCE", 1)
    add_body(doc, "El presente plan aplica a todos los residuos generados en el consultorio [NOMBRE DEL CONSULTORIO], ubicado en [DIRECCIÓN DEL CONSULTORIO], que incluye los servicios de consulta externa de medicina general y procedimientos estéticos no invasivos (aplicación de toxina botulínica, ácido hialurónico, mesoterapia, entre otros). Aplica a todo el personal que labora en el establecimiento.")

    add_heading(doc, "3. MARCO LEGAL", 1)
    normas = [
        ("Decreto 351 de 2014","Por el cual se reglamenta la gestión integral de los residuos generados en la atención en salud y otras actividades. Deroga parcialmente el Decreto 2676 de 2000."),
        ("Resolución 1164 de 2002","Por la cual se adopta el Manual de Procedimientos para la Gestión Integral de los Residuos Hospitalarios y Similares (MPGIRH)."),
        ("Decreto 1076 de 2015","Decreto Único Reglamentario del Sector Ambiente y Desarrollo Sostenible. Compila las normas de residuos peligrosos."),
        ("Resolución 2309 de 1986","Sobre residuos especiales, incluyendo los hospitalarios."),
        ("Ley 9 de 1979","Código Sanitario Nacional. Establece medidas sanitarias sobre manejo de residuos."),
        ("Decreto 4741 de 2005","Por el cual se reglamenta parcialmente la prevención y el manejo de los residuos o desechos peligrosos generados en el marco de la gestión integral."),
        ("Resolución 1402 de 2006","Por la cual se desarrolla parcialmente el Decreto 4741 de 2005 en materia de residuos peligrosos."),
    ]
    t = doc.add_table(rows=len(normas)+1, cols=2)
    t.style = 'Table Grid'
    add_table_header(t, ["Norma","Descripción"])
    for i, (n, d) in enumerate(normas):
        t.rows[i+1].cells[0].text = n
        t.rows[i+1].cells[1].text = d
        set_cell_bg(t.rows[i+1].cells[0], 'EBF3FB')
        for cell in t.rows[i+1].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc, "4. DEFINICIONES", 1)
    defs = [
        ("Residuos hospitalarios y similares","Objetos, elementos o sustancias que se abandonan, botan, desechan, descartan o rechazan y que se generan durante la ejecución o mantenimiento de actividades médicas."),
        ("Residuos peligrosos","Residuo o desecho que por sus características corrosivas, reactivas, explosivas, tóxicas, inflamables, infecciosas o radiactivas puede causar riesgo o daño para la salud humana y el ambiente."),
        ("Biosanitarios","Residuos generados en el diagnóstico, tratamiento, prestación de servicios de salud o en la investigación que han tenido contacto con materia orgánica, sangre o fluidos corporales."),
        ("Cortopunzantes","Residuos provenientes de elementos cortantes o punzantes que han sido usados en la atención a humanos o animales."),
        ("PGIRHS","Plan de Gestión Integral de Residuos Hospitalarios y Similares: instrumento de planeación que promueve prácticas responsables con el entorno."),
        ("Separación en la fuente","Clasificación de los residuos en el sitio donde se generan."),
        ("Ruta sanitaria","Trayecto que recorre el personal de aseo para la recolección interna de residuos."),
        ("Gestor externo","Empresa autorizada por la autoridad ambiental competente para la recolección, transporte, tratamiento y disposición final de residuos hospitalarios."),
    ]
    for term, defi in defs:
        p = doc.add_paragraph()
        run = p.add_run(f"{term}: ")
        run.bold = True; run.font.name = 'Calibri'; run.font.size = Pt(11)
        run2 = p.add_run(defi)
        run2.font.name = 'Calibri'; run2.font.size = Pt(11)

    add_heading(doc, "5. CLASIFICACIÓN DE RESIDUOS GENERADOS EN EL CONSULTORIO", 1)
    add_body(doc, "De acuerdo con el Decreto 351 de 2014 y el Manual MPGIRH, los residuos generados en el consultorio [NOMBRE DEL CONSULTORIO] se clasifican en:")

    add_heading(doc, "5.1 Residuos No Peligrosos", 2)
    t2 = doc.add_table(rows=4, cols=4)
    t2.style = 'Table Grid'
    add_table_header(t2, ["Tipo","Descripción","Ejemplos en el Consultorio","Color Recipiente"])
    rnp = [
        ("Biodegradables","Residuos naturales no peligrosos que se descomponen naturalmente","Residuos de alimentos, papel sucio con alimentos","Verde"),
        ("Reciclables","Residuos que no se descomponen fácilmente y pueden ser reutilizados","Papel, cartón, plástico limpio, vidrio, metales no contaminados","Blanco/Gris"),
        ("Ordinarios e inertes","Residuos que no ofrecen riesgo de infección ni contaminación","Papeles de empaque limpios, servilletas no contaminadas","Negro/Verde"),
    ]
    for i, row in enumerate(rnp):
        for j, val in enumerate(row):
            t2.rows[i+1].cells[j].text = val
            for para in t2.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
    doc.add_paragraph()

    add_heading(doc, "5.2 Residuos Peligrosos", 2)
    t3 = doc.add_table(rows=6, cols=4)
    t3.style = 'Table Grid'
    add_table_header(t3, ["Tipo","Descripción","Ejemplos en el Consultorio","Color Recipiente"])
    rp = [
        ("Biosanitarios","Con materia orgánica, sangre o fluidos corporales","Gasas, algodones, guantes usados en procedimientos, material con sangre","Rojo"),
        ("Cortopunzantes","Elementos que pueden causar cortes o pinchazos","Agujas, jeringas con aguja, lancetas, bisturís, ampolletas rotas","Guardián (rígido rojo)"),
        ("Anatomopatológicos","Tejidos, órganos, partes del cuerpo","Biopsias (si aplica), material de curetaje","Rojo"),
        ("Fármacos parcialmente consumidos","Medicamentos vencidos o parcialmente usados","Medicamentos vencidos, frascos con restos de medicamentos","Rojo"),
        ("Químicos","Sustancias o productos con características de peligrosidad","Desinfectantes concentrados, productos vencidos para estética","Rojo"),
    ]
    for i, row in enumerate(rp):
        for j, val in enumerate(row):
            t3.rows[i+1].cells[j].text = val
            for para in t3.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
        if i % 2 == 0:
            for j in range(4):
                set_cell_bg(t3.rows[i+1].cells[j], 'FCE4D6')
    doc.add_paragraph()

    add_heading(doc, "6. CARACTERIZACIÓN DE RESIDUOS", 1)
    add_body(doc, "La caracterización de residuos permite conocer la composición y cantidad de los residuos generados. Para el consultorio [NOMBRE DEL CONSULTORIO], se estima la siguiente generación mensual:")
    t4 = doc.add_table(rows=7, cols=5)
    t4.style = 'Table Grid'
    add_table_header(t4, ["Tipo de Residuo","Recipiente","Cantidad Estimada/Mes","Peso Estimado (kg/mes)","Gestor Externo"])
    caract = [
        ("Biosanitarios","Bolsa roja en recipiente con tapa","4-6 bolsas","2-4 kg","Empresa gestora autorizada"),
        ("Cortopunzantes","Guardián 2L","1-2 guardianes","0.5-1 kg","Empresa gestora autorizada"),
        ("Ordinarios","Bolsa negra/verde","8-12 bolsas","5-8 kg","Servicio público de aseo"),
        ("Reciclables","Bolsa blanca/gris","2-4 bolsas","1-3 kg","Reciclador/chatarrero autorizado"),
        ("Fármacos vencidos","Bolsa roja sellada","Variable","Variable","Empresa gestora autorizada"),
        ("Residuos químicos","Contenedor especial","Variable","Variable","Empresa gestora autorizada"),
    ]
    for i, row in enumerate(caract):
        for j, val in enumerate(row):
            t4.rows[i+1].cells[j].text = val
            for para in t4.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
    doc.add_paragraph()

    add_heading(doc, "7. SEPARACIÓN EN LA FUENTE", 1)
    add_body(doc, "La separación en la fuente es la primera y más importante etapa del manejo de residuos. Se realiza en el lugar donde se generan los residuos. El personal que genera el residuo es el responsable de su correcta separación.")

    add_heading(doc, "7.1 Sistema de Colores y Recipientes", 2)
    t5 = doc.add_table(rows=6, cols=5)
    t5.style = 'Table Grid'
    add_table_header(t5, ["Color","Tipo de Residuo","Recipiente","Bolsa","Ubicación en el Consultorio"])
    colores = [
        ("ROJO","Peligrosos infecciosos (biosanitarios, anatomopatológicos, fármacos)","Recipiente con tapa de pedal","Bolsa roja","Consultorio, área de procedimientos"),
        ("ROJO (Guardián)","Cortopunzantes","Guardián rígido homologado","N/A - recipiente rígido sellable","Junto al área de trabajo, a la mano"),
        ("NEGRO/VERDE","Ordinarios e inertes","Papelera con bolsa","Bolsa negra/verde","Sala de espera, recepción, baño"),
        ("BLANCO/GRIS","Reciclables (papel, cartón, plástico)","Contenedor con bolsa","Bolsa blanca","Área administrativa"),
        ("VERDE","Biodegradables","Contenedor","Bolsa verde","Cocina/comedor si aplica"),
    ]
    for i, row in enumerate(colores):
        for j, val in enumerate(row):
            t5.rows[i+1].cells[j].text = val
            for para in t5.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
    doc.add_paragraph()

    add_heading(doc, "7.2 Reglas de Separación", 2)
    reglas = [
        "NUNCA mezclar residuos peligrosos con ordinarios.",
        "NUNCA reencapuchar las agujas con las dos manos; usar técnica de una mano.",
        "NUNCA comprimir las bolsas de residuos con las manos.",
        "SIEMPRE descartar agujas con la jeringa completa en el guardián.",
        "SIEMPRE cerrar el guardián cuando esté 3/4 lleno (nunca llenarlo completamente).",
        "SIEMPRE sellar las bolsas de residuos cuando estén 3/4 llenas.",
        "Los medicamentos vencidos deben identificarse claramente y almacenarse separados hasta su entrega al gestor.",
        "Los frascos de toxina botulínica o ácido hialurónico usados son residuos peligrosos (biosanitarios o químicos).",
    ]
    for r in reglas:
        pb = doc.add_paragraph(r, style='List Bullet')
        for run in pb.runs:
            run.font.name = 'Calibri'; run.font.size = Pt(11)

    add_heading(doc, "8. ALMACENAMIENTO INTERNO", 1)
    add_body(doc, "El almacenamiento interno es el depósito temporal de los residuos en el lugar de generación o en el sitio destinado para tal fin dentro del consultorio, hasta su recolección para el almacenamiento central o su entrega al gestor externo.")

    add_heading(doc, "8.1 Condiciones del Área de Almacenamiento Interno", 2)
    condiciones = [
        "El área de almacenamiento debe estar claramente identificada y señalizada.",
        "Debe tener ventilación adecuada para evitar acumulación de olores y gases.",
        "El piso debe ser de material impermeable, lavable y con resistencia a los agentes desinfectantes.",
        "Debe contar con sistema de drenaje si se requiere lavado.",
        "Debe estar cerrada con llave para restringir el acceso a personal no autorizado.",
        "Debe contar con extintor cercano.",
        "Temperatura ambiente controlada (máx. 30°C).",
        "Los residuos peligrosos no deben permanecer más de 7 días en el almacenamiento interno.",
        "Los guardianes llenos deben almacenarse en bolsa roja y en contenedor rígido hasta su entrega al gestor.",
    ]
    for c in condiciones:
        pb = doc.add_paragraph(c, style='List Bullet')
        for run in pb.runs:
            run.font.name = 'Calibri'; run.font.size = Pt(11)

    add_heading(doc, "9. RUTA SANITARIA INTERNA", 1)
    add_body(doc, "La ruta sanitaria interna es el trayecto que recorre el personal encargado de la limpieza para recolectar los residuos de los diferentes puntos de generación y llevarlos al área de almacenamiento central o punto de entrega al gestor externo.")

    add_heading(doc, "9.1 Procedimiento de la Ruta Sanitaria", 2)
    pasos_ruta = [
        ("1","Preparación","El personal se equipa con los EPP: guantes de caucho gruesos, tapabocas, bata impermeable, botas."),
        ("2","Inicio del recorrido","Comienza por las áreas menos contaminadas (administrativa) y termina en las más contaminadas (área clínica/procedimientos)."),
        ("3","Recolección","Se retiran las bolsas cuando estén 3/4 llenas o al final de la jornada. Se cierra la bolsa con nudo doble. Se etiqueta con tipo de residuo, área, fecha y hora."),
        ("4","Transporte interno","Se usa carro de transporte exclusivo (no se usa para transporte de ropa limpia u otros). Las bolsas no se arrastran por el piso."),
        ("5","Depósito","Los residuos se llevan al área de almacenamiento interno separando peligrosos de no peligrosos."),
        ("6","Limpieza del carro","El carro de transporte se lava y desinfecta después de cada uso."),
        ("7","Disposición de EPP","Los guantes y otros EPP desechables se depositan en bolsa roja al finalizar."),
        ("8","Higiene de manos","Lavado de manos con agua y jabón tras retirar los guantes."),
    ]
    t6 = doc.add_table(rows=len(pasos_ruta)+1, cols=3)
    t6.style = 'Table Grid'
    add_table_header(t6, ["Paso","Actividad","Descripción"])
    for i, row in enumerate(pasos_ruta):
        for j, val in enumerate(row):
            t6.rows[i+1].cells[j].text = val
            for para in t6.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
        if i % 2 == 0:
            for j in range(3):
                set_cell_bg(t6.rows[i+1].cells[j], 'F2F7FB')
    doc.add_paragraph()

    add_heading(doc, "9.2 Frecuencia de la Ruta Sanitaria", 2)
    add_body(doc, "La ruta sanitaria se realiza con la siguiente frecuencia:")
    frecuencias = [
        "Residuos biosanitarios y cortopunzantes: al final de cada jornada de atención (mínimo 1 vez al día).",
        "Residuos ordinarios y reciclables: cada vez que los recipientes lleguen a 3/4 de su capacidad o al final de la jornada.",
        "Guardianes: cuando estén llenos en sus 3/4 partes.",
    ]
    for f in frecuencias:
        pb = doc.add_paragraph(f, style='List Bullet')
        for run in pb.runs:
            run.font.name = 'Calibri'; run.font.size = Pt(11)

    add_heading(doc, "10. RECOLECCIÓN Y TRANSPORTE EXTERNO", 1)
    add_body(doc, "La recolección y transporte externo de los residuos peligrosos debe realizarse por un gestor externo autorizado por la autoridad ambiental competente (Corporación Autónoma Regional o autoridad ambiental urbana).")

    add_heading(doc, "10.1 Requisitos del Gestor Externo", 2)
    requisitos = [
        "Licencia ambiental o autorización vigente de la autoridad ambiental.",
        "Plan de gestión aprobado por la autoridad ambiental.",
        "Vehículos de transporte autorizados para residuos peligrosos.",
        "Manifiesto de transporte de residuos peligrosos (diligenciado en cada recolección).",
        "Certificados de disposición final de los residuos recolectados.",
        "Personal capacitado en el manejo de residuos hospitalarios peligrosos.",
    ]
    for r in requisitos:
        pb = doc.add_paragraph(r, style='List Bullet')
        for run in pb.runs:
            run.font.name = 'Calibri'; run.font.size = Pt(11)

    add_heading(doc, "10.2 Documentos de Seguimiento", 2)
    t7 = doc.add_table(rows=5, cols=4)
    t7.style = 'Table Grid'
    add_table_header(t7, ["Documento","Responsable de Emisión","Quién lo Recibe","Tiempo de Archivo"])
    docs_seg = [
        ("Manifiesto de transporte de residuos peligrosos","Gestor externo","Consultorio (copia)","5 años"),
        ("Certificado de disposición final","Gestor externo","Consultorio","5 años"),
        ("Acta de entrega de residuos cortopunzantes","Gestor externo","Consultorio","5 años"),
        ("Pesaje de residuos (bitácora)","Consultorio","Consultorio","5 años"),
    ]
    for i, row in enumerate(docs_seg):
        for j, val in enumerate(row):
            t7.rows[i+1].cells[j].text = val
            for para in t7.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc, "11. REGISTRO PGIRHS", 1)
    add_body(doc, "El Registro del PGIRHS comprende todos los documentos que evidencian el cumplimiento del plan. El consultorio debe mantener los siguientes registros actualizados y disponibles para inspección de las autoridades sanitarias y ambientales:")

    t8 = doc.add_table(rows=9, cols=4)
    t8.style = 'Table Grid'
    add_table_header(t8, ["Registro","Descripción","Periodicidad","Responsable"])
    registros_pgirhs = [
        ("Diagnóstico de la situación actual","Evaluación inicial del manejo de residuos","Anual","[NOMBRE DE LA MÉDICA]"),
        ("Bitácora de generación de residuos","Registro mensual del peso/cantidad de residuos generados por tipo","Mensual","Personal de limpieza / [NOMBRE DE LA MÉDICA]"),
        ("Manifiestos de transporte","Copia de los manifiestos emitidos por el gestor externo en cada recolección","Cada recolección","Gestor externo (copia al consultorio)"),
        ("Certificados de disposición final","Emitidos por el gestor externo","Cada disposición","Gestor externo (copia al consultorio)"),
        ("Actas de capacitación en PGIRHS","Registro de capacitaciones realizadas al personal","Semestral","[NOMBRE DE LA MÉDICA]"),
        ("Informes de indicadores de generación","Análisis de indicadores de generación de residuos","Trimestral","[NOMBRE DE LA MÉDICA]"),
        ("Plan de contingencia de residuos","Procedimiento para emergencias relacionadas con residuos","Cuando ocurra/revisar anual","[NOMBRE DE LA MÉDICA]"),
        ("Contrato con gestor externo","Copia del contrato con la empresa recolectora autorizada","Vigencia del contrato","[NOMBRE DE LA MÉDICA]"),
    ]
    for i, row in enumerate(registros_pgirhs):
        for j, val in enumerate(row):
            t8.rows[i+1].cells[j].text = val
            for para in t8.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
        if i % 2 == 0:
            for j in range(4):
                set_cell_bg(t8.rows[i+1].cells[j], 'EBF3FB')
    doc.add_paragraph()

    add_heading(doc, "12. INDICADORES DE GENERACIÓN DE RESIDUOS", 1)
    add_body(doc, "Los indicadores de generación permiten hacer seguimiento a la cantidad de residuos generados, identificar tendencias y evaluar la efectividad de las medidas de minimización implementadas.")

    t9 = doc.add_table(rows=6, cols=5)
    t9.style = 'Table Grid'
    add_table_header(t9, ["Indicador","Fórmula","Unidad","Meta","Frecuencia"])
    indicadores = [
        ("Tasa de generación de residuos peligrosos","Kg residuos peligrosos / N° consultas realizadas","Kg/consulta","< 0.05 kg/consulta","Mensual"),
        ("Porcentaje de residuos peligrosos vs total","Kg residuos peligrosos / Kg totales x 100","%","< 20%","Mensual"),
        ("Generación de cortopunzantes","N° guardianes llenos / N° procedimientos con agujas","Guardianes/procedim.","Registro y tendencia","Mensual"),
        ("Cumplimiento separación en fuente","N° hallazgos de mezcla / Total verificaciones x 100","%","0% mezclas","Semestral"),
        ("Cumplimiento de entregas al gestor","N° entregas realizadas / N° entregas programadas x 100","%","100%","Trimestral"),
    ]
    for i, row in enumerate(indicadores):
        for j, val in enumerate(row):
            t9.rows[i+1].cells[j].text = val
            for para in t9.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
    doc.add_paragraph()

    add_heading(doc, "13. CAPACITACIÓN DEL PERSONAL", 1)
    add_body(doc, "La capacitación en gestión de residuos es obligatoria para todo el personal que labore en el consultorio. Se establece el siguiente plan de capacitación:")

    t10 = doc.add_table(rows=7, cols=5)
    t10.style = 'Table Grid'
    add_table_header(t10, ["Tema","Contenido Principal","Dirigido a","Periodicidad","Duración"])
    capac = [
        ("Inducción en PGIRHS","Clasificación, separación en fuente, rutas sanitarias, EPP","Todo el personal","Al ingreso","2 horas"),
        ("Separación en la fuente","Colores, recipientes, qué va en cada bolsa/contenedor","Todo el personal","Semestral","1 hora"),
        ("Manejo de cortopunzantes","Guardián, técnica segura, no reencapuchar","Personal asistencial","Semestral","1 hora"),
        ("Ruta sanitaria y almacenamiento","Procedimiento paso a paso, EPP, limpieza del carro","Personal de limpieza","Semestral","1 hora"),
        ("Gestión de residuos estéticos","Clasificación específica de residuos de procedimientos estéticos","Personal asistencial","Anual","1 hora"),
        ("Normatividad vigente","Decreto 351/2014, MPGIRH, actualizaciones normativas","[NOMBRE DE LA MÉDICA]","Anual","2 horas"),
    ]
    for i, row in enumerate(capac):
        for j, val in enumerate(row):
            t10.rows[i+1].cells[j].text = val
            for para in t10.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(9)
        if i % 2 == 0:
            for j in range(5):
                set_cell_bg(t10.rows[i+1].cells[j], 'EBF3FB')
    doc.add_paragraph()

    add_heading(doc, "14. RESPONSABLES", 1)
    t11 = doc.add_table(rows=4, cols=3)
    t11.style = 'Table Grid'
    add_table_header(t11, ["Responsable","Funciones en el PGIRHS","Frecuencia"])
    resp = [
        ("[NOMBRE DE LA MÉDICA]\nDirectora del Consultorio","Elaborar, actualizar y hacer cumplir el PGIRHS. Contratar gestor externo autorizado. Asegurar la capacitación del personal. Verificar el cumplimiento de indicadores. Presentar informes a la autoridad ambiental si se requiere.","Permanente"),
        ("Personal asistencial","Separar correctamente los residuos en la fuente. Usar el guardián para cortopunzantes. Reportar irregularidades o accidentes con residuos.","En cada atención"),
        ("Personal de limpieza / servicios generales","Realizar la ruta sanitaria interna. Limpiar y desinfectar los recipientes. Llevar la bitácora de generación de residuos.","Diario"),
    ]
    for i, row in enumerate(resp):
        for j, val in enumerate(row):
            t11.rows[i+1].cells[j].text = val
            for para in t11.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
    doc.add_paragraph()

    add_heading(doc, "15. REGISTROS", 1)
    t12 = doc.add_table(rows=6, cols=4)
    t12.style = 'Table Grid'
    add_table_header(t12, ["Nombre del Registro","Código","Tiempo de Retención","Responsable Custodia"])
    regs = [
        ("Bitácora mensual de generación de residuos","FOR-PP-006-01","5 años","[NOMBRE DE LA MÉDICA]"),
        ("Acta de capacitación en PGIRHS","FOR-PP-006-02","5 años","[NOMBRE DE LA MÉDICA]"),
        ("Lista de chequeo de separación en fuente","FOR-PP-006-03","2 años","[NOMBRE DE LA MÉDICA]"),
        ("Informe trimestral de indicadores PGIRHS","FOR-PP-006-04","5 años","[NOMBRE DE LA MÉDICA]"),
        ("Copia manifiestos de transporte y certificados","FOR-PP-006-05","5 años","[NOMBRE DE LA MÉDICA]"),
    ]
    for i, row in enumerate(regs):
        for j, val in enumerate(row):
            t12.rows[i+1].cells[j].text = val
            for para in t12.rows[i+1].cells[j].paragraphs:
                for run in para.runs:
                    run.font.name='Calibri'; run.font.size=Pt(10)
        if i % 2 == 0:
            for j in range(4):
                set_cell_bg(t12.rows[i+1].cells[j], 'EBF3FB')

    path = os.path.join(OUT_DIR_5, "PRO-PP-006_Plan_Gestion_Residuos_Hospitalarios.docx")
    doc.save(path)
    print(f"Guardado: {path}")

if __name__ == "__main__":
    gen_pamec()
    gen_residuos()
    print("Documentos 1 y 2 completados.")
