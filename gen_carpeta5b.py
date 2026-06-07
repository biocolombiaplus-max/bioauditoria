#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Carpeta 5 - Parte B: PRO-PP-002 Procedimientos Estéticos"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

BASE5="/home/user/bioauditoria/documentos_habilitacion/CARPETA_5_PROCESOS_PRIORITARIOS"

def margins(doc):
    for s in doc.sections:
        s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5); s.left_margin=Cm(3); s.right_margin=Cm(2.5)

def hf(doc,code,title):
    for s in doc.sections:
        h=s.header; h.paragraphs[0].clear()
        r=h.paragraphs[0].add_run(f"CONSULTORIO [NOMBRE DEL CONSULTORIO]  |  {code}  |  {title}")
        r.font.size=Pt(8); r.font.color.rgb=RGBColor(64,64,64)
        h.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        f=s.footer; f.paragraphs[0].clear()
        r2=f.paragraphs[0].add_run(f"{code} - V1.0 - {datetime.date.today().strftime('%d/%m/%Y')} | Pág. ")
        r2.font.size=Pt(8)
        fld=OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'),'begin')
        f.paragraphs[0].runs[-1]._r.append(fld)
        it=OxmlElement('w:instrText'); it.text='PAGE'
        f.paragraphs[0].runs[-1]._r.append(it)
        fld2=OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'),'end')
        f.paragraphs[0].runs[-1]._r.append(fld2)
        f.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

def portada(doc,code,title):
    doc.add_paragraph(); doc.add_paragraph()
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=t.add_run("CONSULTORIO MÉDICO\n[NOMBRE DEL CONSULTORIO]")
    r.font.size=Pt(18); r.font.bold=True; r.font.color.rgb=RGBColor(31,73,125)
    doc.add_paragraph()
    t2=doc.add_paragraph(); t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=t2.add_run(title); r2.font.size=Pt(14); r2.font.bold=True
    doc.add_paragraph()
    tb=doc.add_table(rows=6,cols=2); tb.style='Table Grid'; tb.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(l,v) in enumerate([("Código:",code),("Versión:","1.0"),("Fecha:",datetime.date.today().strftime('%d/%m/%Y')),
           ("Elaboró:","[NOMBRE DE LA MÉDICA], Médica General"),
           ("Revisó:","[NOMBRE DE LA MÉDICA]"),("Aprobó:","[NOMBRE DE LA MÉDICA]")]):
        tb.rows[i].cells[0].text=l; tb.rows[i].cells[1].text=v
        tb.rows[i].cells[0].paragraphs[0].runs[0].font.bold=True
    doc.add_page_break()
    cv=doc.add_paragraph(); r3=cv.add_run("CONTROL DE VERSIONES")
    r3.font.bold=True; r3.font.size=Pt(13); r3.font.color.rgb=RGBColor(31,73,125)
    cv.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tb2=doc.add_table(rows=3,cols=5); tb2.style='Table Grid'
    hrow(tb2,["Versión","Fecha","Cambio","Elaboró","Aprobó"])
    tb2.rows[1].cells[0].text="1.0"; tb2.rows[1].cells[1].text=datetime.date.today().strftime('%d/%m/%Y')
    tb2.rows[1].cells[2].text="Versión inicial"; tb2.rows[1].cells[3].text="[NOMBRE DE LA MÉDICA]"
    tb2.rows[1].cells[4].text="[NOMBRE DE LA MÉDICA]"
    doc.add_page_break()

def hrow(table,headers,bg="1F497D"):
    row=table.rows[0]
    for i,h in enumerate(headers):
        cell=row.cells[i]; cell.text=""
        run=cell.paragraphs[0].add_run(h)
        run.font.bold=True; run.font.color.rgb=RGBColor(255,255,255); run.font.size=Pt(10)
        cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),bg)
        shd.set(qn('w:color'),'auto'); shd.set(qn('w:val'),'clear'); tcPr.append(shd)

def H(doc,text,lv=1):
    p=doc.add_heading(text,level=lv)
    colors={1:RGBColor(31,73,125),2:RGBColor(46,116,181),3:RGBColor(68,114,196)}
    sizes={1:14,2:12,3:11}
    for run in p.runs:
        run.font.size=Pt(sizes[lv]); run.font.bold=True; run.font.color.rgb=colors[lv]
    return p

def B(doc,text,bold=False,color=None):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=p.add_run(text); r.font.size=Pt(11); r.font.bold=bold; r.font.name='Calibri'
    if color: r.font.color.rgb=color
    return p

def BL(doc,text):
    p=doc.add_paragraph(style='List Bullet'); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=p.add_run(text); r.font.size=Pt(11); r.font.name='Calibri'; return p

def firmas(doc):
    doc.add_paragraph()
    tb=doc.add_table(rows=2,cols=3); tb.style='Table Grid'
    hrow(tb,["ELABORÓ","REVISÓ","APROBÓ"])
    for j in range(3):
        tb.rows[1].cells[j].text="[NOMBRE DE LA MÉDICA]\nMédica General - Propietaria\n\nFirma: _________________________\n\nFecha: "+datetime.date.today().strftime('%d/%m/%Y')
        tb.rows[1].cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

def aviso_legal(doc):
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    # box-like bold red notice
    r=p.add_run("⚠ AVISO LEGAL IMPORTANTE: ")
    r.font.bold=True; r.font.size=Pt(12); r.font.color.rgb=RGBColor(192,0,0)
    r2=p.add_run("Los procedimientos descritos en este documento son ACTOS MÉDICOS de exclusiva competencia del médico legalmente habilitado. La Ley 711 de 2001 y la Resolución 2263 de 2004 establecen con claridad que cualquier procedimiento que involucre la aplicación de medicamentos inyectables (toxina botulínica, ácido hialurónico), agentes químicos con potencial lesivo (ácidos para peelings), procedimientos que atraviesen la epidermis (mesoterapia, PRP) o técnicas de medicina regenerativa, son procedimientos médicos que SOLO pueden ser realizados por médicos con título universitario y tarjeta profesional vigente. Su realización por cosmetólogos, esteticistas u otro personal no médico constituye ejercicio ilegal de la medicina, tipificado como delito en el Código Penal Colombiano.")
    r2.font.size=Pt(11); r2.font.color.rgb=RGBColor(192,0,0)
    doc.add_paragraph()

def crear_pro_pp002():
    doc=Document(); margins(doc)
    hf(doc,"PRO-PP-002","Proceso de Procedimientos Estéticos No Invasivos")
    portada(doc,"PRO-PP-002","PROCESO DE PROCEDIMIENTOS ESTÉTICOS NO INVASIVOS\n(EXCLUSIVO PARA MÉDICOS - ACTOS MÉDICOS)")

    aviso_legal(doc)

    H(doc,"1. OBJETIVO")
    B(doc,"Definir el proceso estandarizado, seguro y de alta calidad para la realización de procedimientos estéticos no invasivos en el Consultorio Médico [NOMBRE DEL CONSULTORIO], estableciendo el paso a paso clínico, los criterios de selección de pacientes, las contraindicaciones, el manejo de complicaciones y los requisitos documentales para cada procedimiento, en el marco de la práctica de la medicina general y la normatividad colombiana vigente.")

    H(doc,"2. ALCANCE")
    B(doc,"Aplica a los siguientes procedimientos estéticos no invasivos realizados exclusivamente por la Dra. [NOMBRE DE LA MÉDICA], Médica General, en las instalaciones del Consultorio [NOMBRE DEL CONSULTORIO]:")
    procs=["Valoración médica estética integral","Aplicación de toxina botulínica tipo A","Aplicación de rellenos dérmicos con ácido hialurónico",
           "Peelings químicos superficiales (ácido glicólico, TCA ≤20%, ácido mandélico)","Mesoterapia facial y corporal","Plasma Rico en Plaquetas (PRP)"]
    for p in procs: BL(doc,p)

    H(doc,"3. MARCO LEGAL Y JUSTIFICACIÓN DE COMPETENCIA MÉDICA")
    H(doc,"3.1 Normatividad que Define los Procedimientos Estéticos como Actos Médicos",2)
    B(doc,"La realización de procedimientos estéticos con los métodos descritos en este documento es de EXCLUSIVA COMPETENCIA MÉDICA, fundamentado en la siguiente normatividad:")

    normas=[
        ("Ley 711 de 2001","Por la cual se reglamenta el ejercicio de la cosmetología y se dictan otras disposiciones en materia de salud estética. Esta ley establece EXPRESAMENTE que la cosmetología comprende actividades de embellecimiento que NO involucren procedimientos médicos. La aplicación de inyecciones, sustancias químicas con potencial de daño o técnicas que impliquen riesgo para la salud están FUERA del ámbito de la cosmetología y son competencia médica exclusiva."),
        ("Resolución 2263 de 2004","Por la cual se establecen los requisitos para la apertura y funcionamiento de los centros de estética y similares. Define los procedimientos permitidos en centros de estética (facial y corporal sin penetración de la piel con inyecciones). Deja claramente establecido que los procedimientos que involucren inyecciones, peelings profundos o técnicas invasivas son de práctica médica."),
        ("Ley 23 de 1981 y Decreto 3380 de 1981","Definen el acto médico: toda actuación sobre el cuerpo humano con fines preventivos, curativos o estéticos que requiera conocimientos médicos especializados. La aplicación de toxina botulínica y rellenos dérmicos son actos médicos por su naturaleza farmacológica y el riesgo de complicaciones graves (oclusión vascular, anafilaxia)."),
        ("Decreto 780 de 2016","Decreto Único Reglamentario del Sector Salud. Ratifica que los procedimientos que impliquen la administración de medicamentos o dispositivos médicos son competencia del personal de salud autorizado."),
        ("Resolución 3100 de 2019","No describe un 'servicio de medicina estética' específico para médicos generales, pero tampoco lo prohíbe cuando los procedimientos son parte de la práctica médica general y el médico tiene la formación adecuada. Los procedimientos estéticos no invasivos realizados por médicos generales se enmarcan en la habilitación del servicio de consulta de medicina general, con las condiciones de talento humano, dotación e infraestructura correspondientes."),
    ]
    tb=doc.add_table(rows=len(normas)+1,cols=2); tb.style='Table Grid'
    hrow(tb,["Norma","Análisis de competencia médica"])
    for i,(n,d) in enumerate(normas):
        tb.rows[i+1].cells[0].text=n; tb.rows[i+1].cells[1].text=d
        for c in tb.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"3.2 Conclusión sobre Competencia Médica",2)
    B(doc,"CONCLUSIÓN: Los procedimientos estéticos descritos en este documento son actos médicos que pueden ser realizados por la Dra. [NOMBRE DE LA MÉDICA] en su condición de Médica General con tarjeta profesional vigente, acreditando la formación específica en cada procedimiento. NINGUNO de estos procedimientos puede ser delegado a personal no médico (esteticistas, cosmetólogos, enfermeras, técnicos o personal auxiliar de cualquier tipo).", bold=True)

    H(doc,"4. VALORACIÓN MÉDICA ESTÉTICA - PROCESO GENERAL")
    H(doc,"4.1 Consulta de Valoración Estética",2)
    B(doc,"TODA solicitud de procedimiento estético debe iniciar con una consulta de valoración médica estética. No se realiza ningún procedimiento sin valoración previa salvo controles documentados de procedimientos anteriores.")

    pasos_val=[("Anamnesis estética","Motivo de consulta estético, expectativas del paciente (realistas vs no realistas), antecedentes médicos relevantes (coagulopatías, autoinmunes, alérgicos), medicamentos actuales (anticoagulantes, antiagregantes, isotretinoína), antecedentes de procedimientos estéticos previos y resultados"),
               ("Historia dermatológica","Fototipo (Fitzpatrick I-VI), antecedentes de queloides, cicatrices hipertróficas, piel sensible, enfermedades dermatológicas activas (acné, rosácea, herpes recurrente, psoriasis)"),
               ("Examen físico facial/corporal","Evaluación del área a tratar: tipo de piel, hidratación, tonicidad, elasticidad, arrugas (estáticas vs dinámicas), volumen, ptosis, asimetría, laxitud. Clasificación de Glogau (envejecimiento). Clasificación de Fitzpatrick (fototipo)."),
               ("Fotografía clínica estandarizada","Toma de fotografías de frente, 3/4 derecho e izquierdo, perfil derecho e izquierdo, y planos específicos según el área a tratar. Con autorización escrita del paciente (incluida en el consentimiento informado)."),
               ("Establecimiento del plan de tratamiento","Diseño del plan individualizado: procedimiento(s) a realizar, productos a usar (con nombre comercial, concentración, registro INVIMA), número de sesiones, resultados esperados realistas, alternativas."),
               ("Consentimiento informado","Entrega, explicación y firma del consentimiento informado específico para cada procedimiento ANTES de proceder. El paciente debe leerlo con tranquilidad, sin presión de tiempo."),
               ("Registro en historia clínica estética","Diligenciamiento completo del FOR-HC-002 con todos los hallazgos de la valoración.")]
    tb2=doc.add_table(rows=len(pasos_val)+1,cols=2); tb2.style='Table Grid'
    hrow(tb2,["Paso","Descripción"])
    for i,(p2,d) in enumerate(pasos_val):
        tb2.rows[i+1].cells[0].text=p2; tb2.rows[i+1].cells[1].text=d
        for c in tb2.rows[i+1].cells:
            for pp in c.paragraphs:
                for r in pp.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"5. TOXINA BOTULÍNICA TIPO A")
    H(doc,"5.1 Descripción del Procedimiento",2)
    B(doc,"La toxina botulínica tipo A es una neurotoxina producida por Clostridium botulinum que bloquea la liberación de acetilcolina en la unión neuromuscular, produciendo una relajación muscular temporal y reversible. En medicina estética se usa para el tratamiento de arrugas de expresión (dinámicas), hiperhidrosis y otras indicaciones médicas.")
    H(doc,"5.2 Indicaciones",2)
    ind_tb=["Arrugas glabelares (líneas del entrecejo)","Arrugas de la frente (frontales)","Arrugas periorbiculares (patas de gallo)",
            "Arrugas nasales (bunny lines)","Arrugas periorales (código de barras)","Platisma (bandas del cuello)","Hiperhidrosis axilar, palmar, plantar",
            "Bruxismo / hipertrofia del masetero","Elevación de comisuras","Corrección de sonrisa gingival"]
    for item in ind_tb: BL(doc,item)
    H(doc,"5.3 Contraindicaciones",2)
    H(doc,"Absolutas:",3)
    contrain_abs=["Hipersensibilidad conocida a la toxina botulínica o a los excipientes del producto",
                  "Enfermedades neuromusculares: miastenia gravis, síndrome de Eaton-Lambert, esclerosis lateral amiotrófica",
                  "Infección activa o inflamación en el sitio de inyección",
                  "Embarazo y lactancia (evidencia insuficiente de seguridad)",
                  "Anticoagulación con alto riesgo de sangrado no controlado"]
    for item in contrain_abs: BL(doc,item)
    H(doc,"Relativas (evaluar riesgo-beneficio):",3)
    contrain_rel=["Uso de aminoglucósidos u otros antibióticos que potencian el efecto de la toxina",
                  "Antecedente de reacciones alérgicas previas a la toxina",
                  "Expectativas poco realistas del paciente",
                  "Cicatrices o alteraciones anatómicas en el área a tratar"]
    for item in contrain_rel: BL(doc,item)

    H(doc,"5.4 Materiales Necesarios",2)
    materiales_tb=["Toxina botulínica tipo A con registro INVIMA vigente ([MARCA COMERCIAL])",
                   "Solución salina 0.9% sin preservantes para reconstitución",
                   "Jeringas de insulina 1ml con aguja 30G o 32G integrada",
                   "Alcohol 70% o clorhexidina al 0.5% para antisepsia",
                   "Gasas estériles","Guantes estériles","Marcador dérmico (opcional, para marcar puntos de inyección)",
                   "Cold pack para crioanalgesiaprevia","Frasco con hielo para mantener la toxina reconstituida","Guardián para residuos cortopunzantes","Kit de emergencias disponible"]
    for item in materiales_tb: BL(doc,item)

    H(doc,"5.5 Paso a Paso del Procedimiento",2)
    pasos_tb=[("PRE-PROCEDIMIENTO","1. Verificar que el consentimiento informado (FOR-PP-002) está firmado\n2. Verificar la identidad del paciente (identificación correcta)\n3. Revisar la historia clínica estética y el plan de tratamiento acordado\n4. Verificar que no hay contraindicaciones activas\n5. Fotografía clínica pre-procedimiento (si no se tomó en la valoración)\n6. Reconstituir la toxina botulínica según el protocolo del fabricante y el plan de dilución elegido\n7. Lavado de manos 6 pasos OMS, colocación de guantes\n8. Limpiar el área a tratar con solución salina o agua micelar\n9. Crio-analgesia previa con cold pack 1-2 minutos si el paciente lo desea\n10. Marcación de los puntos de inyección con marcador dérmico (opcional)"),
              ("PROCEDIMIENTO","11. Aplicar antiséptico (alcohol 70%) en el área de inyección\n12. Esperar 30 segundos y secar\n13. Inyectar en los puntos predeterminados según el mapa de inyección para cada área:\n   - Glabela: 5 puntos estándar (procerus + corrugadores)\n   - Frente: 4-6 puntos según el patrón de contracción\n   - Patas de gallo: 2-3 puntos por lado\n   - Hiperhidrosis axilar: múltiples puntos intradérmicos (10-20 por axila)\n14. Aplicar presión suave post-inyección con gasa\n15. NO masajear (evita migración)\n16. Descartar agujas en guardián inmediatamente"),
              ("POST-PROCEDIMIENTO","17. Registrar: unidades totales aplicadas por área, lote del producto, fecha de reconstitución\n18. Indicaciones al paciente:\n   - Mantener posición vertical 4 horas\n   - No masajear la zona tratada 24 horas\n   - No actividad física intensa el día del procedimiento\n   - No exposición a calor intenso (sauna, vapor) 24 horas\n   - Efecto visible a partir de 3-7 días, máximo efecto a 14 días\n   - Duración esperada: 3-6 meses\n19. Agendar control a los 14 días para evaluación del resultado\n20. Registrar completo en historia clínica estética FOR-HC-002")]
    tb3=doc.add_table(rows=len(pasos_tb)+1,cols=2); tb3.style='Table Grid'
    hrow(tb3,["Fase","Acciones"])
    for i,(fase,acc) in enumerate(pasos_tb):
        tb3.rows[i+1].cells[0].text=fase; tb3.rows[i+1].cells[1].text=acc
        for c in tb3.rows[i+1].cells:
            for pp in c.paragraphs:
                for r in pp.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"5.6 Complicaciones y Manejo",2)
    comps_tb=[("Equimosis/hematoma","Frecuente, generalmente leve","Compresas frías locales, tranquilizar al paciente. Resolución espontánea en 1-2 semanas"),
              ("Ptosis palpebral","Infrecuente, por difusión a músculo elevador del párpado","Colirio de apraclonidina 0.5% (Iopidine) una gota 3 veces/día. Resolución espontánea en 4-6 semanas"),
              ("Ptosis de ceja","Difusión a músculo frontal en zona temporal","Esperar resolución espontánea. Ajustar técnica en próxima sesión"),
              ("Asimetría","Variación en la respuesta muscular","Evaluación a los 14 días, posible retoque con dosis de ajuste"),
              ("Efecto insuficiente","Dosis insuficiente o resistencia","Retoque a los 14 días si el efecto no fue suficiente"),
              ("Cefalea","Efecto secundario transitorio","Analgésicos de venta libre. Resolución en 24-48h"),
              ("Reacción alérgica leve","Eritema, prurito en zona de inyección","Antihistamínicos, compresas frías"),
              ("Reacción anafiláctica","Urticaria generalizada, broncoespasmo, hipotensión","EMERGENCIA: Posición decúbito. Adrenalina 0.3mg IM. Llamar emergencias 123. Trasladar.")]
    tb4=doc.add_table(rows=len(comps_tb)+1,cols=3); tb4.style='Table Grid'
    hrow(tb4,["Complicación","Características","Manejo"])
    for i,row in enumerate(comps_tb):
        for j,v in enumerate(row): tb4.rows[i+1].cells[j].text=v
        for c in tb4.rows[i+1].cells:
            for pp in c.paragraphs:
                for r in pp.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"6. RELLENOS DÉRMICOS CON ÁCIDO HIALURÓNICO")
    H(doc,"6.1 Descripción",2)
    B(doc,"El ácido hialurónico (AH) es un polisacárido natural constituyente de la matriz extracelular. En rellenos dérmicos, se utiliza en forma reticulada (cross-linked) para dar volumen y corregir deficiencias de tejido blando, o en forma no reticulada para hidratación dérmica profunda (skinbooster). Es un dispositivo médico de clase III con registro INVIMA.")

    H(doc,"6.2 Indicaciones",2)
    ind_ah=["Relleno de surcos nasolabiales","Aumento de volumen labial (volumización)","Proyección malar (relleno de mejillas/pómulos)","Relleno de surco lagrimal / ojeras","Perfilado de mandíbula y mentón","Relleno de arrugas estáticas profundas","Hidratación dérmica profunda (skinbooster)","Rinomodelación no quirúrgica","Corrección de asimetrías faciales leves"]
    for item in ind_ah: BL(doc,item)

    H(doc,"6.3 Contraindicaciones",2)
    contrain_ah=["Hipersensibilidad al ácido hialurónico o a los componentes del producto",
                 "Infección activa en la zona a tratar (acné activo severo, herpes activo, impétigo)",
                 "Trastornos de la coagulación no controlados",
                 "Embarazo y lactancia",
                 "Enfermedades autoinmunes activas",
                 "Uso de anticoagulantes (warfarina, acenocumarol) - evaluar suspensión previa con médico tratante",
                 "Antecedente de reacción a rellenos previos",
                 "Cicatrices hipertróficas o queloides activos en la zona",
                 "Expectativas no realistas del paciente"]
    for item in contrain_ah: BL(doc,item)

    H(doc,"6.4 AVISO CRÍTICO: Hialuronidasa Obligatoria",2)
    B(doc,"NUNCA aplicar ácido hialurónico sin tener hialuronidasa disponible en el consultorio. La oclusión vascular por relleno dérmico puede ocurrir incluso en manos expertas y es una emergencia médica que puede producir necrosis tisular o ceguera si no se trata inmediatamente.",
      bold=True, color=RGBColor(192,0,0))

    H(doc,"6.5 Paso a Paso del Procedimiento",2)
    pasos_ah=[("PRE-PROCEDIMIENTO","1. Consentimiento informado FOR-PP-003 firmado\n2. Identificación correcta del paciente\n3. Revisión del plan de tratamiento y fotografías previas\n4. Verificar disponibilidad de hialuronidasa\n5. Evaluar signos vasculares en zona a tratar (blanqueamiento, textura)\n6. Lavado de manos, guantes\n7. Limpieza de la zona con agua micelar o solución salina\n8. Crioanaalgesia o anestésico tópico (lidocaína 4% crema) 20-30 min antes si el paciente lo requiere"),
              ("PROCEDIMIENTO","9. Antisepsia con clorhexidina 0.05% o alcohol 70%\n10. Selección de la técnica según el área:\n   - Surcos nasolabiales: lineal retrótraza o serial puncture\n   - Labios: técnica serial puncture, lineal\n   - Mejillas/pómulos: técnica de bolo profundo (subperióstico o supra-perióstico)\n   - Ojeras: técnica superficial con cánula\n11. ASPIRAR SIEMPRE antes de inyectar (aunque con cánula el riesgo es menor)\n12. Inyectar LENTAMENTE, con volúmenes pequeños por punto\n13. Masajear suavemente para distribuir el producto\n14. VIGILAR constantemente signos de alarma vascular: dolor desproporcionado, blanqueamiento"),
              ("PROTOCOLO DE EMERGENCIA VASCULAR","Si se presenta BLANQUEAMIENTO o dolor desproporcionado:\n1. DETENER inmediatamente la inyección\n2. Administrar hialuronidasa (150-1500 UI según el área) INMEDIATAMENTE\n3. Masaje vigoroso\n4. Calor local\n5. Llamar emergencias 123 si no hay resolución inmediata\n6. Documentar el evento en FOR-PP-005"),
              ("POST-PROCEDIMIENTO","15. Indicaciones al paciente:\n   - Evitar presión sobre la zona 24 horas\n   - No exposición a calor intenso 48 horas\n   - Inflamación y equimosis esperadas los primeros 3-7 días\n   - No masticar excesivamente si se trató zona labial\n   - Signos de alarma que requieren consulta urgente: dolor intenso, cambio de color (blanqueamiento, violáceo), visión borrosa\n16. Control a los 15 días\n17. Registro completo en FOR-HC-002")]
    tb5=doc.add_table(rows=len(pasos_ah)+1,cols=2); tb5.style='Table Grid'
    hrow(tb5,["Fase","Acciones"])
    for i,(fase,acc) in enumerate(pasos_ah):
        tb5.rows[i+1].cells[0].text=fase; tb5.rows[i+1].cells[1].text=acc
        for c in tb5.rows[i+1].cells:
            for pp in c.paragraphs:
                for r in pp.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"7. PEELINGS QUÍMICOS SUPERFICIALES")
    H(doc,"7.1 Definición y Marco",2)
    B(doc,"El peeling químico es la aplicación controlada de agentes químicos sobre la piel para producir exfoliación y regeneración cutánea. Los peelings SUPERFICIALES actúan hasta la capa granulosa de la epidermis y son los únicos permitidos para médicos generales sin formación especializada en dermatología. Los peelings medios y profundos requieren formación especializada específica.")
    B(doc,"Agentes para peelings superficiales usados en este consultorio:")
    agentes_peel=["Ácido glicólico: concentraciones entre 20-70%. El tiempo de aplicación determina la profundidad. Neutralización con bicarbonato de sodio o agua.",
                  "Ácido tricloroacético (TCA): concentraciones ≤20% para superficial. Autolimitado (escarcha blanca indica fin del peeling). NO neutralizable.",
                  "Ácido mandélico: 25-40%. Más suave, indicado para fototipos altos (IV-VI), antiinflamatorio, antibacteriano. Útil en acné."]
    for item in agentes_peel: BL(doc,item)

    H(doc,"7.2 Indicaciones",2)
    ind_peel=["Acné activo leve-moderado y marcas post-acné","Melasma leve a moderado (con fotoprotección estricta)","Hiperpigmentación post-inflamatoria","Piel apagada, sin luminosidad","Fotoenvejecimiento leve","Textura irregular de la piel","Poros dilatados"]
    for item in ind_peel: BL(doc,item)

    H(doc,"7.3 Contraindicaciones",2)
    contrain_peel=["Uso de isotretinoína oral en los últimos 6 meses (contraindicación absoluta - cicatrices anómalas)",
                   "Herpes activo en la zona a tratar",
                   "Infección bacteriana o micótica activa en la piel",
                   "Embarazo y lactancia",
                   "Piel con inflamación activa severa, dermatitis",
                   "Fototipo VI con TCA (alto riesgo de hiperpigmentación post-inflamatoria)",
                   "Exposición solar reciente intensa o bronceado activo",
                   "Antecedente de queloides en la zona",
                   "Uso de retinoides tópicos (suspender 1 semana antes)"]
    for item in contrain_peel: BL(doc,item)

    H(doc,"7.4 Paso a Paso",2)
    pasos_peel=[("PRE-PEELING (1 semana antes)","- Preparación de la piel: retinoides tópicos (suspender 1 semana antes del peeling)\n- Hidratante y fotoprotector diario obligatorio\n- En fototipos ≥IV: crema despigmentante (hidroquinona 4% o azelato) 2 semanas antes para disminuir riesgo de HIpost-inflamatoria"),
                ("DÍA DEL PROCEDIMIENTO","1. Verificar que no hay contraindicaciones activas (preguntar por herpes, infección, exposición solar reciente)\n2. Consentimiento informado FOR-PP-004 firmado\n3. Fotografías pre-peeling\n4. Lavado del rostro con limpiador suave sin aromas\n5. Desengrase con acetona o alcohol isopropílico (crea acceso uniforme del ácido)\n6. Lavado de manos y guantes de nitrilo"),
                ("APLICACIÓN DEL PEELING","7. Aplicar el agente químico con gasa o pincel, capa delgada y uniforme\n8. Iniciar por zona de menor sensibilidad (frente) y terminar en zona más sensible (contorno de ojos y boca)\n9. EVITAR contorno de ojos, mucosas\n10. Observar: eritema difuso (normal), edema moderado (normal), sensación de ardor (normal)\n11. SIGNOS DE ALARMA: eritema intenso no homogéneo, respuesta exagerada, blanqueamiento no esperado\n12. Tiempo de aplicación: según el ácido, la concentración y la respuesta de la piel\n    - Ácido glicólico: 1-5 minutos, neutralizar con bicarbonato de sodio diluido o agua\n    - TCA ≤20%: autolimitado, observar la escarcha blanca (frosting nivel 1)\n    - Ácido mandélico: 3-5 minutos, retirar con agua abundante"),
                ("POST-PEELING","- Aplicar crema calmante (aloe vera, centella asiática, pantenol)\n- Fotoprotector SPF 50+ obligatorio (en zonas expuestas)\n- Indicaciones al paciente:\n  * No frotar ni pelar la piel que se descama\n  * Hidratación constante\n  * Fotoprotector obligatorio hasta 4 semanas post-peeling\n  * No sol directo por 2 semanas\n  * No ácidos exfoliantes ni retinoides hasta resolución completa de la descamación\n  * Control en 2-4 semanas para evaluar resultado")]
    tb6=doc.add_table(rows=len(pasos_peel)+1,cols=2); tb6.style='Table Grid'
    hrow(tb6,["Fase","Acciones"])
    for i,(fase,acc) in enumerate(pasos_peel):
        tb6.rows[i+1].cells[0].text=fase; tb6.rows[i+1].cells[1].text=acc
        for c in tb6.rows[i+1].cells:
            for pp in c.paragraphs:
                for r in pp.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"8. MESOTERAPIA FACIAL Y CORPORAL")
    H(doc,"8.1 Definición",2)
    B(doc,"La mesoterapia es una técnica médica de inyección intradérmica o subdérmica de múltiples microinyecciones de cócteles de principios activos con fines terapéuticos o estéticos. En el ámbito estético, se usa para hidratación, luminosidad, tratamiento de alopecia y reducción localizada de adiposidad. Es un acto médico de competencia exclusiva del médico.")

    H(doc,"8.2 Indicaciones",2)
    ind_meso=["Hidratación y luminosidad facial profunda (mesoterapia hidratante)","Alopecia androgenética y efluvium (mesoterapia capilar con minoxidil, dutasterida, biotina)","Cellulitis y reducción de medidas corporal (fosfatidilcolina, L-carnitina)","Fibrosis y retracción del tejido adiposo"]
    for item in ind_meso: BL(doc,item)

    H(doc,"8.3 Paso a Paso",2)
    pasos_meso=[("PRE-PROCEDIMIENTO","Consentimiento informado (usar FOR-PP-005 o CI general)\nVerificar ausencia de infección activa en zona a tratar\nFotografías\nVerificar registro INVIMA de los principios activos a usar"),
                ("PREPARACIÓN DEL CÓCTEL","Preparar el cóctel en condiciones asépticas\nUsar jeringas estériles, agujas nuevas de mesoterapia 4mm/27G o 13mm/27G\nVerificar vencimiento y aspecto de cada componente"),
                ("APLICACIÓN","Antisepsia de la zona con alcohol 70%\nTécnica de micropápulas intradérmicas o técnica náppage\nDistancia entre inyecciones: 1-2 cm\nProfundidad: intradérmica (4mm) o subcutánea superficial (13mm) según el objetivo\nCantidad por punto: 0.05-0.1 ml por micropápula"),
                ("POST-PROCEDIMIENTO","Antiséptico local\nCompresas frías si hay reacción inflamatoria\nNo exposición solar directa por 24-48 horas\nNo maquillaje las primeras 4 horas\nRegistro en FOR-HC-002")]
    tb7=doc.add_table(rows=len(pasos_meso)+1,cols=2); tb7.style='Table Grid'
    hrow(tb7,["Fase","Acciones"])
    for i,(fase,acc) in enumerate(pasos_meso):
        tb7.rows[i+1].cells[0].text=fase; tb7.rows[i+1].cells[1].text=acc
        for c in tb7.rows[i+1].cells:
            for pp in c.paragraphs:
                for r in pp.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"9. PLASMA RICO EN PLAQUETAS (PRP)")
    H(doc,"9.1 Descripción",2)
    B(doc,"El PRP (Plasma Rico en Plaquetas) es un hemoderivado autólogo obtenido de la centrifugación de la propia sangre del paciente, que concentra los factores de crecimiento plaquetarios. Se usa en medicina regenerativa estética para estimular la producción de colágeno y elastina, tratar la alopecia y rejuvenecer la piel. Al ser autólogo, no requiere registro INVIMA como producto externo, pero los tubos y equipos de procesamiento sí deben tener certificación.")

    H(doc,"9.2 Paso a Paso del PRP",2)
    pasos_prp=[("EXTRACCIÓN SANGUÍNEA","1. Consentimiento informado firmado\n2. Verificar estado de salud actual del paciente (no resfriado, no infección)\n3. Extraer 10-20 ml de sangre venosa con tubos anticoagulantes (citrato de sodio o ACD) certificados\n4. Tubos de vacutainer con citrato de sodio 3.8% (tubos azules) o tubos específicos para PRP"),
               ("PROCESAMIENTO","5. Centrifugar: Primera centrifugación: 1500 RPM por 5 minutos (separación plasma/glóbulos)\n6. Extraer el plasma sobrenadante + la capa leucoplaquetaria\n7. Segunda centrifugación: 3000 RPM por 10 minutos (concentración de plaquetas)\n8. Extraer el PRP concentrado (fracción inferior del plasma)\n9. El PRP resultante debe tener una concentración de plaquetas 3-5 veces la concentración basal de la sangre"),
               ("ACTIVACIÓN (opcional)","Activación con cloruro de calcio (1:10 v/v) o trombina autóloga para liberar los factores de crecimiento. La activación produce un gel de fibrina más resistente."),
               ("APLICACIÓN","10. Antisepsia del área a tratar\n11. Crioanaalgesia o anestesia tópica previa\n12. Para piel: microinyecciones intradérmicas técnica náppage o multipunción con mesoterapia\n13. Para alopecia: inyecciones subcutáneas superficiales en la zona del cuero cabelludo\n14. Para articulaciones (si aplica): inyección intraarticular - precisa entrenamiento específico"),
               ("POST-PROCEDIMIENTO","- Aplicar presión con gasa 2-3 minutos\n- Frío local\n- Eritema, edema y hematomas esperados 24-72 horas\n- No exposición solar directa\n- No antiinflamatorios por 1 semana (disminuyen el efecto del PRP)\n- Resultados visibles a partir de las 4-6 semanas\n- Esquema recomendado: 3-4 sesiones mensuales, mantenimiento cada 6 meses\n- Registro completo en FOR-HC-002")]
    tb8=doc.add_table(rows=len(pasos_prp)+1,cols=2); tb8.style='Table Grid'
    hrow(tb8,["Fase","Acciones"])
    for i,(fase,acc) in enumerate(pasos_prp):
        tb8.rows[i+1].cells[0].text=fase; tb8.rows[i+1].cells[1].text=acc
        for c in tb8.rows[i+1].cells:
            for pp in c.paragraphs:
                for r in pp.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"10. MANEJO GENERAL DE COMPLICACIONES EN PROCEDIMIENTOS ESTÉTICOS")
    B(doc,"Ante cualquier complicación post-procedimiento estético, el paciente debe comunicarse directamente con la médica al número [TELÉFONO DE EMERGENCIAS]. Se establece el siguiente protocolo general:")
    complicaciones_grales=[
        ("Infección post-procedimiento","Eritema, calor, secreción purulenta, fiebre","Antibioticoterapia según germen sospechoso. Si hay absceso: drenaje. Reportar en FOR-PP-005"),
        ("Necrosis por oclusión vascular","Cambio de color (violáceo/negro), dolor intenso, ausencia de relleno capilar","EMERGENCIA: Hialuronidasa IV/local si fue AH. Calor. Nitroglicerina tópica. Remitir a urgencias."),
        ("Granuloma por cuerpo extraño","Nódulo duro, no doloroso, tardío (semanas a meses post-aplicación de AH)","Corticoide intralesional (triamcinolona). Hialuronidasa si es AH. Consulta a dermatólogo."),
        ("Anafilaxia","Urticaria generalizada, broncoespasmo, hipotensión","Adrenalina 0.3mg IM, posición supina con piernas elevadas, O2. Llamar emergencias 123."),
    ]
    tb9=doc.add_table(rows=len(complicaciones_grales)+1,cols=3); tb9.style='Table Grid'
    hrow(tb9,["Complicación","Síntomas","Manejo"])
    for i,row in enumerate(complicaciones_grales):
        for j,v in enumerate(row): tb9.rows[i+1].cells[j].text=v
        for c in tb9.rows[i+1].cells:
            for pp in c.paragraphs:
                for r in pp.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"11. INDICADORES")
    tb10=doc.add_table(rows=5,cols=4); tb10.style='Table Grid'
    hrow(tb10,["Indicador","Fórmula","Meta","Frecuencia"])
    for i,row in enumerate([
        ("% procedimientos con consentimiento informado previo","(CI firmados/Total procedimientos)×100","100%","Mensual"),
        ("% procedimientos con historia clínica completa","(HC completas/Total procedimientos)×100","100%","Mensual"),
        ("Tasa de complicaciones post-procedimiento","N° complicaciones/100 procedimientos","Seguimiento y análisis","Mensual"),
        ("% complicaciones reportadas y documentadas","(Complicaciones documentadas/Total complicaciones)×100","100%","Mensual")]):
        for j,v in enumerate(row): tb10.rows[i+1].cells[j].text=v
        for c in tb10.rows[i+1].cells:
            for pp in c.paragraphs:
                for r in pp.runs: r.font.size=Pt(9)
    doc.add_paragraph(); firmas(doc)
    path=os.path.join(BASE5,"PRO-PP-002_Proceso_Procedimientos_Esteticos_No_Invasivos.docx")
    doc.save(path); print(f"✓ Creado: {path}")

if __name__=="__main__":
    print("Generando Carpeta 5 - Parte B (PRO-PP-002)...")
    crear_pro_pp002()
    print("✅ Parte B completada.")
