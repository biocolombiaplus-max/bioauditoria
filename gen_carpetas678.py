#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Carpetas 6, 7 y 8: Historia Clínica, Referencia/Contrarreferencia, Gestión de Calidad"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

BASE6="/home/user/bioauditoria/documentos_habilitacion/CARPETA_6_HISTORIA_CLINICA"
BASE7="/home/user/bioauditoria/documentos_habilitacion/CARPETA_7_REFERENCIA_CONTRARREFERENCIA"
BASE8="/home/user/bioauditoria/documentos_habilitacion/CARPETA_8_GESTION_CALIDAD"

def margins(doc):
    for s in doc.sections:
        s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5); s.left_margin=Cm(3); s.right_margin=Cm(2.5)

def hf(doc,code,title):
    for s in doc.sections:
        h=s.header; h.paragraphs[0].clear()
        r=h.paragraphs[0].add_run(f"CONSULTORIO [NOMBRE DEL CONSULTORIO]  |  {code}  |  {title}")
        r.font.size=Pt(8); r.font.color.rgb=RGBColor(64,64,64)
        h.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        f=s.footer; f.paragraphs[0].clear()
        r2=f.paragraphs[0].add_run(f"{code} - V1.0 - {datetime.date.today().strftime('%d/%m/%Y')} | Pág. ")
        r2.font.size=Pt(8)
        fld=OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'),'begin')
        f.paragraphs[0].runs[-1]._r.append(fld)
        it=OxmlElement('w:instrText'); it.text='PAGE'
        f.paragraphs[0].runs[-1]._r.append(it)
        fld2=OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'),'end')
        f.paragraphs[0].runs[-1]._r.append(fld2)
        f.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

def portada(doc,code,title):
    doc.add_paragraph(); doc.add_paragraph()
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=t.add_run("CONSULTORIO MÉDICO\n[NOMBRE DEL CONSULTORIO]")
    r.font.size=Pt(18); r.font.bold=True; r.font.color.rgb=RGBColor(31,73,125)
    doc.add_paragraph()
    t2=doc.add_paragraph(); t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=t2.add_run(title); r2.font.size=Pt(14); r2.font.bold=True
    doc.add_paragraph()
    tb=doc.add_table(rows=6,cols=2); tb.style='Table Grid'; tb.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(l,v) in enumerate([("Código:",code),("Versión:","1.0"),("Fecha:",datetime.date.today().strftime('%d/%m/%Y')),
           ("Elaboró:","[NOMBRE DE LA MÉDICA]"),("Revisó:","[NOMBRE DE LA MÉDICA]"),("Aprobó:","[NOMBRE DE LA MÉDICA]")]):
        tb.rows[i].cells[0].text=l; tb.rows[i].cells[1].text=v
        tb.rows[i].cells[0].paragraphs[0].runs[0].font.bold=True
    doc.add_page_break()
    tb2=doc.add_table(rows=3,cols=5); tb2.style='Table Grid'
    hrow(tb2,["Versión","Fecha","Cambio","Elaboró","Aprobó"])
    tb2.rows[1].cells[0].text="1.0"; tb2.rows[1].cells[1].text=datetime.date.today().strftime('%d/%m/%Y')
    tb2.rows[1].cells[2].text="Versión inicial"; tb2.rows[1].cells[3].text="[NOMBRE DE LA MÉDICA]"
    tb2.rows[1].cells[4].text="[NOMBRE DE LA MÉDICA]"
    doc.add_page_break()

def hrow(table,headers,bg="1F497D"):
    row=table.rows[0]
    for i,h in enumerate(headers):
        cell=row.cells[i]; cell.text=""
        run=cell.paragraphs[0].add_run(h)
        run.font.bold=True; run.font.color.rgb=RGBColor(255,255,255); run.font.size=Pt(10)
        cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),bg)
        shd.set(qn('w:color'),'auto'); shd.set(qn('w:val'),'clear'); tcPr.append(shd)

def H(doc,text,lv=1):
    p=doc.add_heading(text,level=lv)
    colors={1:RGBColor(31,73,125),2:RGBColor(46,116,181),3:RGBColor(68,114,196)}
    sizes={1:14,2:12,3:11}
    for run in p.runs:
        run.font.size=Pt(sizes[lv]); run.font.bold=True; run.font.color.rgb=colors[lv]
    return p

def B(doc,text,bold=False):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=p.add_run(text); r.font.size=Pt(11); r.font.bold=bold; r.font.name='Calibri'; return p

def BL(doc,text):
    p=doc.add_paragraph(style='List Bullet'); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=p.add_run(text); r.font.size=Pt(11); r.font.name='Calibri'; return p

def firmas(doc):
    doc.add_paragraph()
    tb=doc.add_table(rows=2,cols=3); tb.style='Table Grid'
    hrow(tb,["ELABORÓ","REVISÓ","APROBÓ"])
    for j in range(3):
        tb.rows[1].cells[j].text="[NOMBRE DE LA MÉDICA]\nMédica General - Propietaria\n\nFirma: _________________________\n\nFecha: "+datetime.date.today().strftime('%d/%m/%Y')
        tb.rows[1].cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

# ===== MAN-HC-001 =====
def crear_man_hc():
    doc=Document(); margins(doc); hf(doc,"MAN-HC-001","Manual de Historia Clínica")
    portada(doc,"MAN-HC-001","MANUAL DE GESTIÓN DE LA HISTORIA CLÍNICA")

    H(doc,"1. OBJETIVO")
    B(doc,"Establecer los lineamientos para la elaboración, diligenciamiento, custodia, conservación, confidencialidad, acceso y auditoría de las historias clínicas del Consultorio [NOMBRE DEL CONSULTORIO], garantizando el cumplimiento de la Resolución 1995 de 1999 y demás normatividad vigente sobre historia clínica en Colombia, y protegiendo los derechos de los pacientes y de la médica propietaria.")

    H(doc,"2. MARCO LEGAL")
    normas=[("Resolución 1995 de 1999","Define la historia clínica: concepto, características, componentes, manejo y custodio. Base fundamental de este manual."),
            ("Ley 23 de 1981","Ética médica. La HC es documento esencial del acto médico."),
            ("Ley 1581 de 2012","Protección de datos personales (Habeas Data). Los datos de salud son datos sensibles con protección especial."),
            ("Decreto 1377 de 2013","Reglamentación de la Ley 1581/2012. Tratamiento de datos personales."),
            ("Ley 594 de 2000","Ley General de Archivos. Principios de gestión documental aplicables a la HC."),
            ("Resolución 3100 de 2019","Estándares de habilitación relacionados con historia clínica."),
            ("Decreto 1011 de 2006","SOGCS. La HC como elemento del sistema de calidad."),
            ("Ley 1751 de 2015","Ley Estatutaria de Salud. Acceso a la HC como derecho del paciente.")]
    tb=doc.add_table(rows=len(normas)+1,cols=2); tb.style='Table Grid'
    hrow(tb,["Norma","Descripción"])
    for i,(n,d) in enumerate(normas):
        tb.rows[i+1].cells[0].text=n; tb.rows[i+1].cells[1].text=d
        for c in tb.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()

    H(doc,"3. CARACTERÍSTICAS DE LA HISTORIA CLÍNICA (Resolución 1995/1999)")
    caract=[("Integralidad","La HC debe recoger la información de todos los aspectos científicos, técnicos y administrativos relativos a la atención en salud en las fases de fomento, prevención, diagnóstico, tratamiento y rehabilitación."),
            ("Secuencialidad","Los registros de la HC deben consignarse en la secuencia cronológica de la atención prestada al usuario."),
            ("Racionalidad científica","Es la aplicación de criterios científicos en el diligenciamiento y registro de las acciones en salud brindadas al usuario, de acuerdo con la naturaleza de cada servicio prestado."),
            ("Disponibilidad","Es la posibilidad de utilizar la HC en el momento en que se necesita, con las limitaciones que impone la ley."),
            ("Oportunidad","Es el diligenciamiento de los registros de atención de la HC, simultánea o inmediatamente después de que ocurra la prestación del servicio.")]
    for nom,desc in caract:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r1=p.add_run("• "+nom+": "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(desc); r2.font.size=Pt(11)

    H(doc,"4. COMPONENTES OBLIGATORIOS DE LA HISTORIA CLÍNICA")
    H(doc,"4.1 Identificación del Usuario",2)
    comp_id=["Apellidos y nombre completo","Número de documento de identidad","Fecha de nacimiento","Sexo","Dirección y teléfono","Aseguradora (EPS, ARS, medicina prepagada, particular)","Nombre y teléfono del contacto de emergencia"]
    for item in comp_id: BL(doc,item)

    H(doc,"4.2 Registro de Consulta (debe incluir)",2)
    comp_cons=["Fecha y hora de la atención","Motivo de consulta (en palabras del paciente)","Anamnesis completa (historia de la enfermedad actual, antecedentes)","Resultados del examen físico","Diagnóstico o diagnósticos (con código CIE-10)","Plan de manejo (medicamentos, exámenes, remisiones, indicaciones)","Nombre, firma y tarjeta profesional del médico tratante"]
    for item in comp_cons: BL(doc,item)

    H(doc,"5. CUSTODIA Y CONSERVACIÓN")
    B(doc,"La Resolución 1995 de 1999 establece que:")
    custod=["El custodio de la HC es el prestador de servicios de salud que la genera",
            "En el Consultorio [NOMBRE DEL CONSULTORIO], la custodio es la Dra. [NOMBRE DE LA MÉDICA] como propietaria",
            "El tiempo mínimo de conservación de la HC es de 15 AÑOS a partir de la fecha de la última atención",
            "Las HC de menores de edad deben conservarse hasta cuando el paciente cumpla 18 años más el período estándar de 15 años",
            "La HC física se conserva en archivo físico bajo llave con acceso restringido",
            "La HC electrónica (si se usa sistema informático) debe tener copia de seguridad periódica y medidas de protección de datos",
            "Al cierre del consultorio, la HC debe ser transferida al prestador de salud que continúe la atención del paciente o a la secretaría de salud territorial"]
    for item in custod: BL(doc,item)

    H(doc,"6. CONFIDENCIALIDAD Y ACCESO A LA HISTORIA CLÍNICA")
    B(doc,"La HC es un documento privado. El acceso está regulado por la Resolución 1995/1999 y la Ley 1581/2012:")
    acceso=[("El paciente o su representante legal","Tiene derecho a acceder a su propia HC en cualquier momento. El consultorio tiene 3 días hábiles para dar respuesta a una solicitud de copia."),
            ("El médico tratante","Accede en el curso del proceso de atención del paciente."),
            ("Las autoridades judiciales y de salud","Cuando exista mandamiento judicial o requerimiento de autoridad sanitaria competente."),
            ("Las entidades aseguradoras","Solo para los procesos de auditoría de calidad de la atención en salud, en los términos definidos por ley."),
            ("Los comités de ética y calidad","En los procesos de auditoría interna y mejoramiento de la calidad.")]
    for titular,desc in acceso:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r1=p.add_run("• "+titular+": "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(desc); r2.font.size=Pt(11)
    B(doc,"PROHIBICIÓN: Está prohibida la entrega de información de la HC a terceros sin autorización expresa del paciente, salvo las excepciones legales. La violación de la confidencialidad de la HC puede dar lugar a sanciones éticas, disciplinarias y penales.")

    H(doc,"7. SISTEMA DE INFORMACIÓN Y TECNOLOGÍA")
    H(doc,"7.1 Historia Clínica en Papel",2)
    B(doc,"Si se utiliza historia clínica en papel:")
    hc_papel=["La escritura debe ser legible, con tinta indeleble (no lápiz)",
              "Las correcciones se realizan con una línea que no borre lo escrito, anotando la corrección y la fecha",
              "No se deben dejar espacios en blanco que puedan ser llenados posteriormente",
              "Cada página de la HC debe tener el nombre del paciente y el número de documento",
              "El archivo físico de HC se organiza por número de historia o por orden alfabético, en folder individual por paciente",
              "El archivo físico se ubica en gabinete bajo llave con acceso restringido"]
    for item in hc_papel: BL(doc,item)

    H(doc,"7.2 Historia Clínica Electrónica",2)
    B(doc,"Si se utiliza software de historia clínica electrónica:")
    hc_elec=["El software debe cumplir con los requisitos técnicos para HC electrónica (integridad, autenticidad, disponibilidad)",
             "Cada registro debe tener la firma electrónica o digital del médico",
             "El sistema debe llevar un registro de auditoría de accesos y modificaciones (log de auditoría)",
             "Se deben realizar copias de seguridad (backup) con mínima frecuencia semanal",
             "El software y los datos deben estar protegidos contra acceso no autorizado (contraseñas, cifrado)",
             "Los datos de los pacientes no deben almacenarse en servicios en la nube sin garantías de seguridad y confidencialidad compatibles con la Ley 1581/2012",
             "Ante cualquier falla del sistema, debe garantizarse la recuperación de los datos"]
    for item in hc_elec: BL(doc,item)

    H(doc,"8. AUDITORÍA DE HISTORIAS CLÍNICAS")
    B(doc,"Se realizará auditoría de historias clínicas con una periodicidad mensual, evaluando una muestra aleatoria del 10% de las HC del período:")
    auditoria=[("Datos de identificación completos","□ Sí  □ No  □ Parcial"),
               ("Motivo de consulta registrado","□ Sí  □ No"),
               ("Anamnesis completa","□ Sí  □ No  □ Parcial"),
               ("Examen físico registrado","□ Sí  □ No  □ Parcial"),
               ("Diagnóstico con código CIE-10","□ Sí  □ No"),
               ("Plan de manejo documentado","□ Sí  □ No  □ Parcial"),
               ("Firma del médico","□ Sí  □ No"),
               ("Legibilidad (si es en papel)","□ Legible  □ Parcialmente  □ Ilegible"),
               ("Consentimiento informado cuando aplica","□ Sí  □ No  □ No aplica")]
    tb2=doc.add_table(rows=len(auditoria)+1,cols=2); tb2.style='Table Grid'
    hrow(tb2,["Criterio de auditoría","Cumplimiento"])
    for i,(crit,cum) in enumerate(auditoria):
        tb2.rows[i+1].cells[0].text=crit; tb2.rows[i+1].cells[1].text=cum
        for c in tb2.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()

    H(doc,"9. INDICADORES")
    tb3=doc.add_table(rows=5,cols=4); tb3.style='Table Grid'
    hrow(tb3,["Indicador","Fórmula","Meta","Frecuencia"])
    for i,row in enumerate([
        ("Calidad de HC (HC completas)","(HC con todos los campos/Total HC auditadas)×100","≥95%","Mensual"),
        ("% HC con CIE-10 registrado","(HC con CIE-10/Total HC)×100","100%","Mensual"),
        ("% HC firmadas por el médico","(HC firmadas/Total HC)×100","100%","Mensual"),
        ("Tiempo de respuesta a solicitud de copia HC","Días hábiles desde solicitud hasta entrega (promedio)","≤3 días hábiles","Cuando aplica")]):
        for j,v in enumerate(row): tb3.rows[i+1].cells[j].text=v
        for c in tb3.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph(); firmas(doc)
    path=os.path.join(BASE6,"MAN-HC-001_Manual_Historia_Clinica.docx")
    doc.save(path); print(f"✓ Creado: {path}")

# ===== PRO-HC-001 =====
def crear_pro_hc():
    doc=Document(); margins(doc); hf(doc,"PRO-HC-001","Proceso de Gestión de Historia Clínica")
    portada(doc,"PRO-HC-001","PROCESO DE GESTIÓN DE LA HISTORIA CLÍNICA")
    H(doc,"1. OBJETIVO")
    B(doc,"Definir el proceso para la apertura, diligenciamiento, archivo, custodia, préstamo y cierre de las historias clínicas del Consultorio [NOMBRE DEL CONSULTORIO], garantizando la integridad, confidencialidad y disponibilidad de la información clínica de los pacientes.")
    H(doc,"2. DESCRIPCIÓN DEL PROCESO")
    pasos=[("Apertura de HC","Paciente nuevo","Crear el expediente con los datos de identificación completos. Asignar número de HC. Crear carpeta física o registro en el sistema electrónico."),
           ("Diligenciamiento","Cada consulta","Registrar la atención en el formato de HC correspondiente: FOR-HC-001 (medicina general) o FOR-HC-002 (estética). Todos los campos obligatorios deben estar completos."),
           ("Firma","Inmediatamente post-consulta","El médico firma y sella (TP) cada registro de atención. En HC electrónica: firmar electrónicamente o con contraseña."),
           ("Archivo","Post-consulta","Archivar la HC en orden cronológico dentro del expediente del paciente. Guardar en el archivo físico bajo llave."),
           ("Préstamo","Cuando se solicita","Registrar el préstamo en el libro de control de archivo: fecha, nombre del solicitante, número de HC, propósito. Devolver en máximo 24 horas."),
           ("Auditoría","Mensual","Revisar aleatoriamente 10% de las HC del período según los criterios del MAN-HC-001."),
           ("Destrucción/baja","Después de 15 años","Solo se pueden destruir HC después de haber cumplido el tiempo mínimo de conservación y con acta de baja. Las HC electrónicas se eliminan de los servidores con método seguro.")]
    tb=doc.add_table(rows=len(pasos)+1,cols=4); tb.style='Table Grid'
    hrow(tb,["Actividad","Cuando","Descripción","Responsable"])
    for i,(act,cdo,desc) in enumerate(pasos):
        tb.rows[i+1].cells[0].text=act; tb.rows[i+1].cells[1].text=cdo
        tb.rows[i+1].cells[2].text=desc; tb.rows[i+1].cells[3].text="[NOMBRE DE LA MÉDICA]"
        for c in tb.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph(); firmas(doc)
    path=os.path.join(BASE6,"PRO-HC-001_Proceso_Gestion_Historia_Clinica.docx")
    doc.save(path); print(f"✓ Creado: {path}")

# ===== FOR-HC-001 =====
def crear_for_hc001():
    doc=Document(); margins(doc); hf(doc,"FOR-HC-001","Historia Clínica Medicina General")
    portada(doc,"FOR-HC-001","HISTORIA CLÍNICA - CONSULTA DE MEDICINA GENERAL")

    H(doc,"I. DATOS DE IDENTIFICACIÓN")
    tb=doc.add_table(rows=6,cols=4); tb.style='Table Grid'
    id_data=[("N° de Historia Clínica:","","Fecha de apertura:",""),
             ("Apellidos:","","Nombres:",""),
             ("Tipo de documento:","□ CC  □ CE  □ TI  □ Pasaporte","N° documento:",""),
             ("Fecha nacimiento:","","Edad:","Sexo: □ M  □ F  □ Otro"),
             ("Dirección:","","Ciudad:","Teléfono:"),
             ("EPS / Aseguradora:","","Tipo de afiliación:","□ Cotizante  □ Beneficiario  □ Particular")]
    for i,row in enumerate(id_data):
        for j,v in enumerate(row): tb.rows[i].cells[j].text=v
        if tb.rows[i].cells[0].paragraphs[0].runs:
            tb.rows[i].cells[0].paragraphs[0].runs[0].font.bold=True
    doc.add_paragraph()

    H(doc,"II. ANTECEDENTES PERSONALES (diligenciar en primera consulta y actualizar)")
    tb2=doc.add_table(rows=8,cols=2); tb2.style='Table Grid'
    hrow(tb2,["Antecedente","Descripción"])
    ants=[("Patológicos (enfermedades crónicas)",""),
          ("Quirúrgicos",""),
          ("Traumáticos",""),
          ("Alérgicos (medicamentos, alimentos, látex)",""),
          ("Farmacológicos (medicamentos actuales)",""),
          ("Familiares (enfermedades hereditarias relevantes)",""),
          ("Gineco-obstétricos (mujeres)","G: __ P: __ A: __ C: __ FUR: ________  ACO: ")]
    for i,(ant,_) in enumerate(ants):
        tb2.rows[i+1].cells[0].text=ant
        for c in tb2.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()

    H(doc,"III. HÁBITOS Y ESTILOS DE VIDA")
    tb3=doc.add_table(rows=1,cols=4); tb3.style='Table Grid'
    habs=["Tabaquismo: □ No  □ Sí (___/día, desde ___)","Alcohol: □ No  □ Ocasional  □ Frecuente","Ejercicio: □ No  □ Sí (___/semana)","Alimentación: □ Adecuada  □ Inadecuada"]
    row=tb3.rows[0]
    for i,h in enumerate(habs):
        row.cells[i].text=h
        for r in row.cells[i].paragraphs[0].runs: r.font.size=Pt(10)
    doc.add_paragraph()

    H(doc,"IV. REGISTRO DE CONSULTAS")
    for n in range(3):
        doc.add_paragraph()
        p_fecha=doc.add_paragraph(); r=p_fecha.add_run(f"CONSULTA N° ___  Fecha: ____________  Hora: ______")
        r.font.bold=True; r.font.size=Pt(11)
        tb_c=doc.add_table(rows=7,cols=2); tb_c.style='Table Grid'
        campos_consulta=[
            ("MOTIVO DE CONSULTA",""),
            ("ANAMNESIS (Historia de la enfermedad actual: inicio, evolución, síntomas, tratamiento previo)",""),
            ("EXAMEN FÍSICO\nT/A: _____/_____ FC: _____ FR: _____ T°: _____ SpO2: _____% Talla: _____ Peso: _____ IMC: _____\n\nHallazgos por sistemas:",""),
            ("DIAGNÓSTICO(S) CIE-10\n1. ___________________ (__.__.__)\n2. ___________________ (__.__.__)\n3. ___________________ (__.__.__)",""),
            ("PLAN DE MANEJO\n□ Medicamentos (ver fórmula)\n□ Exámenes solicitados (ver orden)\n□ Remisión a: ___________________\n□ Incapacidad: ___ días (desde: _____ hasta: _____)\n□ Indicaciones médicas:",""),
            ("EDUCACIÓN AL PACIENTE\n(signos de alarma, cambios de estilo de vida, indicaciones)",""),
            ("Nombre médica: [NOMBRE DE LA MÉDICA]  TP N°: [N°]\nFirma: _________________________  Sello: ",""),
        ]
        for i,(campo,_) in enumerate(campos_consulta):
            tb_c.rows[i].cells[0].text=campo
            for r2 in tb_c.rows[i].cells[0].paragraphs[0].runs: r2.font.bold=True; r2.font.size=Pt(10)
        doc.add_paragraph()

    path=os.path.join(BASE6,"FOR-HC-001_Historia_Clinica_Medicina_General.docx")
    doc.save(path); print(f"✓ Creado: {path}")

# ===== FOR-HC-002 =====
def crear_for_hc002():
    doc=Document(); margins(doc); hf(doc,"FOR-HC-002","Historia Clínica Estética")
    portada(doc,"FOR-HC-002","HISTORIA CLÍNICA ESTÉTICA MÉDICA")

    H(doc,"I. DATOS DE IDENTIFICACIÓN")
    tb=doc.add_table(rows=5,cols=4); tb.style='Table Grid'
    datos=[("N° HC:","","Fecha:",""),
           ("Apellidos y Nombres:","","N° Documento:",""),
           ("Fecha nac.:","","Edad:","Fototipo Fitzpatrick: I  II  III  IV  V  VI"),
           ("Teléfono:","","Correo:",""),
           ("¿Cómo nos conoció?","□ Referido  □ Redes sociales  □ Google  □ Voz a voz  □ Otro:","Fecha primera consulta estética:","")]
    for i,row in enumerate(datos):
        for j,v in enumerate(row): tb.rows[i].cells[j].text=v
    doc.add_paragraph()

    H(doc,"II. ANTECEDENTES MÉDICOS RELEVANTES PARA ESTÉTICA")
    tb2=doc.add_table(rows=10,cols=2); tb2.style='Table Grid'
    hrow(tb2,["Antecedente","Sí / No / Descripción"])
    ants_est=["Embarazo o lactancia actual","Enfermedades autoinmunes (lupus, artritis, etc.)","Trastornos de la coagulación","Tratamiento anticoagulante o antiagregante","Isotretinoína en los últimos 6 meses","Queloides o cicatrices hipertróficas","Herpes oral o labial recurrente","Alergias a medicamentos o productos de uso tópico","Procedimientos estéticos previos (especificar)"]
    for i,ant in enumerate(ants_est):
        tb2.rows[i+1].cells[0].text=ant
        tb2.rows[i+1].cells[1].text="□ No  □ Sí: ___________________________"
        for c in tb2.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()

    H(doc,"III. EXPECTATIVAS Y MOTIVO DE CONSULTA ESTÉTICA")
    doc.add_paragraph("Motivo de consulta en palabras del paciente: ______________________________________________________")
    doc.add_paragraph("Expectativas: ____________________________________________________________________________")
    doc.add_paragraph("¿Son expectativas realistas?: □ Sí  □ No  □ Requieren orientación")
    doc.add_paragraph()

    H(doc,"IV. VALORACIÓN MÉDICA ESTÉTICA")
    H(doc,"4.1 Análisis Facial",2)
    tb3=doc.add_table(rows=9,cols=3); tb3.style='Table Grid'
    hrow(tb3,["Área evaluada","Hallazgo","Tratamiento sugerido"])
    areas=["Tercio superior (frente, cejas)","Tercio medio (nariz, mejillas, surcos)","Tercio inferior (labios, mentón, mandíbula)","Área periorbicular (patas de gallo, ojeras)","Cuello y escote","Textura general de la piel","Hidratación y luminosidad","Tono muscular y laxitud"]
    for i,area in enumerate(areas):
        tb3.rows[i+1].cells[0].text=area
        for c in tb3.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()

    H(doc,"4.2 Clasificación de Glogau (envejecimiento fotográfico)",2)
    B(doc,"□ Tipo I: Sin arrugas. Sin queratosis. 20-30 años, fototipo claro\n□ Tipo II: Arrugas en movimiento. Queratosis incipiente. 30-40 años\n□ Tipo III: Arrugas en reposo. Discromías. 50-60 años\n□ Tipo IV: Arrugas por toda la piel. Queratosis grave. >60 años")

    H(doc,"V. REGISTRO DE PROCEDIMIENTOS ESTÉTICOS")
    for n in range(3):
        doc.add_paragraph()
        p2=doc.add_paragraph(); r=p2.add_run(f"PROCEDIMIENTO N° ___")
        r.font.bold=True; r.font.size=Pt(12); r.font.color.rgb=RGBColor(31,73,125)
        tb_p=doc.add_table(rows=9,cols=2); tb_p.style='Table Grid'
        campos_p=[("Fecha y hora del procedimiento:",""),
                  ("Tipo de procedimiento:","□ Toxina botulínica  □ Relleno HA  □ Peeling  □ Mesoterapia  □ PRP  □ Otro:"),
                  ("Área(s) tratada(s):",""),
                  ("Producto utilizado (nombre, lote, venc, INVIMA):",""),
                  ("Técnica utilizada:",""),
                  ("Dosis/cantidades:",""),
                  ("Complicaciones intra-procedimiento:","□ Ninguna  □ Especificar: ___"),
                  ("Indicaciones post-procedimiento dadas:",""),
                  ("Próxima cita:","Firma médica: _________________________  TP: [N°]")]
        for i,(campo,val) in enumerate(campos_p):
            tb_p.rows[i].cells[0].text=campo; tb_p.rows[i].cells[1].text=val
            for r2 in tb_p.rows[i].cells[0].paragraphs[0].runs: r2.font.bold=True; r2.font.size=Pt(10)

    path=os.path.join(BASE6,"FOR-HC-002_Historia_Clinica_Estetica.docx")
    doc.save(path); print(f"✓ Creado: {path}")

# ===== CARPETA 7: REFERENCIA Y CONTRARREFERENCIA =====
def crear_man_rcr():
    doc=Document(); margins(doc); hf(doc,"MAN-RCR-001","Manual de Referencia y Contrarreferencia")
    portada(doc,"MAN-RCR-001","MANUAL DE REFERENCIA Y CONTRARREFERENCIA")

    H(doc,"1. OBJETIVO")
    B(doc,"Establecer los lineamientos para el proceso de referencia (remisión) y contrarreferencia de pacientes del Consultorio [NOMBRE DEL CONSULTORIO], garantizando la continuidad y oportunidad en la atención de los pacientes que requieran ser valorados por profesionales de mayor especialización o en servicios de mayor complejidad, y asegurando el retorno adecuado de información al médico remitente.")

    H(doc,"2. MARCO LEGAL")
    for n in ["Resolución 3100 de 2019 - Estándares de referencia y contrarreferencia",
              "Resolución 5596 de 2015 - Criterios técnicos para el Sistema de Selección y Clasificación (Triage)",
              "Decreto 780 de 2016 - Decreto Único del Sector Salud",
              "Ley 100 de 1993 - SGSSS, red de prestadores de servicios",
              "Resolución 1044 de 2021 y demás normas de red de servicios de salud"]: BL(doc,n)

    H(doc,"3. DEFINICIONES")
    defs=[("Referencia (Remisión):","Proceso por el cual el prestador de servicios de salud de menor complejidad remite a un paciente a un prestador de mayor complejidad o a otro servicio que no tiene disponible, con toda la información clínica necesaria para la continuidad de la atención."),
          ("Contrarreferencia:","Proceso por el cual el prestador de mayor complejidad devuelve al paciente al médico remitente (o a la red primaria), con el informe de la valoración, el diagnóstico actualizado y las recomendaciones de manejo."),
          ("Urgencia:","Condición clínica que implica daño o riesgo inminente para la vida, y que requiere atención médica inmediata e impostergable."),
          ("Red de prestadores:","Conjunto de prestadores de servicios de salud articulados para garantizar la atención integral a los pacientes en los diferentes niveles de complejidad.")]
    for term,defi in defs:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r1=p.add_run(term+" "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(defi); r2.font.size=Pt(11)

    H(doc,"4. CRITERIOS DE REMISIÓN DESDE EL CONSULTORIO")
    criterios=[("Urgencias médicas","Cualquier condición que represente riesgo inminente para la vida: dolor torácico agudo, dificultad respiratoria severa, alteración del estado de conciencia, convulsiones, sangrado importante. REMISIÓN INMEDIATA a urgencias del hospital/clínica más cercana o activar servicio de emergencias 123."),
               ("Condiciones que requieren especialista","Patologías que superan la capacidad resolutiva de la consulta de medicina general de primer nivel: sospecha de enfermedad neoplásica, condiciones psiquiátricas, problemas cardiovasculares complejos, enfermedades dermatológicas complejas, entre otras."),
               ("Procedimientos no disponibles","Cuando el paciente requiere procedimientos diagnósticos o terapéuticos no disponibles en el consultorio: cirugía, hospitalización, exámenes especializados."),
               ("Criterio del médico","Cuando la médica propietaria considera que la seguridad y bienestar del paciente se optimizan con una valoración complementaria.")]
    for crit,desc in criterios:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r1=p.add_run("• "+crit+": "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(desc); r2.font.size=Pt(11)

    H(doc,"5. PROCESO DE REMISIÓN")
    pasos_rem=[("1. Decisión médica","La médica determina que el paciente requiere valoración o atención en otro nivel."),
               ("2. Explicación al paciente","Se explica al paciente el motivo de la remisión y el proceso a seguir con su asegurador."),
               ("3. Diligenciamiento del formulario","Se diligencia el FOR-RCR-001 con toda la información clínica relevante."),
               ("4. Indicaciones al paciente","Se indica al paciente cómo gestionar la cita con el especialista o cómo ir a urgencias si corresponde."),
               ("5. Para urgencias","Si es urgencia, se acompaña al paciente mientras llega el servicio de emergencias o se coordina el transporte inmediato."),
               ("6. Registro","Se registra la remisión en la historia clínica del paciente."),
               ("7. Seguimiento","En la siguiente consulta del paciente, se verifica si fue valorado por el especialista y se solicita el informe de contrarreferencia.")]
    tb=doc.add_table(rows=len(pasos_rem)+1,cols=2); tb.style='Table Grid'
    hrow(tb,["Paso","Descripción"])
    for i,(p2,d) in enumerate(pasos_rem):
        tb.rows[i+1].cells[0].text=p2; tb.rows[i+1].cells[1].text=d
        for c in tb.rows[i+1].cells:
            for p3 in c.paragraphs:
                for r in p3.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"6. REMISIONES FRECUENTES DESDE MEDICINA GENERAL")
    remisiones=[("Medicina Interna / Cardiología","Hipertensión no controlada, sospecha de cardiopatía, alteraciones del EKG"),
                ("Endocrinología","Diabetes no controlada, hipotiroidismo, obesidad compleja"),
                ("Dermatología","Lesiones sospechosas de malignidad, psoriasis, dermatitis atópica grave, acné severo, melasma refractario"),
                ("Gastroenterología","Dolor abdominal crónico, cambios en el hábito intestinal, sangrado digestivo"),
                ("Ginecología","Alteraciones del ciclo menstrual, miomas, nódulos mamarios, citología anormal"),
                ("Psiquiatría/Psicología","Depresión, ansiedad, trastornos del sueño que no responden al manejo de primer nivel"),
                ("Cirugía","Hernias, hemorroides, quistes, masas palpables que requieren intervención quirúrgica"),
                ("Neurología","Cefaleas crónicas refractarias, sospecha de accidente cerebrovascular, convulsiones"),
                ("Urgencias hospitalarias","Cualquier situación de riesgo vital que supere la capacidad del consultorio")]
    tb2=doc.add_table(rows=len(remisiones)+1,cols=2); tb2.style='Table Grid'
    hrow(tb2,["Especialidad a remitir","Indicaciones frecuentes"])
    for i,(esp,ind) in enumerate(remisiones):
        tb2.rows[i+1].cells[0].text=esp; tb2.rows[i+1].cells[1].text=ind
        for c in tb2.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"7. INDICADORES")
    tb3=doc.add_table(rows=4,cols=4); tb3.style='Table Grid'
    hrow(tb3,["Indicador","Fórmula","Meta","Frecuencia"])
    for i,row in enumerate([
        ("% remisiones con formulario diligenciado completo","(Remisiones con formulario completo/Total remisiones)×100","100%","Mensual"),
        ("% remisiones a urgencias con traslado asegurado","(Traslados confirmados/Remisiones urgentes)×100","100%","Mensual"),
        ("Tasa de remisión","N° remisiones/100 consultas","Seguimiento de tendencia","Mensual")]):
        for j,v in enumerate(row): tb3.rows[i+1].cells[j].text=v
        for c in tb3.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph(); firmas(doc)
    path=os.path.join(BASE7,"MAN-RCR-001_Manual_Referencia_Contrarreferencia.docx")
    doc.save(path); print(f"✓ Creado: {path}")

def crear_for_rcr():
    doc=Document(); margins(doc); hf(doc,"FOR-RCR-001","Formato de Remisión")
    portada(doc,"FOR-RCR-001","FORMULARIO DE REFERENCIA (REMISIÓN)")

    H(doc,"DATOS DEL PRESTADOR REMITENTE")
    tb=doc.add_table(rows=4,cols=4); tb.style='Table Grid'
    datos=[("Institución:","CONSULTORIO [NOMBRE DEL CONSULTORIO]","Código REPS:","[CÓDIGO REPS]"),
           ("Dirección:","[DIRECCIÓN DEL CONSULTORIO]","Ciudad:","[CIUDAD]"),
           ("Teléfono:","[TELÉFONO]","Correo:","[CORREO ELECTRÓNICO]"),
           ("Médico remitente:","Dra. [NOMBRE DE LA MÉDICA]","Tarjeta Prof.:","[N° TP]")]
    for i,row in enumerate(datos):
        for j,v in enumerate(row): tb.rows[i].cells[j].text=v
        tb.rows[i].cells[0].paragraphs[0].runs[0].font.bold=True
        tb.rows[i].cells[2].paragraphs[0].runs[0].font.bold=True
    doc.add_paragraph()

    H(doc,"DATOS DEL PACIENTE")
    tb2=doc.add_table(rows=4,cols=4); tb2.style='Table Grid'
    datos2=[("Nombre completo:","","Fecha nacimiento:",""),
            ("Tipo doc:","□ CC  □ CE  □ TI","N° documento:",""),
            ("Dirección:","","Teléfono:",""),
            ("EPS / Aseguradora:","","N° de afiliación:","")]
    for i,row in enumerate(datos2):
        for j,v in enumerate(row): tb2.rows[i].cells[j].text=v
    doc.add_paragraph()

    H(doc,"DATOS DE LA REMISIÓN")
    tb3=doc.add_table(rows=4,cols=4); tb3.style='Table Grid'
    datos3=[("Fecha de la remisión:","","Hora:",""),
            ("Tipo de remisión:","□ Urgente  □ Prioritaria  □ Programada","Especialidad / servicio:",""),
            ("Institución destino:","","Nivel de atención:","□ I  □ II  □ III  □ IV"),
            ("Diagnóstico principal (CIE-10):","","Código:","")]
    for i,row in enumerate(datos3):
        for j,v in enumerate(row): tb3.rows[i].cells[j].text=v
    doc.add_paragraph()

    H(doc,"INFORMACIÓN CLÍNICA PARA LA CONTINUIDAD DEL CUIDADO")
    B(doc,"MOTIVO DE LA REMISIÓN:",bold=True)
    for i in range(3):
        doc.add_paragraph("___________________________________________________________________________")
    B(doc,"RESUMEN DE LA HISTORIA CLÍNICA RELEVANTE:",bold=True)
    B(doc,"Antecedentes relevantes: ___________________________________________________________________\nMedicamentos actuales: _____________________________________________________________________\nAlergias: __________________________________________________________________________________")
    B(doc,"HALLAZGOS CLÍNICOS ACTUALES:",bold=True)
    for i in range(4):
        doc.add_paragraph("___________________________________________________________________________")
    B(doc,"EXÁMENES REALIZADOS (adjuntar resultados):",bold=True)
    for i in range(3):
        doc.add_paragraph("___________________________________________________________________________")
    B(doc,"TRATAMIENTO ACTUAL:",bold=True)
    for i in range(2):
        doc.add_paragraph("___________________________________________________________________________")
    B(doc,"SOLICITUD ESPECÍFICA AL ESPECIALISTA:",bold=True)
    for i in range(2):
        doc.add_paragraph("___________________________________________________________________________")

    doc.add_paragraph()
    H(doc,"FIRMAS")
    tb4=doc.add_table(rows=2,cols=2); tb4.style='Table Grid'
    hrow(tb4,["MÉDICO REMITENTE","PACIENTE / ACOMPAÑANTE"])
    tb4.rows[1].cells[0].text="Dra. [NOMBRE DE LA MÉDICA]\nMédica General - TP N°[N°]\nFirma: _________________________\nFecha: _________________________"
    tb4.rows[1].cells[1].text="Nombre: _________________________\nFirma: _________________________\nFecha: _________________________\n¿Acepta la remisión?: □ Sí  □ No"
    for j in range(2):
        tb4.rows[1].cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    H(doc,"PARA CONTRARREFERENCIA (DILIGENCIAR POR EL ESPECIALISTA)")
    B(doc,"Institución donde se atendió: ________________________________________________________________\nEspecialista: ________________________________________________________________________________\nFecha de la consulta: ________________________________________________________________________\nDiagnóstico(s): ______________________________________________________________________________\nTratamiento instaurado: _____________________________________________________________________\nRecomendaciones para el médico general: ______________________________________________________\nPróxima cita en especialidad: ________________________________________________________________")
    doc.add_paragraph()
    doc.add_paragraph("Firma del especialista: _________________________  TP: _______________  Fecha: ___________")

    path=os.path.join(BASE7,"FOR-RCR-001_Formato_Remision.docx")
    doc.save(path); print(f"✓ Creado: {path}")

# ===== CARPETA 8: GESTIÓN DE CALIDAD =====
def crear_man_gc():
    doc=Document(); margins(doc); hf(doc,"MAN-GC-001","Manual de Calidad")
    portada(doc,"MAN-GC-001","MANUAL DE GESTIÓN DE LA CALIDAD")

    H(doc,"1. OBJETIVO")
    B(doc,"Establecer el Sistema de Gestión de la Calidad del Consultorio Médico [NOMBRE DEL CONSULTORIO], definiendo la política, los objetivos, los procesos, los indicadores y las herramientas de mejora continua que garanticen la prestación de servicios de salud con los más altos estándares de calidad, seguridad y humanización, en cumplimiento del Sistema Obligatorio de Garantía de Calidad de la Atención en Salud (SOGCS) establecido por el Decreto 1011 de 2006.")

    H(doc,"2. MARCO LEGAL")
    normas=[("Decreto 1011 de 2006","Establece el SOGCS con sus 4 componentes: Sistema Único de Habilitación, PAMEC, Sistema de Información, Acreditación."),
            ("Resolución 3100 de 2019","Sistema Único de Habilitación actualizado."),
            ("Resolución 0256 de 2016","Sistema de Información para la Calidad y los indicadores de calidad."),
            ("Resolución 2082 de 2014","Disposiciones para la operacionalización del PAMEC."),
            ("ISO 9001:2015","Sistema de Gestión de Calidad (referencia internacional)."),
            ("Ley 872 de 2003","Crea el Sistema de Gestión de Calidad en las entidades públicas (referencia)."),
            ("Decreto 780 de 2016","Decreto Único del Sector Salud. Consolida normativa de calidad en salud.")]
    tb=doc.add_table(rows=len(normas)+1,cols=2); tb.style='Table Grid'
    hrow(tb,["Norma","Descripción"])
    for i,(n,d) in enumerate(normas):
        tb.rows[i+1].cells[0].text=n; tb.rows[i+1].cells[1].text=d
        for c in tb.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()

    H(doc,"3. POLÍTICA DE CALIDAD")
    B(doc,"El Consultorio Médico [NOMBRE DEL CONSULTORIO], bajo la dirección de la Dra. [NOMBRE DE LA MÉDICA], se compromete a prestar servicios de salud de alta calidad, seguros y humanizados, orientados a la satisfacción de los pacientes y al mejoramiento continuo de sus procesos, mediante el cumplimiento de los requisitos legales de habilitación, el desarrollo de las competencias del talento humano y la implementación de herramientas sistemáticas de gestión de la calidad.")

    H(doc,"4. OBJETIVOS DE CALIDAD")
    objetivos=["Mantener el 100% de cumplimiento de los estándares de habilitación de la Resolución 3100 de 2019",
               "Lograr una satisfacción de los usuarios ≥90%",
               "Garantizar que el 100% de los eventos adversos sean reportados y analizados",
               "Mantener la calidad de las historias clínicas con un cumplimiento ≥95% de los criterios de auditoría",
               "Cumplir el plan anual de capacitación en ≥80%",
               "Implementar y medir el 100% de los indicadores de calidad definidos",
               "Mejorar continuamente los procesos de atención identificando y gestionando los riesgos"]
    for i,obj in enumerate(objetivos,1):
        p=doc.add_paragraph()
        r1=p.add_run(f"OC-{i:02d}: "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(obj); r2.font.size=Pt(11)

    H(doc,"5. SISTEMA DE GESTIÓN DE CALIDAD - CICLO PHVA")
    B(doc,"El sistema de gestión de calidad del consultorio se basa en el ciclo PHVA (Planear-Hacer-Verificar-Actuar) como herramienta de mejora continua:")
    phva=[("PLANEAR","Definir los procesos, procedimientos, indicadores y metas de calidad.\nPlanificar el plan anual de capacitación.\nIdentificar y evaluar los riesgos de los procesos.\nPlanificar las auditorías internas y el PAMEC."),
          ("HACER","Ejecutar los procesos según los procedimientos documentados.\nImplementar los planes de capacitación y mejoramiento.\nAtender a los pacientes según los protocolos definidos.\nRegistrar la información requerida para los indicadores."),
          ("VERIFICAR","Medir los indicadores de calidad mensualmente.\nRealizar auditorías internas de historias clínicas.\nAnalizar las quejas, reclamos y eventos adversos.\nEvaluar el cumplimiento de los estándares de habilitación."),
          ("ACTUAR","Implementar acciones correctivas ante no conformidades.\nImplementar acciones preventivas ante riesgos identificados.\nDocumentar las lecciones aprendidas.\nActualizar los procesos y procedimientos cuando sea necesario."
           )]
    tb2=doc.add_table(rows=len(phva)+1,cols=2); tb2.style='Table Grid'
    hrow(tb2,["Fase PHVA","Actividades del consultorio"])
    for i,(fase,act) in enumerate(phva):
        tb2.rows[i+1].cells[0].text=fase; tb2.rows[i+1].cells[1].text=act
        for c in tb2.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"6. PAMEC (PROGRAMA DE AUDITORÍA PARA EL MEJORAMIENTO DE LA CALIDAD)")
    B(doc,"En cumplimiento de la Resolución 2082 de 2014, el Consultorio implementa el PAMEC con las siguientes actividades:")
    H(doc,"6.1 Autoevaluación",2)
    B(doc,"Se realiza una autoevaluación anual del cumplimiento de los estándares de habilitación y de los indicadores de calidad. La autoevaluación se documenta y sirve como base para el plan de mejoramiento.")
    H(doc,"6.2 Selección de Procesos a Mejorar",2)
    B(doc,"Los procesos a mejorar se seleccionan con base en:\n• Resultados de los indicadores de calidad (procesos con indicadores por debajo de la meta)\n• Análisis de eventos adversos y quejas recurrentes\n• Hallazgos de la autoevaluación y las auditorías internas\n• Oportunidades de mejora identificadas por el personal")
    H(doc,"6.3 Plan de Mejoramiento",2)
    B(doc,"Para cada proceso a mejorar se elabora un plan de mejoramiento con: descripción del problema, causa raíz, meta de mejoramiento, acciones a implementar, responsable, plazo y forma de seguimiento.")

    H(doc,"7. MAPA DE PROCESOS DEL CONSULTORIO")
    B(doc,"El Consultorio [NOMBRE DEL CONSULTORIO] organiza sus procesos en tres niveles:")
    procesos_mapa=[("Procesos Estratégicos","Gestión de la calidad y seguridad del paciente, planificación estratégica del consultorio, gestión del talento humano"),
                   ("Procesos Misionales (Asistenciales)","Consulta de medicina general, procedimientos estéticos no invasivos, referencia y contrarreferencia, gestión de la historia clínica"),
                   ("Procesos de Apoyo","Gestión de infraestructura, gestión de equipos biomédicos, gestión de medicamentos, gestión documental")]
    for tipo,desc in procesos_mapa:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r1=p.add_run("• "+tipo+": "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(desc); r2.font.size=Pt(11)

    H(doc,"8. INDICADORES DE CALIDAD")
    B(doc,"Los siguientes indicadores de calidad se miden, analizan y reportan mensualmente:")
    indicadores=[
        ("Satisfacción del usuario","(N° usuarios satisfechos / N° usuarios encuestados) × 100","≥90%","Mensual","Encuesta de satisfacción"),
        ("Oportunidad en la asignación de cita de medicina general","% de citas asignadas en ≤3 días hábiles","≥90%","Mensual","Registro de agendamiento"),
        ("Calidad de historia clínica","(HC con todos los componentes requeridos / HC auditadas) × 100","≥95%","Mensual","Auditoría de HC"),
        ("Tasa de eventos adversos","N° de eventos adversos / 100 atenciones","Tendencia decreciente","Mensual","Registro de eventos"),
        ("% eventos adversos analizados","(EA con análisis de causa raíz / Total EA) × 100","100%","Mensual","FOR-PP-005"),
        ("Tasa de quejas y reclamos","N° quejas / 100 consultas","Tendencia decreciente","Mensual","Registro de quejas"),
        ("% cumplimiento plan de capacitación","Capacitaciones realizadas / Programadas × 100","≥80%","Anual","Registro de capacitaciones"),
        ("% cumplimiento estándares habilitación","Estándares cumplidos / Total estándares × 100","100%","Semestral","Autoevaluación"),
        ("Consentimientos informados obtenidos","CI firmados / Procedimientos que requieren CI × 100","100%","Mensual","HC y registros"),
        ("% medicamentos sin vencer","Medicamentos vigentes / Total medicamentos × 100","100%","Mensual","FOR-MED-001"),
    ]
    tb3=doc.add_table(rows=len(indicadores)+1,cols=5); tb3.style='Table Grid'
    hrow(tb3,["Indicador","Fórmula","Meta","Frecuencia","Fuente de datos"])
    for i,row in enumerate(indicadores):
        for j,v in enumerate(row): tb3.rows[i+1].cells[j].text=v
        for c in tb3.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(8)
    doc.add_paragraph()

    H(doc,"9. GESTIÓN DE RIESGOS")
    B(doc,"El consultorio implementa un proceso de identificación y gestión de riesgos que contempla:")
    riesgos=[("Riesgo clínico asistencial","Error diagnóstico, evento adverso por medicamento, complicación de procedimiento","Media","Protocolos clínicos, doble verificación, capacitación"),
             ("Riesgo de seguridad del paciente","Caída del paciente, identificación incorrecta, infección asociada","Media","Protocolos SP, lavado de manos, identificación activa"),
             ("Riesgo de habilitación","No cumplimiento de estándares, documentación desactualizada","Media","Auditorías periódicas, actualización documental"),
             ("Riesgo operacional","Falla de equipos, corte de servicios públicos, ausencia de la médica","Baja","Mantenimiento preventivo, plan de contingencia"),
             ("Riesgo legal/normativo","Cambios en normatividad, sanciones por incumplimiento","Baja","Monitoreo normativo, asesoría jurídica"),
             ("Riesgo de seguridad de la información","Pérdida de historias clínicas, acceso no autorizado","Baja","Archivo seguro, backup de HC electrónica, contraseñas")]
    tb4=doc.add_table(rows=len(riesgos)+1,cols=4); tb4.style='Table Grid'
    hrow(tb4,["Tipo de riesgo","Descripción","Probabilidad","Controles"])
    for i,row in enumerate(riesgos):
        for j,v in enumerate(row): tb4.rows[i+1].cells[j].text=v
        for c in tb4.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph(); firmas(doc)
    path=os.path.join(BASE8,"MAN-GC-001_Manual_Calidad.docx")
    doc.save(path); print(f"✓ Creado: {path}")

def crear_pol_gc():
    doc=Document(); margins(doc); hf(doc,"POL-GC-001","Política de Calidad y Seguridad")
    portada(doc,"POL-GC-001","POLÍTICA DE CALIDAD Y SEGURIDAD DEL PACIENTE")

    H(doc,"POLÍTICA DE CALIDAD")
    B(doc,'"El Consultorio Médico [NOMBRE DEL CONSULTORIO], liderado por la Dra. [NOMBRE DE LA MÉDICA], se compromete a prestar servicios de salud seguros, humanizados y de alta calidad científica y técnica, orientados a la satisfacción integral de nuestros pacientes. Fundamentamos nuestra práctica en los principios de la ética médica, el respeto a los derechos de los usuarios, el trabajo en equipo y la mejora continua, cumpliendo en todo momento con los requisitos legales de habilitación y los estándares del Sistema Obligatorio de Garantía de Calidad de la Atención de Salud (SOGCS) establecidos en el Decreto 1011 de 2006 y la Resolución 3100 de 2019."', bold=True)

    H(doc,"POLÍTICA DE SEGURIDAD DEL PACIENTE")
    B(doc,'"El Consultorio [NOMBRE DEL CONSULTORIO] asume la seguridad del paciente como valor fundamental e irrenunciable de su práctica médica. Nos comprometemos a prevenir activamente la ocurrencia de eventos adversos, a promover una cultura de reporte abierto y aprendizaje de los errores, a garantizar la identificación correcta de cada paciente, a mantener prácticas seguras de bioseguridad y manejo de medicamentos, y a informar a nuestros pacientes sobre cualquier evento adverso que ocurra en el contexto de su atención, con transparencia y respeto."', bold=True)

    H(doc,"PRINCIPIOS QUE GUÍAN NUESTRA POLÍTICA")
    principios=[("Seguridad del paciente","Primum non nocere: lo primero es no hacer daño. Toda decisión clínica busca el mayor beneficio con el menor riesgo para el paciente."),
                ("Calidad científica","Aplicación de evidencia médica actualizada, guías de práctica clínica y estándares técnicos en cada decisión clínica."),
                ("Humanización","El paciente es el centro de la atención. Respeto a su dignidad, autonomía, privacidad y diversidad."),
                ("Mejora continua","Aprendizaje permanente de los resultados, los errores y las mejores prácticas. Nunca dejar de mejorar."),
                ("Ética profesional","Cumplimiento del Código de Ética Médica (Ley 23/1981) en cada acto médico."),
                ("Transparencia","Comunicación honesta con los pacientes sobre diagnósticos, riesgos, resultados y eventos adversos.")]
    for prin,desc in principios:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r1=p.add_run("• "+prin+": "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(desc); r2.font.size=Pt(11)

    doc.add_paragraph()
    B(doc,"Esta política es conocida, entendida y aplicada por todo el personal del consultorio. Es revisada anualmente y actualizada cuando las condiciones del entorno o los resultados de calidad así lo requieran.", bold=False)
    doc.add_paragraph()
    doc.add_paragraph("Adoptada por: Dra. [NOMBRE DE LA MÉDICA] - Médica General - Propietaria\n\nFirma: _________________________\n\nFecha de adopción: " + datetime.date.today().strftime('%d/%m/%Y'))

    path=os.path.join(BASE8,"POL-GC-001_Politica_Calidad_Seguridad.docx")
    doc.save(path); print(f"✓ Creado: {path}")

def crear_pro_auditoria():
    doc=Document(); margins(doc); hf(doc,"PRO-GC-001","Proceso de Auditoría Interna")
    portada(doc,"PRO-GC-001","PROCESO DE AUDITORÍA INTERNA DE CALIDAD")

    H(doc,"1. OBJETIVO")
    B(doc,"Definir el proceso de auditoría interna del Consultorio [NOMBRE DEL CONSULTORIO], como herramienta del PAMEC para evaluar sistemáticamente el cumplimiento de los procesos, los indicadores de calidad y los estándares de habilitación, identificando oportunidades de mejora y generando planes de acción efectivos.")

    H(doc,"2. TIPOS DE AUDITORÍA INTERNA")
    tipos=[("Auditoría de historias clínicas","Mensual","Evalúa el cumplimiento de los criterios de calidad de la HC según Resolución 1995/1999 y los estándares del consultorio."),
           ("Auditoría de indicadores de calidad","Mensual","Recolección, análisis y reporte de todos los indicadores definidos en el MAN-GC-001."),
           ("Auditoría de cumplimiento de estándares de habilitación","Semestral","Autoevaluación del cumplimiento de los estándares de la Resolución 3100/2019."),
           ("Auditoría de eventos adversos","Cuando aplica / Mensual de consolidado","Análisis de todos los eventos adversos e incidentes reportados en el período."),
           ("Auditoría de satisfacción de usuarios","Mensual","Análisis de las encuestas de satisfacción y gestión de quejas.")]
    tb=doc.add_table(rows=len(tipos)+1,cols=3); tb.style='Table Grid'
    hrow(tb,["Tipo de auditoría","Frecuencia","Descripción"])
    for i,row in enumerate(tipos):
        for j,v in enumerate(row): tb.rows[i+1].cells[j].text=v
        for c in tb.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"3. PROCESO DE AUDITORÍA INTERNA")
    pasos=[("Planificación","Al inicio de cada año se elabora el Plan Anual de Auditorías, definiendo: tipo de auditoría, fecha, muestra a auditar, criterios a evaluar y responsable."),
           ("Recolección de información","Se recolectan los datos necesarios según el tipo de auditoría: revisión de HC, registros de indicadores, encuestas, observación directa."),
           ("Análisis","Se analiza la información, se calculan los indicadores, se identifican no conformidades y oportunidades de mejora."),
           ("Informe de auditoría","Se elabora un informe con los hallazgos, las causas identificadas y las recomendaciones de mejora."),
           ("Plan de mejoramiento","Se elabora un plan de mejoramiento con acciones concretas, responsables y fechas comprometidas para cada hallazgo."),
           ("Seguimiento","En la siguiente auditoría se verifica el cumplimiento de las acciones del plan de mejoramiento anterior."),
           ("Archivo","Los informes de auditoría y los planes de mejoramiento se archivan en la carpeta de gestión de calidad del consultorio.")]
    tb2=doc.add_table(rows=len(pasos)+1,cols=3); tb2.style='Table Grid'
    hrow(tb2,["Paso","Actividad","Descripción"])
    for i,(paso,act,desc) in enumerate([(str(i+1),a,d) for i,(a,d) in enumerate(pasos)]):
        tb2.rows[i+1].cells[0].text=paso; tb2.rows[i+1].cells[1].text=act; tb2.rows[i+1].cells[2].text=desc
        for c in tb2.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"4. HERRAMIENTA DE AUDITORÍA DE HISTORIA CLÍNICA")
    B(doc,"Instrucciones: Seleccionar aleatoriamente el 10% de las historias clínicas del período. Para cada HC seleccionada, evaluar los criterios según la escala: C=Conforme / NC=No Conforme / NA=No Aplica")
    criterios_hc=["1. Datos de identificación completos","2. Motivo de consulta registrado",
                  "3. Anamnesis con todos los componentes","4. Examen físico documentado",
                  "5. Diagnóstico con código CIE-10","6. Plan de manejo documentado",
                  "7. Firma del médico","8. Legibilidad","9. Consentimiento informado cuando aplica",
                  "10. Fotografías clínicas cuando aplica (procedimientos estéticos)"]
    cols_hc=["Criterio","HC-01","HC-02","HC-03","HC-04","HC-05","% Cumplimiento"]
    tb3=doc.add_table(rows=len(criterios_hc)+1,cols=len(cols_hc)); tb3.style='Table Grid'
    hrow(tb3,cols_hc)
    for i,crit in enumerate(criterios_hc):
        tb3.rows[i+1].cells[0].text=crit
        for j in range(1,6): tb3.rows[i+1].cells[j].text="C / NC / NA"
        for c in tb3.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(8)
    doc.add_paragraph()

    H(doc,"5. INDICADORES DE LA AUDITORÍA")
    tb4=doc.add_table(rows=4,cols=4); tb4.style='Table Grid'
    hrow(tb4,["Indicador","Fórmula","Meta","Frecuencia"])
    for i,row in enumerate([
        ("% conformidad global HC","(HC conformes/Total HC auditadas)×100","≥95%","Mensual"),
        ("% planes de mejora implementados a tiempo","(Acciones cumplidas/Acciones comprometidas)×100","≥80%","Mensual"),
        ("N° no conformidades recurrentes","Conteo de NC que se repiten en dos auditorías consecutivas","0","Mensual")]):
        for j,v in enumerate(row): tb4.rows[i+1].cells[j].text=v
        for c in tb4.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph(); firmas(doc)
    path=os.path.join(BASE8,"PRO-GC-001_Proceso_Auditoria_Interna.docx")
    doc.save(path); print(f"✓ Creado: {path}")

if __name__=="__main__":
    print("Generando Carpetas 6, 7 y 8...")
    crear_man_hc(); crear_pro_hc(); crear_for_hc001(); crear_for_hc002()
    crear_man_rcr(); crear_for_rcr()
    crear_man_gc(); crear_pol_gc(); crear_pro_auditoria()
    print("✅ Carpetas 6, 7 y 8 completadas.")
