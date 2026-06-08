
# -*- coding: utf-8 -*-
"""Genera MAN-HC-001 Manual de Historia Clínica"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = "/home/user/bioauditoria/documentos_habilitacion/CARPETA_6_HISTORIA_CLINICA/MAN-HC-001_Manual_Historia_Clinica.docx"

AZUL_OSCURO = RGBColor(0, 51, 102)
AZUL_MEDIO = RGBColor(21, 101, 192)
ROJO = RGBColor(180, 0, 0)
BLANCO = RGBColor(255, 255, 255)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_header_footer(doc, code, title):
    section = doc.sections[0]
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
    header = section.header; header.is_linked_to_previous = False
    hp = header.paragraphs[0]; hp.text = f"{code} | {title}"; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.runs[0]; run.font.size = Pt(9); run.font.color.rgb = AZUL_OSCURO; run.font.name = 'Calibri'
    footer = section.footer; footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)

def add_page_number(paragraph):
    paragraph.clear()
    run = paragraph.add_run("Versión 2.0 | Junio 2025    Página ")
    run.font.size = Pt(9); run.font.name = 'Calibri'
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run2 = paragraph.add_run(); run2.font.size = Pt(9); run2.font.name = 'Calibri'
    run2._r.append(fldChar1); run2._r.append(instrText); run2._r.append(fldChar2)

def h1(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text); run.bold = True; run.font.size = Pt(14); run.font.name = 'Calibri'; run.font.color.rgb = AZUL_OSCURO

def h2(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text); run.bold = True; run.font.size = Pt(12); run.font.name = 'Calibri'; run.font.color.rgb = AZUL_MEDIO

def h3(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text); run.bold = True; run.font.size = Pt(11); run.font.name = 'Calibri'; run.font.color.rgb = AZUL_OSCURO

def body(doc, text):
    p = doc.add_paragraph(text); p.paragraph_format.space_after = Pt(6)
    for run in p.runs: run.font.size = Pt(11); run.font.name = 'Calibri'

def alerta(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1); p.paragraph_format.right_indent = Cm(1)
    run = p.add_run("⚠ IMPORTANTE: " + text); run.bold = True
    run.font.size = Pt(11); run.font.name = 'Calibri'; run.font.color.rgb = ROJO

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text); run.font.size = Pt(11); run.font.name = 'Calibri'

def table_header_row(table, headers, bg='003366'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]; cell.text = ''; set_cell_bg(cell, bg)
        run = cell.paragraphs[0].add_run(h); run.bold = True; run.font.color.rgb = BLANCO
        run.font.size = Pt(10); run.font.name = 'Calibri'; cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table_row(table, values, shaded=False):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]; cell.text = str(val)
        if shaded: set_cell_bg(cell, 'F2F2F2')
        for para in cell.paragraphs:
            for run in para.runs: run.font.size = Pt(10); run.font.name = 'Calibri'

doc = Document()
style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
add_header_footer(doc, "MAN-HC-001", "Manual de Historia Clínica")

# PORTADA
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MANUAL DE HISTORIA CLÍNICA")
run.font.size = Pt(20); run.font.name = 'Calibri'; run.font.color.rgb = AZUL_OSCURO; run.bold = True
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[NOMBRE DEL CONSULTORIO]")
run.font.size = Pt(14); run.font.name = 'Calibri'; run.font.color.rgb = AZUL_MEDIO; run.bold = True
doc.add_paragraph()
t = doc.add_table(rows=5, cols=2); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
pd = [("Código:", "MAN-HC-001"), ("Versión:", "2.0"), ("Fecha:", "Junio 2025"),
      ("Elaboró/Aprobó:", "Dra. [NOMBRE DE LA MÉDICA] - Directora Médica"), ("NIT:", "[NIT DEL CONSULTORIO]")]
for i, (k, v) in enumerate(pd):
    t.rows[i].cells[0].text = k; t.rows[i].cells[1].text = v
    set_cell_bg(t.rows[i].cells[0], 'E8EEF4')
    for cell in t.rows[i].cells:
        for para in cell.paragraphs:
            for run in para.runs: run.font.size = Pt(11); run.font.name = 'Calibri'

doc.add_page_break()

h1(doc, "CONTROL DE VERSIONES")
t_ver = doc.add_table(rows=3, cols=5); t_ver.style = 'Table Grid'
table_header_row(t_ver, ["Versión", "Fecha", "Descripción", "Elaboró", "Aprobó"])
add_table_row(t_ver, ["1.0", "Enero 2024", "Versión inicial del manual de historia clínica.", "Dra. [NOMBRE DE LA MÉDICA]", "Dra. [NOMBRE DE LA MÉDICA]"])
add_table_row(t_ver, ["2.0", "Junio 2025", "Ampliación de la sección de historia clínica para estética, instrumento de auditoría con 30 criterios, HC electrónica vs. física, y jurisprudencia colombiana actualizada.", "Dra. [NOMBRE DE LA MÉDICA]", "Dra. [NOMBRE DE LA MÉDICA]"], shaded=True)
doc.add_paragraph()

# 1. INTRODUCCIÓN - HC COMO DOCUMENTO MÉDICO-LEGAL
h1(doc, "1. LA HISTORIA CLÍNICA COMO DOCUMENTO MÉDICO-LEGAL")
body(doc, "La historia clínica es simultáneamente un instrumento clínico, un documento legal y un registro de la relación de confianza entre el médico y el paciente. Como instrumento clínico, es el soporte de la continuidad de la atención médica: permite que cualquier profesional de la salud que atienda al paciente en el futuro conozca sus antecedentes, diagnósticos previos, tratamientos realizados y evolución, garantizando la continuidad, la seguridad y la eficiencia de la atención. Como documento legal, es la principal fuente de prueba en los procesos de responsabilidad médica, ya sean disciplinarios (ante el Tribunal Ético de Medicina), civiles (indemnización de perjuicios) o penales (lesiones personales culposas, homicidio culposo). Una historia clínica bien elaborada es la mejor defensa del médico ante cualquier proceso judicial o disciplinario.")
body(doc, "En Colombia, la historia clínica es un documento de naturaleza jurídica reconocido y regulado expresamente por la Resolución 1995 de 1999 del Ministerio de Salud, que establece sus características obligatorias, sus componentes mínimos, las normas para su diligenciamiento, las reglas de custodia y los plazos de conservación. Esta resolución, que a la fecha de elaboración de este manual continúa vigente, es la norma específica de referencia para todo lo relacionado con la historia clínica en Colombia.")
body(doc, "JURISPRUDENCIA COLOMBIANA RELEVANTE: La Corte Constitucional colombiana, a través de múltiples sentencias, ha reconocido la importancia de la historia clínica como documento protegido por el derecho a la intimidad y el habeas data. La Sentencia T-158 de 1994 reconoció que los datos de salud contenidos en la historia clínica son datos sensibles que gozan de especial protección constitucional. La Sentencia T-696 de 1996 estableció que el paciente tiene derecho a acceder a su propia historia clínica. La Sentencia T-218 de 2014 precisó que la historia clínica debe ser entregada al paciente o su representante legal cuando la solicite, y que negar este acceso sin justificación legal válida constituye una vulneración de derechos fundamentales.")
body(doc, "La Corte Suprema de Justicia colombiana, en sus fallos sobre responsabilidad médica, ha establecido el criterio de que una historia clínica incompleta, ilegible o con omisiones de datos relevantes genera una presunción de falla médica en contra del médico, que este debe desvirtuar con otros medios probatorios. Este criterio jurisprudencial, conocido como 'carga dinámica de la prueba' aplicada a la responsabilidad médica, hace que la calidad y completitud de la historia clínica sea un elemento de primer orden no solo para la calidad de la atención, sino para la protección legal del médico y del establecimiento de salud.")
body(doc, "En el ámbito disciplinario, el Tribunal Ético de Medicina de Colombia ha sancionado en múltiples ocasiones a médicos por omisiones en la historia clínica, por falsedad en los registros clínicos, por no haber registrado el consentimiento informado del paciente, o por haber modificado la historia clínica después de un evento adverso. Estas sanciones pueden ir desde amonestaciones hasta suspensión temporal o definitiva del ejercicio de la medicina. El presente manual tiene como uno de sus propósitos fundamentales prevenir estas situaciones mediante el establecimiento de normas claras y exigentes para el diligenciamiento de la historia clínica en el [NOMBRE DEL CONSULTORIO].")

# 2. RESOLUCIÓN 1995/1999 COMENTADA
h1(doc, "2. RESOLUCIÓN 1995 DE 1999: ANÁLISIS ARTÍCULO POR ARTÍCULO")
body(doc, "La Resolución 1995 de 1999 del Ministerio de Salud es la norma específica que regula la historia clínica en Colombia. A continuación se presenta un análisis comentado de sus artículos más relevantes para la práctica en el [NOMBRE DEL CONSULTORIO]:")

articulos_1995 = [
    ("Art. 1: Definición de Historia Clínica", "Es un documento privado, obligatorio y sometido a reserva, en el cual se registran cronológicamente las condiciones de salud del paciente, los actos médicos y los demás procedimientos ejecutados por el equipo de salud que interviene en su atención. COMENTARIO: Tres elementos son clave: (1) ES OBLIGATORIA: no elaborarla es infracción administrativa y disciplinaria. (2) ES PRIVADA: acceso restringido según las normas de la resolución. (3) ES SOMETIDA A RESERVA: no puede divulgarse sin autorización del paciente."),
    ("Art. 2: Características de la Historia Clínica", "Define las 11 características obligatorias de la HC (ver sección 3 de este manual). COMENTARIO: El incumplimiento de cualquiera de estas características puede tener consecuencias disciplinarias y jurídicas. La 'integralidad' y la 'completitud' son las más frecuentemente incumplidas en la práctica."),
    ("Art. 3: Obligatoriedad del Registro", "Toda institución prestadora de servicios de salud, así como los profesionales independientes que presten servicios de salud, deben llevar historia clínica de los pacientes que atienden. COMENTARIO: El médico de consultorio privado, aunque sea unipersonal, tiene la obligación legal de llevar historia clínica de TODOS sus pacientes, incluyendo aquellos que acuden por certificados médicos u otros trámites."),
    ("Art. 4: Historia Clínica Única", "Cada institución debe contar con un sistema de identificación única del paciente que evite la duplicación de historias clínicas. COMENTARIO: En el [NOMBRE DEL CONSULTORIO] el identificador único es el número de documento de identidad del paciente. No puede haber dos historias clínicas para el mismo número de documento."),
    ("Art. 5: Componentes de la Historia Clínica", "La HC debe tener al menos tres partes: identificación del paciente, registros específicos de cada atención (anamnesis, examen físico, diagnóstico, plan de manejo, evolución) y anexos (consentimientos, resultados de laboratorio, imágenes, remisiones). COMENTARIO: En la práctica de primer nivel, los registros de cada atención son la parte central y deben completarse en su totalidad en cada consulta."),
    ("Art. 6: Apertura y Diligenciamiento", "La HC debe abrirse desde el primer contacto del paciente con el prestador y diligenciarse en el momento de la atención. No debe dejarse para completar después del acto médico. COMENTARIO: Completar la HC horas o días después de la atención es una violación de esta norma y reduce la confiabilidad del registro. La regla de oro es: 'Si no está escrito en el momento, no ocurrió.'"),
    ("Art. 7: Corrección de Errores", "Los errores en la HC no deben borrarse, tacharse con corrector ni eliminarse. Se deben tachar con una línea horizontal, escribir la palabra 'error', la inicial del nombre del profesional, la fecha de la corrección, y la información correcta a continuación. COMENTARIO: El uso de corrector de texto (liquid paper, cinta correctora) es ABSOLUTAMENTE PROHIBIDO en la historia clínica. El descubrimiento de su uso en un proceso judicial es interpretado como intento de alteración de la prueba."),
    ("Art. 8: Custodia y Conservación", "La custodia de la HC es responsabilidad del prestador de servicios de salud. Debe conservarse por un período mínimo de 15 años desde la última atención para adultos, y hasta que el menor alcance 18 años y 5 años más para historias de menores de edad. COMENTARIO: En el [NOMBRE DEL CONSULTORIO], la custodia es responsabilidad directa de la Dra. [NOMBRE DE LA MÉDICA]. La pérdida o deterioro de la HC puede generar responsabilidad civil y disciplinaria."),
    ("Art. 9: Acceso a la Historia Clínica", "El paciente tiene derecho a acceder a su propia HC. También pueden acceder: el equipo de salud que lo atiende, las entidades de vigilancia y control (Secretaría de Salud, Superintendencia), los tribunales y autoridades judiciales mediante orden judicial, y los familiares del paciente fallecido. COMENTARIO: Ninguna otra persona puede acceder a la HC sin autorización escrita del paciente. El médico no puede entregar información de la HC a la empresa del paciente, a la familia sin autorización del paciente adulto, ni a otras personas, salvo las excepciones legales."),
    ("Art. 10: Confidencialidad", "La información de la HC está amparada por el secreto médico y las normas de protección de datos personales (Ley 1581 de 2012). Revelar información de la HC sin autorización del paciente viola el secreto profesional (delito tipificado en el Art. 419 del Código Penal colombiano) y las normas de ética médica. COMENTARIO: Las conversaciones sobre pacientes en lugares públicos (recepción, ascensores, restaurantes) constituyen violaciones a la confidencialidad, aunque no sean intencionales."),
]

for art, comentario in articulos_1995:
    h3(doc, art)
    body(doc, comentario)

# 3. CARACTERÍSTICAS OBLIGATORIAS
h1(doc, "3. CARACTERÍSTICAS OBLIGATORIAS DE LA HISTORIA CLÍNICA")
body(doc, "La Resolución 1995 de 1999 establece once (11) características que debe reunir toda historia clínica en Colombia. A continuación se define y explica cada una con su aplicación práctica en el [NOMBRE DEL CONSULTORIO]:")

caracteristicas = [
    ("1. INTEGRALIDAD", "La HC debe reunir la información de los aspectos científicos, técnicos y administrativos relativos a la atención en salud en las fases de fomento, protección específica, diagnóstico, tratamiento y rehabilitación. En la práctica: esto significa que la HC no es solo el registro de la enfermedad actual y el tratamiento; debe incluir también los datos de identificación completos, los antecedentes completos, los datos de aseguramiento, el registro de actividades preventivas realizadas (vacunación, citología) y las recomendaciones de promoción y prevención dadas al paciente. La ausencia de cualquiera de estas dimensiones la hace incompleta."),
    ("2. SECUENCIALIDAD", "Los registros de la prestación de los servicios de salud deben consignarse en la secuencia cronológica en que ocurrió la atención. En la práctica: todas las entradas de la HC deben tener fecha y hora de registro. Las consultas deben aparecer en orden cronológico (la más antigua primero). No se pueden insertar registros entre consultas previas modificando la cronología."),
    ("3. RACIONALIDAD CIENTÍFICA", "Para los efectos de la presente resolución, es la aplicación de criterios científicos en el diligenciamiento y registro de las acciones en salud. El contenido de la HC debe fundamentarse en principios médicos y científicos reconocidos. En la práctica: el diagnóstico debe corresponder a la CIE-10 con justificación clínica en la anamnesis y el examen físico. El plan de manejo debe ser coherente con el diagnóstico y con las guías de práctica clínica. Un diagnóstico sin sustento clínico documentado, o un tratamiento injustificado, viola esta característica."),
    ("4. DISPONIBILIDAD", "Es la posibilidad de utilizar la historia clínica en el momento en que se necesita, con las limitaciones que impone la ley. En la práctica: la HC debe ser accesible al médico en el momento de la consulta. Si el sistema electrónico falla, debe haber un procedimiento alterno (HC en papel mientras se restaura el sistema). La HC de un paciente que acude en urgencias debe estar disponible inmediatamente."),
    ("5. OPORTUNIDAD", "Es el diligenciamiento de los registros de atención de salud, simultáneamente o dentro de los tiempos establecidos para ello, que permitan una adecuada cronología en la prestación de los servicios de salud. En la práctica: la regla general es que la HC se diligencia DURANTE la atención. En casos excepcionales puede completarse dentro de la primera hora post-atención. Completarla al final del día o días después constituye un incumplimiento de esta característica."),
    ("6. LEGIBILIDAD", "El médico responsable del diligenciamiento de la HC debe asegurarse que los registros sean legibles. Las anotaciones ilegibles no tienen valor probatorio y pueden constituir causa de sanciones disciplinarias. En la práctica: la letra debe ser clara y comprensible por cualquier otro profesional de la salud. Se recomienda la HC electrónica cuando la letra del médico es difícil de leer. Si se usa HC física, debe escribirse con bolígrafo de tinta negra o azul oscura, sin usar lápiz."),
    ("7. LIMPIEZA Y PULCRITUD", "Los registros de la HC deben realizarse en condiciones de limpieza, sin manchas, roturas, líquidos derramados u otras afecciones que afecten su lectura y conservación. En la práctica: no deben archivarse documentos sucios, húmedos o rasgados en la HC. Las manchas de sangre u otros fluidos en la HC deben manejarse con especial cuidado para evitar contaminación biológica."),
    ("8. COMPLETITUD", "La HC debe contener todos los datos, informaciones y registros propios de la atención en salud. No puede haber campos obligatorios en blanco. En la práctica: todos los campos del formato de HC del [NOMBRE DEL CONSULTORIO] son obligatorios. Si un campo no aplica para la consulta específica, debe escribirse 'no aplica' con justificación. Un campo en blanco es inacabado."),
    ("9. AUTENTICIDAD", "La HC debe ser expresión fiel de los hechos ocurridos durante la atención. No pueden registrarse datos falsos, supuestos o hipotéticos como si fueran hechos reales. En la práctica: registrar un examen físico que no se realizó ('examen físico normal' sin haberlo hecho) es una falsedad en documento privado. Registrar datos de signos vitales que no se midieron es igualmente una falsedad. Estas conductas generan responsabilidad disciplinaria, civil y penal."),
    ("10. CONFIDENCIALIDAD", "La HC es un documento sometido a reserva. Solo pueden acceder a ella las personas autorizadas por la norma. En la práctica: la HC debe guardarse en lugar seguro bajo llave (HC física) o con acceso protegido por contraseña (HC electrónica). No puede dejarse sobre el escritorio visible para otros pacientes. No puede comentarse en lugares públicos. Las fotografías del paciente son datos especialmente sensibles."),
    ("11. UNICIDAD", "Para cada paciente solo puede existir una historia clínica por institución prestadora de servicios de salud. En el [NOMBRE DEL CONSULTORIO], el número de documento de identidad del paciente es el identificador único que garantiza que no se abre una segunda HC para el mismo paciente. La duplicación de HC genera fragmentación de la información y riesgo de error médico."),
]

for caract, descripcion in caracteristicas:
    h3(doc, caract)
    body(doc, descripcion)

# 4. COMPONENTES OBLIGATORIOS
h1(doc, "4. COMPONENTES OBLIGATORIOS DE LA HISTORIA CLÍNICA")
body(doc, "La Resolución 1995 de 1999 y la doctrina médica colombiana establecen que toda historia clínica debe contener los siguientes componentes mínimos obligatorios:")
t_comp_hc = doc.add_table(rows=1, cols=3); t_comp_hc.style = 'Table Grid'
table_header_row(t_comp_hc, ["Componente", "Descripción y Contenido Mínimo", "¿Obligatorio en Cada Consulta?"])
comp_hc = [
    ("Identificación del paciente", "Nombre completo, tipo y número de documento de identidad, fecha de nacimiento, edad, sexo, estado civil, ocupación, dirección de residencia, teléfonos de contacto, nombre del contacto de emergencia, tipo de aseguramiento, nombre de la EPS.", "Sí, verificar y actualizar en cada consulta si hay cambios"),
    ("Fecha y hora de la atención", "Fecha completa (día/mes/año) y hora de inicio y finalización de la consulta.", "Sí, obligatorio en cada atención"),
    ("Motivo de consulta", "Razón principal expresada por el paciente en sus propias palabras, entre comillas.", "Sí"),
    ("Anamnesis completa", "Enfermedad actual (inicio, características, evolución, factores modificadores, síntomas asociados), antecedentes personales patológicos, quirúrgicos, traumáticos, farmacológicos, alérgicos, gineco-obstétricos (si aplica), familiares, laborales y sociales, revisión por sistemas.", "Sí (completa en primera vez; resumen actualizado en controles)"),
    ("Signos vitales", "Presión arterial, frecuencia cardíaca, frecuencia respiratoria, temperatura, saturación de oxígeno, peso, talla e IMC. Con unidades y hora de medición.", "Sí, en toda consulta"),
    ("Examen físico", "Descripción sistemática de los hallazgos del examen físico de cada región corporal explorada. Debe especificar los hallazgos POSITIVOS (anormales) y los NEGATIVOS relevantes (ausencia de signos que se buscaron específicamente).", "Sí (completo en primera vez; focalizado en controles)"),
    ("Diagnóstico(s)", "Diagnóstico principal y secundarios con código CIE-10 y descripción. Especificar si es diagnóstico definitivo o presuntivo. Diagnósticos diferenciales si aplica.", "Sí, en toda consulta"),
    ("Plan de manejo", "Tratamiento farmacológico (medicamentos con dosis, frecuencia y duración), tratamiento no farmacológico, órdenes de ayudas diagnósticas, referencias a especialistas, plan educativo para el paciente, plan de seguimiento.", "Sí, en toda consulta"),
    ("Prescripción médica (copia)", "Copia o registro de los medicamentos prescritos con todos los datos de la Resolución 1478/2006.", "Cuando se prescribe"),
    ("Consentimiento informado", "Para procedimientos que lo requieren: constancia de que se obtuvo el CI (puede ser referencia al formato firmado FOR-PP-001, FOR-PP-002, etc.).", "Para consulta (general) y obligatorio para procedimientos"),
    ("Firma, sello y número de tarjeta profesional del médico", "Firma del médico, sello con nombre, especialidad y número de TP al final de cada atención.", "Sí, en toda atención"),
    ("Documentos y resultados anexos", "Resultados de laboratorio, imágenes diagnósticas, resúmenes de hospitalización, remisiones y contraremisiones deben archivarse cronológicamente en la HC.", "Archivar cuando se reciben"),
]
for i, row_data in enumerate(comp_hc):
    add_table_row(t_comp_hc, list(row_data), shaded=(i % 2 == 1))
doc.add_paragraph()

# 5. GUÍA DE DILIGENCIAMIENTO
h1(doc, "5. GUÍA DE DILIGENCIAMIENTO - HISTORIA CLÍNICA DE MEDICINA GENERAL")
body(doc, "La siguiente guía campo por campo orienta el correcto diligenciamiento del formato FOR-HC-001 (Historia Clínica de Medicina General del [NOMBRE DEL CONSULTORIO]):")

h2(doc, "5.1 SECCIÓN I: IDENTIFICACIÓN Y DATOS GENERALES")
body(doc, "FECHA Y HORA: Registrar la fecha en formato DD/MM/AAAA y la hora de inicio de la consulta en formato de 24 horas (ejemplo: 14:30). Este dato es irremplazable para establecer la cronología del acto médico en caso de proceso legal.")
body(doc, "TIPO DE CONSULTA: Marcar el tipo de consulta con claridad: Primera vez / Control / Urgencia menor / Certificado médico / Pre-procedimiento. Esta clasificación determina la extensión requerida de la anamnesis y el examen físico.")
body(doc, "IDENTIFICACIÓN: Registrar el nombre completo del paciente exactamente como aparece en el documento de identidad (sin abreviaturas). Tipo de documento (CC=cédula de ciudadanía, TI=tarjeta de identidad, CE=cédula de extranjería, PA=pasaporte, RC=registro civil, NUIP). Número de documento sin puntos ni espacios. Fecha de nacimiento completa. Edad en años (o meses para menores de 2 años).")
body(doc, "ASEGURAMIENTO: EPS y número de afiliado, tipo de régimen (contributivo/subsidiado), número de autorización si aplica. Si es particular: escribir 'PARTICULAR'. Este dato es esencial para la facturación y para verificar el derecho de atención.")

h2(doc, "5.2 SECCIÓN II: ANAMNESIS")
body(doc, "MOTIVO DE CONSULTA: Escribir textualmente entre comillas lo que el paciente expresa. Ejemplo: 'Me duele la cabeza desde hace 3 días, con náuseas.' El médico NO debe parafrasear ni interpretar en este campo. Incluir el tiempo de evolución si el paciente lo menciona.")
body(doc, "ENFERMEDAD ACTUAL: Descripción cronológica y detallada del problema de salud. Para cada síntoma registrar obligatoriamente: fecha o tiempo de inicio preciso ('desde el 15 de mayo de 2025', 'desde hace 3 semanas'), modo de inicio (súbito/gradual), características cualitativas del síntoma, localización e irradiación para síntomas de dolor, intensidad (EVA 0-10 para dolor), duración y periodicidad, factores que modifican el síntoma (qué lo alivia, qué lo empeora), síntomas acompañantes, consultas y tratamientos previos para el mismo problema con respuesta al tratamiento.")
body(doc, "ANTECEDENTES PERSONALES PATOLÓGICOS: Listar cada enfermedad crónica conocida con año de diagnóstico y estado de control. Ejemplo: 'Hipertensión arterial (diagnóstico 2018), en tratamiento con enalapril 20mg, PA usual 130/80 mmHg'. Si el paciente refiere no tener antecedentes, escribir: 'El paciente refiere no tener antecedentes médicos conocidos' — nunca dejar el campo en blanco.")
body(doc, "ALERGIAS: Este campo es de importancia crítica para la seguridad del paciente. Debe estar visible y destacado en la HC. Si el paciente tiene alergias, se escribe en rojo o en recuadro especial: 'ALERGIA A: [nombre del medicamento/sustancia]. TIPO DE REACCIÓN: [descripción]. FECHA EN QUE SE PRESENTÓ: [fecha]'. Si no tiene alergias conocidas, escribir: 'Paciente refiere no tener alergias conocidas a medicamentos, alimentos ni sustancias'. Este campo NO puede quedar en blanco.")

h2(doc, "5.3 SECCIÓN III: EXAMEN FÍSICO")
body(doc, "SIGNOS VITALES: Se registran en la fila de datos estructurados del formato: PA (mmHg), FC (lpm), FR (rpm), Temperatura (°C con sitio de medición), SpO2 (%), Peso (kg), Talla (cm), IMC (kg/m²). Todos los campos son OBLIGATORIOS. Si un equipo no está disponible (por ejemplo, oxímetro), anotar 'No disponible' con la causa, no dejar en blanco.")
body(doc, "DESCRIPCIÓN DEL EXAMEN FÍSICO: Se describe en formato narrativo o por sistemas. Cada región examinada debe describirse explícitamente. Los hallazgos POSITIVOS (anormales) se describen con detalle morfológico, de localización y de extensión. Los hallazgos negativos relevantes también se registran explícitamente ('ausencia de soplos cardíacos', 'sin signos de irritación peritoneal'). NO son aceptables registros como 'examen físico normal' sin describir qué se examinó y qué se encontró. La descripción debe ser suficientemente detallada como para reconstruir el estado del paciente si fuera necesario.")

h2(doc, "5.4 SECCIÓN IV: DIAGNÓSTICO Y PLAN DE MANEJO")
body(doc, "DIAGNÓSTICO: Se registra el diagnóstico principal primero, seguido de los diagnósticos secundarios. Para cada uno se escribe: código CIE-10 (al mayor nivel de especificidad posible, 4 caracteres cuando aplica), descripción en español del diagnóstico. Si el diagnóstico no está confirmado, se especifica: 'Presuntivo: [diagnóstico] [código CIE-10]' o 'En estudio: [diagnóstico]'. También se registran los diagnósticos diferenciales considerados. Ejemplo: 'Diagnóstico principal: J06.9 Infección aguda de las vías respiratorias superiores, no especificada. Diferencial: J02.9 Faringitis aguda no especificada / J06.0 Laringitis y traqueítis agudas.'")
body(doc, "PLAN DE MANEJO: Se registra cada componente del plan de manera estructurada y completa: (1) Tratamiento farmacológico: medicamento (DCI), concentración, forma farmacéutica, dosis, vía, frecuencia, duración. (2) Tratamiento no farmacológico: todas las medidas no medicamentosas indicadas con suficiente detalle para que el paciente las entienda. (3) Educación: temas de educación proporcionados al paciente y verificación de comprensión. (4) Órdenes: exámenes solicitados con justificación. (5) Remisión: especialidad, motivo, urgencia. (6) Seguimiento: fecha de próxima cita o instrucciones para acudir antes si hay señales de alarma. Debe haber coherencia diagnóstico-terapéutica: el plan de manejo debe corresponder lógicamente a los diagnósticos formulados.")

# 6. HISTORIA CLÍNICA PARA MEDICINA ESTÉTICA
h1(doc, "6. HISTORIA CLÍNICA PARA PROCEDIMIENTOS ESTÉTICOS")
body(doc, "La historia clínica de medicina estética (FOR-HC-002) tiene campos adicionales específicos que no están presentes en la HC de medicina general. Su correcto diligenciamiento es fundamental tanto para la seguridad del procedimiento como para la protección legal del médico. Los campos específicos de la HC de estética incluyen:")
h2(doc, "6.1 Campos Específicos de la HC de Medicina Estética")
bullet(doc, "ÁREA DE PREOCUPACIÓN ESTÉTICA: Descripción de la zona o zonas que el paciente desea tratar, en sus propias palabras. Ejemplo: 'Líneas de expresión en frente y entrecejo. Desea resultados naturales, no 'cara congelada'.'")
bullet(doc, "HISTORIAL DE PROCEDIMIENTOS PREVIOS: Lista detallada de todos los procedimientos estéticos que el paciente ha recibido, en este u otro consultorio: tipo de procedimiento, producto utilizado (si lo conoce), cuándo, dónde, resultado (satisfecho/insatisfecho), complicaciones. Esta información es crítica: saber si hay rellenos previos en el área evita la sobreinyección y ayuda a anticipar planos de inyección alterados.")
bullet(doc, "ANTECEDENTES ESPECÍFICOS DE RELEVANCIA ESTÉTICA: Ver sección 6.2 del PRO-PP-002. Especialmente: uso de anticoagulantes, antecedente de herpes, uso de isotretinoína, historia de queloides.")
bullet(doc, "EVALUACIÓN DE EXPECTATIVAS: Descripción de lo que el paciente espera obtener con el procedimiento, y evaluación por parte del médico de si esas expectativas son realistas y alcanzables. Si las expectativas son irreales, documentar la conversación y la orientación dada.")
bullet(doc, "ESCALA GAIS PRE-PROCEDIMIENTO: Aplicar y registrar la Escala de Mejora Global (GAIS 1-5) antes del procedimiento en el área a tratar, como línea de base para comparación post-procedimiento.")
bullet(doc, "FOTOGRAFÍAS CLÍNICAS: Registrar que se tomaron fotografías (fecha, tomas realizadas: frente/perfil/¾). Las fotografías se almacenan en el sistema digital del consultorio con el número de documento del paciente y la fecha. Es OBLIGATORIO contar con el consentimiento escrito del paciente para fotografías antes de tomarlas.")
bullet(doc, "PRODUCTO UTILIZADO - REGISTRO COMPLETO: Para cada producto inyectable o aplicado se registra: nombre comercial, DCI o composición, laboratorio fabricante, número de lote, fecha de vencimiento (debe ser posterior a la fecha del procedimiento), número de registro INVIMA, volumen o cantidad utilizada. Este registro es esencial para: (1) trazar el producto en caso de reacción adversa, (2) garantizar que se usa producto vigente y con registro INVIMA, (3) estimar el tiempo de duración del efecto."),
bullet(doc, "TÉCNICA DE APLICACIÓN: Descripción sucinta de la técnica utilizada (inyección lineal, bolo, en abanico, cánula vs. aguja, profundidad del plano de inyección) y distribución de puntos de inyección por zona.")
bullet(doc, "ESCALA GAIS POST-PROCEDIMIENTO: Aplicar y registrar la GAIS inmediatamente post-procedimiento. Comparar con el GAIS pre-procedimiento.")
bullet(doc, "INCIDENCIAS DURANTE EL PROCEDIMIENTO: Cualquier evento ocurrido durante el procedimiento (punción vascular, equimosis, dolor intenso, síntoma inusual) debe documentarse con la hora, la descripción del evento y la acción tomada.")

# 7. HC ELECTRÓNICA VS FÍSICA
h1(doc, "7. HISTORIA CLÍNICA ELECTRÓNICA VS. HISTORIA CLÍNICA FÍSICA")
t_hce = doc.add_table(rows=1, cols=3); t_hce.style = 'Table Grid'
table_header_row(t_hce, ["Aspecto", "Historia Clínica Física (en papel)", "Historia Clínica Electrónica (software)"])
hce_data = [
    ("Diligenciamiento", "Manuscrita con bolígrafo negro o azul. Debe ser legible. Firmada y sellada.", "Mecanografiada en sistema. Firma electrónica o impresión y firma física al final de la jornada."),
    ("Corrección de errores", "Tachar con línea, escribir 'error', inicial del médico, fecha. NUNCA corrector.", "El sistema debe registrar el historial de cambios (auditoría de modificaciones). No se borran datos anteriores."),
    ("Seguridad", "Bajo llave en archivador con acceso restringido. Riesgo de pérdida por incendio/agua.", "Acceso con contraseña. Respaldo periódico (mínimo semanal) en servidor externo o nube cifrada. Registro de quién accedió y cuándo."),
    ("Disponibilidad", "Depende de la ubicación física del archivo. Riesgo de extravío.", "Accesible desde cualquier dispositivo autorizado. Búsqueda rápida por número de documento."),
    ("Costo", "Bajo costo inicial (papel, archivador). Alto costo a largo plazo por espacio físico.", "Costo del software y mantenimiento. Ahorro en papel y espacio a largo plazo."),
    ("Cumplimiento normativo", "Cumple si se diligencia correctamente. La ilegibilidad es el principal riesgo de incumplimiento.", "Cumple si el sistema tiene: registro de fecha y hora automático, identificación del usuario que hace el registro, y no permite la eliminación de datos."),
    ("Recomendación para [NOMBRE DEL CONSULTORIO]", "Aceptable mientras el volumen de pacientes sea bajo y el médico tenga letra legible. Implementar archivador con llave.", "Recomendada para mayor seguridad y disponibilidad. Evaluar opciones de HC electrónica para primer nivel con costo accesible."),
]
for i, row_data in enumerate(hce_data):
    add_table_row(t_hce, list(row_data), shaded=(i % 2 == 1))
doc.add_paragraph()

# 8. CUSTODIA Y ARCHIVO
h1(doc, "8. CUSTODIA, ARCHIVO Y CONSERVACIÓN DE LA HISTORIA CLÍNICA")
body(doc, "La Resolución 1995 de 1999 establece que la custodia de la historia clínica es responsabilidad del prestador de servicios de salud. En el [NOMBRE DEL CONSULTORIO], la Dra. [NOMBRE DE LA MÉDICA] es la custodiana legal de todas las historias clínicas. Las normas de custodia y archivo son:")
bullet(doc, "PLAZO MÍNIMO DE CONSERVACIÓN: 15 años contados desde la última atención para historias clínicas de adultos. Para menores de edad: hasta que el paciente cumpla 23 años (18 años + 5 años adicionales).")
bullet(doc, "ORGANIZACIÓN DEL ARCHIVO FÍSICO (si aplica): Las HC físicas se archivan en carpetas individuales identificadas con: número de documento del paciente, nombre completo, fecha de apertura. Se organizan en orden alfanumérico o numérico en archivador bajo llave. El acceso al archivo es restringido: solo personal autorizado del consultorio.")
bullet(doc, "SEGURIDAD FÍSICA: El archivo de HC debe estar en un área con acceso restringido, ventilación adecuada, temperatura controlada (no mayor a 27°C), protegido de humedad, luz directa y roedores. Debe contar con extintor de incendios en el área adyacente.")
bullet(doc, "SEGURIDAD DIGITAL (HC electrónica): El sistema de HC electrónica debe tener: autenticación de usuario con contraseña única e intransferible, registro de auditoría de todos los accesos y modificaciones, respaldo automático diario en servidor secundario o nube cifrada con certificado SSL/TLS, cifrado de los datos en reposo y en tránsito, política de contraseñas seguras (mínimo 8 caracteres, mayúsculas, números, caracteres especiales, cambio cada 90 días).")
bullet(doc, "DESTRUCCIÓN: Cumplido el período mínimo de conservación, las HC pueden destruirse, pero solo mediante proceso documentado que garantice la destrucción total de la información (trituración de papel, borrado seguro de datos digitales). Se debe llevar un registro de las HC destruidas con número de documento, nombre del paciente y fecha de destrucción.")

# 9. ACCESO
h1(doc, "9. CONTROL DE ACCESO A LA HISTORIA CLÍNICA")
t_acceso = doc.add_table(rows=1, cols=4); t_acceso.style = 'Table Grid'
table_header_row(t_acceso, ["Quién Puede Acceder", "Condiciones de Acceso", "Documentos Requeridos", "Restricciones"])
acceso_data = [
    ("El propio paciente", "Derecho fundamental, en cualquier momento. El paciente puede pedir copia de su HC.", "Presentación del documento de identidad. Solicitud escrita (optional pero recomendado).", "Ninguna. No se puede negar el acceso al paciente a su propia HC (sentencia T-696/96 Corte Constitucional)."),
    ("Representante legal del paciente", "Menores de edad: padres o tutor legal. Paciente incapaz: tutor o curador designado judicialmente.", "Documento que acredite la representación legal (registro civil, sentencia judicial).", "Debe acreditarse la representación legal de manera fehaciente."),
    ("Médico tratante del consultorio", "En el ejercicio de sus funciones asistenciales.", "Estar vinculado al consultorio y atender o haber atendido al paciente.", "Solo para fines asistenciales del paciente atendido."),
    ("Otro profesional de salud del mismo establecimiento", "Solo si participa en la atención del paciente.", "Acreditación de su vinculación al consultorio.", "No puede acceder a HC de pacientes que no atiende."),
    ("Secretaría Departamental/Distrital de Salud (inspección)", "En ejercicio de funciones de inspección, vigilancia y control.", "Acto administrativo que ordene la inspección. Identificación del funcionario.", "Solo en el marco de la visita de inspección autorizada."),
    ("Superintendent de Salud", "En ejercicio de funciones de vigilancia y control.", "Auto de inspección o requerimiento oficial.", "Solo en el marco del proceso de vigilancia iniciado."),
    ("Autoridades judiciales y organismos de control", "Mediante orden judicial (juez) o requerimiento de la Fiscalía, Procuraduría o Contraloría.", "Orden judicial o requerimiento oficial debidamente notificado.", "El médico puede solicitar asesoría jurídica antes de entregar la HC ante una orden judicial."),
    ("Familiares del paciente fallecido", "Para conocer las causas de la muerte o con fines sucesorales.", "Registro de defunción. Registro civil que acredite el parentesco.", "Solo información relacionada con la muerte. La información sobre condiciones de salud previas no relacionadas con la muerte puede estar protegida aún."),
    ("EPS (entidades aseguradoras)", "Para auditoría de servicios prestados, reconocimiento de incapacidades, o resolución de glosas.", "Requerimiento escrito de la EPS con justificación. Contrato vigente con el prestador.", "Solo la información relacionada con los servicios facturados o con el objeto específico del requerimiento."),
    ("Investigadores en salud", "Para investigación científica con datos de la HC.", "Concepto favorable de Comité de Ética en Investigación. Consentimiento del paciente (salvo anonimización). Cumplimiento Resolución 8430/93.", "Los datos deben anonimizarse o el paciente debe dar CI específico para uso en investigación."),
]
for i, row_data in enumerate(acceso_data):
    add_table_row(t_acceso, list(row_data), shaded=(i % 2 == 1))
doc.add_paragraph()
alerta(doc, "Ninguna empresa, empleador, compañía de seguros, o familiar del paciente adulto tiene derecho a acceder a la historia clínica sin autorización escrita del propio paciente. Entregar información clínica sin autorización constituye violación al secreto profesional (Art. 419 Código Penal colombiano) y puede generar sanciones disciplinarias, civiles y penales.")

# 10. AUDITORÍA DE HISTORIAS CLÍNICAS
h1(doc, "10. INSTRUMENTO DE AUDITORÍA DE HISTORIAS CLÍNICAS")
body(doc, "El [NOMBRE DEL CONSULTORIO] realiza auditoría de historias clínicas de manera mensual sobre una muestra aleatoria de 10 HC. Para cada HC se evalúan los siguientes 30 criterios de calidad. Cada criterio se califica como: C = Cumple, NC = No Cumple, NA = No Aplica. El porcentaje de cumplimiento global se calcula como: (N° de criterios que cumplen / N° de criterios que aplican) × 100. La meta es 100%.")

t_audit = doc.add_table(rows=1, cols=4); t_audit.style = 'Table Grid'
table_header_row(t_audit, ["N°", "Criterio de Auditoría", "Norma de Referencia", "C/NC/NA"])
audit_criteria = [
    ("1", "La HC tiene registrada la fecha y hora de la atención en formato completo.", "Res. 1995/99 Art. 6"),
    ("2", "La HC tiene registrado el tipo de consulta (primera vez, control, urgencia).", "Res. 1995/99 Art. 5"),
    ("3", "Los datos de identificación del paciente están completos (nombre, documento, fecha de nacimiento, sexo).", "Res. 1995/99 Art. 5"),
    ("4", "Está registrado el tipo de aseguramiento y la EPS o la condición de particular.", "Res. 3100/2019"),
    ("5", "El motivo de consulta está registrado en palabras del paciente, entre comillas.", "Doctrina clínica / PRO-PP-001"),
    ("6", "La enfermedad actual describe inicio, características, evolución y síntomas asociados.", "Res. 1995/99 Art. 5"),
    ("7", "Están registrados los antecedentes patológicos personales o se registra expresamente que el paciente niega tenerlos.", "Res. 1995/99 Art. 5"),
    ("8", "Están registrados los antecedentes farmacológicos (medicamentos actuales) o su ausencia.", "Res. 1995/99 Art. 5"),
    ("9", "Están registradas las alergias del paciente o su ausencia con expresión explícita.", "PRO-PP-001 / Seguridad del paciente"),
    ("10", "Están registrados los antecedentes familiares relevantes o su ausencia.", "Res. 1995/99 Art. 5"),
    ("11", "Están registrados todos los signos vitales con sus valores y unidades (PA, FC, FR, T, SpO2, peso, talla, IMC).", "Res. 3100/2019 / PRO-PP-001"),
    ("12", "El examen físico está descrito por regiones o sistemas con hallazgos explícitos (no solo 'normal').", "Doctrina clínica / Res. 1995/99"),
    ("13", "El diagnóstico está formulado con código CIE-10 y descripción.", "Res. 1995/99 / PRO-PP-001"),
    ("14", "El plan de manejo incluye tratamiento farmacológico completo (medicamento, dosis, frecuencia, duración) cuando se prescribió.", "Res. 1478/2006"),
    ("15", "El plan de manejo incluye medidas no farmacológicas.", "PRO-PP-001 / GPC"),
    ("16", "El plan de manejo incluye información educativa proporcionada al paciente.", "PRO-PP-001"),
    ("17", "Se registran las órdenes de laboratorio o imagen solicitadas con justificación.", "PRO-PP-001"),
    ("18", "Si se realizó remisión, está registrada la especialidad, el motivo y la urgencia.", "Res. 1438/2011"),
    ("19", "Está registrada la fecha de próxima cita o las instrucciones de seguimiento.", "PRO-PP-001"),
    ("20", "La HC tiene firma del médico, sello y número de tarjeta profesional.", "Ley 23/1981 / Res. 1995/99"),
    ("21", "La letra es legible y comprensible.", "Res. 1995/99 Art. 2"),
    ("22", "No se observan correcciones con corrector de texto (liquid paper). Los errores están corregidos con el método legal (tachar-inicial-fecha).", "Res. 1995/99 Art. 7"),
    ("23", "La HC no tiene campos obligatorios en blanco (completitud).", "Res. 1995/99 Art. 2"),
    ("24", "Las fechas de la HC son cronológicamente coherentes (no hay fechas posteriores antes de fechas anteriores).", "Res. 1995/99 Art. 2"),
    ("25", "Los resultados de laboratorio recibidos están archivados en la HC.", "Res. 1995/99 Art. 5"),
    ("26", "El consentimiento informado para procedimientos está firmado, fechado y archivado en la HC cuando el procedimiento lo requiere.", "Ley 23/1981 Art. 15 / Res. 1995/99"),
    ("27", "Para HC de medicina estética: el producto utilizado tiene registrado lote, fecha de vencimiento y número de registro INVIMA.", "PRO-PP-002 / INVIMA"),
    ("28", "Para HC de medicina estética: hay registro fotográfico pre-procedimiento con consentimiento para fotografías.", "PRO-PP-002 / Ley 1581/2012"),
    ("29", "Para HC de medicina estética: la escala GAIS pre-procedimiento está registrada.", "PRO-PP-002"),
    ("30", "No hay evidencia de que la HC fue completada o modificada en fecha posterior a la atención (coherencia interna fecha-contenido).", "Res. 1995/99 Art. 9 / Ética Médica"),
]
for i, (num, crit, norma) in enumerate(audit_criteria):
    add_table_row(t_audit, [num, crit, norma, "___"], shaded=(i % 2 == 1))
doc.add_paragraph()

# 11. SANCIONES
h1(doc, "11. SANCIONES POR INCUMPLIMIENTO DE LAS NORMAS DE HISTORIA CLÍNICA")
body(doc, "Las sanciones por el incumplimiento de las normas de historia clínica pueden producirse en múltiples instancias:")
t_sanc = doc.add_table(rows=1, cols=4); t_sanc.style = 'Table Grid'
table_header_row(t_sanc, ["Tipo de Incumplimiento", "Instancia Sancionatoria", "Tipo de Sanción", "Norma Infringida"])
sanc_data = [
    ("No elaborar historia clínica", "Secretaría de Salud / Tribunal Ético de Medicina", "Multa administrativa. Sanción disciplinaria: amonestación a suspensión.", "Res. 1995/99 Art. 3 / Ley 23/81 Art. 34"),
    ("Historia clínica incompleta", "Tribunal Ético de Medicina / Juzgado civil", "Sanción disciplinaria. En proceso civil: presunción de falla médica (carga dinámica de la prueba).", "Res. 1995/99 Art. 2 / Ley 23/81"),
    ("Historia clínica falsificada o alterada (consignar datos falsos)", "Fiscalía / Juzgado penal", "Delito: falsedad en documento privado (Código Penal, Art. 289) o falsedad ideológica. Pena de prisión.", "Art. 289 Código Penal"),
    ("Alterar la historia clínica después de un evento adverso", "Fiscalía / Juzgado penal / Tribunal Ético", "Delito de falsedad. Agravante en proceso de responsabilidad médica. Sanción disciplinaria severa.", "Código Penal / Ley 23/81"),
    ("Revelar información confidencial de la HC sin autorización", "Tribunal Ético / Fiscalía", "Sanción disciplinaria: amonestación a suspensión. Delito: violación de datos personales (Art. 269F Código Penal) o revelación del secreto profesional (Art. 419).", "Art. 419 Código Penal / Ley 1581/12"),
    ("No conservar la HC por el tiempo mínimo legal", "Secretaría de Salud / Juzgado civil", "Multa administrativa. Pérdida de prueba en proceso de responsabilidad médica.", "Res. 1995/99 Art. 13"),
    ("Historia clínica ilegible", "Tribunal Ético / Juzgado civil", "Sanción disciplinaria. En proceso civil: el médico asume la carga de probar que no hubo falla por medios alternativos.", "Res. 1995/99 Art. 2"),
    ("No entrega de copia de la HC al paciente que la solicita", "Superintendencia de Salud / Juzgado civil", "Multa. Acción de tutela por vulneración de derechos fundamentales.", "Sentencia T-696/96 / Ley 1751/2015"),
]
for i, row_data in enumerate(sanc_data):
    add_table_row(t_sanc, list(row_data), shaded=(i % 2 == 1))
doc.add_paragraph()

# 12. BIBLIOGRAFÍA
h1(doc, "12. BIBLIOGRAFÍA")
refs_hc = [
    "Ministerio de Salud de Colombia. Resolución 1995 de 1999: Por la cual se establecen normas para el manejo de la Historia Clínica. Bogotá: Ministerio de Salud; 1999.",
    "Ministerio de Salud y Protección Social de Colombia. Resolución 3100 de 2019: Procedimientos y condiciones de habilitación de servicios de salud. Bogotá: MINSALUD; 2019.",
    "Congreso de la República de Colombia. Ley 23 de 1981: Código de Ética Médica. Bogotá: Congreso de la República; 1981.",
    "Congreso de la República de Colombia. Ley 1581 de 2012: Protección de datos personales. Bogotá: Congreso de la República; 2012.",
    "Congreso de la República de Colombia. Código Penal Colombiano (Ley 599 de 2000). Arts. 289, 419, 269F. Bogotá: Congreso de la República; 2000.",
    "Corte Constitucional de Colombia. Sentencia T-696 de 1996: Derecho del paciente a acceder a su historia clínica. Bogotá: Corte Constitucional; 1996.",
    "Corte Constitucional de Colombia. Sentencia T-218 de 2014: Acceso a la historia clínica como derecho fundamental. Bogotá: Corte Constitucional; 2014.",
    "Corte Constitucional de Colombia. Sentencia T-158 de 1994: Datos de salud como datos sensibles con protección constitucional. Bogotá: Corte Constitucional; 1994.",
    "Tamayo Lombana A. La Responsabilidad Civil por los Daños Causados con Actos Médicos. 5ª ed. Bogotá: Temis; 2020.",
    "Ministerio de Salud y Protección Social de Colombia. Lineamientos para la implementación de la historia clínica electrónica. Bogotá: MINSALUD; 2022.",
    "Organización Panamericana de la Salud (OPS). Historia Clínica como Documento Médico-Legal: Guía para Profesionales de la Salud. Washington D.C.: OPS; 2016.",
    "Tribunal Ético de Medicina de Colombia. Guía para el diligenciamiento de la historia clínica. Bogotá: Tribunal Ético; 2021.",
]
for i, ref in enumerate(refs_hc):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1); p.paragraph_format.first_line_indent = Cm(-1)
    run = p.add_run(f"{i+1}. {ref}"); run.font.size = Pt(10); run.font.name = 'Calibri'

doc.save(PATH)
print(f"DOC4 guardado: {PATH}")
