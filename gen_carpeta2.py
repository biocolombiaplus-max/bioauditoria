#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

BASE2 = "/home/user/bioauditoria/documentos_habilitacion/CARPETA_2_INFRAESTRUCTURA"

def margins(doc):
    for s in doc.sections:
        s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5)
        s.left_margin=Cm(3); s.right_margin=Cm(2.5)

def hf(doc, code, title):
    for s in doc.sections:
        h=s.header; h.paragraphs[0].clear()
        r=h.paragraphs[0].add_run(f"CONSULTORIO [NOMBRE DEL CONSULTORIO]  |  {code}  |  {title}")
        r.font.size=Pt(8); r.font.color.rgb=RGBColor(64,64,64)
        h.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        f=s.footer; f.paragraphs[0].clear()
        r2=f.paragraphs[0].add_run(f"{code} - Versión 1.0 - {datetime.date.today().strftime('%d/%m/%Y')}  | Pág. ")
        r2.font.size=Pt(8)
        fld=OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'),'begin')
        f.paragraphs[0].runs[-1]._r.append(fld)
        it=OxmlElement('w:instrText'); it.text='PAGE'
        f.paragraphs[0].runs[-1]._r.append(it)
        fld2=OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'),'end')
        f.paragraphs[0].runs[-1]._r.append(fld2)
        f.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

def portada(doc, code, title):
    doc.add_paragraph(); doc.add_paragraph()
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=t.add_run("CONSULTORIO MÉDICO\n[NOMBRE DEL CONSULTORIO]")
    r.font.size=Pt(18); r.font.bold=True; r.font.color.rgb=RGBColor(31,73,125)
    doc.add_paragraph()
    t2=doc.add_paragraph(); t2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=t2.add_run(title); r2.font.size=Pt(16); r2.font.bold=True
    doc.add_paragraph()
    tb=doc.add_table(rows=6,cols=2); tb.style='Table Grid'
    tb.alignment=WD_TABLE_ALIGNMENT.CENTER
    datos=[("Código:",code),("Versión:","1.0"),("Fecha:",datetime.date.today().strftime('%d/%m/%Y')),
           ("Elaboró:","[NOMBRE DE LA MÉDICA], Médica General"),
           ("Revisó:","[NOMBRE DE LA MÉDICA], Propietaria"),
           ("Aprobó:","[NOMBRE DE LA MÉDICA], Propietaria")]
    for i,(l,v) in enumerate(datos):
        tb.rows[i].cells[0].text=l; tb.rows[i].cells[1].text=v
        tb.rows[i].cells[0].paragraphs[0].runs[0].font.bold=True
    doc.add_page_break()
    cv=doc.add_paragraph(); r3=cv.add_run("CONTROL DE VERSIONES")
    r3.font.bold=True; r3.font.size=Pt(13); r3.font.color.rgb=RGBColor(31,73,125)
    cv.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tb2=doc.add_table(rows=3,cols=5); tb2.style='Table Grid'
    hrow(tb2,["Versión","Fecha","Cambio","Elaboró","Aprobó"])
    tb2.rows[1].cells[0].text="1.0"; tb2.rows[1].cells[1].text=datetime.date.today().strftime('%d/%m/%Y')
    tb2.rows[1].cells[2].text="Versión inicial"
    tb2.rows[1].cells[3].text="[NOMBRE DE LA MÉDICA]"
    tb2.rows[1].cells[4].text="[NOMBRE DE LA MÉDICA]"
    doc.add_page_break()

def hrow(table, headers, bg="1F497D"):
    row=table.rows[0]
    for i,h in enumerate(headers):
        cell=row.cells[i]; cell.text=""
        run=cell.paragraphs[0].add_run(h)
        run.font.bold=True; run.font.color.rgb=RGBColor(255,255,255); run.font.size=Pt(10)
        cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),bg)
        shd.set(qn('w:color'),'auto'); shd.set(qn('w:val'),'clear')
        tcPr.append(shd)

def H(doc,text,lv=1):
    p=doc.add_heading(text,level=lv)
    colors={1:RGBColor(31,73,125),2:RGBColor(46,116,181),3:RGBColor(68,114,196)}
    sizes={1:14,2:12,3:11}
    for run in p.runs:
        run.font.size=Pt(sizes[lv]); run.font.bold=True; run.font.color.rgb=colors[lv]
    return p

def B(doc,text,bold=False,indent=False):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.left_indent=Cm(1)
    r=p.add_run(text); r.font.size=Pt(11); r.font.bold=bold; r.font.name='Calibri'
    return p

def BL(doc,text):
    p=doc.add_paragraph(style='List Bullet'); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=p.add_run(text); r.font.size=Pt(11); r.font.name='Calibri'
    return p

def firmas(doc):
    doc.add_paragraph()
    tb=doc.add_table(rows=2,cols=3); tb.style='Table Grid'
    hrow(tb,["ELABORÓ","REVISÓ","APROBÓ"])
    for j in range(3):
        tb.rows[1].cells[j].text="[NOMBRE DE LA MÉDICA]\nMédica General - Propietaria\n\nFirma: _________________________\n\nFecha: "+datetime.date.today().strftime('%d/%m/%Y')
        tb.rows[1].cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

# ========== MAN-INF-001 ==========
def crear_man_inf():
    doc=Document(); margins(doc); hf(doc,"MAN-INF-001","Manual de Gestión de Infraestructura")
    portada(doc,"MAN-INF-001","MANUAL DE GESTIÓN DE INFRAESTRUCTURA")

    H(doc,"1. OBJETIVO")
    B(doc,"Establecer los lineamientos para la gestión, mantenimiento y mejora de la infraestructura física del Consultorio Médico [NOMBRE DEL CONSULTORIO], garantizando que las instalaciones cumplan con los estándares mínimos establecidos en la Resolución 3100 de 2019 del Ministerio de Salud y Protección Social, y brinden condiciones seguras, dignas y adecuadas para la atención de los pacientes y el trabajo del personal.")

    H(doc,"2. ALCANCE")
    B(doc,"Aplica a todas las áreas físicas del Consultorio [NOMBRE DEL CONSULTORIO] ubicado en [DIRECCIÓN DEL CONSULTORIO], [CIUDAD]: sala de espera, consultorio médico, área de procedimientos (si aplica), baño, área de esterilización/desinfección y demás espacios del establecimiento.")

    H(doc,"3. MARCO LEGAL")
    normas=[
        ("Resolución 3100 de 2019","Estándares mínimos de habilitación de infraestructura para servicios de salud, Anexo Técnico 1."),
        ("Decreto 351 de 2014","Gestión integral de los residuos generados en la atención en salud. Aplica para el diseño de áreas de almacenamiento de residuos."),
        ("Resolución 1164 de 2002","Manual de procedimientos para la gestión integral de residuos hospitalarios y similares (MPGIRH). Complemento al Decreto 351/2014."),
        ("NSR-10","Reglamento Colombiano de Construcción Sismo Resistente. Requisitos estructurales de las edificaciones."),
        ("NTC 4595","Norma Técnica Colombiana sobre planeamiento y diseño de instalaciones y ambientes escolares (referencia para espacios accesibles)."),
        ("Ley 361 de 1997","Mecanismos de integración social de personas con discapacidad. Requisitos de accesibilidad en establecimientos."),
        ("Decreto 1072 de 2015","Decreto Único del Sector Trabajo. Condiciones de seguridad en el lugar de trabajo."),
        ("Resolución 0312 de 2019","Estándares mínimos del Sistema de Gestión de Seguridad y Salud en el Trabajo."),
        ("NTC-OHSAS 18001","Sistemas de gestión en seguridad y salud ocupacional (referencia)."),
    ]
    tb=doc.add_table(rows=len(normas)+1,cols=2); tb.style='Table Grid'
    hrow(tb,["Norma","Descripción y aplicabilidad"])
    for i,(n,d) in enumerate(normas):
        tb.rows[i+1].cells[0].text=n; tb.rows[i+1].cells[1].text=d
        for c in tb.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10)
    doc.add_paragraph()

    H(doc,"4. DEFINICIONES")
    defs=[("Infraestructura:","Conjunto de instalaciones físicas, espacios, servicios básicos y condiciones ambientales del establecimiento de salud."),
          ("Área asistencial:","Espacio físico donde se prestan directamente los servicios de salud al paciente."),
          ("Área de apoyo:","Espacios que complementan la atención: sala de espera, baños, áreas de almacenamiento."),
          ("Mantenimiento preventivo:","Actividades programadas para prevenir el deterioro o falla de las instalaciones y equipos."),
          ("Mantenimiento correctivo:","Actividades para restablecer condiciones normales de funcionamiento después de una falla."),
          ("Accesibilidad:","Condiciones que permiten a las personas con discapacidad acceder, circular y usar el establecimiento sin barreras."),
          ("Residuos biosanitarios:","Residuos generados en la atención en salud con potencial de riesgo biológico (Decreto 351/2014)."),]
    for term,defi in defs:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r1=p.add_run(term+" "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(defi); r2.font.size=Pt(11)

    H(doc,"5. ESTÁNDARES FÍSICOS SEGÚN RESOLUCIÓN 3100 DE 2019")
    H(doc,"5.1 Requisitos Generales de la Infraestructura",2)
    B(doc,"La Resolución 3100 de 2019, Anexo Técnico 1, establece los siguientes estándares de infraestructura para el servicio de Consulta de Medicina General:")

    estandares=[
        ("Área mínima del consultorio","12 m² - El consultorio médico debe tener un área mínima de 12 metros cuadrados de área útil, suficiente para contener el escritorio médico, la camilla de examen, área de lavado de manos y espacio para desplazamiento del médico y el paciente."),
        ("Lavamanos","Debe existir lavamanos de uso exclusivo para el consultorio, con agua corriente, jabón líquido y sistema de secado de manos (toallas desechables o secador). El lavamanos debe ser de accionamiento NO manual (codo, rodilla, pie o sensor) para el área de procedimientos."),
        ("Camilla de examen","Camilla o diván de examen con cobertura lavable o con papel desechable, posicionable según necesidad clínica."),
        ("Privacidad visual y auditiva","El consultorio debe garantizar privacidad visual (paredes o biombos) y relativa privacidad auditiva durante la consulta."),
        ("Iluminación","Iluminación natural y/o artificial suficiente para el examen clínico. Mínimo 500 lux en el área de examen físico."),
        ("Ventilación","Ventilación natural o artificial que garantice renovación de aire y confort térmico. Se recomienda ventilación cruzada o sistema de aire acondicionado con filtros adecuados."),
        ("Señalización","El consultorio debe contar con señalización de seguridad, evacuación, riesgo biológico (donde aplique) y accesibilidad según normas colombianas."),
        ("Baño para pacientes","Debe existir baño accesible para pacientes, con lavamanos, inodoro y condiciones de higiene adecuadas. Para atención de personas con discapacidad motora, debe ser accesible según Ley 361/1997."),
        ("Sala de espera","Área de espera cómoda, bien iluminada y ventilada, con sillas suficientes para la demanda de pacientes. Debe incluir señalización y derechos de los pacientes visibles."),
        ("Almacenamiento de residuos","Área o recipientes para la disposición transitoria de residuos ordinarios y biosanitarios, debidamente separados y señalizados según el PGIRH."),
    ]
    for est,desc in estandares:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent=Cm(0.5)
        r1=p.add_run("• "+est+": "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(desc); r2.font.size=Pt(11)

    H(doc,"5.2 Verificación de Cumplimiento de la Infraestructura Actual",2)
    B(doc,"Se realiza la siguiente verificación del estado actual de la infraestructura del consultorio:")
    tb2=doc.add_table(rows=12,cols=4); tb2.style='Table Grid'
    hrow(tb2,["Requisito","¿Cumple?","Observaciones","Acción requerida"])
    reqs=["Área mínima consultorio ≥12 m²","Lavamanos en consultorio","Camilla de examen","Privacidad visual","Iluminación adecuada ≥500 lux",
          "Ventilación adecuada","Señalización de seguridad","Baño para pacientes","Sala de espera","Almacenamiento de residuos",
          "Accesibilidad para discapacitados"]
    for i,req in enumerate(reqs):
        tb2.rows[i+1].cells[0].text=req
        tb2.rows[i+1].cells[1].text="□ Sí  □ No  □ Parcial"
        for c in tb2.rows[i+1].cells:
            for p2 in c.paragraphs:
                for r in p2.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"6. CONDICIONES ESPECÍFICAS PARA PROCEDIMIENTOS ESTÉTICOS NO INVASIVOS")
    B(doc,"Para la realización de procedimientos estéticos no invasivos (toxina botulínica, rellenos dérmicos, peelings, mesoterapia, PRP), el área debe cumplir con condiciones adicionales de asepsia y equipamiento:")
    cond_est=[
        "Iluminación de alta intensidad y ajustable (luz fría preferiblemente, mínimo 1000 lux en el campo de trabajo)",
        "Lavamanos de accionamiento no manual o sensor en el área donde se realizan los procedimientos",
        "Superficie de trabajo de material liso, no poroso y fácilmente desinfectable (acero inoxidable o similar)",
        "Sistema de almacenamiento bajo llave y temperatura controlada para toxina botulínica (2-8°C) y otros biológicos",
        "Kit de emergencias disponible y accesible en el área de procedimientos",
        "Recipientes para residuos cortopunzantes (guardianes) disponibles y seguros",
        "Disponibilidad de oxígeno medicinal y elementos de reanimación básica (ambú, cánulas de Guedel)",
        "Cámara para fotografías clínicas o área con fondo neutro y buena iluminación",
        "Privacidad total durante los procedimientos estéticos (la paciente/paciente debe sentirse cómodo/a)",
    ]
    for item in cond_est: BL(doc,item)

    H(doc,"7. PROGRAMA DE MANTENIMIENTO DE INFRAESTRUCTURA")
    H(doc,"7.1 Mantenimiento Preventivo Programado",2)
    B(doc,"Se establece el siguiente cronograma de mantenimiento preventivo de la infraestructura del consultorio:")
    tb3=doc.add_table(rows=13,cols=4); tb3.style='Table Grid'
    hrow(tb3,["Área/Elemento","Actividad de mantenimiento","Frecuencia","Responsable"])
    mant=[
        ("Paredes y cielos rasos","Inspección de humedad, grietas, pintura. Repintar si necesario","Semestral","Contratista externo"),
        ("Pisos","Verificar estado, sellado de juntas, limpieza profunda","Trimestral","Servicios generales"),
        ("Puertas y ventanas","Revisión de bisagras, cerraduras, empaques, vidrios","Trimestral","Contratista externo"),
        ("Red eléctrica","Verificación de tomacorrientes, interruptores, breakers, tierra física","Anual","Electricista certificado"),
        ("Red hidráulica","Revisión de llaves, sifones, sanitarios, duchas (si aplica)","Semestral","Plomero certificado"),
        ("Aire acondicionado/ventilación","Limpieza de filtros, revisión del compresor, carga de gas","Semestral","Técnico certificado"),
        ("Lavamanos","Verificar sello, grifo, desagüe, dispensador de jabón y papel","Mensual","Servicios generales / médica"),
        ("Iluminación","Cambio de bombillas dañadas, limpieza de luminarias","Mensual","Servicios generales"),
        ("Señalización","Verificar estado y vigencia de señales, reemplazar deterioradas","Semestral","Médica propietaria"),
        ("Área de residuos","Limpieza y desinfección del área de almacenamiento transitorio","Mensual","Servicios generales"),
        ("Extintores","Revisión de carga, vencimiento, accesibilidad","Anual","Empresa certificada"),
        ("Camilla de examen","Verificar estado, limpieza profunda, cambio de recubrimiento si necesario","Mensual","Médica / auxiliar"),
    ]
    for i,(area,act,frec,resp) in enumerate(mant):
        tb3.rows[i+1].cells[0].text=area; tb3.rows[i+1].cells[1].text=act
        tb3.rows[i+1].cells[2].text=frec; tb3.rows[i+1].cells[3].text=resp
        for c in tb3.rows[i+1].cells:
            for p2 in c.paragraphs:
                for r in p2.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"8. GESTIÓN DE RESIDUOS")
    B(doc,"En cumplimiento del Decreto 351 de 2014 y la Resolución 1164 de 2002 (Manual PGIRH), el consultorio implementa el siguiente sistema de gestión de residuos:")
    H(doc,"8.1 Clasificación de Residuos",2)
    residuos=[
        ("Residuos no peligrosos ordinarios","Bolsa NEGRA","Papeles administrativos, envases plásticos limpios, residuos de alimentos"),
        ("Residuos biosanitarios (con riesgo biológico)","Bolsa ROJA","Gasas, apósitos, guantes, material con sangre o fluidos corporales"),
        ("Residuos anatomopatológicos","Bolsa ROJA (especial)","Tejidos, líquidos corporales (si aplica en PRP: sangre residual)"),
        ("Residuos cortopunzantes","Guardián ROJO (contenedor rígido)","Agujas, bisturís, lancetas, ampollas rotas"),
        ("Residuos de medicamentos","Bolsa ROJA o contenedor específico","Frascos con residuos de medicamentos, toxina botulínica vencida"),
        ("Residuos de aparatos eléctricos y electrónicos (RAEE)","Recolección especial","Equipos médicos en desuso, pilas, cables"),
    ]
    tb4=doc.add_table(rows=len(residuos)+1,cols=3); tb4.style='Table Grid'
    hrow(tb4,["Tipo de residuo","Identificación","Ejemplos en el consultorio"])
    for i,(tipo,color,ej) in enumerate(residuos):
        tb4.rows[i+1].cells[0].text=tipo; tb4.rows[i+1].cells[1].text=color; tb4.rows[i+1].cells[2].text=ej
        for c in tb4.rows[i+1].cells:
            for p2 in c.paragraphs:
                for r in p2.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"8.2 Ruta Sanitaria de Residuos",2)
    B(doc,"Los residuos biosanitarios y cortopunzantes deben ser recogidos por una empresa gestora de residuos peligrosos con licencia ambiental vigente. El contrato con la empresa gestora debe estar disponible en la carpeta de habilitación del consultorio.")
    B(doc,"Empresa gestora contratada: [NOMBRE DE LA EMPRESA GESTORA DE RESIDUOS]\nNúmero de contrato: [NÚMERO DE CONTRATO]\nFrecuencia de recolección: [FRECUENCIA]\nLicencia ambiental: [NÚMERO DE LICENCIA]",bold=False)

    H(doc,"9. CONTROL DE PLAGAS Y ROEDORES")
    B(doc,"El consultorio debe contar con un programa de control de vectores y plagas certificado por una empresa autorizada por las autoridades ambientales. Los registros deben ser archivados y estar disponibles para inspección sanitaria.")
    plagas_items=[
        "Contrato vigente con empresa certificada para control de plagas: [NOMBRE EMPRESA] - Contrato N° [N°]",
        "Fumigación y control preventivo con una frecuencia mínima de cada 6 meses",
        "Inspección visual mensual por parte del personal del consultorio",
        "Sellado de grietas, fisuras y posibles vías de ingreso de roedores e insectos",
        "Mantenimiento adecuado de desagües y sifones",
        "Almacenamiento correcto de alimentos (si los hay en el consultorio)",
        "Registro de cada intervención de control de plagas en el formato de inspección de infraestructura (FOR-INF-001)",
    ]
    for item in plagas_items: BL(doc,item)

    H(doc,"10. SEÑALIZACIÓN DEL CONSULTORIO")
    B(doc,"El consultorio debe contar con señalización clara, visible y conforme a las normas colombianas (NTC 1700, NTC 4166):")
    senalizacion=[
        ("Señales de evacuación y emergencia","Verde con blanco","Salidas, rutas de evacuación, punto de encuentro"),
        ("Señales de prohibición","Rojo con blanco","No fumar, no ingerir alimentos en áreas clínicas"),
        ("Señales de advertencia / precaución","Amarillo","Riesgo biológico, piso mojado, alto voltaje"),
        ("Señales de información","Azul con blanco","Consultorio médico, baño, sala de espera, recepción"),
        ("Señales de obligación","Azul","Uso de EPP obligatorio en áreas específicas"),
        ("Derechos de los pacientes","Visible en sala de espera","Resolución 13437/1991"),
    ]
    tb5=doc.add_table(rows=len(senalizacion)+1,cols=3); tb5.style='Table Grid'
    hrow(tb5,["Tipo de señal","Características","Dónde se ubica"])
    for i,(tipo,carac,donde) in enumerate(senalizacion):
        tb5.rows[i+1].cells[0].text=tipo; tb5.rows[i+1].cells[1].text=carac; tb5.rows[i+1].cells[2].text=donde
        for c in tb5.rows[i+1].cells:
            for p2 in c.paragraphs:
                for r in p2.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"11. INDICADORES DE INFRAESTRUCTURA")
    tb6=doc.add_table(rows=5,cols=4); tb6.style='Table Grid'
    hrow(tb6,["Indicador","Fórmula","Meta","Frecuencia"])
    inds=[("% cumplimiento estándares infraestructura","(Requisitos cumplidos/Requisitos totales)×100","100%","Semestral"),
          ("N° mantenimientos preventivos realizados vs programados","Realizados/Programados×100","≥90%","Anual"),
          ("% contratos de servicios de apoyo vigentes","(Contratos vigentes/Contratos requeridos)×100","100%","Trimestral"),
          ("N° no conformidades de infraestructura identificadas en inspección","Conteo absoluto","0","Semestral")]
    for i,row in enumerate(inds):
        for j,v in enumerate(row): tb6.rows[i+1].cells[j].text=v
        for c in tb6.rows[i+1].cells:
            for p2 in c.paragraphs:
                for r in p2.runs: r.font.size=Pt(9)
    doc.add_paragraph()
    firmas(doc)
    path=os.path.join(BASE2,"MAN-INF-001_Manual_Gestion_Infraestructura.docx")
    doc.save(path); print(f"✓ Creado: {path}")

# ========== PRO-INF-001 ==========
def crear_pro_inf():
    doc=Document(); margins(doc); hf(doc,"PRO-INF-001","Programa de Mantenimiento de Infraestructura")
    portada(doc,"PRO-INF-001","PROGRAMA DE MANTENIMIENTO DE INFRAESTRUCTURA")

    H(doc,"1. OBJETIVO")
    B(doc,"Definir el programa sistemático de mantenimiento preventivo y correctivo de la infraestructura física del Consultorio [NOMBRE DEL CONSULTORIO], garantizando que las instalaciones se mantengan en condiciones óptimas de funcionamiento, seguridad, higiene y cumplimiento de los estándares de habilitación establecidos por la Resolución 3100 de 2019.")

    H(doc,"2. ALCANCE")
    B(doc,"Aplica a todas las instalaciones, redes, acabados y elementos físicos del consultorio ubicado en [DIRECCIÓN DEL CONSULTORIO], [CIUDAD], incluyendo: red eléctrica, red hidráulica, sistema de ventilación/climatización, pisos, paredes, cielos rasos, puertas, ventanas, baños, área de espera y área clínica.")

    H(doc,"3. MARCO LEGAL")
    for n in ["Resolución 3100 de 2019 - Estándares de infraestructura para habilitación",
              "NSR-10 - Reglamento de Construcción Sismo Resistente",
              "Decreto 1072 de 2015 - Seguridad en el trabajo",
              "Resolución 0312 de 2019 - Estándares mínimos SST",
              "RETIE - Reglamento Técnico de Instalaciones Eléctricas",
              "NTC 2050 - Código Eléctrico Colombiano"]: BL(doc,n)

    H(doc,"4. DEFINICIONES")
    defs2=[("Mantenimiento preventivo:","Conjunto de actividades programadas orientadas a prevenir fallas y deterioro de la infraestructura, realizadas antes de que ocurra una falla."),
           ("Mantenimiento correctivo:","Actividades realizadas para corregir una falla o deterioro ya ocurrido, restableciendo las condiciones normales de operación."),
           ("Orden de trabajo:","Documento que registra la solicitud, descripción, ejecución y cierre de una actividad de mantenimiento."),
           ("No conformidad de infraestructura:","Condición de la infraestructura que no cumple con los estándares mínimos de habilitación o seguridad establecidos."),]
    for term,defi in defs2:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r1=p.add_run(term+" "); r1.font.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(defi); r2.font.size=Pt(11)

    H(doc,"5. CRONOGRAMA ANUAL DE MANTENIMIENTO PREVENTIVO")
    B(doc,"El siguiente cronograma establece las actividades de mantenimiento preventivo para cada mes del año:")
    meses=["Ene","Feb","Mar","Abr","May","Jun","Jul","Ago","Sep","Oct","Nov","Dic"]
    actividades_mant=[
        ("Limpieza filtros A/A y ventilación","X","","","","","X","","","","","","X"),
        ("Inspección red eléctrica","X","","","","","","","","","","","X"),
        ("Revisión red hidráulica","","","X","","","","","","X","","",""),
        ("Inspección pisos y paredes","","X","","","","X","","","","X","",""),
        ("Revisión puertas y ventanas","","","","X","","","","","","","X",""),
        ("Revisión y recarga extintores","X","","","","","","","","","","","X"),
        ("Control de plagas y vectores","","","X","","","X","","","X","","","X"),
        ("Verificación señalización","","","","","","X","","","","","","X"),
        ("Mantenimiento camilla examen","X","X","X","X","X","X","X","X","X","X","X","X"),
        ("Inspección área residuos","X","X","X","X","X","X","X","X","X","X","X","X"),
        ("Pintura y acabados (si necesario)","","","","","","","X","","","","",""),
        ("Revisión PGIRH y contrato empresa gestora","","","X","","","","","","X","","",""),
    ]
    cols=["Actividad"]+meses
    tb=doc.add_table(rows=len(actividades_mant)+1,cols=13); tb.style='Table Grid'
    hrow(tb,cols)
    for i,row in enumerate(actividades_mant):
        for j,v in enumerate(row):
            tb.rows[i+1].cells[j].text=v
            for p2 in tb.rows[i+1].cells[j].paragraphs:
                for r in p2.runs: r.font.size=Pt(8)
    doc.add_paragraph()

    H(doc,"6. PROCEDIMIENTO PARA MANTENIMIENTO CORRECTIVO")
    pasos=[("Paso 1","Identificación de la falla","El personal del consultorio identifica una condición de la infraestructura que no está funcionando correctamente o que representa un riesgo. La reporta a la médica propietaria verbalmente o por escrito."),
           ("Paso 2","Evaluación de la urgencia","La médica propietaria evalúa si la falla representa: a) Riesgo inmediato para pacientes o personal (acción inmediata), b) Impacto en la operación del consultorio (acción en 24-48h), c) Deterioro cosmético o funcional menor (programar en mantenimiento preventivo)."),
           ("Paso 3","Apertura de orden de trabajo","Se diligencia el formato de orden de trabajo con la descripción de la falla, fecha, área afectada y prioridad asignada."),
           ("Paso 4","Contacto con proveedor","Se contacta al proveedor de mantenimiento calificado para el tipo de falla: electricista, plomero, técnico de aires, contratista de construcción."),
           ("Paso 5","Ejecución del mantenimiento","El proveedor realiza la reparación. El personal del consultorio verifica que el trabajo se realice correctamente."),
           ("Paso 6","Verificación y cierre","Se verifica que la falla fue corregida y que la infraestructura cumple nuevamente con los estándares requeridos. Se cierra la orden de trabajo con la fecha de solución."),
           ("Paso 7","Registro","Se registra la intervención en el formato FOR-INF-001 y se archiva el comprobante del servicio realizado."),]
    tb2=doc.add_table(rows=len(pasos)+1,cols=3); tb2.style='Table Grid'
    hrow(tb2,["Paso","Actividad","Descripción"])
    for i,(paso,act,desc) in enumerate(pasos):
        tb2.rows[i+1].cells[0].text=paso; tb2.rows[i+1].cells[1].text=act; tb2.rows[i+1].cells[2].text=desc
        for c in tb2.rows[i+1].cells:
            for p2 in c.paragraphs:
                for r in p2.runs: r.font.size=Pt(9)
    doc.add_paragraph()

    H(doc,"7. PROVEEDORES DE MANTENIMIENTO")
    B(doc,"El consultorio mantiene los siguientes proveedores calificados para servicios de mantenimiento:")
    tb3=doc.add_table(rows=7,cols=4); tb3.style='Table Grid'
    hrow(tb3,["Tipo de mantenimiento","Proveedor","Teléfono de contacto","Observaciones"])
    provs=[("Eléctrico","[NOMBRE ELECTRICISTA]","[TELÉFONO]","Certificado RETIE vigente"),
           ("Hidráulico/Plomería","[NOMBRE PLOMERO]","[TELÉFONO]",""),
           ("Aire acondicionado","[NOMBRE TÉCNICO A/A]","[TELÉFONO]",""),
           ("Control de plagas","[NOMBRE EMPRESA]","[TELÉFONO]","Licencia ambiental vigente"),
           ("Gestión de residuos peligrosos","[NOMBRE EMPRESA GESTORA]","[TELÉFONO]","Contrato vigente, licencia ambiental"),
           ("Mantenimiento general (pintura, mampostería)","[NOMBRE CONTRATISTA]","[TELÉFONO]",""),]
    for i,(tipo,prov,tel,obs) in enumerate(provs):
        tb3.rows[i+1].cells[0].text=tipo; tb3.rows[i+1].cells[1].text=prov
        tb3.rows[i+1].cells[2].text=tel; tb3.rows[i+1].cells[3].text=obs
        for c in tb3.rows[i+1].cells:
            for p2 in c.paragraphs:
                for r in p2.runs: r.font.size=Pt(10)
    doc.add_paragraph()

    H(doc,"8. INDICADORES")
    tb4=doc.add_table(rows=4,cols=4); tb4.style='Table Grid'
    hrow(tb4,["Indicador","Fórmula","Meta","Frecuencia"])
    for i,row in enumerate([
        ("% actividades mantenimiento preventivo ejecutadas","Realizadas/Programadas×100","≥90%","Anual"),
        ("Tiempo promedio resolución mantenimiento correctivo","Suma días resolución/N° correctivos","≤5 días hábiles","Mensual"),
        ("N° no conformidades infraestructura en visitas de habilitación","Conteo","0","En cada visita")]):
        for j,v in enumerate(row): tb4.rows[i+1].cells[j].text=v
        for c in tb4.rows[i+1].cells:
            for p2 in c.paragraphs:
                for r in p2.runs: r.font.size=Pt(9)
    doc.add_paragraph(); firmas(doc)
    path=os.path.join(BASE2,"PRO-INF-001_Programa_Mantenimiento_Infraestructura.docx")
    doc.save(path); print(f"✓ Creado: {path}")

# ========== FOR-INF-001 ==========
def crear_for_inf():
    doc=Document(); margins(doc); hf(doc,"FOR-INF-001","Formato Inspección de Infraestructura")
    portada(doc,"FOR-INF-001","FORMATO DE INSPECCIÓN DE INFRAESTRUCTURA")

    H(doc,"DATOS DE LA INSPECCIÓN")
    tb=doc.add_table(rows=4,cols=4); tb.style='Table Grid'
    datos=[("Fecha de inspección:","","Inspector(a):",""),
           ("Tipo de inspección:","□ Programada  □ No programada  □ Post-mantenimiento","Próxima inspección:",""),
           ("Área inspeccionada:","□ Todo el consultorio  □ Área específica: _______________","Versión del formato:","1.0"),
           ("Resultado general:","□ Conforme  □ No conforme  □ Con observaciones","N° de inspección:","")]
    for i,row in enumerate(datos):
        for j,v in enumerate(row):
            tb.rows[i].cells[j].text=v
            if j%2==0:
                for run in tb.rows[i].cells[j].paragraphs[0].runs: run.font.bold=True
    doc.add_paragraph()

    areas=[
        ("ÁREA DE RECEPCIÓN / SALA DE ESPERA",[
            ("Pisos: estado, limpieza, sin daños","","",""),
            ("Paredes: sin humedad, pintura en buen estado","","",""),
            ("Iluminación adecuada y funcionando","","",""),
            ("Ventilación adecuada","","",""),
            ("Sillas suficientes y en buen estado","","",""),
            ("Señalización visible (derechos de pacientes, evacuación)","","",""),
            ("Temperatura ambiental confortable","","",""),
        ]),
        ("CONSULTORIO MÉDICO",[
            ("Área ≥12 m² libre y despejada","","",""),
            ("Lavamanos funcionando correctamente con jabón y papel","","",""),
            ("Camilla de examen en buen estado, con cobertura limpia","","",""),
            ("Privacidad visual garantizada","","",""),
            ("Iluminación adecuada para examen clínico","","",""),
            ("Escritorio médico ordenado y funcional","","",""),
            ("Equipo médico disponible y en buen estado","","",""),
            ("Recipientes de residuos con bolsa del color correcto","","",""),
            ("Guardián para cortopunzantes disponible y sin sobrepasar ¾ de capacidad","","",""),
        ]),
        ("ÁREA DE PROCEDIMIENTOS ESTÉTICOS (si aplica)",[
            ("Superficie de trabajo de material liso y desinfectable","","",""),
            ("Lavamanos de accionamiento no manual funcionando","","",""),
            ("Iluminación de alta intensidad disponible","","",""),
            ("Refrigerador para toxina botulínica funcionando (2-8°C)","","",""),
            ("Kit de emergencias completo y accesible","","",""),
            ("Guardián disponible para cortopunzantes","","",""),
            ("Material estéril disponible y vigente","","",""),
        ]),
        ("BAÑO(S)",[
            ("Limpio, sin malos olores","","",""),
            ("Agua funcionando (lavamanos, sanitario)","","",""),
            ("Jabón y papel disponibles","","",""),
            ("Sin humedad excesiva ni goteras","","",""),
            ("Accesible para personas con discapacidad (si aplica)","","",""),
        ]),
        ("INSTALACIONES GENERALES",[
            ("Red eléctrica sin daños visibles (cables, tomacorrientes)","","",""),
            ("Extintores vigentes, accesibles y señalizados","","",""),
            ("Rutas de evacuación despejadas y señalizadas","","",""),
            ("Área de residuos ordenada y con contenedores limpios","","",""),
            ("Sin presencia de plagas o roedores","","",""),
            ("Señalización de riesgo biológico donde aplica","","",""),
        ]),
    ]

    for area_nombre, items in areas:
        H(doc,area_nombre,2)
        tb2=doc.add_table(rows=len(items)+1,cols=4); tb2.style='Table Grid'
        hrow(tb2,["Ítem a verificar","Conforme","No conforme","Observaciones / Acción"])
        for i,(item,c,nc,obs) in enumerate(items):
            tb2.rows[i+1].cells[0].text=item
            tb2.rows[i+1].cells[1].text="□"
            tb2.rows[i+1].cells[2].text="□"
            tb2.rows[i+1].cells[3].text=""
            for c2 in tb2.rows[i+1].cells:
                for p2 in c2.paragraphs:
                    for r in p2.runs: r.font.size=Pt(9)
        doc.add_paragraph()

    H(doc,"HALLAZGOS Y ACCIONES CORRECTIVAS")
    tb3=doc.add_table(rows=5,cols=4); tb3.style='Table Grid'
    hrow(tb3,["N°","Hallazgo / No conformidad","Acción correctiva propuesta","Fecha comprometida"])
    for i in range(1,5):
        tb3.rows[i].cells[0].text=str(i)
    doc.add_paragraph()

    H(doc,"FIRMAS")
    tb4=doc.add_table(rows=2,cols=2); tb4.style='Table Grid'
    hrow(tb4,["INSPECTORA","VISTO BUENO MÉDICA PROPIETARIA"])
    tb4.rows[1].cells[0].text="Nombre: _________________________\nFirma: _________________________\nFecha: _________________________"
    tb4.rows[1].cells[1].text="Nombre: [NOMBRE DE LA MÉDICA]\nFirma: _________________________\nFecha: _________________________"

    path=os.path.join(BASE2,"FOR-INF-001_Formato_Inspeccion_Infraestructura.docx")
    doc.save(path); print(f"✓ Creado: {path}")

if __name__=="__main__":
    print("Generando CARPETA 2 - INFRAESTRUCTURA...")
    crear_man_inf(); crear_pro_inf(); crear_for_inf()
    print("✅ Carpeta 2 completada.")
